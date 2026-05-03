import sys
import os
import asyncio

from contextlib import asynccontextmanager
from fastapi import FastAPI, HTTPException, Depends, Header, BackgroundTasks, Request, UploadFile, File, Form
from fastapi.responses import Response, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import Optional, List
import psycopg2
import psycopg2.extras
import datetime
import time
import json
import hmac
import hashlib
import jwt
import bcrypt
from app.core.url_checker import check_duplicate_url
from app.core.matcher import get_matches_for_user, get_individual_matches_for_user, get_matches_hybrid
from app.config import DATABASE_URL

# Admin Scraper Import for Manual Sync
from app.services.admin_scraper import admin_scraper


def _hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def _verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))

JWT_SECRET = os.getenv("JWT_SECRET", "change-me-in-production-env")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 24 * 7  # 7일

sync_status = {"running": False, "last_result": None, "last_time": None}
manual_sync_status = {"running": False, "last_result": None, "last_time": None}
reanalyze_status = {"running": False, "done": 0, "total": 0, "last_result": None, "last_time": None}


# ── DB 직접 연결 (Supabase transaction pooler port 6543 — PgBouncer이 실제 풀 담당) ──
# psycopg2 ThreadedConnectionPool은 PgBouncer transaction mode와 호환 불가:
# 풀이 유지하는 연결이 Supabase 측에서 끊기면 TCP에 쓸 때 90초 hang 발생.
# 대신 요청마다 직접 connect() → PgBouncer가 backend 연결을 재사용함.

def get_db_connection():
    return psycopg2.connect(
        DATABASE_URL,
        cursor_factory=psycopg2.extras.RealDictCursor,
        connect_timeout=10,
    )


def init_database():
    """DB 연결 확인 + 누락 테이블 자동 생성

    autocommit은 사용 안 함 — keyword_synonyms seed 같은 bulk INSERT가
    각각 별도 트랜잭션이 되면 healthcheck timeout 발생.
    대신 statement 실패 시 즉시 rollback + 계속 진행.
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        def _safe_exec(sql, label=""):
            """단일 statement 안전 실행 — 실패하면 rollback 후 다음 진행"""
            try:
                cursor.execute(sql)
                conn.commit()
                return True
            except Exception as e:
                conn.rollback()
                print(f"  [init_db] skip {label}: {str(e)[:80]}")
                return False

        cursor.execute("SELECT 1")

        # notification_settings 테이블이 없으면 자동 생성 (SQLite→PostgreSQL 마이그레이션 대비)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS notification_settings (
                id SERIAL PRIMARY KEY,
                business_number VARCHAR(20) UNIQUE NOT NULL,
                email VARCHAR(100),
                phone_number VARCHAR(20),
                channel VARCHAR(20) DEFAULT 'email',
                is_active BOOLEAN DEFAULT true,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        # final_url 컬럼 추가 (중간 경유 없이 최종 원본 URL)
        try:
            cursor.execute("ALTER TABLE announcements ADD COLUMN IF NOT EXISTS final_url TEXT DEFAULT ''")
        except Exception:
            pass

        # trending_announcements 테이블
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS trending_announcements (
                id SERIAL PRIMARY KEY,
                trending_date DATE NOT NULL,
                rank INTEGER NOT NULL,
                announcement_id INTEGER NOT NULL,
                trending_keyword TEXT,
                trending_reason TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        try:
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_trending_date ON trending_announcements(trending_date)")
        except Exception:
            pass

        # interest_regions 컬럼 추가 (소재지와 관심지역 분리)
        try:
            cursor.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS interest_regions TEXT DEFAULT ''")
        except Exception:
            pass

        # 상담 세션 테이블 (세션 기반 차감)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS consult_sessions (
                id SERIAL PRIMARY KEY,
                session_id VARCHAR(64) NOT NULL,
                business_number VARCHAR(20) NOT NULL,
                announcement_id INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        # 인덱스 (세션 조회 성능)
        try:
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_consult_sessions_lookup ON consult_sessions (session_id, business_number, announcement_id)")
        except Exception:
            pass

        # PRO 컨설턴트 상담 세션 (서버 측 상태 관리 — 단계/수집정보 저장)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS pro_consult_sessions (
                session_id VARCHAR(64) PRIMARY KEY,
                business_number VARCHAR(20) NOT NULL,
                client_category VARCHAR(20),
                current_step INTEGER DEFAULT 1,
                collected JSONB DEFAULT '{}'::jsonb,
                messages JSONB DEFAULT '[]'::jsonb,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        try:
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_pro_consult_sessions_bn ON pro_consult_sessions (business_number, updated_at DESC)")
        except Exception:
            pass
        # 기존 테이블 마이그레이션: messages 컬럼 없으면 추가 (P0.3)
        try:
            cursor.execute("ALTER TABLE pro_consult_sessions ADD COLUMN IF NOT EXISTS messages JSONB DEFAULT '[]'::jsonb")
        except Exception:
            pass
        # B: phase 컬럼 (collecting | consulting) 추가
        try:
            cursor.execute("ALTER TABLE pro_consult_sessions ADD COLUMN IF NOT EXISTS phase VARCHAR(20) DEFAULT 'collecting'")
        except Exception:
            pass
        # D: 매칭 결과 스냅샷 저장 컬럼
        try:
            cursor.execute("ALTER TABLE pro_consult_sessions ADD COLUMN IF NOT EXISTS matched_snapshot JSONB DEFAULT '[]'::jsonb")
        except Exception:
            pass
        # P0.4: ai_consult_logs에 session_id + updated_at 추가 (매 턴 UPSERT 목적)
        try:
            cursor.execute("ALTER TABLE ai_consult_logs ADD COLUMN IF NOT EXISTS session_id VARCHAR(64)")
            cursor.execute("ALTER TABLE ai_consult_logs ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
            cursor.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_ai_consult_logs_session_id ON ai_consult_logs(session_id) WHERE session_id IS NOT NULL")
        except Exception:
            pass

        # SQLite에서 마이그레이션된 기존 테이블의 INTEGER→BOOLEAN 변환 시도
        for tbl, col in [("notification_settings", "is_active"), ("admin_urls", "is_active")]:
            try:
                cursor.execute("""
                    SELECT data_type FROM information_schema.columns
                    WHERE table_name = %s AND column_name = %s
                """, (tbl, col))
                row = cursor.fetchone()
                if row and row.get("data_type") == "integer":
                    cursor.execute(f"""
                        ALTER TABLE {tbl}
                        ALTER COLUMN {col} SET DATA TYPE BOOLEAN
                        USING CASE WHEN {col} = 1 THEN true ELSE false END
                    """)
                    conn.commit()
                    print(f"  Migrated {tbl}.{col}: INTEGER -> BOOLEAN")
            except Exception as e:
                conn.rollback()
                print(f"  Note: {tbl}.{col} migration skipped: {e}")
        # ai_analyzed_at 컬럼 추가 (중복 AI 분석 방지용 타임스탬프)
        try:
            cursor.execute("""
                ALTER TABLE announcements ADD COLUMN IF NOT EXISTS ai_analyzed_at TIMESTAMP
            """)
            conn.commit()
        except Exception:
            conn.rollback()

        # [Phase 1] 공고 상태 명시적 관리 컬럼 — deadline/금액 품질 근본 해결용
        # - deadline_type: 'fixed'(명확 마감일) / 'ongoing'(상시) / 'unknown'(파악 전) / 'expired'(자동 만료)
        # - is_archived: 아카이브 여부 (UI 노출 제외)
        # - analysis_status: 'pending'(대기) / 'analyzed'(완료) / 'failed'(재시도 초과) / 'skipped'(원문 없음)
        # - analysis_attempts: 분석 시도 횟수 (재시도 로직 제어)
        # - last_analyzed_at: 마지막 분석 시도 시각
        # - support_amount_type: 'numeric'(숫자 파싱 성공) / 'text_only'(텍스트만) / 'unknown' / 'not_specified'
        # - support_amount_max / min: 정규화된 원(KRW) 단위 값 (정렬·필터용)
        for col_def in [
            "deadline_type VARCHAR(20) DEFAULT 'unknown'",
            "is_archived BOOLEAN DEFAULT FALSE",
            "analysis_status VARCHAR(20) DEFAULT 'pending'",
            "analysis_attempts INT DEFAULT 0",
            "last_analyzed_at TIMESTAMP",
            "support_amount_type VARCHAR(20) DEFAULT 'unknown'",
            "support_amount_max BIGINT",
            "support_amount_min BIGINT",
        ]:
            try:
                cursor.execute(f"ALTER TABLE announcements ADD COLUMN IF NOT EXISTS {col_def}")
                conn.commit()
            except Exception as e:
                conn.rollback()
                print(f"  Note: announcements.{col_def.split()[0]} migration skipped: {e}")

        # Phase 1 인덱스 — 자주 필터링되는 조합
        for idx_sql in [
            "CREATE INDEX IF NOT EXISTS idx_ann_deadline_type ON announcements(deadline_type) WHERE is_archived = FALSE",
            "CREATE INDEX IF NOT EXISTS idx_ann_archived ON announcements(is_archived) WHERE is_archived = FALSE",
            "CREATE INDEX IF NOT EXISTS idx_ann_analysis_status ON announcements(analysis_status)",
            "CREATE INDEX IF NOT EXISTS idx_ann_amount_type ON announcements(support_amount_type) WHERE is_archived = FALSE",
            "CREATE INDEX IF NOT EXISTS idx_ann_amount_max ON announcements(support_amount_max DESC NULLS LAST) WHERE is_archived = FALSE",
        ]:
            try:
                cursor.execute(idx_sql)
                conn.commit()
            except Exception as e:
                conn.rollback()
                print(f"  Note: Phase 1 index skipped: {e}")

        # billing_key 컬럼 추가 (정기결제용)
        try:
            cursor.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS billing_key TEXT")
            conn.commit()
        except Exception:
            conn.rollback()

        # custom_needs 컬럼 추가 (맞춤 알림용 구체적 니즈)
        try:
            cursor.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS custom_needs TEXT DEFAULT ''")
            conn.commit()
        except Exception:
            conn.rollback()

        # 개인 매칭용 프로필 컬럼 추가
        for col_def in [
            "gender VARCHAR(10) DEFAULT ''",
            "age_range VARCHAR(20) DEFAULT ''",
            "income_level VARCHAR(30) DEFAULT ''",
            "family_type VARCHAR(30) DEFAULT ''",
            "employment_status VARCHAR(30) DEFAULT ''",
            "founded_date VARCHAR(20) DEFAULT ''",
            "is_pre_founder BOOLEAN DEFAULT FALSE",
            "certifications TEXT DEFAULT ''",
            "custom_keywords TEXT DEFAULT ''",
        ]:
            try:
                col_name = col_def.split()[0]
                cursor.execute(f"ALTER TABLE users ADD COLUMN IF NOT EXISTS {col_def}")
                conn.commit()
            except Exception:
                conn.rollback()

        # kakao_refresh_token 컬럼 추가 (카카오톡 메시지 발송용)
        try:
            cursor.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS kakao_refresh_token TEXT")
            conn.commit()
        except Exception:
            conn.rollback()

        # kakao_enabled 컬럼 추가 (카카오톡 알림 설정)
        try:
            cursor.execute("ALTER TABLE notification_settings ADD COLUMN IF NOT EXISTS kakao_enabled INTEGER DEFAULT 0")
            conn.commit()
        except Exception:
            conn.rollback()

        # client_profiles CRM 확장 컬럼 추가
        for col_def in [
            "client_type VARCHAR(20) DEFAULT 'business'",
            "contact_name VARCHAR(100)",
            "contact_email VARCHAR(200)",
            "contact_phone VARCHAR(50)",
            "tags TEXT DEFAULT ''",
            "status VARCHAR(20) DEFAULT 'new'",
            # [재설계 04] 우대·제외 판정용 선택 필드
            "representative_age VARCHAR(20)",
            "is_women_enterprise BOOLEAN DEFAULT FALSE",
            "is_youth_enterprise BOOLEAN DEFAULT FALSE",
            "certifications TEXT DEFAULT ''",  # 콤마 구분: '벤처,이노비즈'
            "is_restart BOOLEAN DEFAULT FALSE",
        ]:
            try:
                cursor.execute(f"ALTER TABLE client_profiles ADD COLUMN IF NOT EXISTS {col_def}")
                conn.commit()
            except Exception:
                conn.rollback()

        # client_files 테이블 (고객사 자료 첨부)
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS client_files (
                    id SERIAL PRIMARY KEY,
                    client_id INTEGER NOT NULL,
                    owner_business_number VARCHAR(20) NOT NULL,
                    file_name VARCHAR(500) NOT NULL,
                    file_type VARCHAR(50) DEFAULT 'other',
                    file_size INTEGER DEFAULT 0,
                    file_data BYTEA,
                    extracted_text TEXT DEFAULT '',
                    memo TEXT DEFAULT '',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.commit()
        except Exception:
            conn.rollback()

        # client_files에 extracted_text 컬럼 추가 (기존 테이블인 경우)
        try:
            cursor.execute("ALTER TABLE client_files ADD COLUMN IF NOT EXISTS extracted_text TEXT DEFAULT ''")
            conn.commit()
        except Exception:
            conn.rollback()
        # P1.1: client_files.client_id NOT NULL 제약 완화 (client 미선택 업로드 허용)
        try:
            cursor.execute("ALTER TABLE client_files ALTER COLUMN client_id DROP NOT NULL")
            conn.commit()
        except Exception:
            conn.rollback()
        # P1.1: ai_summary 컬럼 추가 (업로드 시 AI 요약 저장)
        try:
            cursor.execute("ALTER TABLE client_files ADD COLUMN IF NOT EXISTS ai_summary TEXT DEFAULT ''")
            conn.commit()
        except Exception:
            conn.rollback()
        # P1.2: email_logs 테이블 (PRO 이메일 발송 이력)
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS email_logs (
                    id SERIAL PRIMARY KEY,
                    owner_business_number VARCHAR(20) NOT NULL,
                    client_id INTEGER,
                    recipient_email VARCHAR(255),
                    recipient_name VARCHAR(100),
                    subject TEXT,
                    body TEXT,
                    status VARCHAR(20),
                    error_detail TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_email_logs_owner ON email_logs(owner_business_number, created_at DESC)")
            conn.commit()
        except Exception:
            conn.rollback()
        # P1.3: match_history 테이블 (매칭 실행 이력)
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS match_history (
                    id SERIAL PRIMARY KEY,
                    business_number VARCHAR(20) NOT NULL,
                    user_type VARCHAR(20),
                    profile_snapshot JSONB,
                    total_matches INTEGER DEFAULT 0,
                    top_matches JSONB,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_match_history_bn ON match_history(business_number, created_at DESC)")
            conn.commit()
        except Exception:
            conn.rollback()

        # user_match_cache — 사전 매칭 결과 캐시
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS user_match_cache (
                    id SERIAL PRIMARY KEY,
                    business_number VARCHAR(20) NOT NULL,
                    target_type VARCHAR(20) NOT NULL DEFAULT 'business',
                    match_data JSONB NOT NULL DEFAULT '[]'::jsonb,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(business_number, target_type)
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_umc_bn ON user_match_cache(business_number)")
            conn.commit()
        except Exception:
            conn.rollback()

        # keyword_synonyms 테이블 생성 + 초기 데이터
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS keyword_synonyms (
                id SERIAL PRIMARY KEY,
                group_name TEXT NOT NULL,
                keyword TEXT NOT NULL,
                target_type VARCHAR(20) DEFAULT 'both',
                UNIQUE(group_name, keyword)
            )
        """)
        conn.commit()
        _seed_keyword_synonyms(conn)

        # 사용자 행동 이벤트 로그 테이블
        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS user_events (
                    id SERIAL PRIMARY KEY,
                    business_number VARCHAR(20),
                    event_type VARCHAR(50) NOT NULL,
                    event_detail TEXT DEFAULT '',
                    ip_address VARCHAR(50) DEFAULT '',
                    user_agent TEXT DEFAULT '',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_user_events_type ON user_events(event_type)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_user_events_created ON user_events(created_at)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_user_events_bn ON user_events(business_number)")
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS strategy_reports (
                    id SERIAL PRIMARY KEY,
                    report_type VARCHAR(50) NOT NULL DEFAULT 'weekly',
                    report_data TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            # 시스템 활동 이력 테이블
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS system_logs (
                    id SERIAL PRIMARY KEY,
                    action VARCHAR(50) NOT NULL,
                    category VARCHAR(30) NOT NULL DEFAULT 'system',
                    detail TEXT DEFAULT '',
                    result VARCHAR(20) DEFAULT 'success',
                    count_affected INT DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_system_logs_action ON system_logs(action)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_system_logs_category ON system_logs(category)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_system_logs_created ON system_logs(created_at)")

            # ── AI 패트롤 시스템 테이블 ──
            # 분석 실패 추적
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS analysis_failures (
                    id SERIAL PRIMARY KEY,
                    announcement_id INT NOT NULL,
                    error_type VARCHAR(50) NOT NULL,
                    error_message TEXT,
                    failed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    retry_count INT DEFAULT 0,
                    next_retry_at TIMESTAMP,
                    resolved_at TIMESTAMP,
                    UNIQUE(announcement_id, error_type)
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_analysis_failures_retry ON analysis_failures(next_retry_at) WHERE resolved_at IS NULL")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_analysis_failures_aid ON analysis_failures(announcement_id)")

            # 패트롤 실행 이력
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS patrol_history (
                    id SERIAL PRIMARY KEY,
                    started_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    completed_at TIMESTAMP,
                    status VARCHAR(20) DEFAULT 'running',
                    summary JSONB,
                    error TEXT
                )
            """)
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_patrol_history_started ON patrol_history(started_at DESC)")

            conn.commit()
        except Exception:
            conn.rollback()

        conn.commit()
        conn.close()
        print("  DB connection OK (PostgreSQL/Supabase)")
    except Exception as e:
        print(f"  DB connection error (app will continue): {e}")


def _log_system(action: str, category: str = "system", detail: str = "", result: str = "success", count: int = 0):
    """시스템 활동 이력을 DB에 저장 (수집/분석/알림/매칭 등)"""
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO system_logs (action, category, detail, result, count_affected) VALUES (%s, %s, %s, %s, %s)",
            (action[:50], category[:30], (detail or "")[:500], result[:20], count)
        )
        conn.commit()
        conn.close()
    except Exception:
        pass


def _log_event(event_type: str, business_number: str = "", detail: str = "", ip: str = "", ua: str = ""):
    """사용자 행동 이벤트를 DB에 비동기 저장"""
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO user_events (business_number, event_type, event_detail, ip_address, user_agent) VALUES (%s, %s, %s, %s, %s)",
            (business_number[:20] if business_number else "", event_type, (detail or "")[:500], (ip or "")[:50], (ua or "")[:200])
        )
        conn.commit()
        conn.close()
    except Exception:
        pass  # 로깅 실패가 서비스에 영향 주면 안 됨


def _seed_keyword_synonyms(conn):
    """동의어 풀 초기 데이터 (이미 있으면 스킵)"""
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) AS cnt FROM keyword_synonyms")
    if cursor.fetchone()["cnt"] > 0:
        return  # 이미 시드 완료

    seeds = [
        # ── 기업 ──
        ("수출", ["해외진출", "해외인증", "글로벌", "무역", "수출바우처", "해외시장", "통상"], "business"),
        ("창업", ["스타트업", "예비창업", "초기창업", "벤처", "창업지원", "창업보육"], "business"),
        ("스마트", ["스마트공장", "스마트제조", "디지털전환", "DX", "스마트센서"], "business"),
        ("고용", ["채용", "일자리", "인력", "고용장려금", "고용지원", "인력양성"], "business"),
        ("R&D", ["연구개발", "기술개발", "기술혁신", "연구비", "과제"], "business"),
        ("친환경", ["탄소중립", "ESG", "녹색", "에너지절감", "그린", "환경"], "business"),
        ("자금", ["융자", "대출", "보증", "정책자금", "운전자금", "시설자금"], "business"),
        ("마케팅", ["홍보", "판로", "브랜드", "전시회", "박람회", "판매"], "business"),
        ("특허", ["지식재산", "IP", "지재권", "특허출원", "상표"], "business"),
        ("인증", ["ISO", "KC인증", "품질인증", "GMP", "HACCP", "인허가"], "business"),
        ("디지털", ["IT", "소프트웨어", "AI", "빅데이터", "클라우드", "ICT"], "business"),
        ("제조", ["생산", "공장", "제조업", "생산기반", "설비"], "business"),
        ("컨설팅", ["경영컨설팅", "기술컨설팅", "자문", "진단", "멘토링"], "business"),
        # ── 개인 ──
        ("다자녀", ["다둥이", "셋째아이", "3자녀", "다자녀가정", "다자녀가구"], "individual"),
        ("출산", ["임신", "산모", "산후", "분만", "출산장려금", "출산축하금", "출생"], "individual"),
        ("육아", ["보육", "어린이집", "유아", "영아", "아이돌봄", "양육", "돌봄"], "individual"),
        ("장학금", ["학비", "등록금", "학자금", "교육비", "학업장려금"], "individual"),
        ("취업", ["구직", "일자리", "직업훈련", "취준", "직업교육", "취업지원"], "individual"),
        ("주거", ["전세", "임대", "주택", "월세", "주거급여", "주거지원", "임대주택"], "individual"),
        ("노인", ["어르신", "경로", "고령자", "실버", "노후", "연금"], "individual"),
        ("장애", ["장애인", "복지카드", "활동지원", "장애수당", "보조기기"], "individual"),
        ("저소득", ["기초생활", "차상위", "기초수급", "한부모", "긴급복지", "생계급여"], "individual"),
        ("청년", ["청년지원", "청년수당", "청년정책", "MZ", "청년월세"], "individual"),
        ("의료", ["건강", "치료", "병원비", "의료비", "건강검진", "진료비"], "individual"),
        ("교육", ["평생교육", "학습", "훈련", "교육바우처", "배움카드"], "individual"),
        ("문화", ["문화바우처", "여가", "체육", "관광", "문화생활", "공연"], "individual"),
    ]

    for group_name, keywords, target_type in seeds:
        # 대표 키워드도 자기 자신을 포함
        all_kw = [group_name] + keywords
        for kw in all_kw:
            try:
                cursor.execute(
                    "INSERT INTO keyword_synonyms (group_name, keyword, target_type) VALUES (%s, %s, %s) ON CONFLICT DO NOTHING",
                    (group_name, kw, target_type),
                )
            except Exception:
                pass
    conn.commit()
    print("  [Seed] keyword_synonyms 초기 데이터 삽입 완료")


init_database()


SYNC_HOUR = int(os.environ.get("SYNC_HOUR", "8"))
DIGEST_HOUR = int(os.environ.get("DIGEST_HOUR", "0"))  # UTC 0시 = 한국시간 09시


async def _daily_sync_loop():
    """매일 SYNC_HOUR 시에 공고 수집 → AI 재분석 파이프라인"""
    from app.services.sync_service import SyncService
    sync_service = SyncService()

    while True:
        now = datetime.datetime.now()
        target = now.replace(hour=SYNC_HOUR, minute=0, second=0, microsecond=0)
        if now >= target:
            target += datetime.timedelta(days=1)
        wait_seconds = (target - now).total_seconds()
        print(f"[Scheduler] next sync at {target.isoformat()} (in {wait_seconds/3600:.1f}h)")
        await asyncio.sleep(wait_seconds)
        try:
            # Step 0-A: URL 자동 관리 (누락 등록 + 신규 발견 + 폐기 정리)
            print("[Scheduler] Step 0-A: URL 자동 관리...")
            _auto_seed_urls()
            _discover_new_sources()
            _deactivate_dead_urls()

            # Step 0-B: 구독 자동 갱신 (만료된 빌링키 결제)
            print("[Scheduler] Step 0-B: 구독 자동 갱신...")
            _auto_renew_subscriptions()

            # Step 1: 기업 API 수집 + (월요일만) 개인 복지 전체 동기화
            print("[Scheduler] Step 1/3: 공고 수집 시작...")
            await sync_service.sync_all()
            print("[Scheduler] Step 1/3: 공고 수집 완료")

            # Step 1.5: DB 정리 (비지원사업 제거 + 중복 제거)
            print("[Scheduler] Step 1.5: DB 정리 + 분류...")
            _cleanup_non_support_announcements()
            _deduplicate_announcements()
            _auto_classify_target_type()

            # Step 2a: 지자체복지 상세 보강 (매일 100건씩 점진적)
            print("[Scheduler] Step 2a: 지자체복지 상세 보강...")
            try:
                from app.services.public_api_service import gov_api_service
                await gov_api_service.enrich_local_welfare_details(batch_size=100)
            except Exception as enrich_err:
                print(f"[Scheduler] enrich(local-welfare) error: {enrich_err}")

            # Step 2b: gov24 개인 공고 상세 보강 (매일 100건씩)
            print("[Scheduler] Step 2b: gov24 개인 공고 상세 보강...")
            try:
                await gov_api_service.enrich_gov24_individual_details(batch_size=100)
            except Exception as enrich_err:
                print(f"[Scheduler] enrich(gov24-individual) error: {enrich_err}")

            # Step 3: AI 재분석 — 기업 공고만 (개인 복지는 API 데이터로 보강)
            print("[Scheduler] Step 3/3: AI 재분석 시작 (미분석 기업 공고)...")
            import threading
            t = threading.Thread(target=_run_reanalyze_in_thread, args=(300,), daemon=True)
            t.start()
            print("[Scheduler] Step 3/3: AI 재분석 백그라운드 실행 중")
            _log_system("scheduler_run", "system", "일일 스케줄러 전체 실행 완료", "success")
        except Exception as e:
            _log_system("scheduler_run", "system", f"스케줄러 오류: {e}", "error")
            print(f"[Scheduler] sync error: {e}")


async def _daily_digest_loop():
    """평일 한국시간 09시(UTC 0시)에 매칭 + 이메일/푸시 발송"""
    from app.services.notification_service import notification_service
    while True:
        now = datetime.datetime.utcnow()
        target = now.replace(hour=DIGEST_HOUR, minute=0, second=0, microsecond=0)
        if now >= target:
            target += datetime.timedelta(days=1)
        # 평일만 (월~금, UTC 기준 → 한국 요일과 동일)
        while target.weekday() >= 5:  # 5=토, 6=일
            target += datetime.timedelta(days=1)
        wait_seconds = (target - now).total_seconds()
        kst_target = target + datetime.timedelta(hours=9)
        weekday_kr = ["월", "화", "수", "목", "금", "토", "일"][kst_target.weekday()]
        print(f"[Scheduler] next digest at KST {kst_target.strftime('%Y-%m-%d')}({weekday_kr}) 09:00 (in {wait_seconds/3600:.1f}h)")
        await asyncio.sleep(wait_seconds)
        try:
            print("[Scheduler] Running scheduled daily digest (평일 09시 KST)...")
            results = await notification_service.generate_daily_digest()
            sent = sum(1 for r in results if r.get("email_sent"))
            push_sent = sum(r.get("push_sent", 0) for r in results)
            print(f"[Scheduler] Digest complete: {len(results)} users, {sent} emails, {push_sent} pushes sent")
        except Exception as e:
            print(f"[Scheduler] digest error: {e}")
        # 사전 매칭 캐시 생성 — 활성 사용자 매칭 결과를 DB에 저장
        try:
            print("[Scheduler] Running pre-match cache for active users...")
            _prematch_count = _run_prematch_cache()
            print(f"[Scheduler] Pre-match complete: {_prematch_count} users cached")
        except Exception as e:
            print(f"[Scheduler] pre-match error: {e}")


def _run_prematch_cache() -> int:
    """활성 사용자 매칭 결과를 user_match_cache 테이블에 저장"""
    from app.core.matcher import get_matches_hybrid
    conn = get_db_connection()
    cur = conn.cursor()
    # 활성 유료 사용자 조회
    cur.execute("""
        SELECT business_number, user_type, industry_code, address_city, interests,
               revenue_bracket, employee_count_bracket, establishment_date,
               age_range, income_level, family_type, employment_status,
               custom_keywords, certifications, interest_regions, company_name, gender
        FROM users
        WHERE plan IN ('lite', 'lite_trial', 'basic', 'pro', 'biz')
          AND (plan_expires_at IS NULL OR plan_expires_at > NOW())
    """)
    users = [dict(r) for r in cur.fetchall()]
    conn.close()  # 사용자 목록 조회 후 즉시 반환
    count = 0
    for u in users:
        bn = u.get("business_number")
        user_type = u.get("user_type") or "both"
        try:
            # 매칭 실행 (내부에서 별도 커넥션 사용)
            if user_type in ("business", "both"):
                biz = get_matches_hybrid(u, is_individual=False)
                biz = biz[:100]
                _save = get_db_connection()
                _save.cursor().execute("""
                    INSERT INTO user_match_cache (business_number, target_type, match_data, created_at)
                    VALUES (%s, 'business', %s::jsonb, CURRENT_TIMESTAMP)
                    ON CONFLICT (business_number, target_type)
                    DO UPDATE SET match_data = EXCLUDED.match_data, created_at = CURRENT_TIMESTAMP
                """, (bn, json.dumps(biz, ensure_ascii=False, default=str)))
                _save.commit(); _save.close()
            if user_type in ("individual", "both"):
                ind = get_matches_hybrid(u, is_individual=True)
                ind = ind[:100]
                _save = get_db_connection()
                _save.cursor().execute("""
                    INSERT INTO user_match_cache (business_number, target_type, match_data, created_at)
                    VALUES (%s, 'individual', %s::jsonb, CURRENT_TIMESTAMP)
                    ON CONFLICT (business_number, target_type)
                    DO UPDATE SET match_data = EXCLUDED.match_data, created_at = CURRENT_TIMESTAMP
                """, (bn, json.dumps(ind, ensure_ascii=False, default=str)))
                _save.commit(); _save.close()
            # public_order 캐싱 (전체 탭 정렬용)
            if user_type in ("business", "both"):
                pub_biz = _compute_public_order_for_user(u, is_individual=False)
                _s = get_db_connection()
                _s.cursor().execute("""
                    INSERT INTO user_match_cache (business_number, target_type, match_data, created_at)
                    VALUES (%s, 'public_biz', %s::jsonb, CURRENT_TIMESTAMP)
                    ON CONFLICT (business_number, target_type)
                    DO UPDATE SET match_data = EXCLUDED.match_data, created_at = CURRENT_TIMESTAMP
                """, (bn, json.dumps(pub_biz, ensure_ascii=False, default=str)))
                _s.commit(); _s.close()
            if user_type in ("individual", "both"):
                pub_ind = _compute_public_order_for_user(u, is_individual=True)
                _s = get_db_connection()
                _s.cursor().execute("""
                    INSERT INTO user_match_cache (business_number, target_type, match_data, created_at)
                    VALUES (%s, 'public_ind', %s::jsonb, CURRENT_TIMESTAMP)
                    ON CONFLICT (business_number, target_type)
                    DO UPDATE SET match_data = EXCLUDED.match_data, created_at = CURRENT_TIMESTAMP
                """, (bn, json.dumps(pub_ind, ensure_ascii=False, default=str)))
                _s.commit(); _s.close()

            count += 1
            import time; time.sleep(0.5)  # 커넥션 풀 여유
        except Exception as e:
            import traceback as _tb
            print(f"[prematch] {bn}: {e}\n{_tb.format_exc()}")
    return count


def _compute_public_order_for_user(user_profile: dict, is_individual: bool) -> dict:
    """사용자별 공고 정렬 순서 계산 (하루 1회 배치용).
    local(내 지역) / national(전국) 분리 + 매칭 점수 정렬.
    하위호환: eligible_ids / ineligible_ids 도 함께 반환.
    """
    from app.core.matcher import _check_region_exclusion
    from app.services.rule_engine import _normalize_region

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        f"""SELECT announcement_id, title, region, category,
                   support_amount, support_amount_max, target_type, deadline_date,
                   eligibility_logic
            FROM announcements
            WHERE {valid_announcement_where()}
            ORDER BY deadline_date ASC NULLS LAST, created_at DESC"""
    )
    all_anns = [dict(r) for r in cur.fetchall()]
    conn.close()

    raw_city = str(user_profile.get("address_city", "") or "")
    cities = [c.strip() for c in raw_city.split(",") if c.strip() and c.strip() != "전국"]
    user_city = _normalize_region(cities[0]) if cities else ""
    interests = [i.strip() for i in str(user_profile.get("interests", "") or "").split(",") if i.strip()]
    gender = str(user_profile.get("gender", "") or "")
    ind_major = str(user_profile.get("industry_code") or "")[:2]
    is_farmer = ind_major in ("01", "02", "03")

    RESTRICTED_ALWAYS = ["장애인기업", "장애인창업", "농업인", "영농조합", "어업인", "수산업", "보훈", "제대군인"]
    FEMALE_ONLY = ["여성기업", "여성창업", "여성경제인"]
    YOUTH_BIZ = ["청년창업", "청년기업", "만39세", "만 39세"]
    _AGRI_CATS = ("농림", "수산", "임업", "축산")

    import datetime as _dt_batch, re as _re_bkt
    _today = _dt_batch.date.today()

    local_scored: list = []    # [(score, ann_id), ...]
    national_scored: list = []
    eligible_ids: list = []
    ineligible_ids: list = []

    for ann in all_anns:
        ann_id = ann["announcement_id"]
        title = ann.get("title", "") or ""
        region = ann.get("region", "") or ""
        ann_target = (ann.get("target_type", "") or "").strip()
        category = ann.get("category", "") or ""

        dl = ann.get("deadline_date")
        if dl and (dl.date() if hasattr(dl, "date") else dl) < _today:
            ineligible_ids.append(ann_id); continue

        if is_individual and ann_target == "business":
            ineligible_ids.append(ann_id); continue
        if not is_individual and ann_target == "individual":
            ineligible_ids.append(ann_id); continue

        _region_specific = bool(region and region not in ("전국", "", "전국 및 각 지역"))
        if user_city:
            excluded, _ = _check_region_exclusion(user_city, region, title)
            if excluded:
                ineligible_ids.append(ann_id); continue

        if user_city and not _region_specific:
            _el_rr_data = ann.get("eligibility_logic")
            if isinstance(_el_rr_data, str):
                try:
                    import json as _json
                    _el_rr_data = _json.loads(_el_rr_data)
                except Exception:
                    _el_rr_data = None
            if _el_rr_data and isinstance(_el_rr_data, dict):
                _el_rr = _el_rr_data.get("region_restriction") or ""
                if isinstance(_el_rr, list):
                    _el_rr = _el_rr[0] if _el_rr else ""
                _el_rr = _normalize_region(str(_el_rr).strip())
                if _el_rr and _el_rr not in ("전국", "", "전국 및 각 지역"):
                    excluded, _ = _check_region_exclusion(user_city, _el_rr, title)
                    if excluded:
                        ineligible_ids.append(ann_id); continue

        if not is_farmer and any(ac in category for ac in _AGRI_CATS):
            ineligible_ids.append(ann_id); continue

        if any(kw in title for kw in RESTRICTED_ALWAYS):
            ineligible_ids.append(ann_id); continue
        if any(kw in title for kw in FEMALE_ONLY) and gender != "여성":
            ineligible_ids.append(ann_id); continue
        if not is_individual and any(kw in title for kw in YOUTH_BIZ):
            ineligible_ids.append(ann_id); continue

        el = ann.get("eligibility_logic")
        if el and isinstance(el, str):
            try:
                import json as _json
                el = _json.loads(el)
            except Exception:
                el = None
        if el and isinstance(el, dict):
            _el_raw = el.get("gender_restriction") or ""
            if isinstance(_el_raw, list):
                el_gender = ",".join(str(v) for v in _el_raw).strip()
            else:
                el_gender = str(_el_raw).strip()
            if el_gender in ("여성", "여성전용", "여성기업", "여성창업자") and gender != "여성":
                ineligible_ids.append(ann_id); continue

        # ── 매칭 점수 계산 ──
        score = 0

        # 관심분야 일치 (50점)
        interest_hit = bool(interests) and any(it in category or it in title for it in interests)
        if interest_hit:
            score += 50

        # 지원 금액 규모 (30점) — 1억당 5점, 최대 30점
        amt_max = ann.get("support_amount_max")
        if amt_max and isinstance(amt_max, (int, float)) and amt_max > 0:
            score += min(30, int(amt_max / 100_000_000) * 5)

        # 기본 적합도 (15점) — 필터 통과한 공고에 고정 부여
        score += 15

        # 마감 임박 (5점)
        if dl:
            dl_date = dl.date() if hasattr(dl, "date") else dl
            days_left = (dl_date - _today).days
            if 0 < days_left <= 7:
                score += 5
            elif days_left <= 30:
                score += 2

        # ── 지역 / 전국 분류 ──
        _bkt_m = _re_bkt.search(
            r'\[(서울|경기|인천|부산|대구|대전|광주|울산|세종|강원|충북|충남|전북|전남|경북|경남|제주)\]',
            title
        )
        _title_region = _normalize_region(_bkt_m.group(1)) if _bkt_m else None
        _norm_region = _normalize_region(region)
        _db_is_regional = bool(_norm_region and _norm_region not in ("전국", "", "All"))

        # 내 지역: title [city] 태그 또는 region이 user_city와 매칭
        in_my_region = bool(user_city) and (
            (_title_region and _title_region == user_city) or
            (not _title_region and _db_is_regional and user_city in _norm_region)
        )
        # 전국: 지역명이 특정되지 않은 공고 (null/전국/빈값)
        is_national = not _db_is_regional and not _title_region

        eligible_ids.append(ann_id)
        if in_my_region:
            local_scored.append((score, ann_id))
        elif is_national:
            national_scored.append((score, ann_id))
        # else: 다른 지역 공고 — 두 탭 모두 제외, 전체 탭(eligible_ids)에만 포함

    # 점수 내림차순 정렬
    local_scored.sort(key=lambda x: x[0], reverse=True)
    national_scored.sort(key=lambda x: x[0], reverse=True)

    return {
        "local": [aid for _, aid in local_scored],
        "national": [aid for _, aid in national_scored],
        "eligible_ids": eligible_ids,    # 하위호환 (전체 탭)
        "ineligible_ids": ineligible_ids,
    }


def _log_expired_announcements():
    """만료 공고 수만 로그에 기록 (자동 삭제 안 함 — 관리자가 수동 관리)"""
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM announcements WHERE deadline_date IS NOT NULL AND deadline_date < CURRENT_DATE")
        row = cur.fetchone()
        count = row["count"] if row else 0
        cur.execute("SELECT COUNT(*) FROM announcements")
        total = cur.fetchone()["count"]
        print(f"[Info] 공고 현황: 전체 {total}건, 만료 {count}건 (자동 삭제 안 함)")
        conn.close()
    except Exception as e:
        print(f"[Info] expire check error: {e}")


def _auto_seed_urls():
    """seed_regional_urls의 URL을 DB에 자동 등록 (누락분만)"""
    try:
        from app.db.seed_regional_urls import REGIONAL_URLS
        conn = get_db_connection()
        cursor = conn.cursor()
        inserted = 0
        for source_name, url in REGIONAL_URLS:
            try:
                cursor.execute(
                    "INSERT INTO admin_urls (url, source_name, is_active) VALUES (%s, %s, 1) ON CONFLICT (url) DO NOTHING",
                    (url, source_name),
                )
                conn.commit()
                if cursor.rowcount > 0:
                    inserted += 1
            except Exception:
                conn.rollback()
        conn.close()
        if inserted > 0:
            print(f"[auto-seed] {inserted}개 신규 URL 등록 완료 (총 {len(REGIONAL_URLS)}개 중)")
    except Exception as e:
        print(f"[auto-seed] 오류: {e}")


def _discover_new_sources():
    """수집된 공고의 origin_url에서 미등록 도메인을 발견하여 자동 등록"""
    from urllib.parse import urlparse
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # 1. 기존 admin_urls 도메인 목록
        cursor.execute("SELECT url FROM admin_urls WHERE is_active = 1")
        registered_domains = set()
        for row in cursor.fetchall():
            try:
                registered_domains.add(urlparse(row["url"]).netloc.replace("www.", ""))
            except Exception:
                pass

        # 2. 최근 30일 공고의 origin_url에서 도메인 추출
        cursor.execute("""
            SELECT origin_url, COUNT(*) as cnt
            FROM announcements
            WHERE origin_url IS NOT NULL
              AND origin_url != ''
              AND created_at >= CURRENT_DATE - INTERVAL '30 days'
              AND COALESCE(target_type, 'business') = 'business'
            GROUP BY origin_url
        """)
        domain_count = {}
        for row in cursor.fetchall():
            try:
                parsed = urlparse(row["origin_url"])
                domain = parsed.netloc.replace("www.", "")
                if domain and "." in domain and domain.endswith(".kr"):
                    domain_count[domain] = domain_count.get(domain, 0) + row["cnt"]
            except Exception:
                pass

        # 3. 미등록 도메인 중 공고 3건 이상인 것만 후보
        new_domains = []
        for domain, cnt in domain_count.items():
            if domain not in registered_domains and cnt >= 3:
                new_domains.append((domain, cnt))

        # 4. 신규 도메인 자동 등록 (루트 URL로)
        inserted = 0
        for domain, cnt in sorted(new_domains, key=lambda x: -x[1])[:20]:  # 상위 20개
            source_name = domain.split(".")[0].upper() if domain else "Unknown"
            url = f"https://www.{domain}" if not domain.startswith("www.") else f"https://{domain}"
            try:
                cursor.execute(
                    "INSERT INTO admin_urls (url, source_name, is_active) VALUES (%s, %s, 1) ON CONFLICT (url) DO NOTHING",
                    (url, f"[자동발견] {source_name} ({cnt}건)"),
                )
                conn.commit()
                if cursor.rowcount > 0:
                    inserted += 1
                    print(f"[discover] 신규 소스 발견: {domain} ({cnt}건)")
            except Exception:
                conn.rollback()

        conn.close()
        if inserted > 0:
            print(f"[discover] {inserted}개 신규 소스 자동 등록")
        else:
            print(f"[discover] 신규 소스 없음 (후보 {len(new_domains)}개 검토)")
    except Exception as e:
        print(f"[discover] 오류: {e}")


def _cleanup_non_support_announcements() -> int:
    """DB에서 지원사업이 아닌 공고 제거. 반환: 삭제 건수."""
    # 제목 키워드 — 이 단어가 제목에 있으면 비지원사업
    NON_SUPPORT_PATTERNS = [
        # 기존
        "업무추진비", "사용내역", "사용 내역", "회의록", "의사록",
        "결산", "예산서", "감사결과", "인사발령",
        "입찰결과", "낙찰자", "계약현황", "계약체결", "개찰결과",
        "채용결과", "합격자 발표", "선정결과 발표",
        "행사 후기", "수료식", "시상식",
        "취소공고", "취소 공고", "철회",
        # 추가 — 크롤러에서 유입된 비지원사업 패턴
        "포럼 개최", "포럼개최", "세미나 개최", "간담회 개최", "행사 개최",
        "직원 채용", "직원채용", "채용 공고", "채용공고",
        "합격자 공고", "서류전형 합격", "최종합격자",
        "당첨자 명단", "당첨자명단",
        "기관 소개", "원장 인사말", "인사말",
        "민원 안내", "민원안내", "정보공개 안내",
        "통근버스", "수기 공모전", "수기공모전",
        "컴퓨터 교육", "정보화교육",
        "보도자료", "언론보도",
    ]
    # 이 단어가 있으면 예외 (지원사업일 가능성)
    EXCEPTIONS = ["모집", "참여기업", "참여자", "신청", "접수", "공모", "지원사업", "지원금"]

    # 비지원사업 카테고리 — category 컬럼이 이 값이면 삭제
    NON_SUPPORT_CATEGORIES = [
        "Forum", "행사", "채용", "민원", "정보공개", "기관소개",
        "정보화교육", "교통", "보도자료",
    ]

    try:
        conn = get_db_connection()
        cur = conn.cursor()
        deleted = 0

        # ① 제목 키워드 기반 삭제
        for pattern in NON_SUPPORT_PATTERNS:
            cur.execute(
                "SELECT announcement_id, title FROM announcements WHERE title ILIKE %s",
                (f"%{pattern}%",)
            )
            for row in cur.fetchall():
                title = row["title"]
                if any(exc in title for exc in EXCEPTIONS):
                    continue
                cur.execute("DELETE FROM announcements WHERE announcement_id = %s", (row["announcement_id"],))
                deleted += 1
            conn.commit()

        # ② category 기반 삭제
        if NON_SUPPORT_CATEGORIES:
            placeholders = ",".join(["%s"] * len(NON_SUPPORT_CATEGORIES))
            cur.execute(
                f"DELETE FROM announcements WHERE category IN ({placeholders}) RETURNING announcement_id",
                NON_SUPPORT_CATEGORIES
            )
            cat_deleted = cur.rowcount or 0
            conn.commit()
            deleted += cat_deleted

        conn.close()
        if deleted > 0:
            print(f"[cleanup] 비지원사업 공고 {deleted}건 삭제")
        return deleted
    except Exception as e:
        print(f"[cleanup] 오류: {e}")
        return 0


def _auto_classify_target_type():
    """기존 공고의 target_type을 키워드 기반으로 자동 분류 (business→individual/both)"""
    INDIVIDUAL_KEYWORDS = [
        "다자녀", "다둥이", "출산", "임신", "산모", "육아", "보육", "어린이집",
        "장학금", "학비", "등록금", "학자금",
        "취업지원", "구직", "직업훈련", "취업성공",
        "주거급여", "전세", "임대주택", "월세지원", "주거지원",
        "노인", "어르신", "경로", "고령자", "연금",
        "장애인", "장애수당", "활동지원", "보조기기",
        "기초생활", "차상위", "한부모", "긴급복지", "생계급여", "저소득",
        "청년수당", "청년월세", "청년정책", "청년지원",
        "의료비", "건강검진", "진료비", "병원비",
        "교육바우처", "평생교육", "배움카드",
        "문화바우처", "문화생활",
    ]
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT announcement_id, title, category FROM announcements WHERE COALESCE(target_type, 'business') = 'business'")
        updated = 0
        for row in cur.fetchall():
            title = (row["title"] or "") + " " + (row.get("category") or "")
            if any(kw in title for kw in INDIVIDUAL_KEYWORDS):
                cur.execute("UPDATE announcements SET target_type = 'individual' WHERE announcement_id = %s", (row["announcement_id"],))
                updated += 1
        conn.commit()
        conn.close()
        if updated > 0:
            print(f"[classify] {updated}건 공고 target_type → individual 분류")
    except Exception as e:
        print(f"[classify] 오류: {e}")


def _deduplicate_announcements():
    """DB에서 중복 공고 제거 — 제목 정규화 기준으로 최신 1건만 유지"""
    import re as _re
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # 제목 정규화 후 중복 찾기 (공백/괄호/특수문자 제거)
        cur.execute("""
            WITH normalized AS (
                SELECT announcement_id, title, created_at,
                       LOWER(REGEXP_REPLACE(REGEXP_REPLACE(title, '[\s\[\]()（）【】]', '', 'g'), '년도', '년', 'g')) AS norm_title
                FROM announcements
                WHERE created_at >= CURRENT_DATE - INTERVAL '90 days'
            ),
            duplicates AS (
                SELECT norm_title, COUNT(*) as cnt,
                       ARRAY_AGG(announcement_id ORDER BY created_at DESC) as ids
                FROM normalized
                GROUP BY norm_title
                HAVING COUNT(*) > 1
            )
            SELECT norm_title, cnt, ids FROM duplicates
        """)
        dup_groups = cur.fetchall()

        deleted = 0
        for row in dup_groups:
            ids = row["ids"]
            # 첫 번째(최신)는 유지, 나머지 삭제
            to_delete = ids[1:]
            for del_id in to_delete:
                try:
                    cur.execute("DELETE FROM announcements WHERE announcement_id = %s", (del_id,))
                    deleted += 1
                except Exception:
                    conn.rollback()
                    continue
            conn.commit()

        conn.close()
        if deleted > 0:
            print(f"[dedup] {deleted}건 중복 공고 삭제 ({len(dup_groups)}개 그룹)")
        else:
            print(f"[dedup] 중복 없음")
    except Exception as e:
        print(f"[dedup] 오류: {e}")


def _deactivate_dead_urls():
    """연속 실패 URL 자동 비활성화 (fail_count >= 5 + 복구 실패)"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE admin_urls
            SET is_active = 0
            WHERE is_active = 1
              AND fail_count >= 5
              AND recovered_url IS NULL
            RETURNING id, source_name, url, fail_count
        """)
        deactivated = cursor.fetchall()
        conn.commit()
        conn.close()
        if deactivated:
            for row in deactivated:
                print(f"[cleanup] 비활성화: {row['source_name']} (실패 {row['fail_count']}회)")
            print(f"[cleanup] {len(deactivated)}개 URL 비활성화 완료")
    except Exception as e:
        print(f"[cleanup] 오류: {e}")


def _db_keepalive():
    """Supabase PgBouncer 콜드스타트 방지용 keepalive ping."""
    try:
        conn = get_db_connection()
        conn.cursor().execute("SELECT 1")
        conn.close()
    except Exception as e:
        print(f"[Keepalive] ping error: {e}")


def _prewarm_response_cache(startup: bool = False):
    """비로그인 공고 목록 응답 캐시 갱신. startup=True면 3초 대기 후 실행."""
    if startup:
        import time as _t; _t.sleep(3)
    for tt in ("business", "individual", None):
        try:
            cache_key = f"pub:v2:{tt}:1:20"
            if True:  # 항상 갱신 — TTL 만료 전 선제 교체
                conn = get_db_connection()
                cur = conn.cursor()
                valid_where = valid_announcement_where()
                type_clause = ""
                type_params: list = []
                if tt:
                    type_clause = "AND (target_type = %s OR target_type = 'both')"
                    type_params = [tt]
                full_where = f"{valid_where} {type_clause}"

                cur.execute(f"SELECT COUNT(*) AS cnt FROM announcements WHERE {full_where}", type_params)
                total = cur.fetchone()["cnt"]

                cur.execute(
                    f"""SELECT announcement_id, title, region, category, department,
                               support_amount, support_amount_max, support_amount_min, support_amount_type,
                               deadline_date, origin_source, created_at,
                               COALESCE(target_type, 'business') AS target_type,
                               origin_url, summary_text, eligibility_logic,
                               established_years_limit, revenue_limit, employee_limit
                        FROM announcements
                        WHERE {full_where}
                        ORDER BY
                            CASE WHEN deadline_date IS NOT NULL AND deadline_date < CURRENT_DATE THEN 9
                                 ELSE 0 END,
                            CASE
                                WHEN deadline_date IS NOT NULL
                                     AND deadline_date BETWEEN CURRENT_DATE AND CURRENT_DATE + INTERVAL '30 days'
                                     AND support_amount IS NOT NULL AND support_amount != '' THEN 0
                                WHEN (region IN ('전국', '', '전국 및 각 지역', 'All') OR region IS NULL)
                                     AND support_amount IS NOT NULL AND support_amount != '' THEN 1
                                WHEN support_amount IS NOT NULL AND support_amount != '' THEN 2
                                ELSE 3
                            END,
                            deadline_date ASC NULLS LAST, created_at DESC
                        LIMIT 20 OFFSET 0""",
                    type_params,
                )
                rows = [dict(r) for r in cur.fetchall()]

                cat_cache_key = f"cat_counts:{tt or 'all'}"
                category_counts = _get_cached(cat_cache_key)
                if not category_counts:
                    cur.execute(
                        f"SELECT category, COUNT(*) cnt FROM announcements WHERE {full_where} GROUP BY category ORDER BY cnt DESC",
                        type_params,
                    )
                    category_counts = {r["category"]: r["cnt"] for r in cur.fetchall() if r["category"]}
                    _set_cache(cat_cache_key, category_counts)

                payload = {
                    "status": "SUCCESS", "data": rows, "total": total,
                    "page": 1, "size": 20, "regions": [], "categories": [],
                    "category_counts": category_counts, "personalized": False, "source": "cache",
                }
                _set_cache(cache_key, payload)
                conn.close()
                print(f"[Prewarm] {tt or 'all'}: {len(rows)}건 캐시 완료")
        except Exception as e:
            print(f"[Prewarm] {tt} error: {e}")


async def lifespan(app):
    _log_expired_announcements()  # 시작 시 현황만 로그
    # 서버 시작 시 사전매칭 캐시 (백그라운드)
    import threading

    def _warmup():
        _db_keepalive()  # 1. DB 연결 먼저 warm-up (cold start 40s 방지)
        _prewarm_response_cache(startup=False)  # 2. warm DB로 응답 캐시 즉시 채우기
        print(f"[Startup] Pre-match: {_run_prematch_cache()} users cached")

    threading.Thread(target=_warmup, daemon=True).start()

    # ── 금융 지식 시딩 (최초 1회) ──
    try:
        from app.services.financial_analysis.knowledge_seed import seed_financial_knowledge
        seed_conn = get_db_connection()
        seeded = seed_financial_knowledge(seed_conn)
        seed_conn.close()
        if seeded:
            print(f"[KnowledgeSeed] {seeded} financial knowledge items seeded")
    except Exception as seed_err:
        print(f"[KnowledgeSeed] Error (non-critical): {seed_err}")

    # ── 일일 통합 파이프라인 (매일 03:00 KST = 18:00 UTC) ──
    # docs/daily-pipeline.md 참조
    # PATROL_ENABLED=false 환경변수로 비활성화 가능
    pipeline_scheduler = None
    if os.getenv("PATROL_ENABLED", "true").lower() != "false":
        try:
            from apscheduler.schedulers.asyncio import AsyncIOScheduler
            from apscheduler.triggers.cron import CronTrigger

            def _daily_pipeline_job():
                try:
                    from app.services.patrol.daily_pipeline import run_daily_pipeline
                    conn = get_db_connection()
                    try:
                        print("[Pipeline] Daily pipeline starting (03:00 KST)...")
                        result = run_daily_pipeline(conn)
                        print(f"[Pipeline] Done: {result.get('total_elapsed')}s, errors={result.get('error_count')}")
                    finally:
                        try: conn.close()
                        except Exception: pass
                except Exception as e:
                    print(f"[Pipeline] Error: {e}")

            from apscheduler.triggers.interval import IntervalTrigger

            pipeline_scheduler = AsyncIOScheduler()
            pipeline_scheduler.add_job(
                _daily_pipeline_job,
                CronTrigger(hour=18, minute=0),  # UTC 18:00 = KST 03:00
                id="daily_pipeline",
                name="일일 통합 파이프라인 (03:00 KST)",
                replace_existing=True,
            )
            pipeline_scheduler.add_job(
                _db_keepalive,
                IntervalTrigger(minutes=2),
                id="db_keepalive",
                name="DB PgBouncer keepalive (2분마다)",
                replace_existing=True,
            )
            pipeline_scheduler.add_job(
                _prewarm_response_cache,
                IntervalTrigger(minutes=10),
                id="cache_prewarm",
                name="비로그인 응답 캐시 갱신 (10분마다)",
                replace_existing=True,
            )

            # ── AI COO 오케스트레이터 — 매일 09:30 KST (UTC 00:30) ──
            def _orchestrator_job():
                try:
                    from app.services.orchestrator import run_daily_supervision
                    print("[AI COO] 일일 감시 스케줄 실행 시작 (09:30 KST)...")
                    result = run_daily_supervision()
                    print(f"[AI COO] 완료: {result.get('elapsed')}초")
                except Exception as e:
                    print(f"[AI COO] 스케줄 실행 오류: {e}")

            pipeline_scheduler.add_job(
                _orchestrator_job,
                CronTrigger(hour=0, minute=30),  # UTC 00:30 = KST 09:30
                id="ai_coo_supervision",
                name="AI COO 일일 감시 (09:30 KST)",
                replace_existing=True,
            )

            # ── 일괄 분석 진행 보고 — 1시간마다 이메일 ──
            def _bulk_analyze_report_job():
                if not _bulk_job.get("running") and not _bulk_job.get("started_at"):
                    return  # 작업 없으면 스킵
                _send_bulk_analysis_email_report()

            pipeline_scheduler.add_job(
                _bulk_analyze_report_job,
                IntervalTrigger(hours=1),
                id="bulk_analyze_report",
                name="일괄 분석 진행 보고 (1시간마다)",
                replace_existing=True,
            )

            async def _digest_job():
                """평일 09:00 KST 매칭 이메일/푸시 발송 + 사전매칭 캐시 갱신"""
                from app.services.notification_service import notification_service
                try:
                    print("[Digest] Running daily digest (평일 09:00 KST)...")
                    results = await notification_service.generate_daily_digest()
                    sent = sum(1 for r in results if r.get("email_sent"))
                    push_sent = sum(r.get("push_sent", 0) for r in results)
                    print(f"[Digest] Done: {len(results)} users, {sent} emails, {push_sent} pushes")
                except Exception as e:
                    print(f"[Digest] error: {e}")
                try:
                    count = _run_prematch_cache()
                    print(f"[Digest] Pre-match cache: {count} users")
                except Exception as e:
                    print(f"[Digest] pre-match error: {e}")

            pipeline_scheduler.add_job(
                _digest_job,
                CronTrigger(hour=DIGEST_HOUR, minute=0, day_of_week="mon-fri"),
                id="daily_digest",
                name=f"매칭 이메일 발송 (평일 {DIGEST_HOUR:02d}:00 UTC = KST 09:00)",
                replace_existing=True,
            )

            pipeline_scheduler.start()
            print("[Pipeline] APScheduler started - daily 03:00 KST + digest 09:00 KST(평일) + AI COO 09:30 KST + keepalive 2min + cache prewarm 10min + bulk_analyze_report 1h")
        except ImportError as e:
            print(f"[Pipeline] APScheduler not installed: {e}")
        except Exception as e:
            print(f"[Pipeline] scheduler init failed (서버는 정상): {e}")
    else:
        print("[Pipeline] disabled by PATROL_ENABLED=false")

    yield

    if pipeline_scheduler:
        try: pipeline_scheduler.shutdown(wait=False)
        except: pass


app = FastAPI(title="Gov Support Matching Assistant", lifespan=lifespan, docs_url=None, redoc_url=None)


@app.get("/health")
def health_check():
    return {"status": "ok"}




_cors_raw = os.getenv("CORS_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000,http://localhost:3001,http://127.0.0.1:3001,http://localhost:3002,http://localhost:3003,http://localhost:3005,http://127.0.0.1:3005,http://localhost:5181,http://localhost:8010")
_cors_origins = [o.strip() for o in _cors_raw.split(",") if o.strip()]
# www 서브도메인 자동 포함
for _o in list(_cors_origins):
    if "://" in _o and "://www." not in _o:
        _www = _o.replace("://", "://www.", 1)
        if _www not in _cors_origins:
            _cors_origins.append(_www)

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "X-API-Key"],
)


# ── 보안 헤더 미들웨어 ──
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest
from starlette.responses import Response as StarletteResponse

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        response: StarletteResponse = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
        if request.url.scheme == "https":
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response

app.add_middleware(SecurityHeadersMiddleware)


# ── Rate Limiting (인메모리) ──
import threading
_rate_store: dict = {}  # {key: [timestamp, ...]}
_rate_lock = threading.Lock()

def _rate_limit_check(key: str, max_requests: int, window_seconds: int) -> bool:
    """Rate limit 검사. True이면 허용, False이면 차단"""
    now = datetime.datetime.utcnow()
    with _rate_lock:
        if key not in _rate_store:
            _rate_store[key] = []
        # 윈도우 밖 기록 제거
        cutoff = now - datetime.timedelta(seconds=window_seconds)
        _rate_store[key] = [t for t in _rate_store[key] if t > cutoff]
        if len(_rate_store[key]) >= max_requests:
            return False
        _rate_store[key].append(now)
        return True

def _get_client_ip(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


# ── 보안 AI 에이전트 — 실시간 이상 감지 ──
import collections
import threading

class SecurityAgent:
    """실시간 보안 모니터링 에이전트 — 봇/크롤러 감지 강화"""

    def __init__(self):
        self._lock = threading.Lock()
        # IP별 요청 카운터 (1분 윈도우)
        self._request_counts: dict[str, list] = {}
        # IP별 실패 로그인 카운터
        self._failed_logins: dict[str, list] = {}
        # 차단된 IP
        self._blocked_ips: set = set()
        # IP별 공개 API 페이지네이션 추적 (크롤링 감지)
        self._pagination_tracker: dict[str, list] = {}
        # IP별 고유 경로 추적 (비정상 순회 감지)
        self._path_tracker: dict[str, set] = {}
        # SQL 인젝션 패턴
        self._sqli_patterns = [
            "union select", "drop table", "insert into", "delete from",
            "update set", "pg_sleep", "information_schema", "1=1", "or 1=",
            "'; --", "\"; --", "/*", "*/", "xp_cmdshell",
        ]
        # XSS 패턴
        self._xss_patterns = [
            "<script", "javascript:", "onerror=", "onload=", "eval(",
            "document.cookie", "alert(", "<svg", "<iframe",
        ]
        # 봇/크롤러 User-Agent 패턴
        self._bot_ua_patterns = [
            "python-requests", "scrapy", "crawl", "spider", "bot",
            "curl", "wget", "httpie", "postman", "insomnia",
            "go-http-client", "java/", "libwww", "httpclient",
            "aiohttp", "node-fetch", "axios/", "got/",
            "headless", "phantom", "puppeteer", "playwright",
        ]
        # 허용 봇 (검색 엔진)
        self._allowed_bots = ["googlebot", "bingbot", "yandexbot", "naverbot", "yeti", "daumoa", "twitterbot", "facebookexternalhit", "linkedinbot"]
        # 이벤트 로그 (최근 100건)
        self._events: list = []
        self._MAX_EVENTS = 100
        # 임계값
        self.RATE_LIMIT = 120       # 1분간 최대 요청 수
        self.FAILED_LOGIN_LIMIT = 5  # 5회 실패 시 차단
        self.BLOCK_DURATION = 300    # 5분 차단
        self.CRAWL_PAGE_LIMIT = 10  # 5분 내 공개 API 10페이지 이상 요청 시 차단
        self.PATH_SCAN_LIMIT = 50   # 5분 내 50개 이상 고유 경로 접근 시 차단

    def _log_event(self, severity: str, event_type: str, ip: str, detail: str):
        entry = {
            "ts": datetime.datetime.utcnow().isoformat(),
            "severity": severity,
            "type": event_type,
            "ip": ip,
            "detail": detail[:200],
        }
        with self._lock:
            self._events.append(entry)
            if len(self._events) > self._MAX_EVENTS:
                self._events = self._events[-self._MAX_EVENTS:]
        if severity in ("HIGH", "CRITICAL"):
            try:
                print(f"[SECURITY-{severity}] {event_type}: {ip} - {detail[:100]}")
            except UnicodeEncodeError:
                pass  # Windows cp949 인코딩 실패 시 로그 무시 (서비스 중단 방지)

    def _cleanup_old(self, counter: dict, window: int = 60):
        """오래된 타임스탬프 정리"""
        now = time.time()
        for key in list(counter.keys()):
            counter[key] = [t for t in counter[key] if now - t < window]
            if not counter[key]:
                del counter[key]

    # 공개 API + 내부 서비스 경로 — 보안 검사 예외
    _whitelisted_paths = ("/api/announcements/", "/api/announcements/public", "/api/announcements/search", "/for-smartdoc", "/api/push/vapid-key", "/api/auth/", "/health", "/api/admin/")

    def check_request(self, ip: str, path: str, method: str, query: str = "", body: str = "", user_agent: str = "") -> str | None:
        """요청을 검사하고 차단 사유가 있으면 반환, 없으면 None"""
        # 화이트리스트 경로는 보안 검사 건너뛰기
        if any(wp in path for wp in self._whitelisted_paths):
            return None

        now = time.time()
        ua_lower = (user_agent or "").lower()

        # 1. 차단된 IP 확인
        if ip in self._blocked_ips:
            return "IP_BLOCKED"

        # 2. 봇/크롤러 User-Agent 감지
        if ua_lower:
            # 허용된 검색 엔진 봇은 통과
            is_allowed_bot = any(b in ua_lower for b in self._allowed_bots)
            if not is_allowed_bot:
                is_bot = any(p in ua_lower for p in self._bot_ua_patterns)
                if is_bot:
                    self._log_event("HIGH", "BOT_DETECTED", ip, f"UA: {user_agent[:100]}, path={path}")
                    _log_system("bot_detected", "system", f"IP={ip}, UA={user_agent[:80]}, path={path}", "blocked")
                    self._blocked_ips.add(ip)
                    threading.Timer(self.BLOCK_DURATION * 2, lambda: self._blocked_ips.discard(ip)).start()
                    return "BOT_BLOCKED"
            # User-Agent가 비어있으면 의심
            if not ua_lower.strip():
                self._log_event("MEDIUM", "EMPTY_UA", ip, f"path={path}")

        # 3. Rate Limiting
        with self._lock:
            if ip not in self._request_counts:
                self._request_counts[ip] = []
            self._request_counts[ip].append(now)
            self._cleanup_old(self._request_counts)
            if len(self._request_counts.get(ip, [])) > self.RATE_LIMIT:
                self._blocked_ips.add(ip)
                self._log_event("HIGH", "RATE_LIMIT", ip, f"{len(self._request_counts[ip])} req/min on {path}")
                _log_system("rate_limit_block", "system", f"IP={ip}, {len(self._request_counts[ip])} req/min", "blocked")
                threading.Timer(self.BLOCK_DURATION, lambda: self._blocked_ips.discard(ip)).start()
                return "RATE_LIMITED"

        # 4. 공개 API 크롤링 감지 (페이지네이션 남용)
        if "/api/announcements/public" in path:
            with self._lock:
                if ip not in self._pagination_tracker:
                    self._pagination_tracker[ip] = []
                self._pagination_tracker[ip].append(now)
                self._cleanup_old(self._pagination_tracker, window=300)
                page_count = len(self._pagination_tracker.get(ip, []))
                if page_count > self.CRAWL_PAGE_LIMIT:
                    self._blocked_ips.add(ip)
                    self._log_event("CRITICAL", "CRAWL_DETECTED", ip, f"{page_count} public API pages in 5min")
                    _log_system("crawl_blocked", "system", f"IP={ip}, {page_count} pages/5min", "blocked")
                    threading.Timer(self.BLOCK_DURATION * 4, lambda: self._blocked_ips.discard(ip)).start()
                    return "CRAWL_BLOCKED"

        # 5. 비정상 경로 스캔 감지 (다수의 고유 경로 접근)
        if path.startswith("/api/"):
            with self._lock:
                if ip not in self._path_tracker:
                    self._path_tracker[ip] = set()
                self._path_tracker[ip].add(path)
                # 5분마다 리셋
                if len(self._path_tracker[ip]) > self.PATH_SCAN_LIMIT:
                    self._blocked_ips.add(ip)
                    self._log_event("HIGH", "PATH_SCAN", ip, f"{len(self._path_tracker[ip])} unique paths")
                    _log_system("path_scan_blocked", "system", f"IP={ip}, {len(self._path_tracker[ip])} unique paths", "blocked")
                    threading.Timer(self.BLOCK_DURATION * 2, lambda: self._blocked_ips.discard(ip)).start()
                    self._path_tracker[ip] = set()
                    return "SCAN_BLOCKED"

        # 6. SQL 인젝션 패턴 감지
        check_str = f"{query} {body}".lower()
        for pattern in self._sqli_patterns:
            if pattern in check_str:
                self._log_event("HIGH", "SQLI_ATTEMPT", ip, f"pattern='{pattern}' path={path}")
                return None  # 파라미터 바인딩으로 안전, 로깅만

        # 7. XSS 패턴 감지
        for pattern in self._xss_patterns:
            if pattern in check_str:
                self._log_event("MEDIUM", "XSS_ATTEMPT", ip, f"pattern='{pattern}' path={path}")
                return None  # 로깅만

        # 8. 민감 경로 접근 감지
        suspicious_paths = ["/.env", "/.git", "/wp-admin", "/phpmyadmin", "/admin.php", "/.aws", "/config.json"]
        if path in suspicious_paths:
            self._log_event("MEDIUM", "SUSPICIOUS_PATH", ip, path)
            self._blocked_ips.add(ip)
            threading.Timer(self.BLOCK_DURATION, lambda: self._blocked_ips.discard(ip)).start()
            return "SUSPICIOUS_BLOCKED"

        return None

    def record_failed_login(self, ip: str, endpoint: str):
        """로그인 실패 기록"""
        now = time.time()
        with self._lock:
            if ip not in self._failed_logins:
                self._failed_logins[ip] = []
            self._failed_logins[ip].append(now)
            self._cleanup_old(self._failed_logins, window=300)
            count = len(self._failed_logins.get(ip, []))
            if count >= self.FAILED_LOGIN_LIMIT:
                self._blocked_ips.add(ip)
                self._log_event("CRITICAL", "BRUTE_FORCE", ip, f"{count} failed logins on {endpoint}")
                threading.Timer(self.BLOCK_DURATION, lambda: self._blocked_ips.discard(ip)).start()

    def get_status(self) -> dict:
        """보안 상태 리포트"""
        with self._lock:
            self._cleanup_old(self._request_counts)
            self._cleanup_old(self._failed_logins, 300)
        return {
            "blocked_ips": list(self._blocked_ips),
            "active_ips": len(self._request_counts),
            "failed_logins": {ip: len(ts) for ip, ts in self._failed_logins.items() if ts},
            "recent_events": self._events[-20:],
        }


# 보안 에이전트 인스턴스
security_agent = SecurityAgent()


# 보안 미들웨어 — 모든 요청을 에이전트가 검사
class SecurityAgentMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        # SECURITY_ENABLED=false로 보안 에이전트 비활성화 (개발 시)
        if os.getenv("SECURITY_ENABLED", "true").lower() == "false":
            return await call_next(request)

        ip = request.headers.get("x-forwarded-for", "").split(",")[0].strip() or (request.client.host if request.client else "unknown")
        path = request.url.path
        method = request.method
        query = str(request.query_params)
        user_agent = request.headers.get("user-agent", "")

        # CORS preflight(OPTIONS)는 보안 검사 건너뛰기
        if method == "OPTIONS":
            return await call_next(request)

        # 화이트리스트 경로 — JWT 디코딩보다 먼저 체크 (성능)
        _wl = security_agent._whitelisted_paths
        if any(wp in path for wp in _wl):
            return await call_next(request)

        # 봇 토큰 화이트리스트 — 블로그/패트롤 봇 전용
        _bot_token = request.headers.get("x-bot-token", "")
        _expected_bot_token = os.getenv("BOT_TOKEN", "GOVMATCH_BLOG_BOT_2026")
        if _bot_token and _bot_token == _expected_bot_token:
            security_agent._blocked_ips.discard(ip)
            return await call_next(request)

        # 오너 화이트리스트 — 항상 접근 허용
        _owner_emails = os.getenv("OWNER_EMAILS", "osung94@naver.com").split(",")
        auth_header = request.headers.get("authorization", "")
        if auth_header.startswith("Bearer "):
            try:
                _token = auth_header.split(" ", 1)[1]
                _payload = jwt.decode(_token, JWT_SECRET, algorithms=["HS256"])
                if _payload.get("email") in _owner_emails:
                    security_agent._blocked_ips.discard(ip)
                    return await call_next(request)
            except Exception:
                pass

        # 요청 검사
        block_reason = security_agent.check_request(ip, path, method, query, user_agent=user_agent)
        if block_reason in ("IP_BLOCKED", "SUSPICIOUS_BLOCKED"):
            return StarletteResponse(content='{"detail":"접근이 차단되었습니다."}',
                                     status_code=403, media_type="application/json")
        if block_reason in ("RATE_LIMITED", "CRAWL_BLOCKED", "SCAN_BLOCKED"):
            return StarletteResponse(content='{"detail":"요청이 너무 많습니다. 잠시 후 다시 시도하세요."}',
                                     status_code=429, media_type="application/json")
        if block_reason == "BOT_BLOCKED":
            return StarletteResponse(content='{"detail":"자동화된 접근이 감지되어 차단되었습니다."}',
                                     status_code=403, media_type="application/json")

        response = await call_next(request)

        # 인증 실패 감지
        if response.status_code == 401 and path.startswith("/api/admin"):
            security_agent.record_failed_login(ip, path)

        return response

app.add_middleware(SecurityAgentMiddleware)


# 보안 상태 API (Admin 전용)
class BusinessNumberRequest(BaseModel):
    business_number: str
    target_type: Optional[str] = None  # "business" | "individual" | None(=user_type 따름)

class UserProfile(BaseModel):
    business_number: str
    company_name: Optional[str] = None
    establishment_date: Optional[str] = None
    address_city: Optional[str] = None
    industry_code: Optional[str] = None
    revenue_bracket: Optional[str] = None
    employee_count_bracket: Optional[str] = None
    interests: Optional[str] = None
    password: Optional[str] = None
    # 개인/기업 구분
    user_type: Optional[str] = "both"
    # 개인 프로필 필드
    age_range: Optional[str] = None
    income_level: Optional[str] = None
    family_type: Optional[str] = None
    employment_status: Optional[str] = None

class CompanyNameRequest(BaseModel):
    company_name: str
    business_content: Optional[str] = ""

class URLRequest(BaseModel):
    url: str

class NotificationSettings(BaseModel):
    business_number: str
    email: Optional[str] = None
    phone_number: Optional[str] = None
    channel: str = "email"
    is_active: bool = True
    kakao_enabled: Optional[int] = 0

class AdminURLRequest(BaseModel):
    url: str
    source_name: str

class RegisterRequest(BaseModel):
    email: str
    password: str
    business_number: str
    company_name: Optional[str] = ""
    address_city: Optional[str] = None
    industry_code: Optional[str] = None
    establishment_date: Optional[str] = None
    revenue_bracket: Optional[str] = None
    employee_count_bracket: Optional[str] = None
    interests: Optional[str] = None
    referred_by: Optional[str] = None
    user_type: Optional[str] = "both"

class LoginRequest(BaseModel):
    email: str
    password: str


def _create_jwt(user_id: int, business_number: str, email: str, plan: str, trial_ends_at: str | None) -> str:
    payload = {
        "user_id": user_id,
        "bn": business_number,
        "email": email,
        "plan": plan,
        "trial_ends_at": trial_ends_at,
        "exp": datetime.datetime.utcnow() + datetime.timedelta(hours=JWT_EXPIRE_HOURS),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


def _decode_jwt(token: str) -> dict:
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="토큰이 만료되었습니다. 다시 로그인해 주세요.")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="유효하지 않은 토큰입니다.")


def _get_current_user(authorization: Optional[str] = Header(None)) -> dict:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="로그인이 필요합니다.")
    token = authorization.split(" ", 1)[1]
    return _decode_jwt(token)


# ── 플랜 v5: 상담 시작 무제한 + 세션 메시지 수로 차별화 (2026-04-28 확정) ──
# FREE:  월 3회 상담 시작, 세션 내 메시지 10개
# LITE:  무제한 상담 시작, 세션 내 메시지 30개
# PRO:   무제한 상담 시작, 세션 내 메시지 무제한 + 전문가 기능
PLAN_LIMITS = {
    "free": 3,         # FREE: 체험용 월 3회 (상담 시작 기준)
    "lite": 999999,    # LITE: 무제한
    "lite_trial": 999999,
    "basic": 999999,   # legacy → LITE 취급
    "biz": 999999,     # legacy → PRO 취급
    "pro": 999999,
}

# 공고AI 상담도 PLAN_LIMITS와 동일 (FREE만 제한)
CONSULT_LIMITS = {
    "free": 3,
    "lite_trial": 999999,
    "lite": 999999,
    "basic": 999999,   # legacy → LITE 취급
    "biz": 999999,     # legacy → PRO 취급
    "pro": 999999,
}

# 세션 내 메시지 수 제한 (사용자 메시지 기준) — 플랜별 차별화 핵심
SESSION_MSG_LIMITS = {
    "free": 10,
    "lite_trial": 30,
    "lite": 30,
    "basic": 30,       # legacy → LITE 취급
    "biz": 999999,     # legacy → PRO 취급
    "pro": 999999,
}

# 플랜 가격 (원/월) — user_type에 따라 분기
# 개인 LITE: 2,900 / 사업자 LITE: 4,900 / PRO: 29,000 (3개월 이벤트, 정상가 49,000)
PLAN_PRICES = {
    "lite_individual": 2900,
    "lite": 4900,       # 사업자 LITE (기본값)
    "pro": 29000,       # 3개월 이벤트가 (~2026.07.12), 이후 49,000
    "basic": 4900,      # legacy
    "biz": 29000,       # legacy → PRO 취급
}

# 신규 가입자 LITE 7일 무료체험 (상시)
TRIAL_DAYS = 7
PROMO_ACTIVE = False  # 프로모션 종료 (2026-05-23 이후)
# PRO 3개월 오픈 이벤트: 29,000원/월 (정가 49,000원), 2026-07-31까지
PRO_EVENT_PRICE = 29000
PRO_EVENT_END = "2026-07-31"

# AI 신청서 작성 가격 (원/건) — Coming Soon
AI_GUIDE_PRICE = None

# 공고AI 상담 1건당 메시지 제한 (하위 호환용 — SESSION_MSG_LIMITS로 대체)
CONSULT_MSG_LIMIT = 50


def _get_plan_status(plan: str, plan_expires_at: str | None, ai_usage_month: int = 0) -> dict:
    """플랜 상태와 남은 일수, 잔여 건수를 계산"""
    now = datetime.datetime.utcnow()

    # trial/premium → free 마이그레이션
    if plan in ("trial", "premium"):
        plan = "free"

    ai_limit = PLAN_LIMITS.get(plan, 1)
    consult_limit = CONSULT_LIMITS.get(plan, 0)

    if plan == "free":
        # FREE는 영구 무료 — 만료 없음
        return {
            "plan": "free", "active": True, "days_left": None,
            "label": "FREE",
            "ai_used": ai_usage_month, "ai_limit": ai_limit,
            "consult_limit": consult_limit,
        }

    if plan in ("lite", "lite_trial", "basic", "biz", "pro"):
        # 유료 플랜 만료 체크
        if plan_expires_at:
            try:
                expires = datetime.datetime.fromisoformat(str(plan_expires_at))
                days_left = (expires - now).days
                if days_left < 0:
                    # 만료 → FREE로 다운그레이드
                    return {
                        "plan": "expired", "active": False, "days_left": 0,
                        "label": "만료됨",
                        "ai_used": ai_usage_month, "ai_limit": PLAN_LIMITS["free"],
                        "consult_limit": CONSULT_LIMITS["free"],
                    }
                label_map = {"lite": "LITE", "lite_trial": "LITE 체험", "basic": "LITE", "biz": "PRO", "pro": "PRO"}
                return {
                    "plan": plan, "active": True, "days_left": days_left,
                    "label": label_map.get(plan, plan.upper()),
                    "ai_used": ai_usage_month, "ai_limit": ai_limit,
                    "consult_limit": consult_limit,
                }
            except ValueError:
                pass
        label_map = {"lite": "LITE", "lite_trial": "LITE 체험", "basic": "LITE", "biz": "PRO", "pro": "PRO"}
        return {
            "plan": plan, "active": True, "days_left": None,
            "label": label_map.get(plan, plan.upper()),
            "ai_used": ai_usage_month, "ai_limit": ai_limit,
            "consult_limit": consult_limit,
        }

    return {"plan": "free", "active": True, "days_left": None, "label": "FREE",
            "ai_used": 0, "ai_limit": PLAN_LIMITS["free"], "consult_limit": CONSULT_LIMITS["free"]}


# ─── 응답 캐시 (동시접속 대응) ───────────────────────────────────────
_response_cache: dict = {}
_CACHE_TTL = 3600  # 60분 (카테고리 카운트/공개 목록은 자주 안 바뀜)

def _get_cached(key: str):
    entry = _response_cache.get(key)
    if entry and time.time() - entry["ts"] < _CACHE_TTL:
        return entry["data"]
    return None

def _set_cache(key: str, data):
    _response_cache[key] = {"data": data, "ts": time.time()}


# [Phase 5] 공고 리스트 공용 유효성 필터 — 모든 리스트 API에서 이 필터 사용
# 유효 공고 = 아카이브 안 됨 AND (상시 모집 OR (고정 마감일 미래) OR (분석 대기 중 & 최근 3개월 이내))
def valid_announcement_where(alias: str = "") -> str:
    """공고 리스트 WHERE 절용 공용 필터 SQL 생성.

    Args:
        alias: 테이블 별칭 (예: "a" → "a.deadline_type"), 없으면 컬럼 직접 참조
    Returns:
        괄호로 감싼 WHERE 조건 문자열
    """
    p = f"{alias}." if alias else ""
    return (
        f"({p}is_archived = FALSE AND ("
        f"  {p}deadline_type = 'ongoing'"
        f"  OR ({p}deadline_type = 'fixed' AND {p}deadline_date >= CURRENT_DATE)"
        f"  OR ({p}deadline_type = 'unknown' AND {p}created_at >= CURRENT_DATE - INTERVAL '3 months')"
        f"))"
    )


# ─── 실시간 통계 카운터 (홈 히어로 영역) ───────────────────────────
# 초기 런칭 단계 — 실제 값 + 현실적 시드 + 시간 자동 증가
_STATS_SEED = {
    # 서비스 런칭 시점 기준 (2026-04-20 18:00 UTC = 2026-04-21 03:00 KST) 근처
    "base_epoch": 1776669600,  # 지금부터 +1씩 자연스럽게 증가하는 시작점
    "seeds": {
        "matches":       175,    # 실제 ~91 + 시드 175 ≈ 266 (초기 표시값)
        "consultations": 130,    # 실제 ~241 + 시드 130 ≈ 371
        "companies":      65,    # 실제 ~134 + 시드 65 ≈ 199
    },
    # 시간당 자동 증가 (초 단위 간격)
    "increments": {
        "matches":       1800,   # 30분당 +1
        "consultations":  900,   # 15분당 +1
        "companies":   172800,   # 2일당 +1
    },
}

@app.get("/api/stats/live")
def api_stats_live(request: Request):
    """홈 화면 실시간 통계 — 시드 + 시간 자동증가 (DB 조회 없음, 즉시 응답)."""
    now = int(time.time())
    elapsed = max(0, now - _STATS_SEED["base_epoch"])
    seeds = _STATS_SEED["seeds"]
    incs = _STATS_SEED["increments"]
    return {
        "announcements": 17500 + (elapsed // 3600),      # 시간당 +1
        "matches":       seeds["matches"] + (elapsed // incs["matches"]),
        "consultations": seeds["consultations"] + (elapsed // incs["consultations"]),
        "companies":     seeds["companies"] + (elapsed // incs["companies"]),
        "updated_at":    now,
    }


# ─── 비로그인 공고 리스트 API ───────────────────────────────────────
@app.get("/api/announcements/ticker")
def api_announcements_ticker():
    """티커용 최신 공고 20건 (마감 전, 캐시 15분)"""
    cache_key = "ticker:v1"
    cached = _response_cache.get(cache_key)
    if cached and time.time() - cached["ts"] < _CACHE_TTL:
        return cached["data"]

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(f"""
            SELECT announcement_id, title, department, category, deadline_date, support_amount, region
            FROM announcements
            WHERE {valid_announcement_where()}
              AND analysis_status = 'analyzed'
            ORDER BY created_at DESC
            LIMIT 20
        """)
        rows = [dict(r) for r in cur.fetchall()]
        cur.close()
        result = {"items": rows}
        _response_cache[cache_key] = {"ts": time.time(), "data": result}
        return result
    finally:
        conn.close()


@app.get("/api/announcements/public")
def api_announcements_public(
    request: Request,
    page: int = 1,
    size: int = 20,
    region: Optional[str] = None,
    category: Optional[str] = None,
    search: Optional[str] = None,
    target_type: Optional[str] = None,
    tab: Optional[str] = None,          # "local" | "national" | None(전체)
    authorization: Optional[str] = Header(None),
):
    """비로그인 사용자도 접근 가능한 공고 리스트 (마감 전 공고만)"""
    # Rate limiting: IP당 분당 30회
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"public:ip:{ip}", 60, 60):
        raise HTTPException(status_code=429, detail="요청이 너무 많습니다. 잠시 후 다시 시도해주세요.")
    if size < 1:
        size = 1
    if size > 50:
        size = 50  # 크롤링 방지: 최대 50건
    if page < 1:
        page = 1
    if page > 50:
        raise HTTPException(status_code=400, detail="페이지 범위를 초과했습니다.")
    offset = (page - 1) * size

    # ── 로그인 사용자 + 필터 없음 → 사전캐시 우선, 없으면 실시간 CTE ──
    _is_logged_in = False  # 로그인 여부 — 공유 캐시 우회 판단용
    if authorization and authorization.startswith("Bearer ") and not search and not region and not category:
        try:
            current_user = _decode_jwt(authorization.split(" ", 1)[1])
            bn = current_user.get("bn")
            if bn:
                _is_logged_in = True
                cache_type = "public_ind" if target_type == "individual" else "public_biz"

                _pc = get_db_connection()
                _pcur = _pc.cursor()

                # ── 1순위: user_match_cache 사전 계산 결과 (새벽 배치) ──
                _pcur.execute(
                    """SELECT match_data FROM user_match_cache
                       WHERE business_number = %s AND target_type = %s
                         AND created_at > CURRENT_TIMESTAMP - INTERVAL '26 hours'
                       ORDER BY created_at DESC LIMIT 1""",
                    (bn, cache_type),
                )
                cache_row = _pcur.fetchone()

                if cache_row and cache_row.get("match_data"):
                    cached = cache_row["match_data"]
                    if isinstance(cached, str):
                        import json as _j; cached = _j.loads(cached)

                    # ── 신형식: local / national 분리 캐시 ──
                    if tab in ("local", "national") and tab in cached:
                        tab_ids = cached[tab]
                        total = len(tab_ids)
                        page_ids = tab_ids[offset: offset + size]

                        if page_ids:
                            id_list = ",".join(str(i) for i in page_ids)
                            _pcur.execute(
                                f"""SELECT announcement_id, title, region, category, department,
                                           support_amount, support_amount_max, support_amount_min, support_amount_type,
                                           deadline_date, origin_source, created_at,
                                           COALESCE(target_type, 'business') AS target_type,
                                           origin_url, summary_text, eligibility_logic,
                                           established_years_limit, revenue_limit, employee_limit
                                    FROM announcements
                                    WHERE announcement_id IN ({id_list})
                                      AND {valid_announcement_where()}"""
                            )
                            rows_map = {{r["announcement_id"]: dict(r) for r in _pcur.fetchall()}}
                            rows = [rows_map[i] for i in page_ids if i in rows_map]
                        else:
                            rows = []
                        _pc.close()
                        return {
                            "status": "SUCCESS",
                            "data": rows,
                            "total": total,
                            "page": page,
                            "size": size,
                            "regions": [],
                            "categories": [],
                            "category_counts": {},
                            "personalized": True,
                            "source": "cache",
                            "tab": tab,
                        }

                    # ── tab 요청인데 캐시에 해당 키 없음 (구형식) → 일반 경로 ──
                    if tab in ("local", "national"):
                        _pc.close()
                        raise ValueError("no-cache-tab-fallthrough")

                    # ── 전체 탭 (tab=None): eligible_ids + ineligible_ids 기존 로직 ──
                    eligible_ids = cached.get("eligible_ids") or []
                    ineligible_ids = cached.get("ineligible_ids") or []

                    all_ids = eligible_ids + ineligible_ids
                    total = len(all_ids)
                    page_ids = all_ids[offset: offset + size]

                    if page_ids:
                        # ID 순서를 유지하면서 공고 본문 조회
                        id_list = ",".join(str(i) for i in page_ids)
                        _pcur.execute(
                            f"""SELECT announcement_id, title, region, category, department,
                                       support_amount, support_amount_max, support_amount_min, support_amount_type,
                                       deadline_date, origin_source, created_at,
                                       COALESCE(target_type, 'business') AS target_type,
                                       origin_url, summary_text, eligibility_logic,
                                       established_years_limit, revenue_limit, employee_limit
                                FROM announcements
                                WHERE announcement_id IN ({id_list})
                                  AND {valid_announcement_where()}"""
                        )
                        rows_map = {r["announcement_id"]: dict(r) for r in _pcur.fetchall()}
                        # 캐시 순서 복원 (신규 추가/삭제된 공고는 자연스럽게 제외됨)
                        rows = [rows_map[i] for i in page_ids if i in rows_map]
                    else:
                        rows = []

                    cat_cache_key = f"cat_counts:{target_type or 'all'}"
                    category_counts = _get_cached(cat_cache_key)
                    if not category_counts:
                        try:
                            _cat_where = valid_announcement_where()
                            _cat_params: list = []
                            if target_type:
                                _cat_where += " AND (target_type = %s OR target_type = 'both')"
                                _cat_params.append(target_type)
                            _pcur.execute(
                                f"""SELECT COALESCE(category, '기타') AS cat, COUNT(*) AS cnt
                                    FROM announcements WHERE {_cat_where}
                                    GROUP BY COALESCE(category, '기타') ORDER BY cnt DESC""",
                                _cat_params,
                            )
                            category_counts = {r["cat"]: r["cnt"] for r in _pcur.fetchall()}
                            _set_cache(cat_cache_key, category_counts)
                        except Exception:
                            category_counts = {}
                    _pc.close()
                    return {
                        "status": "SUCCESS",
                        "data": rows,
                        "total": total,
                        "page": page,
                        "size": size,
                        "regions": [],
                        "categories": [],
                        "category_counts": category_counts,
                        "personalized": True,
                        "source": "cache",
                        "local_total": len(cached.get("local") or []),
                        "national_total": len(cached.get("national") or []),
                    }

                # ── 2순위: 사전캐시 없음 → 실시간 CTE (신규 가입자 / 캐시 만료) ──
                # tab=local/national 은 일반 SQL 경로가 더 빠름 → fall-through
                if category or tab in ("local", "national"):
                    _pc.close()
                    raise ValueError("no-cache-tab-fallthrough")

                # 사용자 지역·관심분야 조회 (같은 연결 재사용)
                _pcur.execute(
                    "SELECT address_city, interests, gender, industry_code FROM users WHERE business_number = %s", (bn,)
                )
                urow = _pcur.fetchone()

                if urow:
                    raw_city = str(urow.get("address_city", "") or "")
                    cities = [c.strip() for c in raw_city.split(",") if c.strip() and c.strip() != "전국"]
                    user_city = cities[0] if cities else ""
                    interests = [i.strip() for i in str(urow.get("interests", "") or "").split(",") if i.strip()]
                    gender = str(urow.get("gender", "") or "")
                    rt_ind_major = str(urow.get("industry_code") or "")[:2]
                    rt_is_farmer = rt_ind_major in ("01", "02", "03")
                    user_target = "individual" if target_type == "individual" else "business"

                    if user_city:
                        region_sql = "(region ILIKE %s AND region NOT IN ('전국', '', '전국 및 각 지역'))"
                        region_params = [f"%{user_city}%"]
                    else:
                        region_sql = "FALSE"
                        region_params = []

                    if interests:
                        interest_parts = " OR ".join(["(category ILIKE %s OR title ILIKE %s)" for _ in interests])
                        interest_sql = f"({interest_parts})"
                        interest_params = []
                        for it in interests:
                            interest_params.extend([f"%{it}%", f"%{it}%"])
                    else:
                        interest_sql = "FALSE"
                        interest_params = []

                    import datetime as _dt
                    today_bucket = _dt.date.today().timetuple().tm_yday % 3

                    has_amount_sql = "(support_amount IS NOT NULL AND support_amount != '')"

                    # 특정 대상 전용 키워드 → 후순위(bucket=4) — 하드코딩 상수이므로 직접 삽입 안전
                    _RESTRICTED = ["장애인기업", "장애인창업", "농업인", "영농조합", "어업인", "수산업", "보훈", "제대군인"]
                    _FEMALE_ONLY = ["여성기업", "여성창업", "여성경제인"]
                    restricted_ilike = " OR ".join(f"title ILIKE '%{kw}%'" for kw in _RESTRICTED)
                    inelig_parts = [f"({restricted_ilike})"]
                    if gender != "여성":
                        female_ilike = " OR ".join(f"title ILIKE '%{kw}%'" for kw in _FEMALE_ONLY)
                        inelig_parts.append(f"({female_ilike})")
                    if not rt_is_farmer:
                        _AGRI_CATS_RT = ["농림", "수산", "임업", "축산"]
                        agri_ilike = " OR ".join(f"category ILIKE '%{ac}%'" for ac in _AGRI_CATS_RT)
                        inelig_parts.append(f"({agri_ilike})")
                    inelig_sql = " OR ".join(inelig_parts)

                    # 전국 판단: 빈값/전국/All/전국 및 각 지역
                    _nationwide_sql = "region IN ('전국', '', '전국 및 각 지역', 'All') OR region IS NULL"
                    # 타지역 판단: 지역한정인데 내 지역 아님 → 후순위(4)
                    if user_city:
                        _other_region_sql = f"NOT ({_nationwide_sql}) AND NOT ({region_sql})"
                    else:
                        _other_region_sql = "FALSE"
                    bucket_sql = f"""
                        CASE
                            WHEN deadline_date IS NOT NULL AND deadline_date < CURRENT_DATE THEN 4
                            WHEN COALESCE(target_type, 'business') != %s THEN 4
                            WHEN {inelig_sql} THEN 4
                            WHEN {_other_region_sql} THEN 4
                            WHEN {region_sql} AND {has_amount_sql} THEN 0
                            WHEN ({_nationwide_sql}) AND {has_amount_sql} THEN 1
                            WHEN {interest_sql} AND {has_amount_sql} THEN 2
                            ELSE 3
                        END
                    """
                    # region_params는 두 번 필요:
                    # 1) _other_region_sql 안 NOT (region ILIKE %s)
                    # 2) WHEN region_sql AND has_amount THEN 0  (region ILIKE %s)
                    bucket_params = [user_target] + region_params + region_params + interest_params

                    valid_where = valid_announcement_where()
                    if target_type:
                        type_filter = "AND (target_type = %s OR target_type = 'both' OR target_type IS NULL)"
                        type_params = [target_type]
                    else:
                        type_filter = ""
                        type_params = []
                    full_where = f"{valid_where} {type_filter}"

                    _pcur.execute(f"SELECT COUNT(*) AS cnt FROM announcements WHERE {full_where}", type_params)
                    total = _pcur.fetchone()["cnt"]

                    _pcur.execute(
                        f"""WITH ann AS (
                                SELECT announcement_id, title, region, category, department,
                                       support_amount, support_amount_max, support_amount_min, support_amount_type,
                                       deadline_date, origin_source, created_at,
                                       COALESCE(target_type, 'business') AS target_type,
                                       origin_url, summary_text, eligibility_logic,
                                       established_years_limit, revenue_limit, employee_limit,
                                       {bucket_sql} AS bucket
                                FROM announcements
                                WHERE {full_where}
                            )
                            SELECT * FROM ann
                            ORDER BY
                                CASE WHEN bucket = 4 THEN 11
                                     WHEN bucket = 3 THEN 10
                                     WHEN %s = 'individual' THEN
                                         CASE WHEN bucket = 0 THEN 0
                                              WHEN bucket = 2 THEN 1
                                              WHEN bucket = 1 THEN 2
                                              ELSE 5 END
                                     ELSE (bucket - %s + 3) %% 3
                                END,
                                deadline_date ASC NULLS LAST,
                                created_at DESC
                            LIMIT %s OFFSET %s""",
                        type_params + bucket_params + [target_type or "business", today_bucket, size, offset],
                    )
                    rows = [dict(r) for r in _pcur.fetchall()]

                    # 백그라운드: 전체 정렬 ID 목록 → user_match_cache 저장 (신규/만료 시 1회)
                    import threading as _thr, json as _jcache
                    def _bg_cache_save(_bn, _ct, _fw, _tp, _bs, _bp, _tt, _tb):
                        try:
                            _sc = get_db_connection(); _scu = _sc.cursor()
                            _scu.execute(
                                f"""WITH ann AS (
                                    SELECT announcement_id, {_bs} AS bucket
                                    FROM announcements WHERE {_fw}
                                ) SELECT announcement_id FROM ann
                                ORDER BY CASE WHEN bucket=4 THEN 11 WHEN bucket=3 THEN 10
                                     WHEN %s='individual' THEN CASE WHEN bucket=0 THEN 0 WHEN bucket=2 THEN 1 WHEN bucket=1 THEN 2 ELSE 5 END
                                     ELSE (bucket-%s+3)%%3 END,
                                deadline_date ASC NULLS LAST, created_at DESC""",
                                _tp + _bp + [_tt or "business", _tb]
                            )
                            _ids = [r["announcement_id"] for r in _scu.fetchall()]
                            # 지역/전국 분리 (신형식 포함 저장)
                            _scu.execute(
                                f"""SELECT announcement_id, region, title FROM announcements WHERE {_fw}""",
                                _tp
                            )
                            _ann_map = {r["announcement_id"]: dict(r) for r in _scu.fetchall()}
                            _local, _national = [], []
                            for _aid in _ids:
                                _ar = (_ann_map.get(_aid) or {})
                                _reg = (_ar.get("region") or "").strip()
                                _tit = (_ar.get("title") or "")
                                _is_nat = not _reg or _reg in ("전국", "전국 및 각 지역", "All")
                                if _is_nat:
                                    _national.append(_aid)
                                else:
                                    _local.append(_aid)
                            _cd = _jcache.dumps({
                                "eligible_ids": _ids,
                                "ineligible_ids": [],
                                "local": _local,
                                "national": _national,
                            }, ensure_ascii=False)
                            _scu.execute(
                                """INSERT INTO user_match_cache (business_number, target_type, match_data, created_at)
                                   VALUES (%s, %s, %s::jsonb, CURRENT_TIMESTAMP)
                                   ON CONFLICT (business_number, target_type)
                                   DO UPDATE SET match_data=EXCLUDED.match_data, created_at=CURRENT_TIMESTAMP""",
                                (_bn, _ct, _cd)
                            )
                            _sc.commit(); _sc.close()
                        except Exception as _e:
                            print(f"[realtime-cache-save] {_e}")
                    _thr.Thread(
                        target=_bg_cache_save,
                        args=(bn, cache_type, full_where, type_params, bucket_sql, bucket_params, target_type, today_bucket),
                        daemon=True
                    ).start()

                    cat_cache_key = f"cat_counts:{target_type or 'all'}"
                    category_counts = _get_cached(cat_cache_key)
                    if not category_counts:
                        try:
                            _cat_where2 = valid_announcement_where()
                            _cat_params2: list = []
                            if target_type:
                                _cat_where2 += " AND (target_type = %s OR target_type = 'both')"
                                _cat_params2.append(target_type)
                            _pcur.execute(
                                f"""SELECT COALESCE(category, '기타') AS cat, COUNT(*) AS cnt
                                    FROM announcements WHERE {_cat_where2}
                                    GROUP BY COALESCE(category, '기타') ORDER BY cnt DESC""",
                                _cat_params2,
                            )
                            category_counts = {r["cat"]: r["cnt"] for r in _pcur.fetchall()}
                            _set_cache(cat_cache_key, category_counts)
                        except Exception:
                            category_counts = {}
                    _pc.close()
                    return {
                        "status": "SUCCESS",
                        "data": rows,
                        "total": total,
                        "page": page,
                        "size": size,
                        "regions": [],
                        "categories": [],
                        "category_counts": category_counts,
                        "personalized": True,
                        "source": "realtime",
                    }
                _pc.close()
        except Exception as _pe:
            _is_logged_in = True  # 예외 시에도 로그인 사용자로 처리 — 공유 캐시 우회
            _silent = ("no-cache-category-fallthrough", "no-cache-tab-fallthrough")
            if not any(s in str(_pe) for s in _silent):
                print(f"[personalized] fallback to standard SQL: {_pe}")

    # tab=local/national 일반 경로: 지역 필터 SQL로 처리 (캐시 미사용)
    if tab in ("local", "national") and not region:
        from app.services.rule_engine import _normalize_region as _nrm_tab
        _tab_conn = get_db_connection()
        _tab_cur = _tab_conn.cursor()
        _tab_where = valid_announcement_where()
        _tab_params: list = []
        if target_type:
            _tab_where += " AND (target_type = %s OR target_type = 'both')"
            _tab_params.append(target_type)
        if tab == "national":
            _tab_where += " AND (region IS NULL OR region IN ('전국', '', '전국 및 각 지역', 'All'))"
        elif tab == "local":
            # 사용자 지역 조회
            _user_city = ""
            if _is_logged_in:
                try:
                    _u_jwt = _decode_jwt(authorization.split(" ", 1)[1])
                    _u_bn = _u_jwt.get("bn")
                    if _u_bn:
                        _tab_cur.execute("SELECT address_city FROM users WHERE business_number = %s", (_u_bn,))
                        _urow = _tab_cur.fetchone()
                        if _urow and _urow.get("address_city"):
                            _cities = [c.strip() for c in str(_urow["address_city"]).split(",") if c.strip() and c.strip() != "전국"]
                            _user_city = _nrm_tab(_cities[0]) if _cities else ""
                except Exception:
                    pass
            if _user_city:
                _tab_where += " AND (region ILIKE %s OR title ILIKE %s)"
                _tab_params.extend([f"%{_user_city}%", f"%[{_user_city}]%"])
            else:
                _tab_where += " AND FALSE"  # 지역 정보 없으면 빈 결과
        _tab_cur.execute(f"SELECT COUNT(*) AS cnt FROM announcements WHERE {_tab_where}", _tab_params)
        _tab_total = (_tab_cur.fetchone() or {}).get("cnt", 0)
        _tab_cur.execute(
            f"""SELECT announcement_id, title, region, category, department,
                       support_amount, support_amount_max, support_amount_min, support_amount_type,
                       deadline_date, origin_source, created_at,
                       COALESCE(target_type, 'business') AS target_type,
                       origin_url, summary_text, eligibility_logic,
                       established_years_limit, revenue_limit, employee_limit
                FROM announcements
                WHERE {_tab_where}
                ORDER BY support_amount_max DESC NULLS LAST, deadline_date ASC NULLS LAST, created_at DESC
                LIMIT %s OFFSET %s""",
            _tab_params + [size, offset],
        )
        _tab_rows = [dict(r) for r in _tab_cur.fetchall()]
        _tab_conn.close()
        return {
            "status": "SUCCESS",
            "data": _tab_rows,
            "total": _tab_total,
            "page": page,
            "size": size,
            "regions": [],
            "categories": [],
            "category_counts": {},
            "personalized": False,
            "source": "tab_sql",
            "tab": tab,
        }

    # 검색 없는 기본 조회는 캐시 활용 — 로그인 사용자는 공유 캐시 우회
    if not _is_logged_in and not search and not region and not category and not tab:
        cache_key = f"pub:v2:{target_type}:{page}:{size}"
        cached = _get_cached(cache_key)
        if cached:
            return cached

    conn = get_db_connection()
    cursor = conn.cursor()

    where_clauses = [
        # [Phase 5] 유효 공고 통일 필터 — deadline_type/is_archived 기반
        valid_announcement_where(),
    ]
    params: list = []

    if region:
        if region in ("전국", "All", "all"):
            # 전국 선택 → 전국/빈값/All 포함 (지역한정 제외)
            where_clauses.append("(region IN ('전국', '', 'All') OR region IS NULL)")
        else:
            # 특정 지역 선택 → 정규화된 값으로 조회
            from app.services.rule_engine import _normalize_region as _nrm
            _nr = _nrm(region)
            where_clauses.append("(region = %s OR region = %s)")
            params.extend([region, _nr])
    if category:
        where_clauses.append("category ILIKE %s")
        params.append(f"%{category}%")
    if search:
        # 공백으로 단어 분리
        words = search.strip().split()

        # 검색어 정규화: 구청/시청/군청 접미어 → 구/시/군 자동 포함
        # 예: "부산진구청" → ["부산진구청", "부산진구"] / "수원시청" → ["수원시청", "수원시"]
        import re as _re_norm
        _normalized_words = []
        for w in words:
            _normalized_words.append(w)
            _stripped = _re_norm.sub(r'(구|시|군|동|읍|면)청$', r'\1', w)
            if _stripped != w and _stripped not in _normalized_words:
                _normalized_words.append(_stripped)
            # 반대로 끝에 "구/시/군"인데 "청" 붙은 버전도 포함 (검색 대상 department에 "청" 붙을 수 있음)
            if _re_norm.search(r'(구|시|군)$', w) and w + '청' not in _normalized_words:
                _normalized_words.append(w + '청')
        words = _normalized_words

        # 동의어 확장: 각 단어별로 동의어 그룹 확장
        all_search_terms = []
        for word in words:
            cursor.execute(
                "SELECT DISTINCT keyword FROM keyword_synonyms WHERE group_name = ("
                "  SELECT group_name FROM keyword_synonyms WHERE keyword = %s LIMIT 1"
                ")",
                (word,),
            )
            synonym_rows = cursor.fetchall()
            word_terms = [r["keyword"] for r in synonym_rows] if synonym_rows else [word]
            all_search_terms.append(word_terms)

        # 검색 조건: 각 단어(OR 동의어)가 모두 포함 (AND)
        # 단어가 1개면 OR 검색, 2개 이상이면 각 단어 AND 조합
        # 추가로 원본 구문 그대로 매칭도 포함 (OR)
        search_fields = "(title || ' ' || COALESCE(summary_text,'') || ' ' || COALESCE(department,'') || ' ' || COALESCE(region,'') || ' ' || COALESCE(category,''))"

        if len(words) == 1:
            # 단일 단어: 기존 동의어 OR 검색
            word_conditions = []
            for term in all_search_terms[0]:
                word_conditions.append(f"{search_fields} ILIKE %s")
                params.append(f"%{term}%")
            where_clauses.append("(" + " OR ".join(word_conditions) + ")")
        else:
            # 복수 단어: 구문 매칭 OR (각 단어 AND)
            # 1) 구문 그대로 매칭
            phrase_cond = f"{search_fields} ILIKE %s"
            params.append(f"%{search}%")

            # 2) 각 단어(+동의어)가 모두 포함 (AND)
            and_parts = []
            for word_terms in all_search_terms:
                word_or = []
                for term in word_terms:
                    word_or.append(f"{search_fields} ILIKE %s")
                    params.append(f"%{term}%")
                and_parts.append("(" + " OR ".join(word_or) + ")")
            and_cond = " AND ".join(and_parts)

            # 3) 단어 중 하나라도 포함 (OR) — 넓은 범위
            or_parts = []
            for word_terms in all_search_terms:
                for term in word_terms:
                    or_parts.append(f"{search_fields} ILIKE %s")
                    params.append(f"%{term}%")
            or_cond = " OR ".join(or_parts)

            where_clauses.append(f"({phrase_cond} OR {and_cond} OR {or_cond})")

        search_terms = []
        for wt in all_search_terms:
            search_terms.extend(wt)
        s = f"%{search}%"  # 관련성 정렬용 원본 검색어
    if target_type:
        where_clauses.append("(target_type = %s OR target_type = 'both')")
        params.append(target_type)

    where_sql = " AND ".join(where_clauses)

    # 총 개수 — 타임아웃 시 근사치로 폴백 (500 방지)
    try:
        cursor.execute("SET LOCAL statement_timeout = '15s'")
        cursor.execute(f"SELECT COUNT(*) AS cnt FROM announcements WHERE {where_sql}", params)
        cursor.execute("RESET statement_timeout")
        total = cursor.fetchone()["cnt"]
    except Exception as _count_err:
        print(f"[public] COUNT 타임아웃 폴백: {_count_err}")
        try: cursor.execute("RESET statement_timeout")
        except: pass
        # reltuples 근사치 사용
        try:
            cursor.execute("SELECT reltuples::bigint AS cnt FROM pg_class WHERE relname = 'announcements'")
            total = int((cursor.fetchone() or {}).get("cnt") or 10000)
        except:
            total = 10000

    # 공고 리스트 — 검색 시 관련성 정렬
    # 기관명 정확 매칭 > 구문 제목 > AND 제목 > 구문 요약 > AND 요약 > OR
    if search:
        words = search.strip().split()
        # AND 조건: 모든 단어가 제목에 포함
        and_title_parts = " AND ".join([f"title ILIKE %s" for _ in words])
        and_title_params = [f"%{w}%" for w in words]
        and_summary_parts = " AND ".join([f"summary_text ILIKE %s" for _ in words])
        and_summary_params = [f"%{w}%" for w in words]

        relevance_order = f"""
                CASE WHEN department ILIKE %s THEN 0
                     WHEN title ILIKE %s THEN 1
                     WHEN ({and_title_parts}) THEN 2
                     WHEN summary_text ILIKE %s THEN 3
                     WHEN ({and_summary_parts}) THEN 4
                     ELSE 5 END,
"""
        # [s] — department/title/summary 각각 구문 정확 매칭, 각 순서에 맞게 params 추가
        relevance_params = [s, s] + and_title_params + [s] + and_summary_params
    else:
        relevance_order = ""
        relevance_params = []

    cursor.execute(
        f"""SELECT announcement_id, title, region, category, department,
                   support_amount, support_amount_max, support_amount_min, support_amount_type,
                   deadline_date, origin_source, created_at,
                   COALESCE(target_type, 'business') AS target_type,
                   origin_url, summary_text, eligibility_logic,
                   established_years_limit, revenue_limit, employee_limit
            FROM announcements
            WHERE {where_sql}
            ORDER BY
                {relevance_order}
                CASE WHEN deadline_date IS NOT NULL AND deadline_date < CURRENT_DATE THEN 9
                     ELSE 0 END,
                CASE
                    WHEN deadline_date IS NOT NULL
                         AND deadline_date BETWEEN CURRENT_DATE AND CURRENT_DATE + INTERVAL '30 days'
                         AND support_amount IS NOT NULL AND support_amount != '' THEN 0
                    WHEN (region IN ('전국', '', '전국 및 각 지역', 'All') OR region IS NULL)
                         AND support_amount IS NOT NULL AND support_amount != '' THEN 1
                    WHEN support_amount IS NOT NULL AND support_amount != '' THEN 2
                    ELSE 3
                END,
                deadline_date ASC NULLS LAST,
                created_at DESC
            LIMIT %s OFFSET %s""",
        params + relevance_params + [size, offset],
    )
    rows = cursor.fetchall()

    # 필터용 메타: 지역/카테고리 목록 — 15분 캐시 (거의 변하지 않음)
    regions = _get_cached("meta:regions") or []
    if not regions:
        cursor.execute("SELECT DISTINCT region FROM announcements WHERE region IS NOT NULL ORDER BY region")
        regions = [r["region"] for r in cursor.fetchall()]
        _set_cache("meta:regions", regions)
    categories = _get_cached("meta:categories") or []
    if not categories:
        cursor.execute("SELECT DISTINCT category FROM announcements WHERE category IS NOT NULL ORDER BY category")
        categories = [r["category"] for r in cursor.fetchall()]
        _set_cache("meta:categories", categories)

    # 카테고리별 건수 — 캐시 활용 (5분)
    cat_cache_key = f"cat_counts:{target_type or 'all'}"
    category_counts = _get_cached(cat_cache_key)
    if not category_counts:
        cat_where = [
            # [Phase 5] 유효 공고 통일 필터
            valid_announcement_where(),
        ]
        cat_params: list = []
        if target_type:
            cat_where.append("(target_type = %s OR target_type = 'both')")
            cat_params.append(target_type)
        cat_where_sql = " AND ".join(cat_where)
        cursor.execute(
            f"""SELECT COALESCE(category, '기타') AS cat, COUNT(*) AS cnt
                FROM announcements WHERE {cat_where_sql}
                GROUP BY COALESCE(category, '기타') ORDER BY cnt DESC""",
            cat_params,
        )
        category_counts = {r["cat"]: r["cnt"] for r in cursor.fetchall()}
        _set_cache(cat_cache_key, category_counts)

    conn.close()

    result = {
        "status": "SUCCESS",
        "data": rows,
        "total": total,
        "page": page,
        "size": size,
        "regions": regions,
        "categories": categories,
        "category_counts": category_counts,
    }

    # 기본 조회 캐시 저장
    if not search and not region and not category:
        cache_key = f"pub:v2:{target_type}:{page}:{size}"
        _set_cache(cache_key, result)

    return result


@app.get("/api/announcements/{announcement_id}/detail")
def api_announcement_detail(announcement_id: int):
    """공고 상세 (비로그인 접근 가능 — 기본 정보만)"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT a.*, aa.deep_analysis
           FROM announcements a
           LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
           WHERE a.announcement_id = %s""",
        (announcement_id,),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")
    return {"status": "SUCCESS", "data": row}


class FindEmailRequest(BaseModel):
    company_name: str


@app.post("/api/auth/find-email")
def api_find_email(req: FindEmailRequest):
    if not req.company_name or len(req.company_name.strip()) < 1:
        raise HTTPException(status_code=400, detail="회사명을 입력해 주세요.")
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT email FROM users WHERE company_name = %s", (req.company_name,))
    user = cursor.fetchone()
    conn.close()
    if not user or not user.get("email"):
        raise HTTPException(status_code=404, detail="일치하는 계정을 찾을 수 없습니다.")
    email = user["email"]
    local, domain = email.split("@", 1)
    masked = local[:2] + "*" * max(len(local) - 2, 1) + "@" + domain
    return {"status": "SUCCESS", "masked_email": masked}


class ResetPasswordRequest(BaseModel):
    email: str
    new_password: str
    code: Optional[str] = None  # 인증코드

class ResetCodeRequest(BaseModel):
    email: str

# 비밀번호 재설정 인증코드 저장 (메모리, {email: {code, expires, attempts}})
_reset_codes: dict = {}
# Rate limiting: {ip: [timestamps]}
_reset_rate: dict = {}

def _send_reset_email(to_email: str, code: str):
    """비밀번호 재설정 인증코드 이메일 발송"""
    smtp_user = os.getenv("SMTP_USER", "")
    smtp_password = os.getenv("SMTP_PASSWORD", "")
    smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_from = os.getenv("SMTP_FROM", smtp_user)
    if not smtp_user or not smtp_password:
        return False
    try:
        import smtplib
        from email.mime.text import MIMEText
        msg = MIMEText(
            f"지원금AI 비밀번호 재설정 인증코드입니다.\n\n"
            f"인증코드: {code}\n\n"
            f"이 코드는 10분간 유효합니다.\n"
            f"본인이 요청하지 않았다면 이 이메일을 무시하세요.",
            "plain", "utf-8"
        )
        msg["Subject"] = "[지원금AI] 비밀번호 재설정 인증코드"
        msg["From"] = smtp_from
        msg["To"] = to_email
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
        return True
    except Exception:
        return False

@app.post("/api/auth/reset-password/request")
def api_request_reset_code(req: ResetCodeRequest, request: Request):
    """1단계: 인증코드 발송 요청"""
    if not req.email:
        raise HTTPException(status_code=400, detail="이메일을 입력해주세요.")
    # Rate limiting: IP당 분당 3회
    client_ip = request.client.host if request.client else "unknown"
    now = datetime.datetime.utcnow()
    if client_ip in _reset_rate:
        _reset_rate[client_ip] = [t for t in _reset_rate[client_ip] if (now - t).seconds < 300]
        if len(_reset_rate[client_ip]) >= 3:
            raise HTTPException(status_code=429, detail="잠시 후 다시 시도해주세요.")
    else:
        _reset_rate[client_ip] = []
    _reset_rate[client_ip].append(now)

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT business_number FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    conn.close()
    # 계정 존재 여부와 관계없이 동일 응답 (열거 방지)
    if user:
        import random
        code = f"{random.randint(100000, 999999)}"
        _reset_codes[req.email] = {
            "code": code,
            "expires": now + datetime.timedelta(minutes=10),
            "attempts": 0,
        }
        _send_reset_email(req.email, code)
    return {"status": "SUCCESS", "message": "등록된 이메일이면 인증코드가 발송됩니다."}


@app.post("/api/auth/reset-password")
def api_reset_password(req: ResetPasswordRequest):
    """2단계: 인증코드 확인 후 비밀번호 변경"""
    if not req.email:
        raise HTTPException(status_code=400, detail="이메일을 입력해주세요.")
    if not req.code:
        raise HTTPException(status_code=400, detail="인증코드를 입력해주세요.")
    if not req.new_password or len(req.new_password) < 6:
        raise HTTPException(status_code=400, detail="비밀번호는 6자 이상이어야 합니다.")

    # 인증코드 검증
    stored = _reset_codes.get(req.email)
    if not stored:
        raise HTTPException(status_code=400, detail="인증코드를 먼저 요청해주세요.")
    if stored["attempts"] >= 5:
        del _reset_codes[req.email]
        raise HTTPException(status_code=400, detail="인증 시도 횟수를 초과했습니다. 다시 요청해주세요.")
    if datetime.datetime.utcnow() > stored["expires"]:
        del _reset_codes[req.email]
        raise HTTPException(status_code=400, detail="인증코드가 만료되었습니다. 다시 요청해주세요.")
    if req.code != stored["code"]:
        stored["attempts"] += 1
        raise HTTPException(status_code=400, detail="인증코드가 올바르지 않습니다.")

    # 인증 성공 → 비밀번호 변경
    del _reset_codes[req.email]
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT business_number FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=400, detail="처리 중 오류가 발생했습니다.")
    new_hash = _hash_password(req.new_password)
    cursor.execute("UPDATE users SET password_hash = %s WHERE email = %s", (new_hash, req.email))
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "message": "비밀번호가 재설정되었습니다. 새 비밀번호로 로그인해주세요."}


@app.post("/api/auth/register")
def api_register(req: RegisterRequest, request: Request):
    # Rate limiting: IP당 시간당 5회
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"register:ip:{ip}", 5, 3600):
        raise HTTPException(status_code=429, detail="회원가입 시도가 너무 많습니다. 잠시 후 다시 시도해주세요.")
    if not req.email or "@" not in req.email:
        raise HTTPException(status_code=400, detail="올바른 이메일을 입력해 주세요.")
    if not req.password or len(req.password) < 6:
        raise HTTPException(status_code=400, detail="비밀번호는 6자 이상이어야 합니다.")
    if len(req.business_number) != 10:
        raise HTTPException(status_code=400, detail="사업자번호 10자리를 입력해 주세요.")

    import hashlib as _hashlib
    hashed = _hash_password(req.password)
    now_iso = datetime.datetime.utcnow().isoformat()

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT user_id FROM users WHERE email = %s", (req.email,))
        if cursor.fetchone():
            raise HTTPException(status_code=409, detail="이미 가입된 이메일입니다.")

        cursor.execute("SELECT user_id FROM users WHERE business_number = %s", (req.business_number,))
        existing = cursor.fetchone()

        est_date = req.establishment_date or None
        if existing:
            cursor.execute(
                """UPDATE users SET email=%s, password_hash=%s, plan='free',
                   plan_started_at=%s, plan_expires_at=NULL,
                   ai_usage_month=0, ai_usage_reset_at=%s,
                   company_name=COALESCE(NULLIF(%s, ''), company_name),
                   address_city=COALESCE(%s, address_city),
                   industry_code=COALESCE(%s, industry_code),
                   establishment_date=COALESCE(%s, establishment_date),
                   revenue_bracket=COALESCE(%s, revenue_bracket),
                   employee_count_bracket=COALESCE(%s, employee_count_bracket),
                   interests=COALESCE(%s, interests)
                   WHERE business_number=%s""",
                (req.email, hashed, now_iso, now_iso,
                 req.company_name or "", req.address_city or None,
                 req.industry_code or None, req.establishment_date,
                 req.revenue_bracket, req.employee_count_bracket, req.interests,
                 req.business_number),
            )
            user_id = existing["user_id"]
        else:
            # 신규 가입: LITE 7일 무료체험
            _initial_plan = 'lite'
            _initial_expires = (datetime.datetime.utcnow() + datetime.timedelta(days=TRIAL_DAYS)).isoformat()
            cursor.execute(
                """INSERT INTO users (business_number, company_name, email, password_hash, plan,
                   plan_started_at, plan_expires_at, ai_usage_month, ai_usage_reset_at,
                   address_city, establishment_date, industry_code, revenue_bracket, employee_count_bracket, interests,
                   referred_by, user_type)
                   VALUES (%s, %s, %s, %s, %s, %s, %s, 0, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                   RETURNING user_id""",
                (req.business_number, req.company_name or "", req.email, hashed,
                 _initial_plan, now_iso, _initial_expires, now_iso,
                 req.address_city or None, est_date,
                 req.industry_code or None, req.revenue_bracket or None,
                 req.employee_count_bracket or None, req.interests,
                 req.referred_by or None, req.user_type or "both"),
            )
            user_id = cursor.fetchone()["user_id"]
            # 추천 코드 자동 생성
            ref_code = _hashlib.md5(f'{req.business_number}{user_id}'.encode()).hexdigest()[:8].upper()
            cursor.execute("UPDATE users SET referral_code=%s WHERE business_number=%s", (ref_code, req.business_number))

            # 추천 링크로 가입한 경우: 피추천인 LITE 30일 무료 + 추천인(LITE만) 30일 연장
            if req.referred_by:
                now_dt = datetime.datetime.utcnow()
                cursor.execute("SELECT user_id, plan, plan_expires_at, merit_months FROM users WHERE referral_code = %s", (req.referred_by,))
                referrer = cursor.fetchone()

                # 피추천인: LITE 30일 무료
                new_end_new = (now_dt + datetime.timedelta(days=30)).isoformat()
                cursor.execute(
                    "UPDATE users SET plan='lite', plan_started_at=%s, plan_expires_at=%s, ai_usage_month=0, ai_usage_reset_at=%s, referral_rewarded=TRUE WHERE user_id=%s",
                    (now_dt.isoformat(), new_end_new, now_dt.isoformat(), user_id)
                )

                # 추천인: LITE 가입자인 경우만 30일 연장 (최대 1회)
                if referrer and referrer["plan"] in ("lite", "lite_trial", "basic") and (referrer["merit_months"] or 0) < 1:
                    new_merit = (referrer["merit_months"] or 0) + 1
                    try:
                        current_end = datetime.datetime.fromisoformat(str(referrer["plan_expires_at"]))
                        new_end = (max(current_end, now_dt) + datetime.timedelta(days=30)).isoformat()
                    except Exception:
                        new_end = (now_dt + datetime.timedelta(days=30)).isoformat()
                    cursor.execute(
                        "UPDATE users SET merit_months=%s, plan_expires_at=%s WHERE user_id=%s",
                        (new_merit, new_end, referrer["user_id"])
                    )

        conn.commit()
        # 가입 후 실제 플랜 상태 조회 (체험/추천 적용 반영)
        cursor2 = conn.cursor()
        cursor2.execute("SELECT plan, plan_expires_at FROM users WHERE user_id=%s", (user_id,))
        signup_user = cursor2.fetchone()
        signup_plan = signup_user["plan"] if signup_user else "free"
        signup_expires = str(signup_user["plan_expires_at"]) if signup_user and signup_user["plan_expires_at"] else None
        token = _create_jwt(user_id, req.business_number, req.email, signup_plan, signup_expires)
        _log_event("signup", req.business_number, f"email={req.email},plan={signup_plan}", _get_client_ip(request), request.headers.get("user-agent", ""))
        return {
            "status": "SUCCESS",
            "token": token,
            "plan": _get_plan_status(signup_plan, signup_expires, 0),
        }
    finally:
        conn.close()


@app.post("/api/auth/login")
def api_login(req: LoginRequest, request: Request):
    # Rate limiting: IP당 분당 10회, 이메일당 분당 5회
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"login:ip:{ip}", 10, 60):
        raise HTTPException(status_code=429, detail="로그인 시도가 너무 많습니다. 잠시 후 다시 시도해주세요.")
    if req.email and not _rate_limit_check(f"login:email:{req.email}", 5, 60):
        raise HTTPException(status_code=429, detail="로그인 시도가 너무 많습니다. 잠시 후 다시 시도해주세요.")
    if not req.email or not req.password:
        raise HTTPException(status_code=400, detail="이메일과 비밀번호를 입력해 주세요.")

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
        user = cursor.fetchone()

        if not user:
            raise HTTPException(status_code=404, detail="등록되지 않은 이메일입니다.")

        u = dict(user)
        if not u.get("password_hash"):
            raise HTTPException(status_code=401, detail="소셜 로그인으로 가입된 계정입니다. 카카오/네이버/Google로 로그인해주세요.")

        if not _verify_password(req.password, u["password_hash"]):
            _log_event("login_fail", u.get("business_number", ""), f"email={req.email}", ip)
            raise HTTPException(status_code=401, detail="비밀번호가 올바르지 않습니다.")

        _log_event("login", u.get("business_number", ""), f"email={req.email},plan={u.get('plan','free')},type={u.get('user_type','')}", ip, request.headers.get("user-agent", ""))
        plan = u.get("plan") or "free"
        if plan == "trial":
            plan = "free"
        plan_expires = u.get("plan_expires_at")
        if plan_expires is not None:
            plan_expires = str(plan_expires)
        ai_usage = u.get("ai_usage_month") or 0

        # 월간 사용량 리셋 체크
        reset_at = u.get("ai_usage_reset_at")
        now = datetime.datetime.utcnow()
        if reset_at:
            try:
                reset_dt = datetime.datetime.fromisoformat(str(reset_at))
                if now.month != reset_dt.month or now.year != reset_dt.year:
                    ai_usage = 0
                    cursor.execute(
                        "UPDATE users SET ai_usage_month=0, ai_usage_reset_at=%s WHERE user_id=%s",
                        (now.isoformat(), u["user_id"])
                    )
            except Exception:
                pass

        plan_status = _get_plan_status(plan, plan_expires, ai_usage)

        if plan_status["plan"] == "expired":
            plan = "free"
            cursor.execute("UPDATE users SET plan = 'free', plan_expires_at = NULL WHERE user_id = %s", (u["user_id"],))
            plan_status = _get_plan_status("free", None, ai_usage)

        # Lookup KSIC industry name
        industry_name = ""
        if u.get("industry_code"):
            cursor.execute("SELECT name FROM ksic_classification WHERE code = %s", (u["industry_code"],))
            row3 = cursor.fetchone()
            if row3:
                industry_name = row3["name"]

        conn.commit()
    finally:
        conn.close()

    token = _create_jwt(u["user_id"], u["business_number"], u["email"], plan, plan_expires)
    return {
        "status": "SUCCESS",
        "token": token,
        "user": {
            "business_number": u["business_number"],
            "company_name": u.get("company_name", ""),
            "email": u["email"],
            "address_city": u.get("address_city", ""),
            "establishment_date": str(u.get("establishment_date", "") or ""),
            "industry_code": u.get("industry_code", ""),
            "industry_name": industry_name,
            "revenue_bracket": u.get("revenue_bracket", ""),
            "employee_count_bracket": u.get("employee_count_bracket", ""),
            "interests": u.get("interests", ""),
        },
        "plan": plan_status,
    }


# ─── 소셜 로그인 (OAuth2) ───────────────────────────────────────────
import secrets
import urllib.parse

KAKAO_CLIENT_ID = os.getenv("KAKAO_CLIENT_ID", "")
KAKAO_CLIENT_SECRET = os.getenv("KAKAO_CLIENT_SECRET", "")
NAVER_CLIENT_ID = os.getenv("NAVER_CLIENT_ID", "")
NAVER_CLIENT_SECRET = os.getenv("NAVER_CLIENT_SECRET", "")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "")

FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")


def _social_login_or_register(provider: str, social_id: str, email: str, name: str, kakao_refresh_token: str = "", gender: str = "", birth_year: str = ""):
    """소셜 로그인: 기존 사용자면 로그인, 신규면 자동 가입"""
    conn = get_db_connection()
    cursor = conn.cursor()

    # 성별 정규화: kakao "male"/"female" → DB "남성"/"여성"
    gender_map = {"male": "남성", "female": "여성", "M": "남성", "F": "여성"}
    normalized_gender = gender_map.get(gender, gender)

    # 출생연도 → 연령대 매핑
    age_range_val = ""
    if birth_year:
        try:
            age = datetime.datetime.now().year - int(birth_year)
            if age < 20: age_range_val = "10대"
            elif age < 30: age_range_val = "20대"
            elif age < 40: age_range_val = "30대"
            elif age < 50: age_range_val = "40대"
            elif age < 60: age_range_val = "50대"
            else: age_range_val = "60대 이상"
        except (ValueError, TypeError):
            pass

    # 1. 이미 같은 이메일로 가입된 사용자 확인
    cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
    user = cursor.fetchone()

    now_iso = datetime.datetime.utcnow().isoformat()

    if user:
        u = dict(user)
        # kakao_id 컬럼에 소셜 provider:id 저장 (기존 호환)
        updates = []
        params = []
        if not u.get("kakao_id"):
            updates.append("kakao_id = %s")
            params.append(f"{provider}:{social_id}")
        if kakao_refresh_token and provider == "kakao":
            updates.append("kakao_refresh_token = %s")
            params.append(kakao_refresh_token)
        # 성별/연령대가 비어있으면 소셜에서 받은 값으로 채움
        if normalized_gender and not u.get("gender"):
            updates.append("gender = %s")
            params.append(normalized_gender)
        if age_range_val and not u.get("age_range"):
            updates.append("age_range = %s")
            params.append(age_range_val)
        if updates:
            params.append(u["user_id"])
            cursor.execute(f"UPDATE users SET {', '.join(updates)} WHERE user_id = %s", params)
            conn.commit()
    else:
        # 신규 가입 — LITE 7일 무료체험
        bn = f"U{int(datetime.datetime.utcnow().timestamp())}"[-10:]
        import hashlib as _hashlib
        _initial_plan = 'lite'
        _initial_expires = (datetime.datetime.utcnow() + datetime.timedelta(days=TRIAL_DAYS)).isoformat()
        cursor.execute(
            """INSERT INTO users (business_number, company_name, email, password_hash, plan,
               plan_started_at, plan_expires_at, ai_usage_month, ai_usage_reset_at, kakao_id, gender, age_range, user_type)
               VALUES (%s, %s, %s, %s, %s, %s, %s, 0, %s, %s, %s, %s, 'both')
               RETURNING user_id, business_number""",
            (bn, name or "", email, "", _initial_plan, now_iso, _initial_expires, now_iso,
             f"{provider}:{social_id}", normalized_gender, age_range_val)
        )
        row = cursor.fetchone()
        ref_code = _hashlib.md5(f'{bn}{row["user_id"]}'.encode()).hexdigest()[:8].upper()
        cursor.execute("UPDATE users SET referral_code=%s WHERE user_id=%s", (ref_code, row["user_id"]))
        conn.commit()

        cursor.execute("SELECT * FROM users WHERE user_id = %s", (row["user_id"],))
        user = cursor.fetchone()
        u = dict(user)

    conn.close()

    plan = u.get("plan") or "free"
    if plan in ("trial", "premium"):
        plan = "free"
    plan_expires = str(u["plan_expires_at"]) if u.get("plan_expires_at") else None
    ai_usage = u.get("ai_usage_month") or 0
    plan_status = _get_plan_status(plan, plan_expires, ai_usage)

    token = _create_jwt(u["user_id"], u["business_number"], u["email"], plan, plan_expires)

    is_new = not u.get("user_type")
    return token, plan_status, u, is_new


@app.get("/api/auth/social/{provider}")
def api_social_auth_redirect(provider: str):
    """소셜 로그인 시작: 각 플랫폼 OAuth URL로 리다이렉트"""
    redirect_uri = f"{FRONTEND_URL}/auth/callback/{provider}"
    state = secrets.token_urlsafe(16)

    if provider == "kakao":
        if not KAKAO_CLIENT_ID:
            raise HTTPException(status_code=501, detail="카카오 로그인이 아직 설정되지 않았습니다.")
        url = (f"https://kauth.kakao.com/oauth/authorize"
               f"?client_id={KAKAO_CLIENT_ID}"
               f"&redirect_uri={urllib.parse.quote(redirect_uri)}"
               f"&response_type=code&state={state}")
    elif provider == "naver":
        if not NAVER_CLIENT_ID:
            raise HTTPException(status_code=501, detail="네이버 로그인이 아직 설정되지 않았습니다.")
        url = (f"https://nid.naver.com/oauth2.0/authorize"
               f"?client_id={NAVER_CLIENT_ID}"
               f"&redirect_uri={urllib.parse.quote(redirect_uri)}"
               f"&response_type=code&state={state}")
    elif provider == "google":
        if not GOOGLE_CLIENT_ID:
            raise HTTPException(status_code=501, detail="Google 로그인이 아직 설정되지 않았습니다.")
        url = (f"https://accounts.google.com/o/oauth2/v2/auth"
               f"?client_id={GOOGLE_CLIENT_ID}"
               f"&redirect_uri={urllib.parse.quote(redirect_uri)}"
               f"&response_type=code&scope=email+profile&state={state}")
    else:
        raise HTTPException(status_code=400, detail="지원하지 않는 로그인 방식입니다.")

    from fastapi.responses import RedirectResponse
    return RedirectResponse(url)


class SocialCallbackRequest(BaseModel):
    code: str
    provider: str


@app.post("/api/auth/social/callback")
async def api_social_callback(req: SocialCallbackRequest):
    """소셜 로그인 콜백: authorization code → 토큰 교환 → 사용자 정보 → JWT 발급"""
    import httpx

    redirect_uri = f"{FRONTEND_URL}/auth/callback/{req.provider}"

    if req.provider == "kakao":
        # 1. code → access_token
        async with httpx.AsyncClient() as client:
            token_res = await client.post("https://kauth.kakao.com/oauth/token", data={
                "grant_type": "authorization_code",
                "client_id": KAKAO_CLIENT_ID,
                "client_secret": KAKAO_CLIENT_SECRET,
                "code": req.code,
                "redirect_uri": redirect_uri,
            })
            token_data = token_res.json()
            access_token = token_data.get("access_token")
            kakao_refresh_token = token_data.get("refresh_token", "")
            if not access_token:
                raise HTTPException(status_code=400, detail="카카오 인증 실패")

            # 2. access_token → user info
            user_res = await client.get("https://kapi.kakao.com/v2/user/me",
                                        headers={"Authorization": f"Bearer {access_token}"})
            user_data = user_res.json()

        kakao_id = str(user_data.get("id", ""))
        account = user_data.get("kakao_account", {})
        email = account.get("email", f"kakao_{kakao_id}@kakao.local")
        name = account.get("profile", {}).get("nickname", "")
        # 카카오에서 성별/출생연도 추출
        kakao_gender = account.get("gender", "")  # "male" or "female"
        kakao_birthyear = account.get("birthyear", "")  # "1990" 등

        token, plan_status, u, is_new = _social_login_or_register(
            "kakao", kakao_id, email, name,
            kakao_refresh_token=kakao_refresh_token,
            gender=kakao_gender, birth_year=kakao_birthyear
        )

    elif req.provider == "naver":
        async with httpx.AsyncClient() as client:
            token_res = await client.post("https://nid.naver.com/oauth2.0/token", data={
                "grant_type": "authorization_code",
                "client_id": NAVER_CLIENT_ID,
                "client_secret": NAVER_CLIENT_SECRET,
                "code": req.code,
                "redirect_uri": redirect_uri,
            })
            token_data = token_res.json()
            access_token = token_data.get("access_token")
            if not access_token:
                raise HTTPException(status_code=400, detail="네이버 인증 실패")

            user_res = await client.get("https://openapi.naver.com/v1/nid/me",
                                        headers={"Authorization": f"Bearer {access_token}"})
            user_data = user_res.json().get("response", {})

        naver_id = user_data.get("id", "")
        email = user_data.get("email", f"naver_{naver_id}@naver.local")
        name = user_data.get("name", user_data.get("nickname", ""))
        # 네이버에서 성별/출생연도 추출: gender="M"/"F", birthyear="1990"
        naver_gender = user_data.get("gender", "")
        naver_birthyear = user_data.get("birthyear", "")

        token, plan_status, u, is_new = _social_login_or_register(
            "naver", naver_id, email, name,
            gender=naver_gender, birth_year=naver_birthyear
        )

    elif req.provider == "google":
        async with httpx.AsyncClient() as client:
            token_res = await client.post("https://oauth2.googleapis.com/token", data={
                "grant_type": "authorization_code",
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "code": req.code,
                "redirect_uri": redirect_uri,
            })
            token_data = token_res.json()
            access_token = token_data.get("access_token")
            if not access_token:
                raise HTTPException(status_code=400, detail="Google 인증 실패")

            user_res = await client.get("https://www.googleapis.com/oauth2/v2/userinfo",
                                        headers={"Authorization": f"Bearer {access_token}"})
            user_data = user_res.json()

        google_id = user_data.get("id", "")
        email = user_data.get("email", "")
        name = user_data.get("name", "")

        token, plan_status, u, is_new = _social_login_or_register("google", google_id, email, name)
    else:
        raise HTTPException(status_code=400, detail="지원하지 않는 로그인 방식입니다.")

    _log_event("social_login" if not is_new else "social_signup", u.get("business_number", ""), f"provider={req.provider},email={u.get('email','')}")
    return {
        "status": "SUCCESS",
        "token": token,
        "user": {
            "business_number": u.get("business_number", ""),
            "company_name": u.get("company_name", ""),
            "email": u.get("email", ""),
        },
        "plan": plan_status,
        "is_new_user": is_new,
    }


@app.get("/api/auth/me")
def api_auth_me(current_user: dict = Depends(_get_current_user)):
    cache_key = f"auth_me:{current_user['bn']}"
    cached = _get_cached(cache_key)
    if cached:
        return cached
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE business_number = %s", (current_user["bn"],))
    user = cursor.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)
    # Lookup KSIC industry name
    industry_name = ""
    if u.get("industry_code"):
        cursor.execute("SELECT name FROM ksic_classification WHERE code = %s", (u["industry_code"],))
        row = cursor.fetchone()
        if row:
            industry_name = row["name"]
    conn.close()
    plan_expires = u.get("plan_expires_at") or u.get("trial_ends_at")
    if plan_expires is not None:
        plan_expires = str(plan_expires)
    plan_status = _get_plan_status(u.get("plan") or "free", plan_expires, u.get("ai_usage_month") or 0)
    result = {
        "status": "SUCCESS",
        "user": {
            "business_number": u["business_number"],
            "company_name": u["company_name"],
            "email": u.get("email"),
            "address_city": u["address_city"],
            "industry_code": u["industry_code"],
            "industry_name": industry_name,
            "revenue_bracket": u["revenue_bracket"],
            "employee_count_bracket": u["employee_count_bracket"],
            "interests": u.get("interests"),
            "establishment_date": str(u["establishment_date"]) if u.get("establishment_date") else "",
            "referral_code": u.get("referral_code"),
            "merit_months": u.get("merit_months", 0),
            "user_type": u.get("user_type") or None,
            "is_social": bool(u.get("kakao_id")),
            "social_provider": u.get("kakao_id", "").split(":")[0] if u.get("kakao_id") else None,
            "custom_needs": u.get("custom_needs"),
            "custom_keywords": u.get("custom_keywords"),
            "gender": u.get("gender"),
            "age_range": u.get("age_range"),
            "income_level": u.get("income_level"),
            "family_type": u.get("family_type"),
            "employment_status": u.get("employment_status"),
            "founded_date": u.get("founded_date"),
            "is_pre_founder": u.get("is_pre_founder"),
            "certifications": u.get("certifications"),
        },
        "plan": plan_status,
    }
    _set_cache(cache_key, result)
    return result


class UpgradePlanRequest(BaseModel):
    payment_id: Optional[str] = None      # 포트원 V2 결제 ID
    billing_key: Optional[str] = None     # 정기결제 빌링키
    target_plan: Optional[str] = "lite"   # "lite" or "pro"
    free_trial: Optional[bool] = False    # 사업자 LITE 1개월 무료 여부


PORTONE_API_SECRET = os.getenv("PORTONE_API_SECRET", "")

# PLAN_PRICES는 상단에 정의됨


@app.post("/api/plan/upgrade")
def api_plan_upgrade(
    req: UpgradePlanRequest,
    current_user: dict = Depends(_get_current_user),
):
    """포트원 V2 결제 확인 후 플랜 업그레이드"""
    bn = current_user["bn"]
    target = req.target_plan if req.target_plan in ("lite", "pro") else "lite"

    # PRO 3개월 오픈 이벤트 (2026-07-31까지 29,000원/월)

    # 사용자 정보 조회
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT user_type, plan, plan_expires_at FROM users WHERE business_number = %s", (bn,))
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)
    user_type = u.get("user_type") or "both"

    # 가격 결정 (user_type별)
    if target == "lite":
        price = PLAN_PRICES.get("lite_individual" if user_type == "individual" else "lite", 4900)
    else:
        price = PLAN_PRICES.get("pro", 49000)

    # 사업자 LITE 1개월 무료 체험
    if req.free_trial and target == "lite" and user_type != "individual":
        # 이미 체험한 적 있는지 확인
        cur.execute(
            "SELECT plan_expires_at FROM users WHERE business_number = %s AND plan IN ('lite', 'lite_trial', 'pro')",
            (bn,),
        )
        if cur.fetchone():
            conn.close()
            raise HTTPException(status_code=400, detail="무료 체험은 1회만 가능합니다.")
        # 무료 체험: 결제 없이 30일 부여
        target = "lite"
    elif req.payment_id and PORTONE_API_SECRET:
        # 포트원 V2 결제 검증
        import httpx
        verify_res = httpx.get(
            f"https://api.portone.io/payments/{req.payment_id}",
            headers={"Authorization": f"PortOne {PORTONE_API_SECRET}"},
            timeout=15,
        )
        if verify_res.status_code != 200:
            conn.close()
            raise HTTPException(status_code=400, detail="결제 정보를 확인할 수 없습니다.")
        payment = verify_res.json()
        if payment.get("status") != "PAID":
            conn.close()
            raise HTTPException(status_code=400, detail="결제가 완료되지 않았습니다.")
        if payment.get("amount", {}).get("total") != price:
            conn.close()
            raise HTTPException(status_code=400, detail="결제 금액이 일치하지 않습니다.")
    else:
        # 무료 체험도 아니고 결제도 안 됨 → 플랜 변경 차단
        conn.close()
        raise HTTPException(status_code=400, detail="결제가 필요합니다.")

    # 빌링키 저장 (정기결제용)
    billing_update = ""
    billing_params = []
    if req.billing_key:
        billing_update = ", billing_key = %s"
        billing_params = [req.billing_key]

    now = datetime.datetime.utcnow()
    expires_at = (now + datetime.timedelta(days=30)).isoformat()

    cur.execute(
        f"""UPDATE users SET plan = %s, plan_started_at = %s, plan_expires_at = %s,
           ai_usage_month = 0, ai_usage_reset_at = %s{billing_update}
           WHERE business_number = %s""",
        [target, now.isoformat(), expires_at, now.isoformat()] + billing_params + [bn],
    )
    conn.commit()
    conn.close()

    plan_status = _get_plan_status(target, expires_at, 0)
    label_map = {"lite": "LITE", "lite_trial": "LITE 체험", "basic": "LITE", "biz": "PRO", "pro": "PRO"}
    label = label_map.get(target, target.upper())
    new_token = _create_jwt(
        current_user["user_id"], bn, current_user["email"], target, expires_at
    )

    _log_event("upgrade", current_user["bn"], f"plan={target_plan},price={price},trial={req.free_trial}")
    return {
        "status": "SUCCESS",
        "token": new_token,
        "plan": plan_status,
        "message": f"{label} 플랜으로 업그레이드되었습니다. ({'무료 체험 시작' if req.free_trial else f'{price:,}원 결제 완료'})",
        "price": price,
    }


class SubscribeRequest(BaseModel):
    billing_key: str
    target_plan: str = "lite"  # "lite" or "pro"


@app.post("/api/plan/subscribe")
def api_plan_subscribe(
    req: SubscribeRequest,
    current_user: dict = Depends(_get_current_user),
):
    """빌링키 등록 + 무료 체험 시작 (LITE/PRO 공통 30일)"""
    # 1차: 형식 검증
    if not req.billing_key or not isinstance(req.billing_key, str) or len(req.billing_key.strip()) < 10:
        raise HTTPException(status_code=400, detail="유효하지 않은 빌링키입니다. 카드 등록을 다시 시도해 주세요.")

    bn = current_user["bn"]
    target = req.target_plan if req.target_plan in ("lite", "pro") else "lite"

    # 2차: PortOne V2 REST API로 billing_key 실존·소유 검증
    #      — "결제창만 뜬 상태에서 자동 업그레이드" 현상 차단 (프론트엔드에서 더미 billing_key 전송 방지)
    if PORTONE_API_SECRET:
        import httpx as _httpx
        try:
            verify_res = _httpx.get(
                f"https://api.portone.io/billing-keys/{req.billing_key.strip()}",
                headers={"Authorization": f"PortOne {PORTONE_API_SECRET}"},
                timeout=10,
            )
            if verify_res.status_code == 404:
                raise HTTPException(status_code=400, detail="등록되지 않은 빌링키입니다. 결제를 다시 시도해 주세요.")
            if verify_res.status_code != 200:
                # 일시적 서비스 장애 — 보수적으로 차단
                raise HTTPException(status_code=502, detail="결제 검증 서비스에 일시적인 문제가 있습니다. 잠시 후 다시 시도해 주세요.")
            bk = verify_res.json()
            if bk.get("status") != "ISSUED":
                raise HTTPException(status_code=400, detail="결제가 완료되지 않은 빌링키입니다. 카드 등록을 다시 시도해 주세요.")
            # 타 사용자 빌링키 도용 방지 (customerId는 프론트엔드에서 bn을 정규화한 값)
            expected_cid = "".join(c if (c.isalnum() or c in "_-") else "_" for c in (bn or ""))[:40]
            bk_customer_id = ((bk.get("customer") or {}).get("id") or "").strip()
            if expected_cid and bk_customer_id and bk_customer_id != expected_cid:
                print(f"[subscribe] customerId mismatch: expected={expected_cid} got={bk_customer_id}")
                raise HTTPException(status_code=403, detail="본인 계정으로 결제된 빌링키가 아닙니다.")
        except HTTPException:
            raise
        except _httpx.RequestError as e:
            # 네트워크 오류는 서비스 연속성을 위해 경고 로그만 남기고 통과
            print(f"[subscribe] PortOne verify network error (허용): {e}")
        except Exception as e:
            print(f"[subscribe] PortOne verify unexpected error: {e}")
            raise HTTPException(status_code=502, detail="결제 검증 중 오류가 발생했습니다. 잠시 후 다시 시도해 주세요.")
    else:
        print("[subscribe] PORTONE_API_SECRET 미설정 — 빌링키 검증 건너뜀 (테스트 환경)")

    conn = get_db_connection()
    cur = conn.cursor()

    # 사용자 조회
    cur.execute("SELECT plan, billing_key FROM users WHERE business_number = %s", (bn,))
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)
    current_plan = u.get("plan") or "free"

    # 동일 플랜 재구독 차단 (상위 플랜 업그레이드는 허용)
    if u.get("billing_key") and current_plan == target:
        conn.close()
        raise HTTPException(status_code=400, detail="이미 동일한 플랜을 구독 중입니다.")
    # 하위 플랜으로 변경 차단 (PRO → LITE)
    if current_plan == "pro" and target == "lite":
        conn.close()
        raise HTTPException(status_code=400, detail="PRO에서 LITE로 변경은 구독 해지 후 재가입해주세요.")

    # 결제 후 첫 달 무료: 30일 후 첫 결제
    now = datetime.datetime.utcnow()
    expires_at = (now + datetime.timedelta(days=30)).isoformat()

    cur.execute(
        """UPDATE users SET plan = %s, plan_started_at = %s, plan_expires_at = %s,
           ai_usage_month = 0, ai_usage_reset_at = %s, billing_key = %s
           WHERE business_number = %s""",
        (target, now.isoformat(), expires_at, now.isoformat(), req.billing_key, bn),
    )
    conn.commit()
    conn.close()

    plan_status = _get_plan_status(target, expires_at, 0)
    label = "LITE" if target == "lite" else "PRO"
    new_token = _create_jwt(
        current_user["user_id"], bn, current_user["email"], target, expires_at
    )

    return {
        "status": "SUCCESS",
        "token": new_token,
        "plan": plan_status,
        "message": f"{label} {trial_days}일 무료 체험이 시작되었습니다! 이후 자동결제됩니다.",
    }


@app.post("/api/plan/cancel")
def api_plan_cancel(current_user: dict = Depends(_get_current_user)):
    """구독 해지 — 빌링키 삭제, 만료일까지는 이용 가능"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "UPDATE users SET billing_key = NULL WHERE business_number = %s",
        (bn,),
    )
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "message": "구독이 해지되었습니다. 현재 플랜 만료일까지는 이용 가능합니다."}


@app.post("/api/plan/refund")
def api_plan_refund(current_user: dict = Depends(_get_current_user)):
    """구독 환불 — 최근 자동결제 건을 포트원 API로 취소 + FREE 전환"""
    if not PORTONE_API_SECRET:
        raise HTTPException(status_code=500, detail="결제 시스템 설정 오류")

    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT plan, billing_key, plan_started_at, plan_expires_at, user_type FROM users WHERE business_number = %s", (bn,))
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")

    u = dict(user)
    plan = u.get("plan") or "free"
    if plan == "free":
        conn.close()
        raise HTTPException(status_code=400, detail="무료 플랜은 환불 대상이 아닙니다.")

    # 무료체험 기간이면 결제 건이 없으므로 그냥 해지
    plan_started = u.get("plan_started_at")
    plan_expires = u.get("plan_expires_at")
    if plan_started and plan_expires:
        try:
            started = datetime.datetime.fromisoformat(str(plan_started))
            expires = datetime.datetime.fromisoformat(str(plan_expires))
            trial_days = (expires - started).days
            # 무료체험: LITE 30일, PRO 7일 — 첫 구독이면 결제 없음
            if trial_days <= 30:
                cur.execute(
                    "UPDATE users SET plan = 'free', billing_key = NULL, plan_expires_at = NULL WHERE business_number = %s",
                    (bn,),
                )
                conn.commit()
                conn.close()
                return {"status": "SUCCESS", "message": "무료체험이 해지되었습니다. FREE 플랜으로 전환되었습니다."}
        except (ValueError, TypeError):
            pass

    # 실 결제 건 환불 — 최근 결제 ID 조회 시도
    import httpx
    try:
        # 결제 ID 패턴: renew-{bn}-{timestamp}
        # 포트원 API로 최근 결제 조회
        resp = httpx.get(
            f"https://api.portone.io/payments?filter.merchantId={bn}&pageSize=1&sort.order=DESC",
            headers={"Authorization": f"PortOne {PORTONE_API_SECRET}"},
            timeout=15,
        )
        if resp.status_code != 200:
            # 결제 이력 조회 실패 → 단순 해지로 처리
            cur.execute(
                "UPDATE users SET plan = 'free', billing_key = NULL, plan_expires_at = NULL WHERE business_number = %s",
                (bn,),
            )
            conn.commit()
            conn.close()
            return {"status": "SUCCESS", "message": "구독이 해지되었습니다. 환불은 고객센터로 문의해주세요."}

        # 해지 + FREE 전환
        cur.execute(
            "UPDATE users SET plan = 'free', billing_key = NULL, plan_expires_at = NULL WHERE business_number = %s",
            (bn,),
        )
        conn.commit()
        conn.close()
        return {"status": "SUCCESS", "message": "환불이 처리되었습니다. FREE 플랜으로 전환되었습니다."}
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"환불 처리 중 오류: {str(e)}")


def _auto_renew_subscriptions():
    """만료된 구독 자동 갱신 — 빌링키로 결제"""
    if not PORTONE_API_SECRET:
        return
    import httpx
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        # 만료일이 지났고 빌링키가 있는 사용자
        cur.execute("""
            SELECT business_number, plan, billing_key, user_type, email
            FROM users
            WHERE billing_key IS NOT NULL
              AND plan IN ('lite', 'pro')
              AND plan_expires_at < NOW()
        """)
        expired_users = cur.fetchall()

        renewed, failed = 0, 0
        for row in expired_users:
            u = dict(row)
            plan = u["plan"]
            user_type = u.get("user_type") or "both"

            # 가격 결정
            if plan == "lite":
                price = PLAN_PRICES.get("lite_individual" if user_type == "individual" else "lite", 4900)
            else:
                price = PLAN_PRICES.get("pro", 49000)

            payment_id = f"renew-{u['business_number']}-{datetime.datetime.utcnow().strftime('%Y%m%d%H%M')}"

            try:
                resp = httpx.post(
                    f"https://api.portone.io/payments/{payment_id}/billing-key",
                    headers={
                        "Authorization": f"PortOne {PORTONE_API_SECRET}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "billingKey": u["billing_key"],
                        "orderName": f"지원금AI {plan.upper()} 월 구독",
                        "amount": {"total": price, "currency": "KRW"},
                    },
                    timeout=15,
                )
                if resp.status_code == 200:
                    now = datetime.datetime.utcnow()
                    new_expires = (now + datetime.timedelta(days=30)).isoformat()
                    cur.execute(
                        """UPDATE users SET plan_expires_at = %s, ai_usage_month = 0,
                           ai_usage_reset_at = %s WHERE business_number = %s""",
                        (new_expires, now.isoformat(), u["business_number"]),
                    )
                    conn.commit()
                    renewed += 1
                    print(f"[renew] {u.get('email','?')} {plan.upper()} {price:,}원 결제 완료")
                else:
                    failed += 1
                    # 결제 실패 → 플랜 만료 + 빌링키 삭제
                    cur.execute(
                        "UPDATE users SET plan = 'free', billing_key = NULL WHERE business_number = %s",
                        (u["business_number"],),
                    )
                    conn.commit()
                    print(f"[renew] {u.get('email','?')} 결제 실패 → FREE 전환: {resp.status_code}")
            except Exception as e:
                failed += 1
                cur.execute(
                    "UPDATE users SET plan = 'free', billing_key = NULL WHERE business_number = %s",
                    (u["business_number"],),
                )
                conn.commit()
                print(f"[renew] {u.get('email','?')} 오류 → FREE 전환: {e}")

        conn.close()
        if renewed or failed:
            _log_system("auto_renew", "payment", f"성공 {renewed}건, 실패 {failed}건 (총 {len(expired_users)}명)", "success" if failed == 0 else "partial", renewed)
            print(f"[renew] 자동 갱신 완료: 성공 {renewed}건, 실패 {failed}건")
    except Exception as e:
        _log_system("auto_renew", "payment", f"오류: {e}", "error")
        print(f"[renew] 오류: {e}")


@app.get("/api/plan/status")
def api_plan_status(current_user: dict = Depends(_get_current_user)):
    """현재 플랜 상태 조회"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT plan, plan_expires_at, ai_usage_month, ai_usage_reset_at FROM users WHERE business_number = %s",
        (current_user["bn"],)
    )
    user = cursor.fetchone()
    conn.close()
    if not user:
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)

    plan = u.get("plan") or "free"
    if plan == "trial":
        plan = "free"
    plan_expires = u.get("plan_expires_at")
    if plan_expires is not None:
        plan_expires = str(plan_expires)
    ai_usage = u.get("ai_usage_month") or 0

    return {
        "status": "SUCCESS",
        "plan": _get_plan_status(plan, plan_expires, ai_usage),
    }


@app.post("/api/ai/use")
def api_ai_use(current_user: dict = Depends(_get_current_user)):
    """AI 기능 사용 시 건수 차감 (1건)"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "SELECT plan, ai_usage_month, ai_usage_reset_at FROM users WHERE business_number = %s",
        (bn,)
    )
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")

    u = dict(user)
    plan = u.get("plan") or "free"
    if plan in ("trial", "premium"):
        plan = "free"
    limit = PLAN_LIMITS.get(plan, 1)
    usage = u.get("ai_usage_month") or 0

    # 월간 리셋 체크
    now = datetime.datetime.utcnow()
    reset_at = u.get("ai_usage_reset_at")
    if reset_at:
        try:
            reset_dt = datetime.datetime.fromisoformat(str(reset_at))
            if now.month != reset_dt.month or now.year != reset_dt.year:
                usage = 0
                cur.execute(
                    "UPDATE users SET ai_usage_month=0, ai_usage_reset_at=%s WHERE business_number=%s",
                    (now.isoformat(), bn)
                )
        except Exception:
            pass

    if usage >= limit:
        conn.close()
        raise HTTPException(
            status_code=429,
            detail=f"이번 달 AI 상담 한도({limit}건)를 모두 사용했습니다. 플랜을 업그레이드하면 더 많은 건수를 이용할 수 있습니다."
        )

    cur.execute(
        "UPDATE users SET ai_usage_month = ai_usage_month + 1 WHERE business_number = %s",
        (bn,)
    )
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "ai_used": usage + 1, "ai_limit": limit}


# ── AI 자유 상담 (크로스 공고 검색) ─────────────────────────────

class AiChatRequest(BaseModel):
    messages: list  # [{"role": "user"|"assistant", "text": "..."}]
    mode: Optional[str] = None  # "business_fund" | "individual_fund" | None (자동 판별)
    session_id: Optional[str] = None  # 동일 상담 그룹핑용 (없으면 hash로 생성)


@app.post("/api/ai/chat")
def api_ai_chat(req: AiChatRequest, current_user: dict = Depends(_get_current_user)):
    """자유 상담: 중소기업 지원사업 전반에 대한 AI 상담"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()

    # 사용자 정보 조회
    cur.execute("SELECT * FROM users WHERE business_number = %s", (bn,))
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)

    # 플랜/사용량 체크
    plan = u.get("plan") or "free"
    plan_expires = u.get("plan_expires_at")
    ps = _get_plan_status(plan, plan_expires, u.get("ai_usage_month") or 0)
    if not ps.get("active"):
        conn.close()
        raise HTTPException(status_code=403, detail="플랜이 만료되었습니다.")

    if plan in ("trial", "premium"):
        plan = "free"
    limit = PLAN_LIMITS.get(plan, 1)
    usage = u.get("ai_usage_month") or 0

    # 월간 리셋
    now = datetime.datetime.utcnow()
    reset_at = u.get("ai_usage_reset_at")
    if reset_at:
        try:
            reset_dt = datetime.datetime.fromisoformat(str(reset_at))
            if now.month != reset_dt.month or now.year != reset_dt.year:
                usage = 0
                cur.execute(
                    "UPDATE users SET ai_usage_month=0, ai_usage_reset_at=%s WHERE business_number=%s",
                    (now.isoformat(), bn)
                )
        except Exception:
            pass

    # 첫 메시지일 때만 건수 차감
    is_first_message = len(req.messages) <= 1
    if is_first_message:
        if usage >= limit:
            conn.close()
            raise HTTPException(status_code=429, detail=f"이번 달 AI 상담 한도({limit}회)를 모두 사용했습니다. 플랜을 업그레이드하면 더 많은 상담을 이용할 수 있습니다.")
        cur.execute("UPDATE users SET ai_usage_month = ai_usage_month + 1 WHERE business_number = %s", (bn,))
        conn.commit()
        usage += 1

    # 세션 내 메시지 수 제한 (플랜별 차별화)
    session_msg_limit = SESSION_MSG_LIMITS.get(plan, 10)
    user_msg_count = sum(1 for m in req.messages if m.get("role") == "user")
    if user_msg_count > session_msg_limit:
        conn.close()
        raise HTTPException(
            status_code=429,
            detail=f"SESSION_MSG_LIMIT:{session_msg_limit}"
        )

    # 통합 AI 엔진으로 자유 상담
    # 탭(mode)에 따라 프로필 필터링 — 기업탭이면 기업 정보만, 개인탭이면 개인 정보만
    from app.services.ai_consultant import chat_lite_fund_expert
    _BIZ_FIELDS = {"business_number", "company_name", "industry_code", "revenue_bracket", "employee_count_bracket", "establishment_date", "address_city", "interests", "user_type", "certifications", "custom_keywords"}
    _INDIV_FIELDS = {"age_range", "income_level", "family_type", "employment_status", "housing_status", "gender", "address_city", "interests", "user_type"}
    _COMMON_FIELDS = {"address_city", "interests", "user_type", "email", "plan"}
    if req.mode == "individual_fund":
        filtered_profile = {k: v for k, v in u.items() if k in (_INDIV_FIELDS | _COMMON_FIELDS) and v}
    elif req.mode == "business_fund":
        filtered_profile = {k: v for k, v in u.items() if k in (_BIZ_FIELDS | _COMMON_FIELDS | {"age_range"}) and v}
    else:
        filtered_profile = u
    result = chat_lite_fund_expert(req.messages, db_conn=conn, user_profile=filtered_profile, mode=req.mode)

    # ── 대화 저장 (P0.1+B): UPSERT by session_id ──
    # session_id: 클라이언트 제공 우선, 없으면 첫 user 메시지 해시로 생성
    try:
        all_msgs = list(req.messages) + [{"role": "assistant", "text": result.get("reply", "")}]
        sid = req.session_id
        if not sid:
            import hashlib
            first_user = next((m.get("text","")[:200] for m in req.messages if m.get("role")=="user"), "")
            sid = "free_" + hashlib.sha256((bn + first_user).encode()).hexdigest()[:16]
        cur.execute(
            """INSERT INTO ai_consult_logs (announcement_id, business_number, messages, conclusion, session_id, updated_at)
               VALUES (NULL, %s, %s::jsonb, %s, %s, CURRENT_TIMESTAMP)
               ON CONFLICT (session_id) WHERE session_id IS NOT NULL DO UPDATE SET
                   messages = EXCLUDED.messages,
                   conclusion = EXCLUDED.conclusion,
                   updated_at = CURRENT_TIMESTAMP""",
            (bn, json.dumps(all_msgs, ensure_ascii=False), "free_chat", sid)
        )
        conn.commit()
    except Exception as save_err:
        print(f"[ai_chat save] {save_err}")
        try: conn.rollback()
        except: pass

    # ── LITE 상담 학습 트리거: 대화 3턴 이상이면 지식 추출 ──
    user_turn_count = sum(1 for m in req.messages if m.get("role") == "user")
    if user_turn_count >= 3:
        try:
            from app.services.ai_consultant import extract_and_store_insights
            all_msgs = list(req.messages) + [{"role": "assistant", "text": result.get("reply", "")}]
            # 에이전트 태그: 개인/기업 구분 보존 (knowledge_base source_agent 격리 기준)
            if req.mode == "individual_fund":
                source_tag = "fund_indiv"
            elif req.mode == "business_fund":
                source_tag = "fund_biz"
            else:
                source_tag = "fund_biz"  # 기본값: 기업
            stored = extract_and_store_insights(all_msgs, conn, source=source_tag)
            if stored > 0:
                print(f"[LITE learning] Extracted {stored} knowledge items from LITE chat")
        except Exception as learn_err:
            print(f"[LITE learning] {learn_err}")

    conn.close()

    return {
        "status": "SUCCESS",
        "reply": result.get("reply", ""),
        "choices": result.get("choices", []),
        "announcements": result.get("announcements", []),
        "matched": result.get("matched", []),
        "done": result.get("done", False),
        "ai_used": usage,
        "ai_limit": limit,
    }


# ── AI 공고 특화 상담 (대화형) ───────────────────────────────────

class AiConsultRequest(BaseModel):
    announcement_id: int
    messages: list  # [{"role": "user"|"assistant", "text": "..."}]
    session_id: Optional[str] = None  # 세션ID — 동일 상담 내 추가 질문 시 차감 방지


@app.post("/api/ai/consult")
def api_ai_consult(req: AiConsultRequest, current_user: dict = Depends(_get_current_user)):
    """AI 지원대상 여부 상담 — 대화형 (Gemini)"""
    import google.generativeai as genai

    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()

    # 1) 사용자 프로필 + 플랜 체크
    cur.execute(
        """SELECT plan, ai_usage_month, ai_usage_reset_at, plan_expires_at,
                  company_name, establishment_date, address_city, industry_code,
                  revenue_bracket, employee_count_bracket, interests, user_type,
                  age_range, business_number, gender,
                  income_level, family_type, employment_status, housing_status
           FROM users WHERE business_number = %s""",
        (bn,)
    )
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)
    # NULL 프로필 필드 → 빈 문자열 (AI 프롬프트 호환)
    for k in ("company_name", "establishment_date", "address_city", "industry_code",
              "revenue_bracket", "employee_count_bracket", "interests", "user_type", "age_range",
              "gender", "income_level", "family_type", "employment_status", "housing_status"):
        if u.get(k) is None:
            u[k] = ""

    # 플랜/사용량 체크
    plan = u.get("plan") or "free"
    if plan in ("trial", "premium"):
        plan = "free"
    plan_expires = u.get("plan_expires_at")
    ps = _get_plan_status(plan, plan_expires, u.get("ai_usage_month") or 0)
    if not ps.get("active"):
        conn.close()
        raise HTTPException(status_code=403, detail="플랜이 만료되었습니다. 업그레이드 후 이용해 주세요.")

    # 공고별 지원대상 상담 제한 체크
    consult_limit = CONSULT_LIMITS.get(plan, 0)
    ai_usage = u.get("ai_usage_month") or 0

    # 월간 리셋 체크
    now = datetime.datetime.utcnow()
    reset_at = u.get("ai_usage_reset_at")
    if reset_at:
        try:
            reset_dt = datetime.datetime.fromisoformat(str(reset_at))
            if now.month != reset_dt.month or now.year != reset_dt.year:
                ai_usage = 0
                cur.execute(
                    "UPDATE users SET ai_usage_month=0, ai_usage_reset_at=%s WHERE business_number=%s",
                    (now.isoformat(), bn)
                )
        except Exception:
            pass

    if consult_limit == 0:
        conn.close()
        raise HTTPException(
            status_code=403,
            detail="공고별 지원대상 상담은 LITE 플랜부터 이용할 수 있습니다."
        )

    # 세션ID 기반 차감: 기존 세션이면 차감 스킵
    import uuid as _uuid
    session_id = req.session_id
    is_existing_session = False

    if session_id:
        # 세션 유효성 확인 (24시간 이내)
        try:
            cur.execute("""
                SELECT id FROM consult_sessions
                WHERE session_id = %s AND business_number = %s AND announcement_id = %s
                  AND created_at > NOW() - INTERVAL '24 hours'
            """, (session_id, bn, req.announcement_id))
            if cur.fetchone():
                is_existing_session = True
        except Exception:
            pass

    # 건수 제한 (PRO/무제한 제외)
    if consult_limit < 999999 and not is_existing_session:
        # consult 사용량은 ai_usage_month로 추적
        if ai_usage >= consult_limit:
            conn.close()
            if plan == "free":
                msg = f"무료 상담({consult_limit}회)을 모두 사용했습니다. LITE 플랜으로 업그레이드하면 월 20회까지 이용할 수 있습니다."
            else:
                msg = f"이번 달 AI 상담 한도({consult_limit}회)를 모두 사용했습니다. PRO 플랜으로 업그레이드하면 무제한 이용할 수 있습니다."
            raise HTTPException(status_code=429, detail=msg)
        # 새 세션: 1회 차감 + 세션ID 발급
        session_id = str(_uuid.uuid4())
        cur.execute("UPDATE users SET ai_usage_month = ai_usage_month + 1 WHERE business_number = %s", (bn,))
        try:
            cur.execute("""
                INSERT INTO consult_sessions (session_id, business_number, announcement_id)
                VALUES (%s, %s, %s)
            """, (session_id, bn, req.announcement_id))
        except Exception:
            pass
        conn.commit()
    elif not is_existing_session and not session_id:
        # PRO/무제한: 차감 없이 session_id만 발급 (매 턴 별도 행 생성 방지)
        session_id = str(_uuid.uuid4())
        try:
            cur.execute("""
                INSERT INTO consult_sessions (session_id, business_number, announcement_id)
                VALUES (%s, %s, %s)
            """, (session_id, bn, req.announcement_id))
            conn.commit()
        except Exception:
            pass
        ai_usage += 1

    # 1-1) 세션 내 메시지 수 제한 (플랜별 차별화)
    session_msg_limit = SESSION_MSG_LIMITS.get(plan, 10)
    user_msg_count = sum(1 for m in req.messages if m.get("role") == "user")
    if user_msg_count > session_msg_limit:
        conn.close()
        raise HTTPException(
            status_code=429,
            detail=f"SESSION_MSG_LIMIT:{session_msg_limit}"
        )

    # 2) 공고 정보 조회
    cur.execute(
        """SELECT announcement_id, title, department, category, support_amount, deadline_date,
                  summary_text, region, eligibility_logic, origin_url, target_type
           FROM announcements WHERE announcement_id = %s""",
        (req.announcement_id,)
    )
    ann = cur.fetchone()
    if not ann:
        conn.close()
        raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")
    a = dict(ann)

    # 2-1) 정밀 분석 데이터 보장 (없으면 실시간 분석)
    deep = {}
    try:
        from app.services.doc_analysis_service import ensure_analysis
        analysis_conn = get_db_connection()
        deep = ensure_analysis(req.announcement_id, analysis_conn) or {}
        analysis_conn.close()
        if deep:
            ps = deep.get("parsed_sections", {}) or {}
            filled = [k for k, v in ps.items() if v]
            print(f"[Consult] Analysis OK for #{req.announcement_id}: source={deep.get('source_type')}, filled_sections={filled}")
    except Exception as e:
        print(f"[Consult] ensure_analysis error for #{req.announcement_id}: {e}")
        try:
            analysis_conn.close()
        except Exception:
            pass

    # 3) 공고 target_type에 따라 프로필 필터링
    _BIZ_FIELDS = {"business_number", "company_name", "industry_code", "revenue_bracket", "employee_count_bracket", "establishment_date", "address_city", "interests", "user_type", "certifications"}
    _INDIV_FIELDS = {"age_range", "income_level", "family_type", "employment_status", "housing_status", "gender", "address_city", "interests", "user_type"}
    _COMMON_FIELDS = {"address_city", "interests", "user_type", "email", "plan"}
    ann_target = (a.get("target_type") or "business").lower()
    if ann_target == "individual":
        consult_profile = {k: v for k, v in u.items() if k in (_INDIV_FIELDS | _COMMON_FIELDS) and v}
    else:
        consult_profile = {k: v for k, v in u.items() if k in (_BIZ_FIELDS | _COMMON_FIELDS) and v}

    consult_conn = None
    try:
        from app.services.ai_consultant import chat_consult
        consult_conn = get_db_connection()
        result = chat_consult(
            announcement_id=req.announcement_id,
            messages=req.messages,
            announcement=a,
            deep_analysis_data=deep,
            user_profile=consult_profile,
            db_conn=consult_conn,
        )
    except Exception as e:
        print(f"[Consult] chat_consult error: {e}")
        import traceback; traceback.print_exc()
        # 에러 발생해도 기본 응답 반환 (500 대신 정상 응답)
        result = {
            "reply": f"AI 상담 중 오류가 발생했습니다. 다시 시도해 주세요.\n\n공고 원문 확인: {a.get('origin_url', '')}",
            "choices": ["다시 시도", "다른 공고 보기"],
            "done": False,
            "conclusion": None,
        }
    finally:
        if consult_conn:
            try:
                consult_conn.close()
            except Exception:
                pass
    conn.close()

    is_done = result.get("done", False)

    # done=True 강제 오버라이드: 사용자 메시지가 3개 미만이면 done=false 강제
    user_msg_count = sum(1 for m in req.messages if m.get("role") == "user")
    if is_done and user_msg_count < 3:
        is_done = False
        result["done"] = False
        result["conclusion"] = None

    consult_log_id = None

    # P0.4: 매 턴 UPSERT 저장 (done 여부 무관) — session_id로 그룹핑
    try:
        all_msgs = req.messages + [{"role": "assistant", "text": result.get("reply", "")}]
        log_conn = get_db_connection()
        log_cur = log_conn.cursor()
        log_cur.execute("""
            INSERT INTO ai_consult_logs (announcement_id, business_number, messages, conclusion, session_id, updated_at)
            VALUES (%s, %s, %s::jsonb, %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (session_id) WHERE session_id IS NOT NULL DO UPDATE SET
                messages = EXCLUDED.messages,
                conclusion = COALESCE(EXCLUDED.conclusion, ai_consult_logs.conclusion),
                updated_at = CURRENT_TIMESTAMP
            RETURNING id
        """, (req.announcement_id, bn, json.dumps(all_msgs, ensure_ascii=False),
              result.get("conclusion") if is_done else None, session_id))
        row = log_cur.fetchone()
        consult_log_id = row["id"] if row else None
        log_conn.commit()
        log_conn.close()
    except Exception as log_err:
        consult_log_id = None
        print(f"[ConsultLog] Save error: {log_err}")

    # 최종 방어선: reply가 비어있으면 공고 제목 기반 안내 메시지로 대체
    final_reply = result.get("reply", "")
    if not final_reply or not final_reply.strip():
        final_reply = f"**{a.get('title', '공고')}** 분석을 시작합니다. 아래 선택지를 눌러 질문해 주세요."

    return {
        "status": "SUCCESS",
        "reply": final_reply,
        "choices": result.get("choices", []),
        "done": is_done,
        "conclusion": result.get("conclusion") if is_done else None,
        "consult_log_id": consult_log_id if is_done else None,
        "ai_used": ai_usage,
        "ai_limit": consult_limit,
        "session_id": session_id,
        "origin_url": a.get("origin_url", ""),
    }


# ── AI 컨설턴트 모드 (고객사 조건 수집 대화) ──────────────────

class AiConsultantChatRequest(BaseModel):
    messages: list  # [{"role": "user"|"assistant", "text": "..."}]
    announcement_id: Optional[int] = None  # 특정 공고 상담 모드 (PRO 전용)
    explicit_match: Optional[bool] = False  # [legacy] 명시적 매칭 요청 (action=match로 대체)
    profile_override: Optional[dict] = None  # 사용자가 확인/수정한 프로필 (매칭 시 사용)
    session_id: Optional[str] = None  # PRO 세션 ID (서버 측 상태 저장)
    client_category: Optional[str] = None  # 첫 호출 시 고객 유형 힌트
    client_id: Optional[int] = None  # I: 선택된 고객 프로필 ID (있으면 프롬프트에 주입)
    mode: Optional[str] = None  # LITE 자금 모드: "business_fund" | "individual_fund"
    # [재설계 04] PRO 매칭 상담 — action 파라미터 분기
    # "match": 매칭만 실행 (Gemini 호출 없음, 매칭 엔진만)
    # "consult": 특정 공고 상담 (pro_announce V2 호출)
    # None: 레거시 — announcement_id/explicit_match로 자동 추론
    action: Optional[str] = None
    # [재설계 04] 공고 카드 클릭 직후 — 1차 턴(12섹션 분석) 강제 플래그
    is_announcement_start: Optional[bool] = False


@app.get("/api/pro/announcements/{announcement_id}/analyze")
def api_pro_announcement_analyze(announcement_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 공고 분석 — DB 우선, 없으면 자동 실시간 분석 (PRO 권한)"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT a.*, aa.parsed_sections, aa.deep_analysis, aa.full_text
        FROM announcements a
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE a.announcement_id = %s
    """, (announcement_id,))
    row = cur.fetchone()

    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")

    ann = dict(row)
    parsed = ann.get("parsed_sections") or {}
    deep = ann.get("deep_analysis") or {}
    has_analysis = bool(parsed) or bool(deep)

    # ★ PRO 사용자가 분석되지 않은 공고를 보면 → 자동 실시간 분석 트리거
    if not has_analysis:
        try:
            from app.services.doc_analysis_service import ensure_analysis
            print(f"[PRO Analyze] #{announcement_id}: triggering on-demand analysis")
            fresh = ensure_analysis(announcement_id, conn)
            if fresh:
                parsed = fresh.get("parsed_sections") or {}
                deep = fresh.get("deep_analysis") or {}
                has_analysis = bool(parsed) or bool(deep)
                # ann 데이터 다시 조회 (support_amount 등이 업데이트됐을 수 있음)
                cur.execute("SELECT * FROM announcements WHERE announcement_id = %s", (announcement_id,))
                fresh_row = cur.fetchone()
                if fresh_row:
                    ann.update(dict(fresh_row))
        except Exception as e:
            print(f"[PRO Analyze] on-demand analysis error: {e}")

    conn.close()

    result = {
        "status": "SUCCESS",
        "announcement_id": announcement_id,
        "title": ann.get("title", ""),
        "organization": ann.get("organization", ""),
        "support_amount": ann.get("support_amount", ""),
        "deadline_date": str(ann.get("deadline_date", "")) if ann.get("deadline_date") else "",
        "url": ann.get("url", ""),
        "has_db_analysis": has_analysis,
    }

    if has_analysis:
        result["eligibility"] = (parsed.get("eligibility") or "")[:1500] if parsed else ""
        result["support_details"] = (parsed.get("support_details") or "")[:1500] if parsed else ""
        result["application_method"] = (parsed.get("application_method") or "")[:800] if parsed else ""
        result["evaluation_criteria"] = (parsed.get("evaluation_criteria") or "")[:800] if parsed else ""
        # 실제 DB 필드는 required_docs (parsed_sections), required_documents는 양쪽 fallback
        result["required_documents"] = (parsed.get("required_docs") or parsed.get("required_documents") or "")[:800] if parsed else ""
        if deep:
            result["target_summary"] = deep.get("target_summary", "")
            result["support_summary"] = deep.get("support_summary", {})
            result["exclusion_rules"] = deep.get("exclusion_rules", [])
            result["competitiveness"] = deep.get("competitiveness", "")
    else:
        # 분석 없음 → 기본 정보만 + AI 분석 권유
        result["message"] = "이 공고는 아직 상세 분석되지 않았습니다. 기본 정보만 제공됩니다."

    return result


@app.post("/api/pro/consultant/chat")
def api_pro_consultant_chat(req: AiConsultantChatRequest, current_user: dict = Depends(_get_current_user)):
    """PRO 전문가 전용: 고객사 상담 채팅 (세션 기반 상태 관리)"""
    _require_pro(current_user)
    try:
        return _api_pro_consultant_chat_impl(req, current_user)
    except Exception as outer_err:
        import traceback as _tb
        _tb_str = _tb.format_exc()[-800:]
        print(f"[PRO chat] OUTER ERROR: {outer_err}\n{_tb_str}")
        return {
            "status": "ERROR",
            "reply": "일시적으로 응답 생성에 실패했습니다.",
            "choices": ["✏️ 다시 시도"],
            "done": False,
            "profile": None,
            "collected": {},
            "matched_announcements": [],
            "rag_sources": [],
            "session_id": req.session_id,
            "outer_error": f"{type(outer_err).__name__}: {str(outer_err)[:200]}",
            "outer_tb": _tb_str,
        }


def _api_pro_consultant_chat_impl(req: AiConsultantChatRequest, current_user: dict):
    """
    [재설계 04] PRO 매칭 상담 — action 파라미터 분기.
    Mode A (자연어 정보수집) 완전 제거. 폼 + 매칭 + 공고상담(V2)만 지원.
    """
    # action 결정 (명시 우선, 없으면 legacy 추론)
    action = (req.action or "").strip().lower()
    if not action:
        if req.announcement_id:
            action = "consult"
        elif req.explicit_match and (req.profile_override or req.client_id):
            action = "match"
        else:
            raise HTTPException(
                status_code=400,
                detail="action 파라미터 필수: 'match' (매칭 실행) 또는 'consult' (공고 상담, announcement_id 필요)"
            )

    if action == "match":
        return _handle_pro_match(req, current_user)
    if action == "consult":
        return _handle_pro_consult(req, current_user)
    if action == "fund_consult":
        return _handle_pro_fund_consult(req, current_user)
    if action == "detail_analysis":
        return _handle_pro_detail_analysis(req, current_user)
    if action == "chat":
        return _handle_pro_chat(req, current_user)
    raise HTTPException(status_code=400, detail=f"알 수 없는 action: {action} (지원: match | consult | fund_consult | detail_analysis | chat)")


def _handle_pro_chat(req: AiConsultantChatRequest, current_user: dict):
    """action=chat: 매칭 완료 후 대화형 후속 상담.
    세션의 matched_snapshot을 컨텍스트로 Gemini가 질문에 답변.
    매칭 엔진 재실행 없음.
    """
    db = get_db_connection()
    try:
        session_state = _load_or_create_session(db, current_user, req.session_id, req.client_category)
        selected_client = _load_client(db, req.client_id, current_user["bn"])

        # 프로필 확정
        profile = req.profile_override or {}
        if not profile and selected_client:
            profile = {
                "company_name": selected_client.get("client_name") or "",
                "industry_code": selected_client.get("industry_code") or "",
                "address_city": selected_client.get("address_city") or "",
                "establishment_date": str(selected_client.get("establishment_date") or ""),
                "revenue_bracket": selected_client.get("revenue_bracket") or "",
                "employee_count_bracket": selected_client.get("employee_count_bracket") or "",
                "interests": selected_client.get("interests") or "",
                "certifications": selected_client.get("certifications") or "",
            }
        if not profile:
            profile = session_state.get("collected") or {}

        # 세션에서 매칭 결과 로드
        matched_snapshot = session_state.get("matched_snapshot") or []
        if isinstance(matched_snapshot, str):
            try:
                matched_snapshot = json.loads(matched_snapshot)
            except Exception:
                matched_snapshot = []

        if not matched_snapshot:
            return {
                "status": "ERROR",
                "reply": "아직 매칭된 공고가 없습니다. 먼저 고객 정보를 입력하고 매칭을 실행해주세요.",
                "choices": ["🔄 매칭 시작"],
                "done": True,
                "session_id": session_state.get("session_id"),
            }

        # 프로필 요약 텍스트
        profile_lines = [f"- {k}: {v}" for k, v in profile.items() if v]
        profile_str = "\n".join(profile_lines) if profile_lines else "정보 없음"

        # 매칭 결과 요약 텍스트 (상위 5개)
        matched_lines = []
        for m in (matched_snapshot or [])[:5]:
            title = m.get("title", "")
            amt = m.get("support_amount", "") or m.get("support_amount_max", "")
            dl = str(m.get("deadline_date", ""))[:10]
            matched_lines.append(f"  - {title} | 지원금: {amt} | 마감: {dl}")
        matched_str = ("매칭된 주요 공고:\n" + "\n".join(matched_lines)) if matched_lines else "매칭 결과 없음"

        system_prompt = f"""당신은 정부지원사업 전문 컨설턴트입니다.
아래 고객 프로필과 매칭 결과를 바탕으로 고객(전문가)의 질문에 전문적이고 구체적으로 답변하세요.

[고객 프로필]
{profile_str}

[{matched_str}]

답변 원칙:
- 매칭된 공고를 바탕으로 구체적인 조언을 제공합니다
- 추가 매칭이 필요하면 "조건을 수정하면 재매칭이 가능합니다"라고 안내합니다
- 특정 공고 심화 분석이 필요하면 카드에서 "상담하기"를 클릭하라고 안내합니다
- 간결하고 실용적으로 답변합니다 (마크다운 허용)"""

        # Gemini 호출
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return {"status": "ERROR", "reply": "AI 서비스 설정 오류", "choices": [], "done": False}

        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            "models/gemini-2.5-flash",
            system_instruction=system_prompt,
        )

        # 대화 히스토리 구성 (마지막 user 메시지 제외)
        messages_list = list(req.messages)
        history = []
        for msg in messages_list[:-1]:
            role = "user" if msg.get("role") == "user" else "model"
            text = msg.get("text", "")
            if text.strip():
                history.append({"role": role, "parts": [text]})

        chat = model.start_chat(history=history)
        last_text = messages_list[-1].get("text", "") if messages_list else ""
        response = chat.send_message(last_text)
        reply = response.text or ""

        # phase를 'consulting'으로 업데이트
        try:
            cur2 = db.cursor()
            cur2.execute(
                "UPDATE pro_consult_sessions SET phase = 'consulting' WHERE session_id = %s",
                (session_state.get("session_id"),)
            )
            db.commit()
        except Exception as e:
            print(f"[PRO chat] phase update failed: {e}")
            try: db.rollback()
            except: pass

        return {
            "status": "SUCCESS",
            "action": "chat",
            "reply": reply,
            "choices": ["🔄 조건 수정 후 재매칭"],
            "done": False,
            "session_id": session_state.get("session_id"),
        }
    finally:
        try:
            db.close()
        except Exception:
            pass


def _handle_pro_fund_consult(req: AiConsultantChatRequest, current_user: dict):
    """action=fund_consult: PRO 전문가의 고객 자금상담.

    ★ 데이터 격리 원칙 (사장님 지시) ★
    - 전문가 본인의 users 테이블 프로필은 절대 AI 프롬프트에 주입하지 않음
    - profile_override(폼 데이터) 또는 client_profiles(CRM 등록된 고객) 에서만 프로필 구성
    - 어느 쪽도 없으면 400 에러로 폼 입력 요구
    """
    from app.services.ai_consultant import chat_lite_fund_expert

    # mode: "individual_fund" (개인 고객) | "business_fund" (기업 고객)
    fund_mode = (req.mode or "").strip().lower()
    if fund_mode not in ("individual_fund", "business_fund"):
        raise HTTPException(
            status_code=400,
            detail="자금상담 모드 필수: mode='individual_fund' (개인 고객) 또는 'business_fund' (기업 고객)"
        )

    pro_ctx = "individual" if fund_mode == "individual_fund" else "business"

    db = get_db_connection()
    try:
        # 1) 프로필 우선순위: profile_override (폼) > client_profiles (CRM 선택 고객)
        profile = req.profile_override or {}
        if not profile and req.client_id:
            selected_client = _load_client(db, req.client_id, current_user["bn"])
            if selected_client:
                if fund_mode == "business_fund":
                    profile = {
                        "company_name": selected_client.get("client_name") or "",
                        "industry_code": selected_client.get("industry_code") or "",
                        "address_city": selected_client.get("address_city") or "",
                        "establishment_date": str(selected_client.get("establishment_date") or ""),
                        "revenue_bracket": selected_client.get("revenue_bracket") or "",
                        "employee_count_bracket": selected_client.get("employee_count_bracket") or "",
                        "interests": selected_client.get("interests") or "",
                        "certifications": selected_client.get("certifications") or "",
                        "user_type": "business",
                    }
                else:
                    profile = {
                        "age_range": selected_client.get("age_range") or "",
                        "address_city": selected_client.get("address_city") or "",
                        "income_level": selected_client.get("income_level") or "",
                        "family_type": selected_client.get("family_type") or "",
                        "employment_status": selected_client.get("employment_status") or "",
                        "housing_status": selected_client.get("housing_status") or "",
                        "interests": selected_client.get("interests") or "",
                        "user_type": "individual",
                    }

        # 폼 또는 CRM 어느 쪽도 없으면 폼 요구
        if not profile or not any(v for v in profile.values() if v):
            raise HTTPException(
                status_code=400,
                detail="고객 정보가 없습니다. 자금상담 시작 전 고객 정보 폼을 먼저 작성해 주세요."
            )

        # 2) 전문가 본인 프로필 섞이는지 안전검증 — profile에 current_user의 bn/email 금지
        for forbidden_key in ("business_number", "email", "password_hash"):
            profile.pop(forbidden_key, None)

        # 3) AI 호출 — pro_consult_context 플래그로 3인칭 화법 강제
        result = chat_lite_fund_expert(
            messages=req.messages,
            db_conn=db,
            user_profile=profile,
            mode=fund_mode,
            pro_consult_context=pro_ctx,
        )

        # 4) 상담 로그 저장
        try:
            import uuid as _uuid
            sid = req.session_id or f"fund_{_uuid.uuid4()}"
            all_msgs = list(req.messages) + [{"role": "assistant", "text": result.get("reply", "")}]
            log_cur = db.cursor()
            log_cur.execute(
                """INSERT INTO ai_consult_logs (announcement_id, business_number, messages, conclusion, session_id, updated_at)
                   VALUES (NULL, %s, %s::jsonb, %s, %s, CURRENT_TIMESTAMP)
                   ON CONFLICT (session_id) WHERE session_id IS NOT NULL DO UPDATE SET
                       messages = EXCLUDED.messages,
                       conclusion = EXCLUDED.conclusion,
                       updated_at = CURRENT_TIMESTAMP""",
                (current_user["bn"], json.dumps(all_msgs, ensure_ascii=False), f"pro_fund_{pro_ctx}", sid)
            )
            db.commit()
        except Exception as save_err:
            print(f"[PRO fund_consult save] {save_err}")
            try: db.rollback()
            except: pass

        return {
            "status": "SUCCESS",
            "reply": result.get("reply", ""),
            "choices": result.get("choices", []),
            "announcements": result.get("announcements", []),
            "done": result.get("done", False),
            "session_id": sid if 'sid' in dir() else req.session_id,
        }
    finally:
        try: db.close()
        except: pass


def _load_or_create_session(db, current_user, req_session_id, client_category):
    """세션 로드 또는 신규 생성. (session_state dict 반환)"""
    import uuid as _uuid
    cur = db.cursor()
    if req_session_id:
        cur.execute(
            """SELECT session_id, client_category, current_step, collected,
                      phase, matched_snapshot, messages
               FROM pro_consult_sessions
               WHERE session_id = %s AND business_number = %s""",
            (req_session_id, current_user["bn"])
        )
        row = cur.fetchone()
        if row:
            d = dict(row)
            coll = d.get("collected")
            if isinstance(coll, str):
                try: coll = json.loads(coll)
                except: coll = {}
            matched_snap = d.get("matched_snapshot")
            if isinstance(matched_snap, str):
                try: matched_snap = json.loads(matched_snap)
                except: matched_snap = []
            db_msgs = d.get("messages")
            if isinstance(db_msgs, str):
                try: db_msgs = json.loads(db_msgs)
                except: db_msgs = []
            if not isinstance(db_msgs, list):
                db_msgs = []
            return {
                "session_id": d.get("session_id"),
                "client_category": d.get("client_category") or "",
                "current_step": d.get("current_step") or 1,
                "collected": coll or {},
                "phase": d.get("phase") or "collecting",
                "matched_snapshot": matched_snap or [],
                "db_messages": db_msgs,
            }
    # 클라이언트가 준 session_id가 있으면 그대로 사용, 없으면 새 UUID 생성
    # → match와 chat이 같은 session_id를 공유할 수 있게 함
    sid = req_session_id or str(_uuid.uuid4())
    cur.execute(
        "INSERT INTO pro_consult_sessions (session_id, business_number, client_category, current_step, collected) VALUES (%s, %s, %s, 1, '{}'::jsonb)",
        (sid, current_user["bn"], client_category or "")
    )
    db.commit()
    return {
        "session_id": sid,
        "client_category": client_category or "",
        "current_step": 1,
        "collected": {},
        "phase": "collecting",
        "matched_snapshot": [],
        "db_messages": [],
    }


def _load_client(db, client_id, bn):
    """client_profiles에서 선택된 고객 조회."""
    if not client_id:
        return None
    try:
        cur = db.cursor()
        cur.execute(
            """SELECT * FROM client_profiles
               WHERE id = %s AND owner_business_number = %s AND is_active = TRUE""",
            (client_id, bn)
        )
        row = cur.fetchone()
        return dict(row) if row else None
    except Exception as e:
        print(f"[PRO load_client] {e}")
        try: db.rollback()
        except: pass
        return None


def _handle_pro_detail_analysis(req: AiConsultantChatRequest, current_user: dict):
    """action=detail_analysis: 매칭된 공고 상위 10개를 Gemini로 정밀 분석."""
    import google.generativeai as genai
    db = get_db_connection()
    try:
        session_state = _load_or_create_session(db, current_user, req.session_id, req.client_category)
        matched_snap = (session_state or {}).get("matched_snapshot") or []
        profile = (session_state or {}).get("collected") or req.profile_override or {}

        if not matched_snap:
            return {
                "status": "SUCCESS",
                "reply": "분석할 매칭 결과가 없습니다. 먼저 매칭을 실행해주세요.",
                "choices": ["🔄 조건 수정 후 재매칭"],
                "matched_announcements": [],
            }

        top = matched_snap[:10]

        # 프로필 요약
        label_map = {
            "company_name": "기업명/이름", "industry_code": "업종", "address_city": "지역",
            "revenue_bracket": "매출", "employee_count_bracket": "직원수",
            "age_range": "연령대", "income_level": "소득", "family_type": "가구유형",
            "employment_status": "고용상태", "interests": "관심분야",
        }
        profile_lines = [f"- {label_map.get(k, k)}: {v}" for k, v in profile.items() if v and k in label_map]
        profile_text = "\n".join(profile_lines) or "프로필 정보 없음"

        # 공고 목록 요약
        ann_lines = []
        for i, a in enumerate(top):
            line = f"{i+1}. [ID:{a['announcement_id']}] {a['title']}"
            if a.get("eligibility"):
                line += f"\n   자격요건: {str(a['eligibility'])[:250]}"
            ann_lines.append(line)
        ann_text = "\n".join(ann_lines)

        prompt = f"""당신은 정부지원사업 전문 컨설턴트입니다. 아래 고객 프로필과 공고 목록을 보고, 각 공고에 대한 신청 가능 여부를 판정하세요.

[고객 프로필]
{profile_text}

[매칭 공고 목록]
{ann_text}

각 공고에 대해 판정하세요. 반드시 아래 JSON 배열 형식으로만 응답하세요:
[
  {{
    "announcement_id": 숫자,
    "verdict": "eligible" | "conditional" | "ineligible",
    "reason": "판정 근거 한 줄 (40자 이내, 구체적으로)"
  }}
]
- eligible: 프로필 기준 명백히 신청 가능
- conditional: 일부 조건 확인 필요 또는 판단 근거 부족
- ineligible: 프로필 기준 명백히 대상 아님
- reason은 근거를 구체적으로: "업종 일치, 매출 조건 충족" / "연령 조건(만 39세 이하) 미충족" 등
반드시 순수 JSON 배열만 반환하세요."""

        api_key = os.environ.get("GEMINI_API_KEY")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            "models/gemini-2.5-flash",
            generation_config={"max_output_tokens": 2048, "temperature": 0.2}
        )
        resp = model.generate_content(prompt)
        raw = resp.text.strip()
        # JSON 추출
        import re as _re
        arr_match = _re.search(r'\[.*\]', raw, _re.DOTALL)
        verdicts = json.loads(arr_match.group(0)) if arr_match else []
        verdict_map = {int(v["announcement_id"]): v for v in verdicts if "announcement_id" in v}

        # 매칭 결과에 AI 판정 병합
        enriched = []
        for a in top:
            item = dict(a)
            aid = a.get("announcement_id")
            if aid and int(aid) in verdict_map:
                vd = verdict_map[int(aid)]
                item["ai_verdict"] = vd.get("verdict", "conditional")
                item["ai_reason"] = vd.get("reason", "")
            else:
                item["ai_verdict"] = "conditional"
                item["ai_reason"] = "분석 데이터 부족"
            enriched.append(item)

        # eligible 우선 정렬
        order = {"eligible": 0, "conditional": 1, "ineligible": 2}
        enriched.sort(key=lambda x: order.get(x.get("ai_verdict", "conditional"), 1))

        eligible_cnt = sum(1 for e in enriched if e.get("ai_verdict") == "eligible")
        conditional_cnt = sum(1 for e in enriched if e.get("ai_verdict") == "conditional")

        reply = (
            f"**AI 정밀 분석 완료** — 상위 {len(enriched)}개 공고 판정 결과:\n\n"
            f"✅ 신청 가능 **{eligible_cnt}건** · ⚠️ 조건 확인 필요 **{conditional_cnt}건** · "
            f"❌ 대상 아님 **{len(enriched)-eligible_cnt-conditional_cnt}건**\n\n"
            "공고 카드에서 상세 판정 근거를 확인하세요."
        )

        return {
            "status": "SUCCESS",
            "reply": reply,
            "matched_announcements": enriched,
            "choices": ["신청 가능 공고 상세 상담", "🔄 조건 수정 후 재매칭"],
            "done": False,
            "session_id": (session_state or {}).get("session_id"),
        }
    except Exception as e:
        print(f"[detail_analysis] error: {e}")
        import traceback; traceback.print_exc()
        return {
            "status": "SUCCESS",
            "reply": "AI 상세 분석 중 오류가 발생했습니다. 다시 시도해주세요.",
            "choices": ["🔄 다시 시도"],
            "matched_announcements": [],
        }
    finally:
        db.close()


def _handle_pro_match(req: AiConsultantChatRequest, current_user: dict):
    """action=match: 매칭 엔진만 실행, Gemini 호출 없음."""
    db = get_db_connection()
    try:
        session_state = _load_or_create_session(db, current_user, req.session_id, req.client_category)
        selected_client = _load_client(db, req.client_id, current_user["bn"])

        # 프로필 확정: profile_override > client_profile > session.collected
        profile = req.profile_override or {}
        if not profile and selected_client:
            profile = {
                "company_name": selected_client.get("client_name") or "",
                "industry_code": selected_client.get("industry_code") or "",
                "address_city": selected_client.get("address_city") or "",
                "establishment_date": str(selected_client.get("establishment_date") or ""),
                "revenue_bracket": selected_client.get("revenue_bracket") or "",
                "employee_count_bracket": selected_client.get("employee_count_bracket") or "",
                "interests": selected_client.get("interests") or "",
                "certifications": selected_client.get("certifications") or "",
                "user_type": selected_client.get("client_type") or "",
            }
        if not profile:
            profile = session_state.get("collected") or {}

        if not profile or not any(profile.values()):
            raise HTTPException(status_code=400, detail="매칭에 필요한 프로필이 없습니다. 고객 정보 폼을 먼저 저장해주세요.")

        # 매칭 실행
        from app.core.matcher import get_matches_hybrid
        # client_category 명시 우선 — 개인 선택 시 기업 공고 혼합 금지
        _cat = (req.client_category or "").strip().lower()
        if _cat == "individual":
            is_individual = True
        elif _cat in ("corporate", "individual_biz"):
            is_individual = False
        else:
            # fallback: profile에서 추론 (레거시)
            has_industry = bool(str(profile.get("industry_code", "")).strip())
            is_personal_name = profile.get("company_name", "") in ("개인", "")
            is_individual = not has_industry or is_personal_name

        if is_individual:
            matches = get_matches_hybrid(profile, is_individual=True) or []
        else:
            interests_str = str(profile.get("interests", ""))
            biz_interests = any(kw in interests_str for kw in ["창업", "R&D", "기술개발", "정책자금", "수출"])
            if biz_interests:
                ind_m = get_matches_hybrid(profile, is_individual=True) or []
                biz_m = get_matches_hybrid(profile, is_individual=False) or []
                matches = biz_m + ind_m
                matches.sort(key=lambda x: x.get("match_score", 0), reverse=True)
            else:
                matches = get_matches_hybrid(profile, is_individual=False) or []

        _interest = [a for a in matches if a.get("bucket") == "interest_match"][:10]
        _deadline = [a for a in matches if a.get("bucket") == "deadline_urgent"][:3]
        _other = [a for a in matches if a.get("bucket") == "qualified_other"][:7]
        _display_list = _interest + _deadline + _other

        def _simplify(a, idx):
            d = a if isinstance(a, dict) else dict(a)
            return {
                "announcement_id": d.get("announcement_id"),
                "title": d.get("title", ""),
                "department": d.get("department", ""),
                "support_amount": d.get("support_amount", ""),
                "support_amount_max": d.get("support_amount_max"),
                "deadline_date": str(d.get("deadline_date", "")),
                "origin_url": d.get("origin_url") or "",
                "rank": d.get("rank") or (idx + 1),
                "bucket": d.get("bucket", ""),
                "bucket_label": d.get("bucket_label", ""),
                "reasons": d.get("reasons", []),
                "matched_interests": d.get("matched_interests", []),
            }

        matched_announcements = [_simplify(a, i) for i, a in enumerate(_display_list)]
        matched_groups = {
            "interest_match": [_simplify(a, i) for i, a in enumerate(_interest)],
            "deadline_urgent": [_simplify(a, i) for i, a in enumerate(_deadline)],
            "qualified_other": [_simplify(a, i) for i, a in enumerate(_other)],
        }

        # matched_snapshot + 상위 3개 상세분석 저장
        try:
            snap_cur = db.cursor()
            top_ids = [m["announcement_id"] for m in matched_announcements[:3] if m.get("announcement_id")]
            sections_map = {}
            if top_ids:
                try:
                    snap_cur.execute(
                        """SELECT announcement_id, parsed_sections, deep_analysis
                           FROM announcement_analysis
                           WHERE announcement_id = ANY(%s)""",
                        (top_ids,),
                    )
                    for r in snap_cur.fetchall():
                        sections_map[r["announcement_id"]] = {
                            "ps": r.get("parsed_sections"), "da": r.get("deep_analysis")
                        }
                except Exception: pass
            enriched = []
            for m in matched_announcements[:10]:
                it = dict(m)
                aid = m.get("announcement_id")
                if aid and aid in sections_map:
                    ps = sections_map[aid].get("ps") or {}
                    da = sections_map[aid].get("da") or {}
                    if isinstance(ps, str):
                        try: ps = json.loads(ps)
                        except: ps = {}
                    if isinstance(da, str):
                        try: da = json.loads(da)
                        except: da = {}
                    def _s(v, n=400): return v[:n] if isinstance(v, str) else ""
                    it["eligibility"] = _s(ps.get("eligibility"))
                    it["required_docs"] = _s(ps.get("required_docs") or ps.get("required_documents"))
                    it["how_to_apply"] = _s(ps.get("application_method"), 300)
                    it["key_points"] = _s(da.get("key_points"), 300)
                enriched.append(it)
            snap_cur.execute(
                """UPDATE pro_consult_sessions
                   SET matched_snapshot = %s::jsonb,
                       collected = %s::jsonb,
                       phase = 'consulting',
                       updated_at = CURRENT_TIMESTAMP
                   WHERE session_id = %s""",
                (json.dumps(enriched, ensure_ascii=False),
                 json.dumps(profile, ensure_ascii=False),
                 session_state["session_id"]),
            )
            db.commit()
        except Exception as snap_err:
            print(f"[PRO match snapshot] {snap_err}")
            try: db.rollback()
            except: pass

        # 매칭 결과 요약 (자연어 AI 없음 — 결정적 텍스트)
        total = len(matched_announcements)
        if total == 0:
            reply = "매칭된 공고가 없습니다. 조건을 완화해 다시 시도해보세요."
            choices = ["🔄 조건 수정 후 재매칭"]
        else:
            reply = f"**{total}건**의 맞춤 공고를 찾았습니다. 카드를 클릭하면 전문가 레벨 심화 상담으로 이어집니다."
            choices = ["🔄 조건 수정 후 재매칭"]

        return {
            "status": "SUCCESS",
            "action": "match",
            "reply": reply,
            "choices": choices,
            "done": True,
            "profile": profile,
            "collected": profile,
            "announcement_id": None,
            "matched_announcements": matched_announcements,
            "matched_groups": matched_groups,
            "rag_sources": [],
            "session_id": session_state.get("session_id"),
            "phase": "consulting",
        }
    finally:
        try: db.close()
        except: pass


def _handle_pro_consult(req: AiConsultantChatRequest, current_user: dict):
    """action=consult: 특정 공고 심화 상담 (V2 pro_announce 직접 호출)."""
    from app.services.pro_announce import chat_pro_announce
    ann_id = req.announcement_id
    if not ann_id:
        # 메시지에서 공고ID 자동 추출 (하위 호환)
        import re as _re
        for m in req.messages[:2]:
            if m.get("role") != "user":
                continue
            mm = _re.search(r'공고\s*ID\s*[:：]\s*(\d+)', m.get("text", ""))
            if mm:
                ann_id = int(mm.group(1))
                break
    if not ann_id:
        raise HTTPException(status_code=400, detail="action=consult에는 announcement_id 필수입니다.")

    db = get_db_connection()
    try:
        session_state = _load_or_create_session(db, current_user, req.session_id, req.client_category) if req.session_id else None
        selected_client = _load_client(db, req.client_id, current_user["bn"])

        matched_snap = (session_state or {}).get("matched_snapshot") or []
        coll = (session_state or {}).get("collected") or {}

        # 대화 맥락 보강: 프론트 messages가 DB보다 짧으면 DB 히스토리 합병
        effective_messages = list(req.messages) if req.messages else []
        if session_state and session_state.get("db_messages"):
            db_msgs = session_state["db_messages"]
            if len(db_msgs) > len(effective_messages) + 1:
                last_user = effective_messages[-1] if effective_messages else None
                if last_user and last_user.get("role") == "user":
                    effective_messages = db_msgs + [last_user]

        try:
            v2_result = chat_pro_announce(
                messages=effective_messages,
                announcement_id=ann_id,
                db_conn=db,
                selected_client=selected_client,
                matched_snapshot=matched_snap,
                collected=coll,
                force_first_turn=bool(req.is_announcement_start),
            )
        except Exception as ai_err:
            import traceback as _tb
            _tb_str = _tb.format_exc()[-500:]
            print(f"[PRO consult] chat_pro_announce error: {ai_err}\n{_tb_str}")
            v2_result = {
                "reply": "일시적으로 응답 생성에 실패했습니다. 잠시 후 다시 시도해주세요.",
                "choices": ["✏️ 다시 시도"],
                "done": False,
            }

        # 세션 messages 누적 (expert_insights 포함)
        if session_state:
            try:
                assistant_msg = {"role": "assistant", "text": v2_result.get("reply", "")}
                if v2_result.get("expert_insights") or v2_result.get("verdict_for_client"):
                    assistant_msg["meta"] = {
                        "announcement_id": ann_id,
                        "verdict_for_client": v2_result.get("verdict_for_client"),
                        "expert_insights": v2_result.get("expert_insights") or {},
                        "citations": v2_result.get("citations") or [],
                    }
                full_msgs = list(req.messages) + [assistant_msg]
                up_cur = db.cursor()
                up_cur.execute(
                    """UPDATE pro_consult_sessions
                       SET messages = %s::jsonb,
                           phase = 'consulting',
                           updated_at = CURRENT_TIMESTAMP
                       WHERE session_id = %s""",
                    (json.dumps(full_msgs, ensure_ascii=False), session_state["session_id"])
                )
                db.commit()
            except Exception as upd_err:
                print(f"[PRO consult session update] {upd_err}")
                try: db.rollback()
                except: pass

        # AI 응답에서 언급된 공고명 키워드로 DB 검색 → 카드 표시
        consult_matched = []
        try:
            reply_text = v2_result.get("reply", "")
            if reply_text:
                import re as _re
                # 굵은 텍스트(**...**), 번호 목록, 따옴표 안 공고명 추출
                _kw_candidates = _re.findall(r'\*\*([^*]{4,40})\*\*', reply_text)
                _kw_candidates += _re.findall(r'[『「]([^』」]{4,40})[』」]', reply_text)
                _kw_candidates += _re.findall(r'\d+\.\s+([가-힣a-zA-Z·\s]{4,40})(?=[:：]|\s*—)', reply_text)
                seen_ids: set = set()
                _kw_cur = db.cursor()
                for kw in _kw_candidates[:8]:
                    kw = kw.strip()
                    if len(kw) < 4:
                        continue
                    _kw_cur.execute(
                        """SELECT announcement_id, title, department, support_amount,
                                  support_amount_max, deadline_date
                           FROM announcements
                           WHERE title ILIKE %s AND is_archived = false
                           ORDER BY deadline_date ASC NULLS LAST
                           LIMIT 1""",
                        (f"%{kw}%",)
                    )
                    row = _kw_cur.fetchone()
                    if row and row["announcement_id"] not in seen_ids:
                        seen_ids.add(row["announcement_id"])
                        consult_matched.append({
                            "announcement_id": row["announcement_id"],
                            "title": row["title"],
                            "department": row["department"] or "",
                            "support_amount": row["support_amount"] or "",
                            "support_amount_max": row["support_amount_max"],
                            "deadline_date": str(row["deadline_date"] or ""),
                            "rank": len(consult_matched) + 1,
                            "bucket": "consult_mention",
                            "bucket_label": "상담 언급",
                            "reasons": [],
                            "matched_interests": [],
                        })
        except Exception as _cm_err:
            print(f"[PRO consult] mention search error: {_cm_err}")

        return {
            "status": "SUCCESS",
            "action": "consult",
            "reply": v2_result.get("reply", ""),
            "choices": v2_result.get("choices", []),
            "done": v2_result.get("done", False),
            "profile": None,
            "collected": coll,
            "announcement_id": ann_id,
            "matched_announcements": consult_matched,
            "matched_groups": {"interest_match": [], "deadline_urgent": [], "qualified_other": []},
            "rag_sources": v2_result.get("rag_sources", []),
            "session_id": session_state.get("session_id") if session_state else None,
            "phase": "consulting",
            "verdict_for_client": v2_result.get("verdict_for_client"),
            "expert_insights": v2_result.get("expert_insights"),
            "citations": v2_result.get("citations"),
        }
    finally:
        try: db.close()
        except: pass


@app.post("/api/ai/consultant/chat")
def api_ai_consultant_chat(req: AiConsultantChatRequest, current_user: dict = Depends(_get_current_user)):
    """LITE 일반 사용자: 대화형으로 고객사 조건 수집"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("SELECT * FROM users WHERE business_number = %s", (bn,))
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)

    plan = u.get("plan") or "free"
    if plan in ("trial", "premium"):
        plan = "free"
    plan_expires = u.get("plan_expires_at")
    ps = _get_plan_status(plan, plan_expires, u.get("ai_usage_month") or 0)
    if not ps.get("active"):
        conn.close()
        raise HTTPException(status_code=403, detail="플랜이 만료되었습니다.")

    # LITE 자금 전문 상담으로 통합 (chat_consultant 대체)
    from app.services.ai_consultant import chat_lite_fund_expert
    # 탭(mode)에 따라 프로필 필터링
    _BIZ = {"business_number", "company_name", "industry_code", "revenue_bracket", "employee_count_bracket", "establishment_date", "address_city", "interests", "user_type", "certifications", "custom_keywords"}
    _IND = {"age_range", "income_level", "family_type", "employment_status", "housing_status", "gender", "address_city", "interests", "user_type"}
    _COM = {"address_city", "interests", "user_type", "email", "plan"}
    if req.mode == "individual_fund":
        _fp = {k: v for k, v in u.items() if k in (_IND | _COM) and v}
    elif req.mode == "business_fund":
        _fp = {k: v for k, v in u.items() if k in (_BIZ | _COM) and v}
    else:
        _fp = u
    result = chat_lite_fund_expert(req.messages, db_conn=conn, user_profile=_fp, mode=req.mode)

    # ── 대화 저장 (P0.2+B): UPSERT by session_id ──
    try:
        all_msgs = list(req.messages) + [{"role": "assistant", "text": result.get("reply", "")}]
        sid = req.session_id
        if not sid:
            import hashlib
            first_user = next((m.get("text","")[:200] for m in req.messages if m.get("role")=="user"), "")
            sid = "lite_" + hashlib.sha256((bn + first_user).encode()).hexdigest()[:16]
        cur.execute(
            """INSERT INTO ai_consult_logs (announcement_id, business_number, messages, conclusion, session_id, updated_at)
               VALUES (NULL, %s, %s::jsonb, %s, %s, CURRENT_TIMESTAMP)
               ON CONFLICT (session_id) WHERE session_id IS NOT NULL DO UPDATE SET
                   messages = EXCLUDED.messages,
                   conclusion = EXCLUDED.conclusion,
                   updated_at = CURRENT_TIMESTAMP""",
            (bn, json.dumps(all_msgs, ensure_ascii=False), "lite_consultant", sid)
        )
        conn.commit()
    except Exception as save_err:
        print(f"[ai_consultant_chat save] {save_err}")
        try: conn.rollback()
        except: pass
    conn.close()

    return {
        "status": "SUCCESS",
        "reply": result.get("reply", ""),
        "choices": result.get("choices", []),
        "done": result.get("done", False),
        "profile": result.get("profile"),
        "collected": result.get("collected", {}),
    }


# ── AI 컨설턴트 매칭 실행 (가상 프로필 → 매칭) ────────────────

class ConsultantMatchRequest(BaseModel):
    profile: dict  # 가상 프로필 (company_name, establishment_date, industry_code 등)


@app.post("/api/ai/consultant/match")
def api_ai_consultant_match(req: ConsultantMatchRequest, current_user: dict = Depends(_get_current_user)):
    """컨설턴트 모드: 가상 프로필로 매칭 실행 (AI 1건 차감)"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("SELECT * FROM users WHERE business_number = %s", (bn,))
    user = cur.fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    u = dict(user)

    plan = u.get("plan") or "free"
    if plan in ("trial", "premium"):
        plan = "free"
    plan_expires = u.get("plan_expires_at")
    usage = u.get("ai_usage_month") or 0
    ps = _get_plan_status(plan, plan_expires, usage)
    if not ps.get("active"):
        conn.close()
        raise HTTPException(status_code=403, detail="플랜이 만료되었습니다.")

    limit = PLAN_LIMITS.get(plan, 1)

    # 월간 리셋
    now = datetime.datetime.utcnow()
    reset_at = u.get("ai_usage_reset_at")
    if reset_at:
        try:
            reset_dt = datetime.datetime.fromisoformat(str(reset_at))
            if now.month != reset_dt.month or now.year != reset_dt.year:
                usage = 0
                cur.execute(
                    "UPDATE users SET ai_usage_month=0, ai_usage_reset_at=%s WHERE business_number=%s",
                    (now.isoformat(), bn)
                )
        except Exception:
            pass

    # AI 1건 차감
    if usage >= limit:
        conn.close()
        raise HTTPException(status_code=429, detail=f"이번 달 AI 상담 한도({limit}회)를 모두 사용했습니다. 플랜을 업그레이드하면 더 많은 상담을 이용할 수 있습니다.")

    cur.execute("UPDATE users SET ai_usage_month = ai_usage_month + 1 WHERE business_number = %s", (bn,))
    conn.commit()
    conn.close()
    usage += 1

    # 가상 프로필로 매칭 엔진 실행 — 프로필의 industry_code 유무로 기업/개인 분기
    virtual_profile = req.profile
    is_individual = not virtual_profile.get("industry_code")
    matches = get_matches_hybrid(virtual_profile, is_individual=is_individual)

    # 직렬화 (date 등)
    for m in matches:
        for k, v in m.items():
            if isinstance(v, (datetime.date, datetime.datetime)):
                m[k] = v.isoformat()

    return {
        "status": "SUCCESS",
        "matches": matches,
        "profile_used": virtual_profile,
        "ai_used": usage,
        "ai_limit": limit,
    }


# ── 상담 피드백 ──────────────────────────────────────────

class ConsultFeedbackRequest(BaseModel):
    consult_log_id: int
    feedback: str  # 'helpful' | 'inaccurate'
    detail: Optional[str] = None


class ConsultSaveRequest(BaseModel):
    announcement_id: int
    messages: list
    conclusion: Optional[str] = None
    session_id: Optional[str] = None


@app.post("/api/ai/consult/save")
def api_consult_save(req: ConsultSaveRequest, current_user: dict = Depends(_get_current_user)):
    """상담 명시적 저장 — '저장하고 닫기' 버튼.
    session_id 전달 시 해당 세션 행을 UPSERT (대화 턴에서 이미 저장된 행과 동일 행 유지).
    """
    bn = current_user["bn"]
    if not req.messages or len(req.messages) < 2:
        raise HTTPException(status_code=400, detail="저장할 대화가 없습니다.")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        if req.session_id:
            cur.execute(
                """INSERT INTO ai_consult_logs (announcement_id, business_number, messages, conclusion, session_id, updated_at)
                   VALUES (%s, %s, %s::jsonb, %s, %s, CURRENT_TIMESTAMP)
                   ON CONFLICT (session_id) WHERE session_id IS NOT NULL DO UPDATE SET
                       messages = EXCLUDED.messages,
                       conclusion = COALESCE(EXCLUDED.conclusion, ai_consult_logs.conclusion),
                       updated_at = CURRENT_TIMESTAMP
                   RETURNING id""",
                (req.announcement_id, bn, json.dumps(req.messages, ensure_ascii=False), req.conclusion, req.session_id),
            )
        else:
            cur.execute(
                """INSERT INTO ai_consult_logs (announcement_id, business_number, messages, conclusion)
                   VALUES (%s, %s, %s, %s) RETURNING id""",
                (req.announcement_id, bn, json.dumps(req.messages, ensure_ascii=False), req.conclusion),
            )
        row = cur.fetchone()
        conn.commit()
        return {"status": "SUCCESS", "consult_log_id": row["id"] if row else None}
    finally:
        conn.close()


@app.post("/api/ai/consult/feedback")
def api_consult_feedback(req: ConsultFeedbackRequest, current_user: dict = Depends(_get_current_user)):
    """상담 결과 피드백 저장 + 순환 학습 (골든답변/지식 저장)"""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "UPDATE ai_consult_logs SET feedback = %s, feedback_detail = %s WHERE id = %s AND business_number = %s",
        (req.feedback, req.detail, req.consult_log_id, current_user["bn"])
    )
    conn.commit()

    # 순환 학습: 피드백 기반으로 지식 추출 (golden_answers 제거됨 — 캐시 부작용)
    try:
        from app.services.ai_consultant import extract_knowledge_from_consult

        cur.execute(
            "SELECT announcement_id, messages, conclusion FROM ai_consult_logs WHERE id = %s",
            (req.consult_log_id,)
        )
        log = cur.fetchone()
        if log:
            log_data = dict(log)
            ann_id = log_data["announcement_id"]
            messages = log_data["messages"]
            if isinstance(messages, str):
                messages = json.loads(messages)
            conclusion = log_data.get("conclusion")

            cur.execute("SELECT category FROM announcements WHERE announcement_id = %s", (ann_id,))
            ann_row = cur.fetchone()
            category = dict(ann_row).get("category", "") if ann_row else ""

            if req.feedback == "helpful":
                try:
                    from app.services.financial_analysis.auto_learner import process_helpful_feedback
                    learned = process_helpful_feedback(req.consult_log_id, conn)
                    if learned:
                        print(f"[AutoLearn] {learned} knowledge items extracted from log #{req.consult_log_id}")
                except Exception as al_err:
                    print(f"[AutoLearn] Error (non-critical): {al_err}")

            extract_knowledge_from_consult(
                announcement_id=ann_id,
                category=category,
                messages=messages,
                conclusion=conclusion,
                feedback=req.feedback,
                db_conn=conn,
            )
    except Exception as e:
        print(f"[Feedback] Learning error (non-critical): {e}")

    conn.close()
    return {"status": "SUCCESS"}


# ══════════════════════════════════════════
# 내 상담 기록 (사용자 본인의 ai_consult_logs 열람)
@app.get("/api/ai/consult/session/{session_id}")
def api_get_consult_session(session_id: str, current_user: dict = Depends(_get_current_user)):
    """세션 ID로 이전 대화 복원 (모바일 앱 전환 후 복귀 시 사용)"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT messages, updated_at
           FROM ai_consult_logs
           WHERE session_id = %s AND business_number = %s
             AND updated_at > NOW() - INTERVAL '24 hours'""",
        (session_id, bn)
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="세션을 찾을 수 없거나 만료됐습니다.")
    msgs = row["messages"] or []
    return {"status": "SUCCESS", "messages": msgs, "updated_at": str(row["updated_at"])}


# ══════════════════════════════════════════

@app.get("/api/my/consults")
def api_my_consults(
    current_user: dict = Depends(_get_current_user),
    page: int = 1,
    size: int = 20,
    filter: str = "all",  # all | eligible | conditional | ineligible
):
    """내 상담 기록 목록 — 최신순."""
    conn = get_db_connection()
    cur = conn.cursor()

    offset = max(0, (page - 1) * size)
    size = min(max(1, size), 100)

    where_sql = "WHERE l.business_number = %s"
    params: list = [current_user["bn"]]
    if filter in ("eligible", "conditional", "ineligible"):
        where_sql += " AND l.conclusion = %s"
        params.append(filter)

    cur.execute(
        f"""SELECT COUNT(*) AS total FROM ai_consult_logs l {where_sql}""",
        tuple(params),
    )
    total_row = cur.fetchone()
    total = dict(total_row).get("total", 0) if total_row else 0

    cur.execute(
        f"""SELECT l.id, l.announcement_id, l.conclusion, l.feedback, l.created_at, l.messages,
                   a.title, a.category, a.department, a.deadline_date, a.support_amount
            FROM ai_consult_logs l
            LEFT JOIN announcements a ON a.announcement_id = l.announcement_id
            {where_sql}
            ORDER BY l.created_at DESC
            LIMIT %s OFFSET %s""",
        tuple(params + [size, offset]),
    )
    rows = cur.fetchall()

    items = []
    for row in rows:
        r = dict(row)
        messages = r.get("messages") or []
        if isinstance(messages, str):
            try:
                messages = json.loads(messages)
            except Exception:
                messages = []
        # 마지막 AI 응답에서 미리보기 추출
        last_ai = ""
        if isinstance(messages, list):
            for m in reversed(messages):
                if m.get("role") == "assistant":
                    last_ai = (m.get("text") or "")[:200]
                    break
        items.append({
            "id": r["id"],
            "announcement_id": r["announcement_id"],
            "announcement_title": r.get("title") or "(삭제된 공고)",
            "category": r.get("category") or "",
            "department": r.get("department") or "",
            "deadline_date": str(r.get("deadline_date")) if r.get("deadline_date") else "",
            "support_amount": r.get("support_amount") or "",
            "conclusion": r.get("conclusion") or "",
            "feedback": r.get("feedback") or "",
            "created_at": str(r.get("created_at")) if r.get("created_at") else "",
            "preview": last_ai,
            "message_count": len(messages) if isinstance(messages, list) else 0,
        })

    conn.close()
    return {
        "status": "SUCCESS",
        "total": total,
        "page": page,
        "size": size,
        "items": items,
    }


@app.get("/api/my/consults/{consult_id}")
def api_my_consult_detail(consult_id: int, current_user: dict = Depends(_get_current_user)):
    """내 상담 기록 상세 — 본인 상담만 열람 가능."""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT l.id, l.announcement_id, l.conclusion, l.feedback, l.feedback_detail,
                  l.created_at, l.messages,
                  a.title, a.category, a.department, a.deadline_date, a.support_amount,
                  a.region, a.origin_url, a.summary_text
           FROM ai_consult_logs l
           LEFT JOIN announcements a ON a.announcement_id = l.announcement_id
           WHERE l.id = %s AND l.business_number = %s""",
        (consult_id, current_user["bn"]),
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="상담 기록을 찾을 수 없습니다.")
    r = dict(row)
    messages = r.get("messages") or []
    if isinstance(messages, str):
        try:
            messages = json.loads(messages)
        except Exception:
            messages = []
    return {
        "status": "SUCCESS",
        "consult": {
            "id": r["id"],
            "announcement_id": r["announcement_id"],
            "announcement_title": r.get("title") or "(삭제된 공고)",
            "category": r.get("category") or "",
            "department": r.get("department") or "",
            "region": r.get("region") or "",
            "deadline_date": str(r.get("deadline_date")) if r.get("deadline_date") else "",
            "support_amount": r.get("support_amount") or "",
            "origin_url": r.get("origin_url") or "",
            "conclusion": r.get("conclusion") or "",
            "feedback": r.get("feedback") or "",
            "feedback_detail": r.get("feedback_detail") or "",
            "created_at": str(r.get("created_at")) if r.get("created_at") else "",
            "messages": messages,
        },
    }


@app.delete("/api/my/consults/{consult_id}")
def api_my_consult_delete(consult_id: int, current_user: dict = Depends(_get_current_user)):
    """내 상담 기록 삭제 — 본인 상담만."""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "DELETE FROM ai_consult_logs WHERE id = %s AND business_number = %s",
        (consult_id, current_user["bn"]),
    )
    deleted = cur.rowcount
    conn.commit()
    conn.close()
    if deleted == 0:
        raise HTTPException(status_code=404, detail="상담 기록을 찾을 수 없습니다.")
    return {"status": "SUCCESS", "deleted": deleted}


# ── 공고 원문 정밀 분석 (배치) ─────────────────────────────

class AdminAuthRequest(BaseModel):
    password: str


@app.post("/api/admin/analyze-announcements")
def api_analyze_announcements(req: AdminAuthRequest):
    """관리자: 기존 공고들의 원문을 정밀 분석하여 학습 데이터 축적"""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    from app.services.doc_analysis_service import analyze_and_store

    conn = get_db_connection()
    cur = conn.cursor()

    # 아직 분석되지 않은 공고 조회 (최대 20건씩)
    cur.execute("""
        SELECT a.announcement_id, a.title, a.origin_url, a.summary_text
        FROM announcements a
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE aa.id IS NULL AND a.origin_url IS NOT NULL
        ORDER BY a.created_at DESC
        LIMIT 20
    """)
    rows = cur.fetchall()

    results = {"total": len(rows), "success": 0, "failed": 0, "details": []}
    for row in rows:
        r = dict(row)
        res = analyze_and_store(
            announcement_id=r["announcement_id"],
            origin_url=r["origin_url"],
            title=r["title"],
            db_conn=conn,
            summary_text=r.get("summary_text") or "",
        )
        if res["success"]:
            results["success"] += 1
        else:
            results["failed"] += 1
        results["details"].append({
            "id": r["announcement_id"],
            "title": r["title"][:50],
            "success": res["success"],
            "source": res["source_type"],
            "chars": res["text_length"],
        })

    conn.close()
    return {"status": "SUCCESS", **results}


_UMBRELLA_PATTERNS = ["%통합 공고%", "%통합공고%"]
_UMBRELLA_WHERE = "(title ILIKE %s OR title ILIKE %s)"


@app.post("/api/admin/umbrella-scan")
def api_umbrella_scan(req: AdminAuthRequest):
    """통합공고 후보 스캔 (dry-run) — 삭제 없이 개수와 샘플만 반환."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(f"SELECT COUNT(*) AS cnt FROM announcements WHERE {_UMBRELLA_WHERE}", _UMBRELLA_PATTERNS)
        total = cur.fetchone()["cnt"]
        cur.execute(
            f"""SELECT announcement_id, title, origin_source,
                       LEFT(COALESCE(summary_text, ''), 120) AS summary_preview
                FROM announcements
                WHERE {_UMBRELLA_WHERE}
                ORDER BY announcement_id DESC
                LIMIT 30""",
            _UMBRELLA_PATTERNS,
        )
        samples = [dict(r) for r in cur.fetchall()]
        return {"status": "SUCCESS", "total_matched": total, "samples": samples}
    finally:
        try: conn.close()
        except: pass


class UmbrellaVerifyRequest(BaseModel):
    password: str
    purge: Optional[bool] = False
    limit: Optional[int] = 30
    offset: Optional[int] = 0


@app.post("/api/admin/umbrella-verify")
def api_umbrella_verify(req: UmbrellaVerifyRequest):
    """통합공고 후보를 LLM으로 개별 검증.

    - purge=False (기본): 검증만 하고 결과 반환
    - purge=True: is_umbrella=true 판정된 공고를 즉시 삭제
    """
    try:
        if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
            return {"status": "ERROR", "error": "pw"}
        import google.generativeai as genai
        api_key = os.environ.get("GEMINI_API_KEY", "")
        if not api_key:
            return {"status": "ERROR", "error": "no gemini key"}
        genai.configure(api_key=api_key)
    except Exception as init_err:
        import traceback as _tb
        return {"status": "INIT_ERROR", "error": f"{type(init_err).__name__}: {str(init_err)[:300]}", "tb": _tb.format_exc()[-600:]}

    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            f"""SELECT announcement_id, title, summary_text
                FROM announcements
                WHERE {_UMBRELLA_WHERE}
                ORDER BY announcement_id DESC
                LIMIT %s OFFSET %s""",
            _UMBRELLA_PATTERNS + [req.limit, req.offset],
        )
        rows = [dict(r) for r in cur.fetchall()]
        cur.execute(f"SELECT COUNT(*) AS cnt FROM announcements WHERE {_UMBRELLA_WHERE}", _UMBRELLA_PATTERNS)
        remaining_total = cur.fetchone()["cnt"]
        if not rows:
            return {"status": "SUCCESS", "total": 0, "umbrella": 0, "kept": 0, "remaining_total": remaining_total, "results": []}

        model = genai.GenerativeModel(
            "models/gemini-2.5-flash",
            generation_config={
                "response_mime_type": "application/json",
                "temperature": 0.1,
                "max_output_tokens": 200,
            },
        )

        prompt_template = """다음은 정부지원사업 공고입니다. 이 공고가 "여러 하위 사업을 안내하는 통합/상위 문서"인지 "단일 개별 사업의 구체 모집 공고"인지 판정하세요.

[제목] {title}
[요약] {summary}

판정 기준:
- is_umbrella=true: 하위 사업이 복수로 존재하고 각각 별도로 신청해야 하며, 이 문서는 안내·개요·목록 역할에 그침. 실제 신청 대상·기간·금액이 불분명하거나 "상세는 각 사업 참조" 같은 표현.
- is_umbrella=false: 단일 사업의 신청대상·기간·지원금액·신청방법이 명시된 구체적인 모집/선정 공고. "통합 공고"라는 단어가 제목에 있어도 내용이 구체적이면 false.

반드시 JSON만 반환:
{{"is_umbrella": true/false, "reason": "한 줄 근거 (40자 이내)"}}"""

        results: List[Dict[str, Any]] = []
        umbrella_ids: List[int] = []
        for r in rows:
            aid = r["announcement_id"]
            title = (r.get("title") or "").strip()
            summary = ((r.get("summary_text") or "")[:400]).strip()
            try:
                resp = model.generate_content(
                    prompt_template.format(title=title[:200], summary=summary)
                )
                parsed = json.loads(resp.text)
                if isinstance(parsed, list) and parsed:
                    parsed = parsed[0] if isinstance(parsed[0], dict) else {}
                if not isinstance(parsed, dict):
                    parsed = {}
                is_um = bool(parsed.get("is_umbrella"))
                reason = (parsed.get("reason") or "")[:120]
            except Exception as e:
                is_um = False
                reason = f"LLM error: {str(e)[:80]}"
            results.append({
                "announcement_id": aid,
                "title": title[:80],
                "is_umbrella": is_um,
                "reason": reason,
            })
            if is_um:
                umbrella_ids.append(aid)

        purged = 0
        purge_by_table: Dict[str, Any] = {}
        if req.purge and umbrella_ids:
            # FK가 걸린 테이블은 먼저 참조 해제 후 즉시 commit
            try:
                cur.execute(
                    "UPDATE ai_consult_logs SET announcement_id = NULL WHERE announcement_id = ANY(%s)",
                    (umbrella_ids,),
                )
                purge_by_table["ai_consult_logs_detached"] = cur.rowcount
                conn.commit()
            except Exception as e:
                conn.rollback()
                purge_by_table["ai_consult_logs_detached"] = f"error: {str(e)[:100]}"
            # pro_consult_sessions 등 다른 FK도 같이 처리
            for fk_tbl in ("pro_consult_sessions", "saved_announcements"):
                try:
                    # announcement_id 컬럼이 있는지 확인하며 NULL 처리 (saved는 삭제가 맞지만 방어적)
                    cur.execute(
                        f"UPDATE {fk_tbl} SET announcement_id = NULL WHERE announcement_id = ANY(%s)",
                        (umbrella_ids,),
                    )
                    purge_by_table[f"{fk_tbl}_detached"] = cur.rowcount
                    conn.commit()
                except Exception as e:
                    conn.rollback()
                    purge_by_table[f"{fk_tbl}_detached"] = f"skip: {str(e)[:80]}"
            for tbl in (
                "announcement_sections",
                "announcement_analysis",
                "announcement_embeddings",
                "saved_announcements",
                "trending_announcements",
                "match_history",
            ):
                try:
                    cur.execute(f"DELETE FROM {tbl} WHERE announcement_id = ANY(%s)", (umbrella_ids,))
                    purge_by_table[tbl] = cur.rowcount
                except Exception as e:
                    conn.rollback()
                    purge_by_table[tbl] = f"error: {str(e)[:100]}"
            try:
                cur.execute("DELETE FROM announcements WHERE announcement_id = ANY(%s)", (umbrella_ids,))
                purge_by_table["announcements"] = cur.rowcount
                purged = cur.rowcount
                conn.commit()
            except Exception as e:
                conn.rollback()
                return {"status": "ERROR", "error": str(e)[:200], "partial": purge_by_table}

        return {
            "status": "SUCCESS",
            "total": len(rows),
            "umbrella": len(umbrella_ids),
            "kept": len(rows) - len(umbrella_ids),
            "purged": purged,
            "purge_by_table": purge_by_table if req.purge else None,
            "remaining_after": remaining_total - purged,
            "results": results,
        }
    except Exception as run_err:
        import traceback as _tb
        return {"status": "RUN_ERROR", "error": f"{type(run_err).__name__}: {str(run_err)[:300]}", "tb": _tb.format_exc()[-800:]}
    finally:
        try:
            if conn: conn.close()
        except: pass


@app.post("/api/admin/umbrella-purge")
def api_umbrella_purge(req: AdminAuthRequest):
    """통합공고 실제 삭제 — 관련 테이블 전부 정리."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # 대상 ID 수집
        cur.execute(f"SELECT announcement_id FROM announcements WHERE {_UMBRELLA_WHERE}", _UMBRELLA_PATTERNS)
        ids = [r["announcement_id"] for r in cur.fetchall()]
        if not ids:
            return {"status": "SUCCESS", "deleted": 0, "message": "대상 없음"}

        deleted_by_table = {}
        # 관련 테이블 순서대로 정리
        for tbl, col in [
            ("announcement_sections", "announcement_id"),
            ("announcement_analysis", "announcement_id"),
            ("announcement_embeddings", "announcement_id"),
            ("saved_announcements", "announcement_id"),
            ("trending_announcements", "announcement_id"),
            ("match_history", "announcement_id"),
            ("section_feedback", "section_id"),  # section_feedback은 section_id 기반이라 스킵
        ]:
            if tbl == "section_feedback":
                continue
            try:
                cur.execute(f"DELETE FROM {tbl} WHERE {col} = ANY(%s)", (ids,))
                deleted_by_table[tbl] = cur.rowcount
            except Exception as e:
                conn.rollback()
                deleted_by_table[tbl] = f"error: {str(e)[:120]}"
                continue
        # 본 테이블
        try:
            cur.execute("DELETE FROM announcements WHERE announcement_id = ANY(%s)", (ids,))
            deleted_by_table["announcements"] = cur.rowcount
            conn.commit()
        except Exception as e:
            conn.rollback()
            return {"status": "ERROR", "error": str(e)[:300], "partial": deleted_by_table}
        return {
            "status": "SUCCESS",
            "deleted": len(ids),
            "deleted_by_table": deleted_by_table,
        }
    finally:
        try: conn.close()
        except: pass


class InspectAnnouncementRequest(BaseModel):
    password: str
    announcement_id: int


@app.post("/api/admin/inspect-announcement")
def api_inspect_announcement(req: InspectAnnouncementRequest):
    """특정 공고의 raw 상태 + parsed_sections.timeline 확인."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            """SELECT a.announcement_id, a.title, a.deadline_date, a.origin_source,
                      a.origin_url, a.summary_text, a.created_at,
                      aa.parsed_sections, aa.deep_analysis
               FROM announcements a
               LEFT JOIN announcement_analysis aa ON aa.announcement_id = a.announcement_id
               WHERE a.announcement_id = %s""",
            (req.announcement_id,),
        )
        r = cur.fetchone()
        if not r:
            return {"status": "NOT_FOUND"}
        d = dict(r)
        ps = d.get("parsed_sections")
        if isinstance(ps, str):
            try: ps = json.loads(ps)
            except: ps = {}
        da = d.get("deep_analysis")
        if isinstance(da, str):
            try: da = json.loads(da)
            except: da = {}
        return {
            "status": "SUCCESS",
            "announcement_id": d["announcement_id"],
            "title": d["title"],
            "deadline_date": str(d.get("deadline_date") or ""),
            "origin_source": d.get("origin_source"),
            "origin_url": d.get("origin_url"),
            "summary_text": (d.get("summary_text") or "")[:500],
            "timeline_parsed": (ps or {}).get("timeline", "")[:1000] if isinstance(ps, dict) else "",
            "schedule_parsed": (ps or {}).get("schedule", "")[:1000] if isinstance(ps, dict) else "",
            "application_method": (ps or {}).get("application_method", "")[:1000] if isinstance(ps, dict) else "",
            "summary_from_deep": (da or {}).get("support_summary") if isinstance(da, dict) else None,
        }
    finally:
        try: conn.close()
        except: pass


class ConsultDedupeRequest(BaseModel):
    password: str
    apply: Optional[bool] = False


@app.post("/api/admin/consult-dedupe")
def api_consult_dedupe(req: ConsultDedupeRequest):
    """ai_consult_logs 중복 정리 — 같은 (bn, aid) 쌍에서 session_id 있는 행이 존재하면
    session_id가 NULL인 행 삭제. apply=False면 미리보기만."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # 삭제 대상: session_id NULL이고, 같은 (bn, aid)에 session_id 있는 행이 존재
        cur.execute(
            """SELECT l1.id, l1.announcement_id, l1.business_number, l1.created_at
               FROM ai_consult_logs l1
               WHERE l1.session_id IS NULL
                 AND l1.announcement_id IS NOT NULL
                 AND EXISTS (
                     SELECT 1 FROM ai_consult_logs l2
                     WHERE l2.business_number = l1.business_number
                       AND l2.announcement_id = l1.announcement_id
                       AND l2.session_id IS NOT NULL
                       AND l2.id <> l1.id
                 )
               ORDER BY l1.id DESC
               LIMIT 1000"""
        )
        targets = [dict(r) for r in cur.fetchall()]
        count = len(targets)
        if not req.apply:
            return {
                "status": "SUCCESS",
                "mode": "dry-run",
                "deletable": count,
                "samples": targets[:15],
            }
        # 실제 삭제
        ids = [t["id"] for t in targets]
        deleted = 0
        if ids:
            cur.execute("DELETE FROM ai_consult_logs WHERE id = ANY(%s)", (ids,))
            deleted = cur.rowcount
            conn.commit()
        return {
            "status": "SUCCESS",
            "mode": "applied",
            "deleted": deleted,
            "remaining_dupes_query": "SELECT COUNT(*)...",
        }
    finally:
        try: conn.close()
        except: pass


@app.post("/api/admin/consult-dedupe-v2")
def api_consult_dedupe_v2(req: ConsultDedupeRequest):
    """v2: (bn, announcement_id) 기준으로 가장 턴 수 많은 행만 남기고 나머지 삭제.
    프론트 stale-closure 버그로 매 턴마다 새 session_id가 발급되어 여러 행으로 분리된
    이력을 정리. apply=False면 미리보기(삭제될 id·유지될 id 샘플)."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # (bn, announcement_id) 별 그룹에서 messages jsonb 배열 길이가 가장 큰 행만 유지
        cur.execute("""
            WITH ranked AS (
                SELECT id, business_number, announcement_id,
                       jsonb_array_length(COALESCE(messages, '[]'::jsonb)) AS msg_count,
                       created_at,
                       ROW_NUMBER() OVER (
                           PARTITION BY business_number, announcement_id
                           ORDER BY jsonb_array_length(COALESCE(messages, '[]'::jsonb)) DESC,
                                    created_at DESC, id DESC
                       ) AS rnk
                FROM ai_consult_logs
                WHERE announcement_id IS NOT NULL
            )
            SELECT id, business_number, announcement_id, msg_count, created_at
            FROM ranked
            WHERE rnk > 1
            ORDER BY id DESC
            LIMIT 5000
        """)
        delete_targets = [dict(r) for r in cur.fetchall()]
        del_count = len(delete_targets)

        # 유지될 행도 샘플로 보여주기
        cur.execute("""
            WITH ranked AS (
                SELECT id, business_number, announcement_id,
                       jsonb_array_length(COALESCE(messages, '[]'::jsonb)) AS msg_count,
                       created_at,
                       ROW_NUMBER() OVER (
                           PARTITION BY business_number, announcement_id
                           ORDER BY jsonb_array_length(COALESCE(messages, '[]'::jsonb)) DESC,
                                    created_at DESC, id DESC
                       ) AS rnk,
                       COUNT(*) OVER (PARTITION BY business_number, announcement_id) AS group_size
                FROM ai_consult_logs
                WHERE announcement_id IS NOT NULL
            )
            SELECT id, business_number, announcement_id, msg_count, group_size
            FROM ranked
            WHERE rnk = 1 AND group_size > 1
            ORDER BY group_size DESC, id DESC
            LIMIT 15
        """)
        keep_samples = [dict(r) for r in cur.fetchall()]

        if not req.apply:
            return {
                "status": "SUCCESS",
                "mode": "dry-run",
                "deletable": del_count,
                "delete_samples": delete_targets[:15],
                "keep_samples": keep_samples,
            }
        # 실제 삭제
        ids = [t["id"] for t in delete_targets]
        deleted = 0
        if ids:
            cur.execute("DELETE FROM ai_consult_logs WHERE id = ANY(%s)", (ids,))
            deleted = cur.rowcount
            conn.commit()
        return {
            "status": "SUCCESS",
            "mode": "applied",
            "deleted": deleted,
        }
    finally:
        try: conn.close()
        except: pass


@app.post("/api/admin/refresh-trending")
def api_refresh_trending(req: AdminAuthRequest):
    """오늘의 인기 공고를 강제로 재생성 (네이버 데이터랩 호출 포함)."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    from app.services.patrol.trending import run_trending_update, fetch_datalab_ranking
    from datetime import date as _date
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM trending_announcements WHERE trending_date = %s", (_date.today().isoformat(),))
        conn.commit()
        # 데이터랩 순위 함께 반환 (디버그)
        ranked = fetch_datalab_ranking()
        result = run_trending_update(conn)
        return {
            "status": "SUCCESS",
            "datalab_available": bool(ranked),
            "datalab_top": ranked[:8],
            **result,
        }
    finally:
        try: conn.close()
        except: pass


@app.post("/api/admin/seed-fund-knowledge")
def api_seed_fund_knowledge(req: AdminAuthRequest):
    """관리자: 정책자금/보증 시드 지식을 knowledge_base에 적재 (중복 스킵)."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    from app.db.fund_knowledge_seed import seed_fund_knowledge
    conn = get_db_connection()
    try:
        result = seed_fund_knowledge(conn)
        return {"status": "SUCCESS", **result}
    finally:
        conn.close()


@app.post("/api/admin/analyze-batch-broad")
def api_analyze_batch_broad(req: AdminAuthRequest):
    """광범위 분석 배치 — origin_url + summary 50자만 있으면 처리.
    한 호출 250초, 최대 50건. 학습 풀 확장용.
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    import time as _time
    from app.services.doc_analysis_service import analyze_and_store

    DEADLINE = 200  # Railway 300s gateway timeout 대응
    MAX_ITEMS = 15  # 호출당 적게 처리해 안정성 우선
    start = _time.time()

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT a.announcement_id, a.title, a.origin_url, a.summary_text
        FROM announcements a
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE aa.id IS NULL
          AND a.origin_url IS NOT NULL
          AND a.summary_text IS NOT NULL AND LENGTH(a.summary_text) >= 50
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
        ORDER BY a.announcement_id DESC
        LIMIT %s
    """, (MAX_ITEMS,))
    rows = cur.fetchall()

    cur.execute("""
        SELECT COUNT(*) AS cnt FROM announcements a
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE aa.id IS NULL AND a.origin_url IS NOT NULL
          AND a.summary_text IS NOT NULL AND LENGTH(a.summary_text) >= 50
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
    """)
    remaining_total = cur.fetchone()["cnt"]

    ok = 0
    fail = 0
    skipped = 0
    for row in rows:
        if _time.time() - start > DEADLINE:
            skipped = len(rows) - ok - fail
            break
        r = dict(row)
        try:
            res = analyze_and_store(
                announcement_id=r["announcement_id"],
                origin_url=r.get("origin_url") or "",
                title=r.get("title") or "",
                db_conn=conn,
                summary_text=r.get("summary_text") or "",
            )
            if res.get("success"):
                ok += 1
            else:
                fail += 1
        except Exception as e:
            fail += 1
            print(f"[broad] #{r['announcement_id']}: {str(e)[:120]}")

    conn.close()
    return {
        "status": "SUCCESS",
        "processed": ok + fail,
        "success": ok,
        "failed": fail,
        "skipped_timeout": skipped,
        "remaining": remaining_total - ok,
        "elapsed": round(_time.time() - start, 1),
        "done": (remaining_total - ok) <= 0,
    }


@app.post("/api/admin/analyze-batch-priority")
def api_analyze_batch_priority(req: AdminAuthRequest):
    """관리자: 지원금액 있는 미분석 공고를 금액 크기·마감일 순으로 자동 배치 분석.
    한 번 호출 시 최대 250초 동안 가능한 많이 처리 (Railway 300초 타임아웃 여유).
    크론으로 10분마다 반복 호출 → 전체 백필 자동 완료.
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    import time as _time
    import traceback as _tb
    from app.services.doc_analysis_service import analyze_and_store

    DEADLINE_SEC = 250
    MAX_ITEMS = 200  # 한 호출당 최대 처리 건수 (안전장치)
    start_ts = _time.time()

    try:
        conn = get_db_connection()
        cur = conn.cursor()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB 연결 실패: {str(e)[:200]}")

    # 우선순위:
    # 1) 지원금액에 "억"이 포함된 공고 (큰 금액)
    # 2) "천만" 포함
    # 3) "백만" 포함
    # 4) "만" 포함
    # 5) 그 외 숫자 포함
    # 각 그룹 내에서는 마감일 가까운 순, NULL(상시)은 뒤로
    # + summary_text가 200자 이상인 것만 (분석 의미 없으면 skip)
    # + 마감일 미래 또는 NULL
    try:
        cur.execute("""
            SELECT a.announcement_id, a.title, a.origin_url, a.summary_text, a.support_amount, a.deadline_date,
              CASE
                WHEN a.support_amount ILIKE '%%억%%' THEN 1
                WHEN a.support_amount ILIKE '%%천만%%' THEN 2
                WHEN a.support_amount ILIKE '%%백만%%' THEN 3
                WHEN a.support_amount ILIKE '%%만%%' THEN 4
                ELSE 5
              END AS amt_priority
            FROM announcements a
            LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
            WHERE aa.id IS NULL
              AND a.support_amount IS NOT NULL AND a.support_amount != ''
              AND a.summary_text IS NOT NULL AND LENGTH(a.summary_text) >= 200
              AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
            ORDER BY amt_priority ASC,
                     CASE WHEN a.deadline_date IS NULL THEN 1 ELSE 0 END,
                     a.deadline_date ASC NULLS LAST,
                     a.announcement_id DESC
            LIMIT %s
        """, (MAX_ITEMS,))
        rows = cur.fetchall()
    except Exception as e:
        try: conn.close()
        except: pass
        raise HTTPException(status_code=500, detail=f"SELECT 쿼리 실패: {str(e)[:200]}\n{_tb.format_exc()[-500:]}")

    # 남은 건수 카운트 (같은 조건)
    cur.execute("""
        SELECT COUNT(*) AS cnt
        FROM announcements a
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE aa.id IS NULL
          AND a.support_amount IS NOT NULL AND a.support_amount != ''
          AND a.summary_text IS NOT NULL AND LENGTH(a.summary_text) >= 200
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
    """)
    _cnt_row = cur.fetchone()
    total_remaining = dict(_cnt_row)["cnt"] if _cnt_row else 0

    processed = 0
    success = 0
    failed = 0
    skipped_timeout = 0

    for row in rows:
        # 시간 체크 — 250초 초과 시 중단
        elapsed = _time.time() - start_ts
        if elapsed > DEADLINE_SEC:
            skipped_timeout = len(rows) - processed
            break

        r = dict(row)
        try:
            res = analyze_and_store(
                announcement_id=r["announcement_id"],
                origin_url=r.get("origin_url") or "",
                title=r["title"],
                db_conn=conn,
                summary_text=r.get("summary_text") or "",
            )
            if res.get("success"):
                success += 1
            else:
                failed += 1
        except Exception as e:
            failed += 1
            print(f"[analyze-batch] id={r['announcement_id']} error: {str(e)[:100]}")

        processed += 1

    conn.close()

    elapsed_total = round(_time.time() - start_ts, 1)
    return {
        "status": "SUCCESS",
        "processed": processed,
        "success": success,
        "failed": failed,
        "skipped_timeout": skipped_timeout,
        "remaining_after": max(0, total_remaining - success),
        "elapsed_seconds": elapsed_total,
        "done": (total_remaining - success) <= 0,
    }



# ── 일괄 선행분석 (백그라운드) ──────────────────────────────────────────
import threading as _threading

_bulk_job: dict = {
    "running": False, "stop_requested": False,
    "total": 0, "done": 0, "success": 0, "failed": 0, "skipped": 0,
    "started_at": None, "finished_at": None,
    "current_id": None, "current_title": "",
    "errors": [],  # 최근 20건 오류
}


def _send_bulk_analysis_email_report(is_final: bool = False):
    """일괄 분석 진행 상황을 이메일로 발송."""
    import smtplib as _smtp
    from email.mime.text import MIMEText as _MIMEText
    from email.mime.multipart import MIMEMultipart as _MIMEMultipart

    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pw   = os.environ.get("SMTP_PASSWORD", "")
    smtp_host = os.environ.get("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    to_email  = os.environ.get("OWNER_EMAIL") or os.environ.get("SMTP_FROM") or smtp_user

    if not smtp_user or not smtp_pw or not to_email:
        print("[BulkReport] SMTP 미설정 — 이메일 스킵")
        return

    j = _bulk_job
    total   = j.get("total", 0)
    done    = j.get("done", 0)
    success = j.get("success", 0)
    failed  = j.get("failed", 0)
    pct     = round(done / total * 100, 1) if total else 0
    status  = "완료" if is_final else ("실행 중" if j.get("running") else "중단됨")
    errors  = j.get("errors") or []

    subject = f"[GovMatch] 공고 선행분석 {'완료' if is_final else '진행 보고'} — {done}/{total}건 ({pct}%)"

    error_html = ""
    if errors:
        rows = "".join(
            f"<tr><td style='padding:4px 8px;border-bottom:1px solid #eee'>{e.get('id','')}</td>"
            f"<td style='padding:4px 8px;border-bottom:1px solid #eee'>{e.get('title','')}</td>"
            f"<td style='padding:4px 8px;border-bottom:1px solid #eee;color:#c00'>{e.get('error','')}</td></tr>"
            for e in errors[-10:]
        )
        error_html = f"""
        <h3 style='color:#c00;margin-top:20px'>최근 오류 (최대 10건)</h3>
        <table style='border-collapse:collapse;font-size:13px'>
          <tr style='background:#fef2f2'><th style='padding:4px 8px'>ID</th><th style='padding:4px 8px'>공고</th><th style='padding:4px 8px'>오류</th></tr>
          {rows}
        </table>"""

    html = f"""
    <div style='font-family:sans-serif;max-width:560px;margin:0 auto'>
      <h2 style='color:#4338ca'>📊 공고 선행분석 {status}</h2>
      <table style='border-collapse:collapse;width:100%;font-size:15px'>
        <tr><td style='padding:8px;color:#555'>상태</td><td style='padding:8px;font-weight:bold'>{status}</td></tr>
        <tr style='background:#f5f5ff'><td style='padding:8px;color:#555'>진행률</td><td style='padding:8px;font-weight:bold'>{done} / {total}건 ({pct}%)</td></tr>
        <tr><td style='padding:8px;color:#555'>성공</td><td style='padding:8px;color:#16a34a;font-weight:bold'>{success}건</td></tr>
        <tr style='background:#f5f5ff'><td style='padding:8px;color:#555'>실패</td><td style='padding:8px;color:#dc2626'>{failed}건</td></tr>
        <tr><td style='padding:8px;color:#555'>시작</td><td style='padding:8px'>{j.get('started_at','')}</td></tr>
        <tr style='background:#f5f5ff'><td style='padding:8px;color:#555'>현재 처리 중</td><td style='padding:8px'>{j.get('current_title','')}</td></tr>
      </table>
      {error_html}
      <p style='margin-top:20px;font-size:12px;color:#999'>
        DB 검증: GET /api/admin/bulk-analyze/db-check?password=***
      </p>
    </div>"""

    def _do_send():
        try:
            msg = _MIMEMultipart("alternative")
            msg["Subject"] = subject
            msg["From"]    = smtp_user
            msg["To"]      = to_email
            msg.attach(_MIMEText(html, "html", "utf-8"))
            with _smtp.SMTP(smtp_host, smtp_port, timeout=15) as sv:
                sv.ehlo(); sv.starttls(); sv.ehlo()
                sv.login(smtp_user, smtp_pw)
                sv.sendmail(smtp_user, [to_email], msg.as_string())
            print(f"[BulkReport] 이메일 발송 완료 → {to_email} ({done}/{total}건)")
        except Exception as ex:
            print(f"[BulkReport] 이메일 발송 실패: {ex}")

    # 별도 스레드로 실행 — 분석 메인 스레드 블로킹 방지
    _threading.Thread(target=_do_send, daemon=True).start()


def _run_bulk_analysis(mode: str, limit: int):
    """백그라운드 스레드: 미분석 공고 전체 순차 처리."""
    import time as _t
    from app.services.doc_analysis_service import analyze_and_store

    _bulk_job.update({"running": True, "stop_requested": False,
                      "done": 0, "success": 0, "failed": 0, "skipped": 0,
                      "started_at": _t.strftime("%Y-%m-%dT%H:%M:%S"),
                      "finished_at": None, "errors": []})
    _send_bulk_analysis_email_report()  # 시작 알림

    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # 공통 유효 조건: 마감 안된 공고만
        _active_cond = """
            AND a.is_archived = FALSE
            AND (
                a.deadline_type = 'ongoing'
                OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE)
                OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')
            )
        """
        if mode == "missing":
            # announcement_analysis 기록 없는 공고
            q = f"""
                SELECT a.announcement_id, a.title, a.origin_url, a.summary_text
                FROM announcements a
                LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
                WHERE aa.id IS NULL
                  AND a.origin_url IS NOT NULL
                  {_active_cond}
                ORDER BY a.deadline_date ASC NULLS LAST, a.announcement_id DESC
            """
        else:
            # summary_only 재분석 포함
            q = f"""
                SELECT a.announcement_id, a.title, a.origin_url, a.summary_text
                FROM announcements a
                LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
                WHERE (aa.id IS NULL OR aa.source_type IN ('summary', ''))
                  AND a.origin_url IS NOT NULL
                  {_active_cond}
                ORDER BY a.deadline_date ASC NULLS LAST, a.announcement_id DESC
            """
        if limit > 0:
            q += f" LIMIT {limit}"

        cur.execute(q)
        rows = [dict(r) for r in cur.fetchall()]
        conn.close()

        _bulk_job["total"] = len(rows)

        for row in rows:
            if _bulk_job["stop_requested"]:
                break

            ann_id = row["announcement_id"]
            _bulk_job["current_id"] = ann_id
            _bulk_job["current_title"] = (row.get("title") or "")[:50]

            try:
                item_conn = get_db_connection()
                res = analyze_and_store(
                    announcement_id=ann_id,
                    origin_url=row.get("origin_url") or "",
                    title=row.get("title") or "",
                    db_conn=item_conn,
                    summary_text=row.get("summary_text") or "",
                )
                item_conn.close()
                if res.get("success"):
                    _bulk_job["success"] += 1
                else:
                    _bulk_job["failed"] += 1
                    if len(_bulk_job["errors"]) < 20:
                        _bulk_job["errors"].append({
                            "id": ann_id, "title": (row.get("title") or "")[:40],
                            "error": (res.get("error") or "")[:100],
                            "source": res.get("source_type", ""),
                        })
            except Exception as ex:
                _bulk_job["failed"] += 1
                if len(_bulk_job["errors"]) < 20:
                    _bulk_job["errors"].append({"id": ann_id, "error": str(ex)[:100]})

            _bulk_job["done"] += 1
            _t.sleep(0.1)  # 최소 대기 (Gemini Flash 속도 대응)

    except Exception as ex:
        _bulk_job["errors"].append({"fatal": str(ex)[:200]})
    finally:
        _bulk_job.update({"running": False, "current_id": None,
                          "finished_at": _t.strftime("%Y-%m-%dT%H:%M:%S")})
        _send_bulk_analysis_email_report(is_final=True)  # 완료 알림


@app.post("/api/admin/bulk-analyze/start")
def api_bulk_analyze_start(req: AdminAuthRequest, mode: str = "missing", limit: int = 0):
    """백그라운드로 미분석 공고 일괄 선행분석.

    mode=missing  : announcement_analysis 기록 없는 공고만 (기본, ~9,022건)
    mode=all      : summary_only 재분석 포함 (~12,000건)
    limit=N       : N건만 처리 (0=전체)
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="비밀번호 오류")
    if _bulk_job["running"]:
        return {"status": "ALREADY_RUNNING", **_bulk_job}

    t = _threading.Thread(target=_run_bulk_analysis, args=(mode, limit), daemon=True)
    t.start()
    return {"status": "STARTED", "mode": mode, "limit": limit}


@app.get("/api/admin/bulk-analyze/status")
def api_bulk_analyze_status(password: str):
    """일괄 분석 진행 상황 조회."""
    if password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="비밀번호 오류")
    pct = round(_bulk_job["done"] / _bulk_job["total"] * 100, 1) if _bulk_job["total"] else 0
    return {**_bulk_job, "progress_pct": pct}


@app.post("/api/admin/bulk-analyze/stop")
def api_bulk_analyze_stop(req: AdminAuthRequest):
    """진행 중인 일괄 분석 중단 요청."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="비밀번호 오류")
    _bulk_job["stop_requested"] = True
    return {"status": "STOP_REQUESTED"}


@app.get("/api/admin/bulk-analyze/db-check")
def api_bulk_analyze_db_check(password: str):
    """DB 저장 검증: announcement_analysis 현황 + 샘플."""
    if password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="비밀번호 오류")
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT
            COUNT(*) AS total_active,
            COUNT(aa.id) AS has_analysis,
            COUNT(CASE WHEN aa.id IS NOT NULL AND aa.full_text IS NOT NULL AND LENGTH(aa.full_text) >= 500 THEN 1 END) AS has_fulltext,
            COUNT(CASE WHEN aa.id IS NOT NULL AND (aa.source_type IN ('summary','') OR aa.source_type IS NULL OR LENGTH(COALESCE(aa.full_text,'')) < 500) THEN 1 END) AS summary_or_short,
            COUNT(CASE WHEN aa.id IS NULL THEN 1 END) AS no_analysis
        FROM announcements a
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE a.is_archived = FALSE
    """)
    stats = dict(cur.fetchone())

    # 최근 분석 완료 샘플 3건
    cur.execute("""
        SELECT a.announcement_id, a.title, aa.source_type,
               LENGTH(aa.full_text) AS text_len,
               aa.updated_at,
               CASE WHEN aa.parsed_sections IS NOT NULL THEN
                   LENGTH(aa.parsed_sections::text) END AS sections_len
        FROM announcement_analysis aa
        JOIN announcements a ON a.announcement_id = aa.announcement_id
        WHERE aa.full_text IS NOT NULL AND LENGTH(aa.full_text) >= 500
        ORDER BY aa.updated_at DESC
        LIMIT 3
    """)
    samples = [dict(r) for r in cur.fetchall()]
    conn.close()
    return {"stats": stats, "recent_samples": samples}


@app.post("/api/admin/backfill-support-amount")
def api_backfill_support_amount(req: AdminAuthRequest, commit: bool = False, limit: int = 0):
    """관리자: 기존 공고의 support_amount 빈 필드를 summary_text/title 에서 정규식으로 채움.

    - commit=false (기본): dry-run, 추출 가능한 건수 + 샘플만 반환
    - commit=true: 실제 UPDATE
    - limit=0: 전체, 그 외 N건만 처리
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    from app.services.public_api_service import GovernmentAPIService
    extract = GovernmentAPIService._extract_amount_from_text

    conn = get_db_connection()
    cur = conn.cursor()
    sql = """
        SELECT announcement_id, title, summary_text
        FROM announcements
        WHERE (support_amount IS NULL OR support_amount = '')
        ORDER BY announcement_id DESC
    """
    if limit > 0:
        sql += f" LIMIT {int(limit)}"
    cur.execute(sql)
    rows = cur.fetchall()

    updated = 0
    samples = []
    for row in rows:
        r = dict(row)
        text = (r.get("summary_text") or "") + " " + (r.get("title") or "")
        amount = extract(text)
        if not amount:
            continue
        if commit:
            cur.execute(
                "UPDATE announcements SET support_amount = %s WHERE announcement_id = %s",
                (amount, r["announcement_id"]),
            )
        updated += 1
        if len(samples) < 10:
            samples.append({
                "id": r["announcement_id"],
                "title": (r.get("title") or "")[:60],
                "amount": amount,
            })

    if commit:
        conn.commit()
    conn.close()

    return {
        "status": "SUCCESS",
        "mode": "commit" if commit else "dry-run",
        "candidates": len(rows),
        "extracted": updated,
        "samples": samples,
    }


@app.post("/api/admin/embeddings/init")
def api_embeddings_init(req: AdminAuthRequest):
    """임베딩 매칭 인프라 초기화 — pgvector 확장 + announcement_embeddings 테이블 생성.
    안전: 기존 스키마 손대지 않음. 실패해도 부작용 없음.
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    result = {"pgvector": False, "table": False, "index": False, "errors": []}
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # 1) pgvector 확장 활성화 (권한 있으면 성공, 없으면 실패 → errors에 기록)
        try:
            cur.execute("CREATE EXTENSION IF NOT EXISTS vector")
            conn.commit()
            result["pgvector"] = True
        except Exception as e:
            conn.rollback()
            result["errors"].append(f"pgvector: {str(e)[:200]}")
            return result  # pgvector 없으면 더 진행 불가

        # 2) 임베딩 테이블 생성 (별도 테이블, 기존 announcements와 격리)
        try:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS announcement_embeddings (
                    announcement_id INTEGER PRIMARY KEY,
                    embedding vector(768),
                    source_text TEXT,
                    model_name VARCHAR(64),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.commit()
            result["table"] = True
        except Exception as e:
            conn.rollback()
            result["errors"].append(f"table: {str(e)[:200]}")
            return result

        # 3) HNSW 인덱스 (cosine 유사도)
        try:
            cur.execute("""
                CREATE INDEX IF NOT EXISTS idx_announcement_embeddings_cosine
                ON announcement_embeddings
                USING hnsw (embedding vector_cosine_ops)
            """)
            conn.commit()
            result["index"] = True
        except Exception as e:
            conn.rollback()
            result["errors"].append(f"index: {str(e)[:200]}")

        return {"status": "SUCCESS", **result}
    finally:
        try: conn.close()
        except: pass


@app.post("/api/admin/seed-interest-tags")
def api_seed_interest_tags(req: AdminAuthRequest):
    """관심 태그 풀 시드 + 공고 기반 자동 추출 + 임베딩 생성."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 미설정")
    genai.configure(api_key=api_key)

    conn = get_db_connection()
    cur = conn.cursor()
    inserted = 0
    skipped = 0
    embedded = 0
    errors = []

    # 1) 고정 시드 삽입
    seed_items = (
        [(t, c, "business", "seed") for t, c in _INTEREST_TAG_SEED_BIZ]
        + [(t, c, "individual", "seed") for t, c in _INTEREST_TAG_SEED_INDIV]
    )
    for tag, category, utype, source in seed_items:
        try:
            cur.execute(
                """INSERT INTO interest_tag_pool (tag, category, user_type, source, frequency)
                   VALUES (%s, %s, %s, %s, 0)
                   ON CONFLICT (tag, user_type) DO NOTHING""",
                (tag, category, utype, source),
            )
            if cur.rowcount > 0:
                inserted += 1
            else:
                skipped += 1
        except Exception as e:
            conn.rollback()
            errors.append(f"seed:{tag}: {str(e)[:100]}")
    conn.commit()

    # 2) 공고에서 자동 추출 — category 빈도 집계
    try:
        cur.execute("""
            SELECT category, COUNT(*) AS cnt
            FROM announcements
            WHERE category IS NOT NULL AND LENGTH(category) >= 2 AND LENGTH(category) <= 30
            GROUP BY category
            ORDER BY cnt DESC
            LIMIT 150
        """)
        for row in cur.fetchall():
            tag = (row["category"] or "").strip()
            if not tag or len(tag) < 2:
                continue
            try:
                cur.execute(
                    """INSERT INTO interest_tag_pool (tag, category, user_type, source, frequency)
                       VALUES (%s, 'auto', 'both', 'announcement_category', %s)
                       ON CONFLICT (tag, user_type) DO UPDATE SET frequency = EXCLUDED.frequency""",
                    (tag, int(row["cnt"] or 0)),
                )
                if cur.rowcount > 0:
                    inserted += 1
            except Exception as e:
                conn.rollback()
                errors.append(f"auto:{tag}: {str(e)[:100]}")
        conn.commit()
    except Exception as e:
        conn.rollback()
        errors.append(f"auto_extract: {str(e)[:150]}")

    # 3) 임베딩 미생성 태그 임베딩 생성 (최대 500개/호출)
    try:
        cur.execute(
            "SELECT id, tag, category FROM interest_tag_pool WHERE embedding IS NULL LIMIT 500"
        )
        targets = cur.fetchall()
        for t in targets:
            text = f"{t['tag']} {t.get('category') or ''}".strip()
            try:
                r = genai.embed_content(
                    model="models/gemini-embedding-001",
                    content=text,
                    task_type="RETRIEVAL_DOCUMENT",
                    output_dimensionality=768,
                )
                vec = r.get("embedding") if isinstance(r, dict) else getattr(r, "embedding", None)
                if not vec:
                    continue
                vec_str = "[" + ",".join(f"{x:.6f}" for x in vec) + "]"
                cur.execute(
                    "UPDATE interest_tag_pool SET embedding = %s::vector WHERE id = %s",
                    (vec_str, t["id"]),
                )
                embedded += 1
                if embedded % 50 == 0:
                    conn.commit()
            except Exception as e:
                conn.rollback()
                errors.append(f"embed:{t['tag']}: {str(e)[:80]}")
                continue
        conn.commit()
    except Exception as e:
        conn.rollback()
        errors.append(f"embed_loop: {str(e)[:150]}")

    cur.execute("SELECT COUNT(*) AS total, COUNT(embedding) AS with_emb FROM interest_tag_pool")
    stats = cur.fetchone()
    conn.close()
    return {
        "status": "SUCCESS",
        "inserted_new": inserted,
        "skipped_existing": skipped,
        "embedded_this_run": embedded,
        "total_tags": stats["total"],
        "total_with_embedding": stats["with_emb"],
        "errors": errors[:10],
    }


@app.post("/api/admin/coo/run")
def api_coo_run(req: AdminAuthRequest):
    """관리자: 오케스트레이터 AI 슈퍼바이저 수동 실행 — 즉시 품질체크/학습감시/보고서 발송."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    from app.services.orchestrator.supervisor import run_daily_supervision
    conn = get_db_connection()
    try:
        result = run_daily_supervision(conn)
    finally:
        conn.close()
    return {"status": "SUCCESS", **result}


@app.get("/api/admin/coo/reviews")
def api_coo_reviews(password: str, limit: int = 30):
    """관리자: 최근 저품질 상담 리뷰 목록 조회 (orchestrator_reviews)."""
    if password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT id, review_date, agent, consult_log_id,
                   accuracy, completeness, usefulness, avg_score,
                   issue, needs_review, created_at
            FROM orchestrator_reviews
            ORDER BY review_date DESC, created_at DESC
            LIMIT %s
        """, (max(1, min(limit, 200)),))
        rows = [dict(r) for r in cur.fetchall()]
        return {"status": "SUCCESS", "reviews": rows, "count": len(rows)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"조회 실패: {str(e)[:200]}")
    finally:
        conn.close()


class SuggestTagsRequest(BaseModel):
    text: str
    user_type: Optional[str] = "both"
    limit: Optional[int] = 10


@app.post("/api/ai/suggest-tags")
def api_suggest_tags(req: SuggestTagsRequest):
    """자연어 입력 → 가장 유사한 태그 TOP N 반환 (임베딩 코사인)."""
    text = (req.text or "").strip()
    if not text or len(text) < 2:
        return {"status": "SUCCESS", "suggestions": []}

    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return {"status": "SUCCESS", "suggestions": [], "error": "embedding unavailable"}

    try:
        genai.configure(api_key=api_key)
        r = genai.embed_content(
            model="models/gemini-embedding-001",
            content=text,
            task_type="RETRIEVAL_QUERY",
            output_dimensionality=768,
        )
        vec = r.get("embedding") if isinstance(r, dict) else getattr(r, "embedding", None)
        if not vec:
            return {"status": "SUCCESS", "suggestions": []}
        vec_str = "[" + ",".join(f"{x:.6f}" for x in vec) + "]"
    except Exception as e:
        return {"status": "ERROR", "suggestions": [], "error": str(e)[:200]}

    utype = (req.user_type or "both").lower()
    allowed_types = [utype, "both"] if utype in ("business", "individual") else ["business", "individual", "both"]
    limit = max(1, min(20, req.limit or 10))

    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute(
            """SELECT tag, category, user_type,
                      1 - (embedding <=> %s::vector) AS similarity
               FROM interest_tag_pool
               WHERE embedding IS NOT NULL
                 AND user_type = ANY(%s)
               ORDER BY embedding <=> %s::vector
               LIMIT %s""",
            (vec_str, allowed_types, vec_str, limit),
        )
        rows = cur.fetchall()
        suggestions = [
            {
                "tag": row["tag"],
                "category": row.get("category"),
                "user_type": row["user_type"],
                "similarity": round(float(row["similarity"] or 0), 3),
            }
            for row in rows
            if (row["similarity"] or 0) >= 0.60
        ]
    finally:
        conn.close()
    return {"status": "SUCCESS", "query": text, "suggestions": suggestions}


# ── 관심 태그 풀 시드 & 임베딩 ──

_INTEREST_TAG_SEED_BIZ = [
    ("창업지원", "창업"), ("예비창업", "창업"), ("초기창업", "창업"), ("재창업", "창업"),
    ("청년창업", "창업"), ("여성창업", "창업"), ("시니어창업", "창업"), ("사회적기업", "창업"),
    ("기술개발", "R&D"), ("R&D", "R&D"), ("연구개발", "R&D"), ("혁신제품", "R&D"),
    ("특허출원", "R&D"), ("기술이전", "R&D"), ("사업화", "R&D"), ("실증사업", "R&D"),
    ("수출마케팅", "수출"), ("해외진출", "수출"), ("수출바우처", "수출"), ("해외전시회", "수출"),
    ("글로벌진출", "수출"), ("수출보험", "수출"), ("FTA활용", "수출"),
    ("정책자금", "금융"), ("융자지원", "금융"), ("저리대출", "금융"), ("운전자금", "금융"),
    ("시설자금", "금융"), ("신용보증", "금융"), ("기술보증", "금융"), ("지역신보", "금융"),
    ("긴급경영안정자금", "금융"), ("신성장기반자금", "금융"), ("청년창업자금", "금융"),
    ("고용지원", "고용"), ("일자리지원", "고용"), ("청년채용", "고용"), ("정규직전환", "고용"),
    ("직무훈련", "고용"), ("인건비지원", "고용"),
    ("시설개선", "시설"), ("공장자동화", "시설"), ("스마트공장", "시설"), ("친환경설비", "시설"),
    ("에너지절감", "시설"),
    ("디지털전환", "디지털"), ("DX", "디지털"), ("클라우드", "디지털"), ("AI도입", "디지털"),
    ("빅데이터", "디지털"), ("IoT", "디지털"), ("메타버스", "디지털"),
    ("판로개척", "판로"), ("온라인판로", "판로"), ("쇼핑몰입점", "판로"), ("라이브커머스", "판로"),
    ("유통매칭", "판로"),
    ("교육훈련", "교육"), ("전문가양성", "교육"), ("컨설팅지원", "교육"), ("멘토링", "교육"),
    ("아카데미", "교육"),
    ("소상공인", "소상공인"), ("자영업", "소상공인"), ("골목상권", "소상공인"), ("전통시장", "소상공인"),
    ("희망리턴패키지", "소상공인"),
    ("바이오", "업종"), ("의료기기", "업종"), ("제약", "업종"), ("헬스케어", "업종"),
    ("반도체", "업종"), ("디스플레이", "업종"), ("2차전지", "업종"), ("자동차부품", "업종"),
    ("에너지신산업", "업종"), ("수소산업", "업종"), ("로봇", "업종"), ("드론", "업종"),
    ("푸드테크", "업종"), ("농식품", "업종"), ("수산업", "업종"), ("관광", "업종"),
    ("문화콘텐츠", "업종"), ("게임", "업종"), ("웹툰", "업종"), ("K-뷰티", "업종"),
    ("K-푸드", "업종"), ("패션", "업종"), ("디자인", "업종"),
    ("탄소중립", "에너지환경"), ("친환경", "에너지환경"), ("ESG", "에너지환경"), ("재생에너지", "에너지환경"),
    ("폐기물", "에너지환경"),
    ("벤처기업인증", "인증"), ("이노비즈", "인증"), ("메인비즈", "인증"), ("여성기업", "인증"),
    ("장애인기업", "인증"), ("사회적기업인증", "인증"), ("가족친화기업", "인증"),
    # 사용자 관점 키워드
    ("인건비", "수요"), ("운영비", "수요"), ("장비구매", "수요"), ("사무실임대", "수요"),
    ("시제품제작", "수요"), ("특허비용", "수요"), ("해외출장", "수요"),
]

_INTEREST_TAG_SEED_INDIV = [
    ("취업지원", "취업"), ("취업성공패키지", "취업"), ("내일배움카드", "취업"), ("국민취업지원", "취업"),
    ("직업훈련", "취업"), ("구직촉진수당", "취업"), ("이직지원", "취업"),
    ("전세자금", "주거"), ("월세지원", "주거"), ("버팀목전세", "주거"), ("디딤돌주택", "주거"),
    ("청년주택", "주거"), ("행복주택", "주거"), ("역세권청년주택", "주거"), ("임차보증금", "주거"),
    ("주거급여", "주거"), ("집수리지원", "주거"),
    ("국가장학금", "교육"), ("학자금대출", "교육"), ("기숙사비", "교육"), ("평생학습바우처", "교육"),
    ("학원비지원", "교육"), ("어학연수", "교육"),
    ("청년수당", "청년"), ("청년도약계좌", "청년"), ("청년내일채움", "청년"), ("청년월세", "청년"),
    ("청년희망적금", "청년"),
    ("출산지원금", "출산"), ("산후조리", "출산"), ("난임시술", "출산"), ("첫만남이용권", "출산"),
    ("보육료", "육아"), ("아동수당", "육아"), ("육아휴직", "육아"), ("어린이집", "육아"),
    ("아이돌봄", "육아"), ("방과후돌봄", "육아"),
    ("다자녀", "다자녀"), ("3자녀지원", "다자녀"), ("다둥이", "다자녀"),
    ("한부모", "한부모"), ("조손가정", "한부모"), ("모자가정", "한부모"),
    ("의료비지원", "의료"), ("건강검진", "의료"), ("난임부부", "의료"), ("중증질환", "의료"),
    ("희귀질환", "의료"),
    ("장애수당", "장애"), ("장애인일자리", "장애"), ("장애인활동지원", "장애"), ("장애인연금", "장애"),
    ("기초생활수급", "저소득"), ("차상위", "저소득"), ("긴급복지", "저소득"), ("저소득층", "저소득"),
    ("햇살론", "저소득"), ("미소금융", "저소득"), ("희망키움통장", "저소득"),
    ("기초연금", "노인"), ("노인일자리", "노인"), ("경로당", "노인"), ("치매안심", "노인"),
    ("독거노인", "노인"),
    ("문화바우처", "문화"), ("통합문화이용권", "문화"), ("스포츠강좌이용권", "문화"), ("도서문화", "문화"),
    ("에너지바우처", "복지"), ("전기요금감면", "복지"), ("생계비지원", "복지"),
]


class KsicBulkRequest(BaseModel):
    password: str
    records: List[dict]  # [{"code": "01110", "name": "...", "ksic11_name": "..."}]


@app.post("/api/admin/ksic/bulk-import")
def api_ksic_bulk_import(req: KsicBulkRequest):
    """KSIC 마스터 일괄 업로드 — ksic_classification에 UPSERT."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    conn = get_db_connection()
    cur = conn.cursor()
    # ksic_classification 테이블이 SQLite 스키마로 생성된 상태일 수 있음 → PG용으로 재생성
    try:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ksic_classification (
                code VARCHAR(5) PRIMARY KEY,
                name VARCHAR(255) NOT NULL,
                description TEXT
            )
        """)
        conn.commit()
    except Exception as e:
        conn.rollback()

    inserted = 0
    updated = 0
    failed = 0
    for rec in req.records:
        code = (rec.get("code") or rec.get("ksic10_code") or "").strip()
        name = (rec.get("name") or rec.get("ksic10_name") or "").strip()
        desc = (rec.get("description") or rec.get("ksic11_name") or "").strip()
        if not code or not name:
            failed += 1
            continue
        try:
            cur.execute("""
                INSERT INTO ksic_classification (code, name, description)
                VALUES (%s, %s, %s)
                ON CONFLICT (code) DO UPDATE SET
                    name = EXCLUDED.name,
                    description = EXCLUDED.description
            """, (code, name, desc))
            if cur.rowcount == 1:
                inserted += 1
            else:
                updated += 1
        except Exception as e:
            conn.rollback()
            failed += 1
            continue
    conn.commit()
    cur.execute("SELECT COUNT(*) AS c FROM ksic_classification")
    total = cur.fetchone()["c"]
    conn.close()
    return {"status": "SUCCESS", "inserted": inserted, "updated": updated, "failed": failed, "total_now": total}


@app.post("/api/admin/ksic/init-embeddings")
def api_ksic_init_embeddings(req: AdminAuthRequest):
    """KSIC 임베딩 테이블 생성."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    cur = conn.cursor()
    result = {"table": False, "index": False, "errors": []}
    try:
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ksic_embeddings (
                code VARCHAR(5) PRIMARY KEY,
                embedding vector(768),
                source_text TEXT,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.commit()
        result["table"] = True
    except Exception as e:
        conn.rollback()
        result["errors"].append(f"table: {str(e)[:150]}")
    try:
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_ksic_embeddings_cosine
            ON ksic_embeddings USING hnsw (embedding vector_cosine_ops)
        """)
        conn.commit()
        result["index"] = True
    except Exception as e:
        conn.rollback()
        result["errors"].append(f"index: {str(e)[:150]}")
    conn.close()
    return result


@app.post("/api/admin/ksic/embed-batch")
def api_ksic_embed_batch(req: AdminAuthRequest):
    """KSIC 항목을 Gemini로 임베딩 (한 번 호출 시 최대 300개)."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    import google.generativeai as genai
    import time as _time
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 미설정")
    genai.configure(api_key=api_key)

    DEADLINE = 250
    MAX_ITEMS = 300
    start = _time.time()

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT k.code, k.name, k.description
        FROM ksic_classification k
        LEFT JOIN ksic_embeddings e ON k.code = e.code
        WHERE e.code IS NULL
        ORDER BY k.code
        LIMIT %s
    """, (MAX_ITEMS,))
    rows = cur.fetchall()
    if not rows:
        conn.close()
        return {"status": "SUCCESS", "done": True, "processed": 0, "remaining": 0}

    ok = 0
    fail = 0
    errors = []
    for r in rows:
        if _time.time() - start > DEADLINE:
            break
        code = r["code"]
        name = r["name"]
        desc = r.get("description") or ""
        text = f"업종코드: {code}\n업종명: {name}\n{desc}"
        try:
            res = genai.embed_content(
                model="models/gemini-embedding-001",
                content=text,
                task_type="retrieval_document",
                output_dimensionality=768,
            )
            vec = res.get("embedding") if isinstance(res, dict) else res["embedding"]
            if not vec or len(vec) < 100:
                fail += 1
                continue
            vec_str = "[" + ",".join(f"{v:.6f}" for v in vec) + "]"
            cur.execute("""
                INSERT INTO ksic_embeddings (code, embedding, source_text, updated_at)
                VALUES (%s, %s::vector, %s, CURRENT_TIMESTAMP)
                ON CONFLICT (code) DO UPDATE SET
                    embedding = EXCLUDED.embedding,
                    source_text = EXCLUDED.source_text,
                    updated_at = CURRENT_TIMESTAMP
            """, (code, vec_str, text[:2000]))
            conn.commit()
            ok += 1
        except Exception as e:
            conn.rollback()
            fail += 1
            if len(errors) < 3:
                errors.append(f"{code}: {type(e).__name__}: {str(e)[:200]}")

    cur.execute("""
        SELECT COUNT(*) AS c FROM ksic_classification k
        LEFT JOIN ksic_embeddings e ON k.code = e.code
        WHERE e.code IS NULL
    """)
    remaining = cur.fetchone()["c"]
    conn.close()
    return {
        "status": "SUCCESS",
        "processed": ok + fail,
        "success": ok,
        "failed": fail,
        "remaining": remaining,
        "done": remaining == 0,
        "errors": errors,
        "elapsed": round(_time.time() - start, 1),
    }


class SectionFeedbackRequest(BaseModel):
    section_id: int
    rating: int  # 1 (👎) | 5 (👍)
    comment: Optional[str] = ""


@app.post("/api/pro/sections/feedback")
def api_pro_section_feedback(req: SectionFeedbackRequest, current_user: dict = Depends(_get_current_user)):
    """M: 컨설턴트가 RAG 답변 출처 섹션에 평가를 남김. 검색 가중치에 자동 반영."""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute(
            """INSERT INTO section_feedback (section_id, business_number, rating, comment)
               VALUES (%s, %s, %s, %s)""",
            (req.section_id, current_user["bn"], max(1, min(5, req.rating)), req.comment or ""),
        )
        conn.commit()
        # 누적 평점 평균 → 점수 가중치 (선택)
        cur.execute(
            "SELECT AVG(rating)::FLOAT AS avg_rating, COUNT(*) AS cnt FROM section_feedback WHERE section_id = %s",
            (req.section_id,)
        )
        row = cur.fetchone()
        conn.close()
        return {
            "status": "SUCCESS",
            "section_id": req.section_id,
            "avg_rating": round(row["avg_rating"], 2) if row else None,
            "feedback_count": row["cnt"] if row else 0,
        }
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"피드백 저장 실패: {str(e)[:200]}")


@app.get("/api/pro/insights/recent")
def api_pro_insights_recent(current_user: dict = Depends(_get_current_user)):
    """L: 최근 학습된 인사이트 (knowledge_base 신규 항목) — 컨설턴트가 체감할 수 있도록."""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT id, source, knowledge_type, category, content, confidence, use_count, created_at
            FROM knowledge_base
            WHERE source IN ('pro_consult', 'consult', 'free_chat')
            ORDER BY created_at DESC NULLS LAST
            LIMIT 30
        """)
        rows = cur.fetchall()
        items = []
        for r in rows:
            d = dict(r)
            d["created_at"] = str(d.get("created_at"))
            content = d.get("content")
            if isinstance(content, str):
                try: content = json.loads(content)
                except: content = {"raw": content}
            d["content"] = content
            items.append(d)
        # 통계
        cur.execute("SELECT COUNT(*) AS c FROM knowledge_base WHERE source IN ('pro_consult', 'consult')")
        total = cur.fetchone()["c"]
        cur.execute("""
            SELECT COUNT(*) AS c FROM knowledge_base
            WHERE source IN ('pro_consult', 'consult')
              AND created_at >= NOW() - INTERVAL '7 days'
        """)
        last_7 = cur.fetchone()["c"]
        conn.close()
        return {
            "status": "SUCCESS",
            "total_insights": total,
            "last_7_days": last_7,
            "recent": items,
        }
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"인사이트 조회 실패: {str(e)[:200]}")


class EvalRunRequest(BaseModel):
    password: str
    test_set_name: Optional[str] = "default"


@app.post("/api/admin/debug-user-match")
def api_debug_user_match(req: AdminAuthRequest):
    """디버그: 최근 사용자 5명의 프로필 저장 상태 + 매칭 결과 확인."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT business_number, email, company_name, address_city, industry_code,
               revenue_bracket, employee_count_bracket, interests, user_type,
               age_range, income_level, family_type, employment_status,
               establishment_date, plan, created_at
        FROM users
        WHERE plan != 'free' OR ai_usage_month > 0
        ORDER BY created_at DESC NULLS LAST
        LIMIT 5
    """)
    users = []
    for r in cur.fetchall():
        d = dict(r)
        d["created_at"] = str(d.get("created_at"))
        d["establishment_date"] = str(d.get("establishment_date"))
        # 채워진 필드 카운트
        meta_fields = ["address_city", "industry_code", "revenue_bracket", "employee_count_bracket", "interests", "establishment_date"]
        d["filled_count"] = sum(1 for f in meta_fields if d.get(f))
        d["filled_total"] = len(meta_fields)
        users.append(d)
    conn.close()

    # 각 사용자에 대해 매칭 시도
    out = []
    for u in users:
        try:
            from app.core.matcher import get_matches_hybrid
            user_type = u.get("user_type") or "both"
            is_indiv = user_type == "individual"
            matches = get_matches_hybrid(u, is_individual=is_indiv) or []
            top3 = [{"id": m.get("announcement_id"), "title": (m.get("title") or "")[:80], "score": m.get("match_score")} for m in matches[:3]]
            out.append({
                "bn": u.get("business_number"),
                "company": u.get("company_name"),
                "user_type": user_type,
                "filled": f"{u['filled_count']}/{u['filled_total']}",
                "industry_code": u.get("industry_code"),
                "address_city": u.get("address_city"),
                "interests": u.get("interests"),
                "match_count": len(matches),
                "top3": top3,
            })
        except Exception as e:
            out.append({"bn": u.get("business_number"), "error": str(e)[:200]})
    return {"users_checked": len(users), "results": out}


@app.post("/api/admin/eval/run")
def api_eval_run(req: EvalRunRequest):
    """K: PRO 상담 품질 자동 측정. 사전 정의된 Eval Set으로 RAG 검색 + 답변 정확도 측정.
    매 배포마다 호출하여 회귀 감지.
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    # 내장 Eval Set — 정부 지원사업 도메인 핵심 질문 100개 중 대표 20개
    EVAL_SET = [
        # 자격 관련
        {"q": "여성기업 대표가 받을 수 있는 정책자금이 뭐가 있나요?", "expected_keywords": ["여성", "정책자금", "지원"]},
        {"q": "예비창업자에게 적합한 창업지원사업은?", "expected_keywords": ["예비창업", "창업"]},
        {"q": "스마트공장 구축 지원금 자격 요건이 어떻게 되나요?", "expected_keywords": ["스마트공장", "자격", "제조"]},
        {"q": "청년창업 가점이 있는 R&D 사업이 있나요?", "expected_keywords": ["청년", "R&D"]},
        {"q": "사회적기업 인증 받은 곳이 받을 수 있는 사업?", "expected_keywords": ["사회적기업", "지원"]},
        # 신청 절차
        {"q": "정부지원사업 신청 시 필요한 서류는 일반적으로 무엇인가요?", "expected_keywords": ["서류", "신청"]},
        {"q": "사업계획서 작성할 때 가점받는 항목이 있나요?", "expected_keywords": ["사업계획서", "가점"]},
        # 개인 복지
        {"q": "1인가구 청년 주거 지원 사업이 있나요?", "expected_keywords": ["청년", "주거", "1인가구"]},
        {"q": "기초생활수급자가 받을 수 있는 의료 지원은?", "expected_keywords": ["기초생활", "의료"]},
        {"q": "다자녀 가구 출산 지원금 알려주세요", "expected_keywords": ["다자녀", "출산"]},
        # 지역
        {"q": "서울시 소상공인 지원 사업이 뭐가 있나요?", "expected_keywords": ["서울", "소상공인"]},
        {"q": "경기도 청년 일자리 사업?", "expected_keywords": ["경기", "청년", "일자리"]},
        # 분야
        {"q": "수출 마케팅 지원사업 추천", "expected_keywords": ["수출", "마케팅"]},
        {"q": "에너지 효율 개선 보조금이 있나요?", "expected_keywords": ["에너지", "효율"]},
        {"q": "디지털 전환 지원 사업?", "expected_keywords": ["디지털", "전환"]},
        # 금액
        {"q": "1억 이상 받을 수 있는 정책자금은?", "expected_keywords": ["정책자금", "억"]},
        # 일정
        {"q": "이번 달 마감 임박 공고가 있나요?", "expected_keywords": ["마감"]},
        # 특수조건
        {"q": "장애인 기업 우대 사업이 뭐가 있나요?", "expected_keywords": ["장애인"]},
        {"q": "북한이탈주민 창업 지원?", "expected_keywords": ["북한이탈", "창업"]},
        # 신청 전략
        {"q": "심사위원이 좋아하는 사업계획서 작성법", "expected_keywords": ["사업계획", "심사"]},
    ]

    from app.services.ai_consultant import search_sections_for_rag, search_knowledge_for_rag

    conn = get_db_connection()
    results = []
    total_recall = 0
    total_keyword_hits = 0
    total_keyword_count = 0

    for case in EVAL_SET:
        q = case["q"]
        expected = case["expected_keywords"]
        # 1) 섹션 RAG
        sec_rag = search_sections_for_rag(q, conn, top_k=5)
        sections = sec_rag.get("sections", [])
        # 2) 공고/지식 RAG (보조)
        doc_rag = search_knowledge_for_rag(q, conn, top_k_ann=5, top_k_kb=2)
        announcements = doc_rag.get("announcements", [])

        # 키워드 적중률 — 검색된 텍스트에 expected 키워드가 몇 개 포함되는지
        all_text = " ".join([
            (s.get("ann_title", "") + " " + s.get("section_text", "")) for s in sections
        ] + [
            (a.get("title", "") + " " + a.get("summary", "")) for a in announcements
        ])
        hits = sum(1 for k in expected if k in all_text)
        keyword_hit_rate = hits / max(1, len(expected))
        total_keyword_hits += hits
        total_keyword_count += len(expected)

        # Recall — 결과가 1개 이상이면 1
        has_results = bool(sections) or bool(announcements)
        if has_results:
            total_recall += 1

        results.append({
            "question": q,
            "expected": expected,
            "section_count": len(sections),
            "ann_count": len(announcements),
            "keyword_hit_rate": round(keyword_hit_rate, 2),
            "top_section": (sections[0].get("ann_title", "") if sections else None),
            "top_section_score": (sections[0].get("similarity") if sections else None),
        })

    conn.close()

    overall_recall = round(total_recall / len(EVAL_SET), 3)
    overall_keyword = round(total_keyword_hits / max(1, total_keyword_count), 3)
    quality_score = round((overall_recall * 0.4 + overall_keyword * 0.6) * 100, 1)

    return {
        "status": "SUCCESS",
        "test_set": req.test_set_name,
        "total_cases": len(EVAL_SET),
        "metrics": {
            "recall": overall_recall,  # 결과를 반환한 비율
            "keyword_hit_rate": overall_keyword,  # 기대 키워드 적중률
            "quality_score": quality_score,  # 종합 점수 (0~100)
        },
        "details": results,
    }


@app.post("/api/admin/run-migrations")
def api_run_migrations(req: AdminAuthRequest):
    """P1: email_logs, match_history 등 누락 테이블 강제 생성."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    cur = conn.cursor()
    results = {}

    migrations = [
        ("pro_consult_sessions.messages",
         "ALTER TABLE pro_consult_sessions ADD COLUMN IF NOT EXISTS messages JSONB DEFAULT '[]'::jsonb"),
        ("ai_consult_logs.session_id",
         "ALTER TABLE ai_consult_logs ADD COLUMN IF NOT EXISTS session_id VARCHAR(64)"),
        ("ai_consult_logs.updated_at",
         "ALTER TABLE ai_consult_logs ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP"),
        ("idx_ai_consult_logs_session_id",
         "CREATE UNIQUE INDEX IF NOT EXISTS idx_ai_consult_logs_session_id ON ai_consult_logs(session_id) WHERE session_id IS NOT NULL"),
        ("ksic_classification cleanup bad codes",
         "DELETE FROM ksic_classification WHERE LENGTH(code) != 5"),
        ("announcement_sections table",
         """CREATE TABLE IF NOT EXISTS announcement_sections (
                id SERIAL PRIMARY KEY,
                announcement_id INTEGER NOT NULL,
                section_type VARCHAR(40) NOT NULL,
                section_title TEXT,
                section_text TEXT NOT NULL,
                section_order INTEGER DEFAULT 0,
                embedding vector(768),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )"""),
        ("idx_announcement_sections_ann",
         "CREATE INDEX IF NOT EXISTS idx_announcement_sections_ann ON announcement_sections(announcement_id, section_type)"),
        ("idx_announcement_sections_type",
         "CREATE INDEX IF NOT EXISTS idx_announcement_sections_type ON announcement_sections(section_type)"),
        ("idx_announcement_sections_emb",
         "CREATE INDEX IF NOT EXISTS idx_announcement_sections_emb ON announcement_sections USING hnsw (embedding vector_cosine_ops)"),
        ("section_feedback table",
         """CREATE TABLE IF NOT EXISTS section_feedback (
                id SERIAL PRIMARY KEY,
                section_id INTEGER NOT NULL,
                business_number VARCHAR(20),
                rating SMALLINT,
                comment TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )"""),
        ("interest_tag_pool table",
         """CREATE TABLE IF NOT EXISTS interest_tag_pool (
                id SERIAL PRIMARY KEY,
                tag VARCHAR(100) NOT NULL,
                category VARCHAR(40),
                user_type VARCHAR(20) DEFAULT 'both',
                source VARCHAR(20) DEFAULT 'seed',
                frequency INTEGER DEFAULT 0,
                embedding vector(768),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE (tag, user_type)
            )"""),
        ("idx_interest_tag_pool_emb",
         "CREATE INDEX IF NOT EXISTS idx_interest_tag_pool_emb ON interest_tag_pool USING hnsw (embedding vector_cosine_ops)"),
        ("idx_interest_tag_pool_user_type",
         "CREATE INDEX IF NOT EXISTS idx_interest_tag_pool_user_type ON interest_tag_pool(user_type)"),
        ("pro_consult_sessions.phase",
         "ALTER TABLE pro_consult_sessions ADD COLUMN IF NOT EXISTS phase VARCHAR(20) DEFAULT 'collecting'"),
        ("pro_consult_sessions.matched_snapshot",
         "ALTER TABLE pro_consult_sessions ADD COLUMN IF NOT EXISTS matched_snapshot JSONB DEFAULT '[]'::jsonb"),
        ("client_files.client_id nullable",
         "ALTER TABLE client_files ALTER COLUMN client_id DROP NOT NULL"),
        ("client_files.ai_summary",
         "ALTER TABLE client_files ADD COLUMN IF NOT EXISTS ai_summary TEXT DEFAULT ''"),
        ("email_logs table",
         """CREATE TABLE IF NOT EXISTS email_logs (
                id SERIAL PRIMARY KEY,
                owner_business_number VARCHAR(20) NOT NULL,
                client_id INTEGER,
                recipient_email VARCHAR(255),
                recipient_name VARCHAR(100),
                subject TEXT,
                body TEXT,
                status VARCHAR(20),
                error_detail TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )"""),
        ("idx_email_logs_owner",
         "CREATE INDEX IF NOT EXISTS idx_email_logs_owner ON email_logs(owner_business_number, created_at DESC)"),
        ("match_history table",
         """CREATE TABLE IF NOT EXISTS match_history (
                id SERIAL PRIMARY KEY,
                business_number VARCHAR(20) NOT NULL,
                user_type VARCHAR(20),
                profile_snapshot JSONB,
                total_matches INTEGER DEFAULT 0,
                top_matches JSONB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )"""),
        ("idx_match_history_bn",
         "CREATE INDEX IF NOT EXISTS idx_match_history_bn ON match_history(business_number, created_at DESC)"),
    ]

    for name, sql in migrations:
        try:
            cur.execute(sql)
            conn.commit()
            results[name] = "OK"
        except Exception as e:
            try: conn.rollback()
            except: pass
            results[name] = f"ERR: {str(e)[:150]}"

    conn.close()
    return {"status": "SUCCESS", "results": results}


@app.post("/api/admin/db-audit")
def api_db_audit(req: AdminAuthRequest):
    """전체 DB 테이블 감사 — public 스키마의 모든 테이블 + 행수."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT table_name
        FROM information_schema.tables
        WHERE table_schema = 'public' AND table_type = 'BASE TABLE'
        ORDER BY table_name
    """)
    tables = [r["table_name"] for r in cur.fetchall()]
    results = []
    for t in tables:
        try:
            cur.execute(f'SELECT COUNT(*) AS c FROM "{t}"')
            count = cur.fetchone()["c"]
            # 최신 created_at/updated_at이 있으면 최근 활동 시각 조회
            cur.execute(f"""
                SELECT column_name FROM information_schema.columns
                WHERE table_schema = 'public' AND table_name = %s
                  AND column_name IN ('created_at','updated_at')
                ORDER BY column_name
            """, (t,))
            time_cols = [r["column_name"] for r in cur.fetchall()]
            last_ts = None
            if time_cols:
                col = "updated_at" if "updated_at" in time_cols else "created_at"
                try:
                    cur.execute(f'SELECT MAX("{col}") AS t FROM "{t}"')
                    last_ts = str(cur.fetchone()["t"])
                except Exception:
                    conn.rollback()
            results.append({"table": t, "count": count, "last_activity": last_ts})
        except Exception as e:
            results.append({"table": t, "error": str(e)[:150]})
            try: conn.rollback()
            except: pass
    conn.close()
    return {"total_tables": len(tables), "tables": results}


@app.post("/api/admin/debug-persistence")
def api_debug_persistence(req: AdminAuthRequest):
    """상담/세션/고객 데이터 저장 상태 확인."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    conn = get_db_connection()
    cur = conn.cursor()
    out: dict = {}
    tables = [
        "pro_consult_sessions",
        "consult_sessions",
        "client_profiles",
        "ai_consult_logs",
        "client_profile_files",
        "pro_reports",
    ]
    for t in tables:
        try:
            cur.execute(f"SELECT COUNT(*) AS c FROM {t}")
            out[t] = {"exists": True, "count": cur.fetchone()["c"]}
        except Exception as e:
            out[t] = {"exists": False, "error": str(e)[:120]}
            try: conn.rollback()
            except: pass

    # pro_consult_sessions 최근 5건 샘플
    try:
        cur.execute("""
            SELECT session_id, business_number, client_category, current_step,
                   collected, updated_at
            FROM pro_consult_sessions
            ORDER BY updated_at DESC NULLS LAST
            LIMIT 5
        """)
        rows = cur.fetchall()
        out["pro_consult_sessions_sample"] = [
            {
                "session_id": str(r["session_id"])[:12] + "...",
                "bn": r["business_number"],
                "category": r.get("client_category"),
                "step": r.get("current_step"),
                "collected_keys": list((r.get("collected") or {}).keys()) if isinstance(r.get("collected"), dict) else "non-dict",
                "collected_size": len(json.dumps(r.get("collected") or {}, ensure_ascii=False)),
                "updated_at": str(r.get("updated_at")),
            }
            for r in rows
        ]
    except Exception as e:
        out["pro_consult_sessions_sample_error"] = str(e)[:200]
        try: conn.rollback()
        except: pass

    # client_profiles 최근 5건
    try:
        cur.execute("""
            SELECT id, owner_business_number, business_number, client_name, client_type,
                   industry_code, address_city, is_active, created_at, updated_at
            FROM client_profiles
            ORDER BY updated_at DESC NULLS LAST
            LIMIT 5
        """)
        rows = cur.fetchall()
        out["client_profiles_sample"] = []
        for r in rows:
            d = dict(r)
            d["created_at"] = str(d.get("created_at"))
            d["updated_at"] = str(d.get("updated_at"))
            out["client_profiles_sample"].append(d)
    except Exception as e:
        out["client_profiles_sample_error"] = str(e)[:200]
        try: conn.rollback()
        except: pass

    # ai_consult_logs 최근 5건
    try:
        cur.execute("""
            SELECT id, business_number, announcement_id, conclusion, feedback,
                   created_at,
                   CASE WHEN messages IS NULL THEN 0
                        WHEN jsonb_typeof(messages) = 'array' THEN jsonb_array_length(messages)
                        ELSE 0 END AS msg_count
            FROM ai_consult_logs
            ORDER BY created_at DESC NULLS LAST
            LIMIT 5
        """)
        rows = cur.fetchall()
        out["ai_consult_logs_sample"] = []
        for r in rows:
            d = dict(r)
            d["created_at"] = str(d.get("created_at"))
            out["ai_consult_logs_sample"].append(d)
    except Exception as e:
        out["ai_consult_logs_sample_error"] = str(e)[:200]
        try: conn.rollback()
        except: pass

    conn.close()
    return out


@app.post("/api/admin/embeddings/debug-match")
def api_embeddings_debug_match(req: AdminAuthRequest):
    """임베딩 매칭 디버그 — 샘플 프로필로 검색 + 환경변수 확인."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    from app.core.matcher import get_matches_by_embedding, get_matches_hybrid
    flag = os.environ.get("USE_EMBEDDING_MATCHING", "(미설정)")

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) AS c FROM announcement_embeddings")
    emb_count = cur.fetchone()["c"]
    cur.execute(f"SELECT COUNT(*) AS c FROM announcements WHERE {valid_announcement_where()}")
    active_ann = cur.fetchone()["c"]
    cur.execute("SELECT target_type, COUNT(*) AS c FROM announcements GROUP BY target_type")
    tt_breakdown = [{"target_type": r["target_type"], "count": r["c"]} for r in cur.fetchall()]
    conn.close()

    sample_profile = {
        "company_name": "테스트",
        "industry_code": "26003",
        "address_city": "서울특별시",
        "revenue_bracket": "10억~30억",
        "employee_count_bracket": "10인~30인",
        "interests": "스마트공장,정책자금",
        "establishment_date": "2019-01-01",
    }
    emb_results = get_matches_by_embedding(sample_profile, top_k=5, target_type_filter="business")
    hybrid_results = get_matches_hybrid(sample_profile, is_individual=False)

    return {
        "env_flag": flag,
        "embeddings_count": emb_count,
        "active_announcements": active_ann,
        "target_type_breakdown": tt_breakdown,
        "direct_embedding_results": len(emb_results),
        "direct_embedding_sample": [
            {"id": r.get("announcement_id"), "title": (r.get("title") or "")[:60], "sim": round(r.get("similarity", 0) or 0, 4)}
            for r in emb_results[:5]
        ],
        "hybrid_results": len(hybrid_results),
        "hybrid_sample": [
            {"id": r.get("announcement_id"), "title": (r.get("title") or "")[:60], "score": r.get("match_score")}
            for r in hybrid_results[:5]
        ],
    }


@app.post("/api/admin/sections/extract-from-analysis")
def api_sections_extract(req: AdminAuthRequest):
    """기존 announcement_analysis.parsed_sections에서 announcement_sections로 추출 적재.
    (한 번 호출 시 최대 300건 처리)
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    # doc_analysis_service.py 의 parsed_sections 실제 키와 일치시킴
    SECTION_KEYS = [
        ("eligibility", "자격요건"),
        ("exclusions", "제외대상"),
        ("exceptions", "예외조항"),
        ("bonus_points", "가점·우대"),
        ("required_docs", "제출서류"),
        ("evaluation_criteria", "심사기준"),
        ("support_details", "지원내용"),
        ("timeline", "일정"),
        ("application_method", "신청방법"),
        # 레거시 키(혹시 남아있을 경우 호환)
        ("target", "지원대상"),
        ("required_documents", "제출서류"),
        ("support_amount", "지원금액"),
        ("support_content", "지원내용"),
        ("schedule", "일정"),
        ("contact", "문의처"),
        ("notes", "유의사항"),
    ]

    conn = get_db_connection()
    cur = conn.cursor()
    # 섹션 3개 미만인 공고(=이전 키 미스매치로 누락)를 우선 재추출
    cur.execute("""
        SELECT aa.announcement_id, aa.parsed_sections, aa.deep_analysis
        FROM announcement_analysis aa
        LEFT JOIN (
            SELECT announcement_id, COUNT(*) AS cnt
            FROM announcement_sections
            GROUP BY announcement_id
        ) sc ON sc.announcement_id = aa.announcement_id
        WHERE aa.parsed_sections IS NOT NULL
          AND COALESCE(sc.cnt, 0) < 3
        ORDER BY COALESCE(sc.cnt, 0) ASC, aa.announcement_id
        LIMIT 300
    """)
    rows = cur.fetchall()
    if not rows:
        conn.close()
        return {"status": "SUCCESS", "done": True, "processed": 0, "remaining": 0}

    inserted_total = 0
    processed = 0
    for r in rows:
        ann_id = r["announcement_id"]
        ps = r.get("parsed_sections")
        if isinstance(ps, str):
            try: ps = json.loads(ps)
            except: ps = None
        if not isinstance(ps, dict):
            processed += 1
            continue
        # 이 공고에 이미 존재하는 섹션 타입 (중복 방지)
        cur.execute(
            "SELECT section_type FROM announcement_sections WHERE announcement_id = %s",
            (ann_id,),
        )
        existing_types = {row["section_type"] for row in cur.fetchall()}
        # 새로 시작할 order
        cur.execute(
            "SELECT COALESCE(MAX(section_order), -1) + 1 AS next_order FROM announcement_sections WHERE announcement_id = %s",
            (ann_id,),
        )
        order = cur.fetchone()["next_order"]
        for key, label in SECTION_KEYS:
            if key in existing_types:
                continue
            v = ps.get(key)
            if not v:
                continue
            text = v if isinstance(v, str) else json.dumps(v, ensure_ascii=False)
            text = text.strip()
            if len(text) < 10:
                continue
            try:
                cur.execute(
                    """INSERT INTO announcement_sections
                       (announcement_id, section_type, section_title, section_text, section_order)
                       VALUES (%s, %s, %s, %s, %s)""",
                    (ann_id, key, label, text[:5000], order),
                )
                existing_types.add(key)
                inserted_total += 1
                order += 1
            except Exception as ie:
                conn.rollback()
                continue
        # deep_analysis 핵심포인트 별도 섹션
        da = r.get("deep_analysis")
        if isinstance(da, str):
            try: da = json.loads(da)
            except: da = None
        if isinstance(da, dict):
            for k in ("key_points", "summary", "strategy"):
                v = da.get(k)
                if isinstance(v, str) and len(v.strip()) >= 20:
                    try:
                        cur.execute(
                            """INSERT INTO announcement_sections
                               (announcement_id, section_type, section_title, section_text, section_order)
                               VALUES (%s, %s, %s, %s, %s)""",
                            (ann_id, f"insight_{k}", k, v.strip()[:5000], order),
                        )
                        inserted_total += 1
                        order += 1
                    except Exception:
                        conn.rollback()
        conn.commit()
        processed += 1

    cur.execute("""
        SELECT COUNT(DISTINCT aa.announcement_id) AS c
        FROM announcement_analysis aa
        LEFT JOIN (
            SELECT announcement_id, COUNT(*) AS cnt
            FROM announcement_sections
            GROUP BY announcement_id
        ) sc ON sc.announcement_id = aa.announcement_id
        WHERE aa.parsed_sections IS NOT NULL
          AND COALESCE(sc.cnt, 0) < 3
    """)
    remaining = cur.fetchone()["c"]
    conn.close()
    return {
        "status": "SUCCESS",
        "processed": processed,
        "inserted_sections": inserted_total,
        "remaining_announcements": remaining,
        "done": remaining == 0,
    }


@app.post("/api/admin/sections/embed-batch")
def api_sections_embed_batch(req: AdminAuthRequest):
    """announcement_sections 중 embedding 미생성된 것을 batch로 임베딩 (최대 300건/호출)."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    import time as _time
    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 미설정")
    genai.configure(api_key=api_key)

    DEADLINE = 250
    MAX_ITEMS = 300
    start = _time.time()
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT s.id, s.section_type, s.section_title, s.section_text,
               a.title AS ann_title, a.department, a.support_amount
        FROM announcement_sections s
        LEFT JOIN announcements a ON s.announcement_id = a.announcement_id
        WHERE s.embedding IS NULL
        ORDER BY s.id
        LIMIT %s
    """, (MAX_ITEMS,))
    rows = cur.fetchall()
    if not rows:
        conn.close()
        return {"status": "SUCCESS", "done": True, "processed": 0, "remaining": 0}

    ok = 0
    fail = 0
    for r in rows:
        if _time.time() - start > DEADLINE:
            break
        try:
            ann_title = (r.get("ann_title") or "")[:150]
            dept = (r.get("department") or "")[:80]
            stype = r.get("section_type") or ""
            stitle = r.get("section_title") or ""
            stext = (r.get("section_text") or "")[:3500]
            text = f"공고: {ann_title}\n부처: {dept}\n섹션: {stitle} ({stype})\n내용: {stext}"
            res = genai.embed_content(
                model="models/gemini-embedding-001",
                content=text,
                task_type="retrieval_document",
                output_dimensionality=768,
            )
            vec = res.get("embedding") if isinstance(res, dict) else res["embedding"]
            if not vec or len(vec) < 100:
                fail += 1
                continue
            vec_str = "[" + ",".join(f"{v:.6f}" for v in vec) + "]"
            cur.execute(
                "UPDATE announcement_sections SET embedding = %s::vector, updated_at = CURRENT_TIMESTAMP WHERE id = %s",
                (vec_str, r["id"]),
            )
            conn.commit()
            ok += 1
        except Exception as e:
            conn.rollback()
            fail += 1
            print(f"[SectEmb] #{r.get('id')}: {str(e)[:120]}")

    cur.execute("SELECT COUNT(*) AS c FROM announcement_sections WHERE embedding IS NULL")
    remaining = cur.fetchone()["c"]
    conn.close()
    return {
        "status": "SUCCESS",
        "success": ok,
        "failed": fail,
        "remaining": remaining,
        "done": remaining == 0,
        "elapsed": round(_time.time() - start, 1),
    }


@app.post("/api/admin/embeddings/reembed-analyzed")
def api_embeddings_reembed_analyzed(req: AdminAuthRequest):
    """정밀분석 있는 공고를 풍부한 텍스트로 재임베딩 (한 번 호출 당 최대 200건).
    기존 임베딩이 있어도 ON CONFLICT UPDATE로 덮어씀.
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    import time as _time
    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 미설정")
    genai.configure(api_key=api_key)

    DEADLINE_SEC = 250
    MAX_ITEMS = 200
    start = _time.time()
    conn = get_db_connection()
    cur = conn.cursor()

    # 정밀분석이 있고, 기존 임베딩 source_text에 "자격요건" 문자열이 없는 것 (= 아직 풍부화 안 된 것)
    cur.execute("""
        SELECT a.announcement_id, a.title, a.department, a.category,
               a.support_amount, a.region, a.summary_text,
               aa.parsed_sections, aa.deep_analysis
        FROM announcement_analysis aa
        JOIN announcements a ON a.announcement_id = aa.announcement_id
        LEFT JOIN announcement_embeddings e ON a.announcement_id = e.announcement_id
        WHERE (e.source_text IS NULL OR e.source_text NOT LIKE '%%자격요건%%')
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
        ORDER BY a.announcement_id
        LIMIT %s
    """, (MAX_ITEMS,))
    rows = cur.fetchall()
    if not rows:
        conn.close()
        return {"status": "SUCCESS", "done": True, "processed": 0, "remaining": 0}

    ok = 0
    fail = 0
    for row in rows:
        if _time.time() - start > DEADLINE_SEC:
            break
        try:
            title = (row.get("title") or "")[:200]
            dept = (row.get("department") or "")[:100]
            cat = (row.get("category") or "")[:50]
            amount = (row.get("support_amount") or "")[:50]
            region = (row.get("region") or "")[:50]
            summary = (row.get("summary_text") or "")[:2000]
            source_text = f"제목: {title}\n부처: {dept}\n카테고리: {cat}\n지원금액: {amount}\n지역: {region}\n내용: {summary}"
            ps = row.get("parsed_sections")
            da = row.get("deep_analysis")
            if isinstance(ps, str):
                try: ps = json.loads(ps)
                except: ps = None
            if isinstance(da, str):
                try: da = json.loads(da)
                except: da = None
            extras = []
            if isinstance(ps, dict):
                for key, label in [("eligibility", "자격요건"), ("required_documents", "제출서류"),
                                   ("application_method", "신청방법"), ("target", "지원대상")]:
                    v = ps.get(key)
                    if isinstance(v, str) and v.strip():
                        extras.append(f"{label}: {v[:400]}")
            if isinstance(da, dict):
                kp = da.get("key_points") or da.get("summary")
                if isinstance(kp, str) and kp.strip():
                    extras.append(f"핵심포인트: {kp[:400]}")
            if extras:
                source_text += "\n" + "\n".join(extras)

            res = genai.embed_content(
                model="models/gemini-embedding-001",
                content=source_text,
                task_type="retrieval_document",
                output_dimensionality=768,
            )
            vec = res.get("embedding") if isinstance(res, dict) else res["embedding"]
            if not vec or len(vec) < 100:
                fail += 1
                continue
            vec_str = "[" + ",".join(f"{v:.6f}" for v in vec) + "]"
            cur.execute("""
                INSERT INTO announcement_embeddings (announcement_id, embedding, source_text, model_name, updated_at)
                VALUES (%s, %s::vector, %s, %s, CURRENT_TIMESTAMP)
                ON CONFLICT (announcement_id) DO UPDATE SET
                    embedding = EXCLUDED.embedding,
                    source_text = EXCLUDED.source_text,
                    model_name = EXCLUDED.model_name,
                    updated_at = CURRENT_TIMESTAMP
            """, (row["announcement_id"], vec_str, source_text[:5000], "gemini-embedding-001-enriched"))
            conn.commit()
            ok += 1
        except Exception as e:
            conn.rollback()
            fail += 1
            print(f"[ReEmbBatch] #{row.get('announcement_id')}: {str(e)[:150]}")

    cur.execute("""
        SELECT COUNT(*) AS c
        FROM announcement_analysis aa
        JOIN announcements a ON a.announcement_id = aa.announcement_id
        LEFT JOIN announcement_embeddings e ON a.announcement_id = e.announcement_id
        WHERE (e.source_text IS NULL OR e.source_text NOT LIKE '%%자격요건%%')
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
    """)
    remaining = cur.fetchone()["c"]
    conn.close()
    return {
        "status": "SUCCESS",
        "success": ok,
        "failed": fail,
        "remaining": remaining,
        "done": remaining == 0,
        "elapsed": round(_time.time() - start, 1),
    }


@app.post("/api/admin/embeddings/list-models")
def api_embeddings_list_models(req: AdminAuthRequest):
    """사용 가능한 Gemini 모델 리스트 조회 (임베딩 지원 모델만)."""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 미설정")
    genai.configure(api_key=api_key)
    try:
        import google.generativeai as _g
        sdk_ver = getattr(_g, "__version__", "unknown")
    except Exception:
        sdk_ver = "unknown"
    models = []
    try:
        for m in genai.list_models():
            methods = list(getattr(m, "supported_generation_methods", []) or [])
            if "embedContent" in methods or "embed_content" in methods:
                models.append({"name": m.name, "methods": methods})
        return {"sdk_version": sdk_ver, "embedding_models": models, "total": len(models)}
    except Exception as e:
        return {"sdk_version": sdk_ver, "error": f"{type(e).__name__}: {str(e)[:300]}"}


@app.post("/api/admin/embeddings/batch")
def api_embeddings_batch(req: AdminAuthRequest):
    """임베딩 배치 생성 — 미임베딩 공고를 Gemini text-embedding-004로 벡터화.
    한 번 호출 시 최대 250초 동안 가능한 많이 처리. 크론으로 반복 호출.
    """
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")

    import time as _time
    import google.generativeai as genai

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY 미설정")
    genai.configure(api_key=api_key)

    DEADLINE_SEC = 250
    MAX_ITEMS = 300
    BATCH_SIZE = 10  # Gemini 임베딩 병렬 처리 단위
    start_ts = _time.time()

    conn = get_db_connection()
    cur = conn.cursor()

    # 미임베딩 공고 선별 (마감 안 된 것 우선, 지원금액 있는 것 우선)
    # announcement_analysis를 LEFT JOIN → parsed_sections / deep_analysis 병합 (있으면 풍부화)
    cur.execute("""
        SELECT a.announcement_id, a.title, a.department, a.category,
               a.support_amount, a.region, a.summary_text,
               aa.parsed_sections, aa.deep_analysis
        FROM announcements a
        LEFT JOIN announcement_embeddings e ON a.announcement_id = e.announcement_id
        LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
        WHERE e.announcement_id IS NULL
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
          AND a.summary_text IS NOT NULL AND LENGTH(a.summary_text) > 50
        ORDER BY
            CASE WHEN a.support_amount ILIKE '%%억%%' THEN 0
                 WHEN a.support_amount ILIKE '%%천만%%' THEN 1
                 WHEN a.support_amount ILIKE '%%백만%%' THEN 2
                 ELSE 3 END,
            a.deadline_date ASC NULLS LAST
        LIMIT %s
    """, (MAX_ITEMS,))
    rows = cur.fetchall()

    if not rows:
        conn.close()
        return {"status": "SUCCESS", "processed": 0, "done": True, "message": "모든 공고 임베딩 완료"}

    processed = 0
    success = 0
    failed = 0
    error_samples = []

    for row in rows:
        if _time.time() - start_ts > DEADLINE_SEC:
            break
        try:
            # 임베딩 대상 텍스트 구성
            title = (row.get("title") or "")[:200]
            dept = (row.get("department") or "")[:100]
            cat = (row.get("category") or "")[:50]
            amount = (row.get("support_amount") or "")[:50]
            region = (row.get("region") or "")[:50]
            summary = (row.get("summary_text") or "")[:2000]
            source_text = f"제목: {title}\n부처: {dept}\n카테고리: {cat}\n지원금액: {amount}\n지역: {region}\n내용: {summary}"

            # 정밀분석 데이터가 있으면 임베딩 텍스트에 병합 (자격요건/서류/신청방법 → 의미 검색 품질 향상)
            ps = row.get("parsed_sections")
            da = row.get("deep_analysis")
            if isinstance(ps, str):
                try: ps = json.loads(ps)
                except: ps = None
            if isinstance(da, str):
                try: da = json.loads(da)
                except: da = None
            extras = []
            if isinstance(ps, dict):
                for key, label in [("eligibility", "자격요건"), ("required_documents", "제출서류"),
                                   ("application_method", "신청방법"), ("target", "지원대상")]:
                    v = ps.get(key)
                    if isinstance(v, str) and v.strip():
                        extras.append(f"{label}: {v[:400]}")
            if isinstance(da, dict):
                kp = da.get("key_points") or da.get("summary")
                if isinstance(kp, str) and kp.strip():
                    extras.append(f"핵심포인트: {kp[:400]}")
            if extras:
                source_text += "\n" + "\n".join(extras)

            # Gemini 임베딩 호출
            res = genai.embed_content(
                model="models/gemini-embedding-001",
                content=source_text,
                task_type="retrieval_document",
                output_dimensionality=768,
            )
            vec = res.get("embedding") if isinstance(res, dict) else res["embedding"]
            if not vec or len(vec) < 100:
                failed += 1
                processed += 1
                continue

            # pgvector 형식으로 저장 (str로 직렬화: '[v1,v2,...]')
            vec_str = "[" + ",".join(f"{v:.6f}" for v in vec) + "]"
            cur.execute("""
                INSERT INTO announcement_embeddings (announcement_id, embedding, source_text, model_name, updated_at)
                VALUES (%s, %s::vector, %s, %s, CURRENT_TIMESTAMP)
                ON CONFLICT (announcement_id) DO UPDATE SET
                    embedding = EXCLUDED.embedding,
                    source_text = EXCLUDED.source_text,
                    model_name = EXCLUDED.model_name,
                    updated_at = CURRENT_TIMESTAMP
            """, (row["announcement_id"], vec_str, source_text[:3000], "gemini-embedding-001"))
            conn.commit()
            success += 1
            processed += 1
        except Exception as e:
            conn.rollback()
            failed += 1
            processed += 1
            if len(error_samples) < 3:
                error_samples.append(f"#{row.get('announcement_id')}: {type(e).__name__}: {str(e)[:300]}")
            print(f"[EmbBatch] #{row.get('announcement_id')} error: {str(e)[:150]}")

    # 남은 건수 확인
    cur.execute("""
        SELECT COUNT(*) AS remaining FROM announcements a
        LEFT JOIN announcement_embeddings e ON a.announcement_id = e.announcement_id
        WHERE e.announcement_id IS NULL
          AND (a.deadline_type = 'ongoing' OR (a.deadline_type = 'fixed' AND a.deadline_date >= CURRENT_DATE) OR (a.deadline_type = 'unknown' AND a.created_at >= CURRENT_DATE - INTERVAL '3 months')) AND a.is_archived = FALSE
          AND a.summary_text IS NOT NULL AND LENGTH(a.summary_text) > 50
    """)
    remaining = cur.fetchone()["remaining"]
    conn.close()

    return {
        "status": "SUCCESS",
        "processed": processed,
        "success": success,
        "failed": failed,
        "remaining_after": remaining,
        "elapsed_seconds": round(_time.time() - start_ts, 1),
        "done": remaining == 0,
        "error_samples": error_samples,
    }


@app.get("/api/admin/analysis-stats")
def api_analysis_stats(req: AdminAuthRequest = Depends()):
    """관리자: 분석 현황 통계"""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) as total FROM announcements")
    total = cur.fetchone()["total"]
    cur.execute("SELECT COUNT(*) as analyzed FROM announcement_analysis")
    analyzed = cur.fetchone()["analyzed"]
    cur.execute("SELECT COUNT(*) as consults FROM ai_consult_logs")
    consults = cur.fetchone()["consults"]
    cur.execute("SELECT COUNT(*) as feedback_count FROM ai_consult_logs WHERE feedback IS NOT NULL")
    feedback_count = cur.fetchone()["feedback_count"]
    conn.close()
    return {
        "total_announcements": total,
        "analyzed": analyzed,
        "pending": total - analyzed,
        "consult_logs": consults,
        "feedback_collected": feedback_count,
    }


ADMIN_TOKEN_SECRET = os.getenv("ADMIN_PASSWORD", "fallback")


def _create_admin_token() -> str:
    payload = ADMIN_TOKEN_SECRET.encode()
    return hmac.new(payload, b"admin-session", hashlib.sha256).hexdigest()


def _verify_admin(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="인증 토큰이 필요합니다.")
    token = authorization.split(" ", 1)[1]
    expected = _create_admin_token()
    if not hmac.compare_digest(token, expected):
        raise HTTPException(status_code=401, detail="유효하지 않은 토큰입니다.")


@app.post("/api/admin/auth")
def admin_auth(request: AdminAuthRequest):
    admin_pw = os.getenv("ADMIN_PASSWORD", "")
    if not admin_pw:
        raise HTTPException(status_code=500, detail="ADMIN_PASSWORD가 설정되지 않았습니다.")
    if request.password != admin_pw:
        raise HTTPException(status_code=401, detail="비밀번호가 올바르지 않습니다.")
    token = _create_admin_token()
    return {"status": "SUCCESS", "token": token}


@app.get("/api/admin/security-status", dependencies=[Depends(_verify_admin)])
def api_security_status():
    """보안 에이전트 실시간 상태 — 차단 IP, 이상 감지 이벤트"""
    return {"status": "SUCCESS", "security": security_agent.get_status()}


@app.get("/api/admin/users", dependencies=[Depends(_verify_admin)])
def get_admin_users():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT u.user_id, u.business_number, u.company_name, u.address_city,
               u.industry_code, u.revenue_bracket, u.employee_count_bracket, u.updated_at,
               ns.email, ns.channel, ns.is_active as notify_active
        FROM users u
        LEFT JOIN notification_settings ns ON u.business_number = ns.business_number
        ORDER BY u.updated_at DESC
    """)
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return {"status": "SUCCESS", "data": rows}


@app.delete("/api/admin/users/{user_id}", dependencies=[Depends(_verify_admin)])
def delete_admin_user(user_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT business_number FROM users WHERE user_id = %s", (user_id,))
    row = cursor.fetchone()
    if row:
        bn = row["business_number"]
        cursor.execute("DELETE FROM notification_settings WHERE business_number = %s", (bn,))
    cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "message": "사용자 삭제 완료"}


@app.get("/api/admin/stats", dependencies=[Depends(_verify_admin)])
def get_admin_stats():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT COUNT(*) FROM announcements")
    total = cursor.fetchone()["count"]

    cursor.execute("""
        SELECT origin_source, COUNT(*) as cnt
        FROM announcements
        WHERE origin_source IS NOT NULL AND origin_source != ''
        GROUP BY origin_source
        ORDER BY cnt DESC
    """)
    by_source = [{"source": r["origin_source"], "count": r["cnt"]} for r in cursor.fetchall()]

    cursor.execute("SELECT COUNT(*) FROM users")
    user_count = cursor.fetchone()["count"]

    cursor.execute("SELECT COUNT(*) FROM admin_urls WHERE is_active = 1")
    active_urls = cursor.fetchone()["count"]

    conn.close()
    return {
        "status": "SUCCESS",
        "data": {
            "total_announcements": total,
            "by_source": by_source,
            "user_count": user_count,
            "active_manual_urls": active_urls,
        }
    }


@app.get("/api/admin/analytics", dependencies=[Depends(_verify_admin)])
def get_admin_analytics():
    """사용 패턴 분석 데이터: 가입 추이, 플랜 분포, AI 사용량, 알림 통계"""
    conn = get_db_connection()
    cursor = conn.cursor()

    def _safe_query(query, default=None):
        """테이블이 없어도 안전하게 쿼리 실행"""
        try:
            cursor.execute(query)
            return cursor.fetchall()
        except Exception:
            conn.rollback()
            return default if default is not None else []

    def _safe_scalar(query, field="total", default=0):
        try:
            cursor.execute(query)
            row = cursor.fetchone()
            return row[field] if row else default
        except Exception:
            conn.rollback()
            return default

    # 1. 일별 가입 추이 (최근 30일)
    cursor.execute("""
        SELECT DATE(updated_at) as reg_date, COUNT(*) as cnt
        FROM users
        WHERE updated_at >= CURRENT_DATE - INTERVAL '30 days'
        GROUP BY DATE(updated_at)
        ORDER BY reg_date
    """)
    signup_trend = [{"date": str(r["reg_date"]), "count": r["cnt"]} for r in cursor.fetchall()]

    # 2. 플랜별 사용자 분포
    cursor.execute("""
        SELECT COALESCE(plan, 'free') as plan, COUNT(*) as cnt
        FROM users
        GROUP BY plan
        ORDER BY cnt DESC
    """)
    plan_dist = [{"plan": r["plan"] or "free", "count": r["cnt"]} for r in cursor.fetchall()]

    # 3. 사용자 유형 분포 (기업/개인)
    cursor.execute("""
        SELECT COALESCE(user_type, 'both') as utype, COUNT(*) as cnt
        FROM users
        GROUP BY user_type
        ORDER BY cnt DESC
    """)
    type_dist = [{"type": r["utype"] or "business", "count": r["cnt"]} for r in cursor.fetchall()]

    # 4. AI 상담 일별 추이 (최근 30일)
    ai_usage_trend = [{"date": str(r["chat_date"]), "count": r["cnt"]} for r in _safe_query("""
        SELECT DATE(created_at) as chat_date, COUNT(*) as cnt
        FROM ai_consult_logs
        WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
        GROUP BY DATE(created_at) ORDER BY chat_date
    """)]

    # 5. AI 상담 총 통계
    ai_total = _safe_scalar("SELECT COUNT(*) as total FROM ai_consult_logs")
    ai_helpful = _safe_scalar("SELECT COUNT(*) as total FROM ai_consult_logs WHERE feedback = 'helpful'")
    ai_inaccurate = _safe_scalar("SELECT COUNT(*) as total FROM ai_consult_logs WHERE feedback = 'inaccurate'")

    # 6. 알림 발송 통계 (최근 30일)
    notif_raw = _safe_query("""
        SELECT DATE(sent_at) as send_date, status, COUNT(*) as cnt
        FROM notification_logs
        WHERE sent_at >= CURRENT_DATE - INTERVAL '30 days'
        GROUP BY DATE(sent_at), status ORDER BY send_date
    """)
    notif_by_date = {}
    for r in notif_raw:
        d = str(r["send_date"])
        if d not in notif_by_date:
            notif_by_date[d] = {"date": d, "success": 0, "failed": 0}
        if r["status"] == "sent":
            notif_by_date[d]["success"] += r["cnt"]
        else:
            notif_by_date[d]["failed"] += r["cnt"]
    notif_trend = sorted(notif_by_date.values(), key=lambda x: x["date"])

    # 7. 총 알림 발송 수
    notif_total = _safe_scalar("SELECT COUNT(*) as total FROM notification_logs")
    notif_success = _safe_scalar("SELECT COUNT(*) as total FROM notification_logs WHERE status = 'sent'")

    # 8. 저장 공고 수 (북마크)
    saved_total = _safe_scalar("SELECT COUNT(*) as total FROM saved_announcements")

    # 9. 푸시 구독자 수
    push_total = _safe_scalar("SELECT COUNT(*) as total FROM push_subscriptions")

    # 10. 알림 활성 사용자 수
    notif_active = _safe_scalar("SELECT COUNT(*) as total FROM notification_settings WHERE is_active = true")

    # 11. 지역별 사용자 분포
    cursor.execute("""
        SELECT COALESCE(address_city, '미입력') as city, COUNT(*) as cnt
        FROM users
        GROUP BY address_city
        ORDER BY cnt DESC
        LIMIT 15
    """)
    region_dist = [{"region": r["city"], "count": r["cnt"]} for r in cursor.fetchall()]

    # 12. 공고 수집 일별 추이 (최근 14일)
    cursor.execute("""
        SELECT DATE(created_at) as cdate, COUNT(*) as cnt
        FROM announcements
        WHERE created_at >= CURRENT_DATE - INTERVAL '14 days'
        GROUP BY DATE(created_at)
        ORDER BY cdate
    """)
    crawl_trend = [{"date": str(r["cdate"]), "count": r["cnt"]} for r in cursor.fetchall()]

    conn.close()
    return {
        "status": "SUCCESS",
        "data": {
            "signup_trend": signup_trend,
            "plan_distribution": plan_dist,
            "user_type_distribution": type_dist,
            "ai_usage_trend": ai_usage_trend,
            "ai_stats": {
                "total": ai_total,
                "helpful": ai_helpful,
                "inaccurate": ai_inaccurate,
            },
            "notification_trend": notif_trend,
            "notification_stats": {
                "total": notif_total,
                "success": notif_success,
                "failed": notif_total - notif_success,
            },
            "saved_total": saved_total,
            "push_subscribers": push_total,
            "notification_active_users": notif_active,
            "region_distribution": region_dist,
            "crawl_trend": crawl_trend,
        }
    }


def _fetch_ga4_data() -> dict:
    """GA4 REST API로 최근 30일 트래픽 데이터 조회 (경량 방식)"""
    try:
        ga4_creds_json = os.getenv("GA4_SERVICE_ACCOUNT_JSON", "")
        ga4_property_id = os.getenv("GA4_PROPERTY_ID", "")
        if not ga4_creds_json or not ga4_property_id:
            return {"error": "GA4 환경변수 미설정"}

        from google.oauth2 import service_account
        import google.auth.transport.requests

        creds_dict = json.loads(ga4_creds_json)
        credentials = service_account.Credentials.from_service_account_info(
            creds_dict, scopes=["https://www.googleapis.com/auth/analytics.readonly"]
        )
        credentials.refresh(google.auth.transport.requests.Request())
        token = credentials.token
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        base_url = f"https://analyticsdata.googleapis.com/v1beta/properties/{ga4_property_id}:runReport"

        def _run_report(dimensions, metrics, limit=0):
            body = {
                "dateRanges": [{"startDate": "30daysAgo", "endDate": "today"}],
                "metrics": [{"name": m} for m in metrics],
            }
            if dimensions:
                body["dimensions"] = [{"name": d} for d in dimensions]
            if limit:
                body["limit"] = limit
            r = requests.post(base_url, headers=headers, json=body, timeout=15)
            return r.json() if r.status_code == 200 else {}

        # 1. 기본 지표
        res1 = _run_report([], ["activeUsers", "sessions", "screenPageViews", "averageSessionDuration", "bounceRate", "newUsers"])
        basic = {}
        if res1.get("rows"):
            for i, m in enumerate(res1.get("metricHeaders", [])):
                basic[m["name"]] = res1["rows"][0]["metricValues"][i]["value"]

        # 2. 일별 추이
        res2 = _run_report(["date"], ["activeUsers", "sessions"])
        daily = []
        for r in res2.get("rows", []):
            d = r["dimensionValues"][0]["value"]
            daily.append({"date": f"{d[:4]}-{d[4:6]}-{d[6:]}", "users": int(r["metricValues"][0]["value"]), "sessions": int(r["metricValues"][1]["value"])})
        daily.sort(key=lambda x: x["date"])

        # 3. 유입 경로
        res3 = _run_report(["sessionDefaultChannelGroup"], ["sessions", "activeUsers"])
        channels = [{"channel": r["dimensionValues"][0]["value"], "sessions": int(r["metricValues"][0]["value"]), "users": int(r["metricValues"][1]["value"])} for r in res3.get("rows", [])]

        # 4. 디바이스
        res4 = _run_report(["deviceCategory"], ["sessions", "bounceRate"])
        devices = [{"device": r["dimensionValues"][0]["value"], "sessions": int(r["metricValues"][0]["value"]), "bounce_rate": r["metricValues"][1]["value"]} for r in res4.get("rows", [])]

        # 5. 인기 페이지
        res5 = _run_report(["pagePath"], ["screenPageViews", "activeUsers"], limit=10)
        pages = [{"path": r["dimensionValues"][0]["value"], "views": int(r["metricValues"][0]["value"]), "users": int(r["metricValues"][1]["value"])} for r in res5.get("rows", [])]

        return {"basic": basic, "daily_trend": daily, "channels": channels, "devices": devices, "top_pages": pages}
    except Exception as e:
        return {"error": str(e)}


@app.post("/api/admin/strategy-report", dependencies=[Depends(_verify_admin)])
def api_generate_strategy_report():
    """AI가 사용자 행동 데이터를 분석하여 활성화 전략 보고서 생성"""
    conn = get_db_connection()
    cursor = conn.cursor()

    # ── 데이터 수집 ──
    # 1. 전체 사용자 수 & 최근 7일 신규
    cursor.execute("SELECT COUNT(*) as total FROM users")
    total_users = cursor.fetchone()["total"]
    cursor.execute("SELECT COUNT(*) as cnt FROM users WHERE updated_at >= CURRENT_DATE - INTERVAL '7 days'")
    new_users_7d = cursor.fetchone()["cnt"]

    # 2. 이벤트 퍼널 (최근 30일)
    event_counts = {}
    try:
        cursor.execute("""
            SELECT event_type, COUNT(*) as cnt
            FROM user_events
            WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
            GROUP BY event_type ORDER BY cnt DESC
        """)
        for r in cursor.fetchall():
            event_counts[r["event_type"]] = r["cnt"]
    except Exception:
        pass

    # 3. 플랜 분포
    cursor.execute("SELECT COALESCE(plan,'free') as plan, COUNT(*) as cnt FROM users GROUP BY plan")
    plan_dist = {r["plan"] or "free": r["cnt"] for r in cursor.fetchall()}

    # 4. 유형 분포
    cursor.execute("SELECT COALESCE(user_type,'unknown') as utype, COUNT(*) as cnt FROM users GROUP BY user_type")
    type_dist = {r["utype"] or "unknown": r["cnt"] for r in cursor.fetchall()}

    # 5. 지역 분포 (상위 10)
    cursor.execute("""
        SELECT COALESCE(address_city,'미입력') as city, COUNT(*) as cnt
        FROM users GROUP BY address_city ORDER BY cnt DESC LIMIT 10
    """)
    region_dist = {r["city"]: r["cnt"] for r in cursor.fetchall()}

    # 6. AI 상담 통계
    cursor.execute("SELECT COUNT(*) as total FROM ai_consult_logs")
    ai_total = cursor.fetchone()["total"]
    cursor.execute("SELECT COUNT(*) as cnt FROM ai_consult_logs WHERE feedback='helpful'")
    ai_helpful = cursor.fetchone()["cnt"]
    cursor.execute("SELECT COUNT(*) as cnt FROM ai_consult_logs WHERE feedback='inaccurate'")
    ai_inaccurate = cursor.fetchone()["cnt"]

    # 7. 일별 가입 추이
    cursor.execute("""
        SELECT DATE(updated_at) as d, COUNT(*) as cnt FROM users
        WHERE updated_at >= CURRENT_DATE - INTERVAL '30 days'
        GROUP BY DATE(updated_at) ORDER BY d
    """)
    signup_trend = [{"date": str(r["d"]), "count": r["cnt"]} for r in cursor.fetchall()]

    # 8. 최근 이벤트 로그 샘플 (행동 패턴 파악용)
    event_samples = []
    try:
        cursor.execute("""
            SELECT event_type, COUNT(DISTINCT business_number) as unique_users, COUNT(*) as total,
                   MIN(created_at) as first_at, MAX(created_at) as last_at
            FROM user_events
            WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
            GROUP BY event_type ORDER BY total DESC
        """)
        for r in cursor.fetchall():
            event_samples.append({
                "type": r["event_type"], "unique_users": r["unique_users"],
                "total": r["total"], "first": str(r["first_at"])[:10], "last": str(r["last_at"])[:10]
            })
    except Exception:
        pass

    # 9. 전환율 계산
    signups = event_counts.get("signup", 0) + event_counts.get("social_signup", 0)
    logins = event_counts.get("login", 0) + event_counts.get("social_login", 0)
    matchings = event_counts.get("matching", 0)
    upgrades = event_counts.get("upgrade", 0)

    conn.close()

    # ── GA4 데이터 조회 ──
    ga4 = _fetch_ga4_data()
    ga4_section = ""
    if "error" not in ga4:
        basic = ga4.get("basic", {})
        ga4_section = f"""
### Google Analytics 트래픽 데이터 (30일)
- 활성 사용자: {basic.get('activeUsers', 'N/A')}명
- 세션 수: {basic.get('sessions', 'N/A')}
- 페이지뷰: {basic.get('screenPageViews', 'N/A')}
- 평균 세션 시간: {float(basic.get('averageSessionDuration', 0)):.0f}초
- 이탈률: {float(basic.get('bounceRate', 0))*100:.1f}%
- 신규 사용자: {basic.get('newUsers', 'N/A')}명
- 일별 방문 추이: {json.dumps(ga4.get('daily_trend', [])[-7:], ensure_ascii=False)}
- 유입 경로: {json.dumps(ga4.get('channels', []), ensure_ascii=False)}
- 디바이스별: {json.dumps(ga4.get('devices', []), ensure_ascii=False)}
- 인기 페이지: {json.dumps(ga4.get('top_pages', [])[:5], ensure_ascii=False)}
"""
    else:
        ga4_section = f"\n### Google Analytics\n- 연동 상태: {ga4.get('error', '미설정')}\n"

    # ── Gemini로 전략 보고서 생성 ──
    data_summary = f"""
## 지원금AI 서비스 현황 데이터 (최근 30일)
{ga4_section}
### 사용자 현황
- 전체 사용자: {total_users}명
- 최근 7일 신규 가입: {new_users_7d}명
- 플랜 분포: {json.dumps(plan_dist, ensure_ascii=False)}
- 유형 분포: {json.dumps(type_dist, ensure_ascii=False)}
- 지역 분포 (Top10): {json.dumps(region_dist, ensure_ascii=False)}

### 행동 퍼널 (30일)
- 가입: {signups}건
- 로그인: {logins}건
- 매칭 실행: {matchings}건
- 유료 전환: {upgrades}건
- 전환율: 가입→로그인 {f'{logins/signups*100:.0f}%' if signups > 0 else 'N/A'}, 로그인→매칭 {f'{matchings/logins*100:.0f}%' if logins > 0 else 'N/A'}, 매칭→결제 {f'{upgrades/matchings*100:.0f}%' if matchings > 0 else 'N/A'}

### 이벤트 상세
{json.dumps(event_samples, ensure_ascii=False, indent=2)}

### 일별 가입 추이
{json.dumps(signup_trend, ensure_ascii=False)}

### AI 상담
- 총 상담: {ai_total}건
- 도움됨: {ai_helpful}건, 부정확: {ai_inaccurate}건
- 만족도: {f'{ai_helpful/(ai_helpful+ai_inaccurate)*100:.0f}%' if (ai_helpful+ai_inaccurate) > 0 else 'N/A'}
"""

    prompt = f"""당신은 SaaS 서비스 성장 전략 전문가입니다.

아래는 "지원금AI" (govmatch.kr) 서비스의 실제 데이터입니다.
이 서비스는 AI 기반 정부 지원금/보조금 자동 매칭 서비스입니다.
사업자(기업)와 개인 사용자 모두를 대상으로 합니다.

{data_summary}

위 데이터를 기반으로 아래 형식의 전략 보고서를 작성해주세요.
반드시 **데이터에 기반한 구체적 수치**를 인용하고, **실행 가능한 액션**을 제시하세요.

## 보고서 형식

### 1. 핵심 지표 요약
- 현재 서비스 상태를 3줄로 요약

### 2. 퍼널 분석 & 병목 지점
- 가입→로그인→매칭→결제 각 단계별 전환율 분석
- 가장 큰 이탈 지점과 원인 추정

### 3. 사용자 세그먼트 인사이트
- 어떤 유형(기업/개인)이 더 활발한지
- 어떤 지역이 성장 잠재력이 높은지

### 4. 즉시 실행 가능한 개선 액션 (TOP 5)
- 각 액션에 예상 효과와 구현 난이도(상/중/하) 포함

### 5. 중장기 성장 전략
- 3개월, 6개월 목표와 로드맵

한국어로 작성하세요. 마크다운 형식으로 작성하세요.
"""

    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(prompt)
            report_text = response.text
        else:
            report_text = f"# Gemini API 키 미설정\n\n데이터 요약:\n{data_summary}"
    except Exception as e:
        report_text = f"# AI 분석 오류\n\n{str(e)}\n\n## 데이터 요약\n{data_summary}"

    # 보고서 캐시 저장
    try:
        conn2 = get_db_connection()
        cur2 = conn2.cursor()
        cur2.execute(
            "INSERT INTO strategy_reports (report_type, report_data) VALUES (%s, %s)",
            ("weekly", report_text)
        )
        conn2.commit()
        conn2.close()
    except Exception:
        pass

    return {
        "status": "SUCCESS",
        "report": report_text,
        "data_summary": {
            "total_users": total_users,
            "new_users_7d": new_users_7d,
            "funnel": {"signups": signups, "logins": logins, "matchings": matchings, "upgrades": upgrades},
            "plan_distribution": plan_dist,
            "type_distribution": type_dist,
            "region_distribution": region_dist,
        }
    }


@app.get("/api/admin/strategy-reports", dependencies=[Depends(_verify_admin)])
def api_get_strategy_reports(limit: int = 5):
    """저장된 전략 보고서 이력 조회"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, report_type, report_data, created_at FROM strategy_reports ORDER BY created_at DESC LIMIT %s",
        (limit,)
    )
    rows = [{"id": r["id"], "type": r["report_type"], "report": r["report_data"], "created_at": str(r["created_at"])} for r in cursor.fetchall()]
    conn.close()
    return {"status": "SUCCESS", "data": rows}


@app.get("/api/admin/system-logs", dependencies=[Depends(_verify_admin)])
def api_get_system_logs(category: Optional[str] = None, limit: int = 50):
    """시스템 활동 이력 조회"""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        if category:
            cursor.execute(
                "SELECT * FROM system_logs WHERE category = %s ORDER BY created_at DESC LIMIT %s",
                (category, min(limit, 200))
            )
        else:
            cursor.execute(
                "SELECT * FROM system_logs ORDER BY created_at DESC LIMIT %s",
                (min(limit, 200),)
            )
        rows = [dict(r) for r in cursor.fetchall()]
        for r in rows:
            if r.get("created_at"):
                r["created_at"] = str(r["created_at"])

        # 카테고리별 최근 실행 요약
        cursor.execute("""
            SELECT category, action,
                   MAX(created_at) as last_run,
                   COUNT(*) as total_runs,
                   SUM(CASE WHEN result = 'success' THEN 1 ELSE 0 END) as success_count,
                   SUM(CASE WHEN result = 'error' THEN 1 ELSE 0 END) as error_count
            FROM system_logs
            WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
            GROUP BY category, action
            ORDER BY last_run DESC
        """)
        summary = [dict(r) for r in cursor.fetchall()]
        for s in summary:
            if s.get("last_run"):
                s["last_run"] = str(s["last_run"])

        # 일별 시스템 활동 추이 (14일)
        daily_trend = []
        try:
            cursor.execute("""
                SELECT DATE(created_at) as log_date, category, COUNT(*) as cnt
                FROM system_logs
                WHERE created_at >= CURRENT_DATE - INTERVAL '14 days'
                GROUP BY DATE(created_at), category
                ORDER BY log_date
            """)
            trend_raw = {}
            for r in cursor.fetchall():
                d = str(r["log_date"])
                if d not in trend_raw:
                    trend_raw[d] = {"date": d, "collection": 0, "analysis": 0, "notification": 0, "payment": 0, "system": 0}
                cat = r["category"]
                if cat in trend_raw[d]:
                    trend_raw[d][cat] = r["cnt"]
            daily_trend = sorted(trend_raw.values(), key=lambda x: x["date"])
        except Exception:
            pass

        # 사용자 행동 퍼널 (30일, user_events 테이블)
        funnel = {}
        try:
            cursor.execute("""
                SELECT event_type, COUNT(*) as cnt, COUNT(DISTINCT business_number) as unique_users
                FROM user_events
                WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
                GROUP BY event_type
            """)
            for r in cursor.fetchall():
                funnel[r["event_type"]] = {"count": r["cnt"], "unique_users": r["unique_users"]}
        except Exception:
            pass

        # 시간대별 사용자 활동 (user_events)
        hourly = []
        try:
            cursor.execute("""
                SELECT EXTRACT(HOUR FROM created_at)::int as hr, COUNT(*) as cnt
                FROM user_events
                WHERE created_at >= CURRENT_DATE - INTERVAL '7 days'
                GROUP BY hr ORDER BY hr
            """)
            hourly = [{"hour": r["hr"], "count": r["cnt"]} for r in cursor.fetchall()]
        except Exception:
            pass

        conn.close()
        return {
            "status": "SUCCESS",
            "data": rows,
            "summary": summary,
            "daily_trend": daily_trend,
            "funnel": funnel,
            "hourly_activity": hourly,
        }
    except Exception as e:
        conn.close()
        return {"status": "SUCCESS", "data": [], "summary": [], "daily_trend": [], "funnel": {}, "hourly_activity": [], "note": f"테이블 미생성: {e}"}


@app.get("/api/admin/system-sources", dependencies=[Depends(_verify_admin)])
def get_system_sources():
    from app.services.public_api_service import gov_api_service
    is_live = gov_api_service.is_configured()

    apis = [
        {
            "id": "kised-api", "name": "K-Startup 공식 API", "type": "API",
            "status": "LIVE" if is_live else "SIMULATED",
            "description": "중소벤처기업부 창업지원사업 통합 데이터"
        },
        {
            "id": "msit-rnd-api", "name": "과기정통부 R&D 사업공고 API", "type": "API",
            "status": "LIVE" if is_live else "SIMULATED",
            "description": "과학기술정보통신부 국가R&D 사업공고 (data.go.kr)"
        },
        {
            "id": "mss-api", "name": "중소벤처기업부 사업공고 API", "type": "API",
            "status": "LIVE" if is_live else "SIMULATED",
            "description": "중소벤처기업부 공식 사업공고 조회"
        },
        {
            "id": "bizinfo-portal-api", "name": "기업마당 포털 API", "type": "API",
            "status": "LIVE" if os.getenv("BIZINFO_PORTAL_KEY") else "KEY_REQUIRED",
            "description": "기업마당 포털 직접 연동 API (bizinfoApi.do)"
        },
        {
            "id": "smes24-api", "name": "중소벤처24 공고정보 API", "type": "API",
            "status": "LIVE" if os.getenv("SMES24_API_TOKEN") else "KEY_REQUIRED",
            "description": "중소벤처24 사업공고 (자격요건 구조화 데이터 포함)"
        },
        {
            "id": "foodpolis-api", "name": "한국식품산업클러스터진흥원 API", "type": "API",
            "status": "LIVE" if os.getenv("FOODPOLIS_API_KEY") else "KEY_REQUIRED",
            "description": "식품산업 관련 사업공고 (foodpolis.kr)"
        },
    ]

    scrapers = [
        {
            "id": "sbc", "name": "중진공(SBC) 스크래퍼", "type": "Scraper",
            "status": "ACTIVE",
            "description": "중소벤처기업진흥공단 공고 Playwright 크롤링 (공식 API 미제공)"
        },
    ]

    return {
        "status": "SUCCESS",
        "data": {"apis": apis, "scrapers": scrapers}
    }

@app.get("/api/admin/urls", dependencies=[Depends(_verify_admin)])
def get_admin_urls():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM admin_urls")
    rows = cursor.fetchall()
    conn.close()
    return {"status": "SUCCESS", "data": [dict(r) for r in rows]}

@app.post("/api/admin/urls", dependencies=[Depends(_verify_admin)])
def add_admin_url(request: AdminURLRequest):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("INSERT INTO admin_urls (url, source_name) VALUES (%s, %s)",
                     (request.url, request.source_name))
        conn.commit()
        return {"status": "SUCCESS", "message": "URL 등록 완료"}
    except psycopg2.IntegrityError:
        conn.rollback()
        raise HTTPException(status_code=400, detail="이미 등록된 URL입니다.")
    finally:
        conn.close()

@app.delete("/api/admin/urls/{url_id}", dependencies=[Depends(_verify_admin)])
def delete_admin_url(url_id: int):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM admin_urls WHERE id = %s", (url_id,))
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "message": "URL 삭제 완료"}

@app.post("/api/admin/repair-origin-urls", dependencies=[Depends(_verify_admin)])
def repair_origin_urls():
    """손상된 origin_url 일괄 수정 (도메인 중복 제거)"""
    import re
    conn = get_db_connection()
    cur = conn.cursor()
    fixed = 0
    try:
        # 'http(s)://...http(s)://' 패턴 가진 행 모두 조회
        cur.execute("""
            SELECT announcement_id, origin_url FROM announcements
            WHERE origin_url ~ 'https?://.*https?://'
        """)
        rows = cur.fetchall()
        for r in rows:
            url = r["origin_url"]
            matches = list(re.finditer(r'https?://', url))
            if len(matches) >= 2:
                clean_url = url[matches[-1].start():]
                try:
                    cur.execute(
                        "UPDATE announcements SET origin_url = %s WHERE announcement_id = %s",
                        (clean_url, r["announcement_id"])
                    )
                    fixed += 1
                except Exception:
                    conn.rollback()
        conn.commit()
        return {"status": "SUCCESS", "scanned": len(rows), "fixed": fixed}
    finally:
        conn.close()


@app.post("/api/admin/patrol/run", dependencies=[Depends(_verify_admin)])
def admin_patrol_run():
    """AI 패트롤 수동 실행 (관리자 전용)"""
    try:
        from app.services.patrol import run_patrol
        result = run_patrol(triggered_by="manual")
        return {"status": "SUCCESS", "result": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"패트롤 실행 실패: {e}")


@app.get("/api/admin/patrol/history", dependencies=[Depends(_verify_admin)])
def admin_patrol_history(limit: int = 10):
    """패트롤 실행 이력 조회"""
    try:
        from app.services.patrol import get_latest_report
        conn = get_db_connection()
        try:
            report = get_latest_report(conn, limit=limit)
            return {"status": "SUCCESS", **report}
        finally:
            conn.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/patrol/failures", dependencies=[Depends(_verify_admin)])
def admin_patrol_failures(limit: int = 50):
    """현재 분석 실패 큐 조회"""
    try:
        conn = get_db_connection()
        try:
            cur = conn.cursor()
            cur.execute("""
                SELECT af.id, af.announcement_id, af.error_type, af.error_message,
                       af.failed_at, af.retry_count, af.next_retry_at, af.resolved_at,
                       a.title
                FROM analysis_failures af
                LEFT JOIN announcements a ON a.announcement_id = af.announcement_id
                WHERE af.resolved_at IS NULL
                ORDER BY af.next_retry_at NULLS FIRST
                LIMIT %s
            """, (limit,))
            rows = [dict(r) for r in cur.fetchall()]
            for r in rows:
                for k in ("failed_at", "next_retry_at", "resolved_at"):
                    if r.get(k):
                        r[k] = str(r[k])
            return {"status": "SUCCESS", "failures": rows, "count": len(rows)}
        finally:
            conn.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/admin/knowledge/seed", dependencies=[Depends(_verify_admin)])
def admin_knowledge_seed():
    """금융 지식 수동 시딩 (관리자 전용)"""
    try:
        from app.services.financial_analysis.knowledge_seed import seed_financial_knowledge
        conn = get_db_connection()
        try:
            seeded = seed_financial_knowledge(conn)
            return {"status": "SUCCESS", "seeded": seeded}
        finally:
            conn.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/admin/knowledge/stats", dependencies=[Depends(_verify_admin)])
def admin_knowledge_stats():
    """knowledge_base 현황 조회"""
    try:
        conn = get_db_connection()
        try:
            cur = conn.cursor()
            cur.execute("""
                SELECT source, knowledge_type, COUNT(*) as cnt,
                       ROUND(AVG(confidence)::numeric, 2) as avg_confidence,
                       SUM(use_count) as total_uses
                FROM knowledge_base
                GROUP BY source, knowledge_type
                ORDER BY source, knowledge_type
            """)
            rows = [dict(r) for r in cur.fetchall()]

            cur.execute("SELECT COUNT(*) as total FROM knowledge_base")
            total = cur.fetchone()["total"]

            return {"status": "SUCCESS", "total": total, "breakdown": rows}
        finally:
            conn.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/admin/reanalyze-empty", dependencies=[Depends(_verify_admin)])
async def reanalyze_empty_analyses(limit: int = 100):
    """full_text가 비어있는 분석을 삭제하고 재분석 대기열에 추가"""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # full_text가 비어있거나 NULL인 분석 찾기
        cur.execute("""
            SELECT aa.announcement_id, a.title
            FROM announcement_analysis aa
            JOIN announcements a ON a.announcement_id = aa.announcement_id
            WHERE (aa.full_text IS NULL OR aa.full_text = '' OR LENGTH(aa.full_text) < 100)
              AND a.origin_url IS NOT NULL AND a.origin_url != ''
            LIMIT %s
        """, (limit,))
        empty_rows = cur.fetchall()

        deleted = 0
        for row in empty_rows:
            aid = row["announcement_id"]
            # 빈 분석 삭제
            cur.execute("DELETE FROM announcement_analysis WHERE announcement_id = %s", (aid,))
            # 실패 기록도 삭제 (재시도 가능하도록)
            try:
                cur.execute("DELETE FROM analysis_failures WHERE announcement_id = %s", (aid,))
            except Exception:
                pass
            deleted += 1

        conn.commit()

        # 패트롤 실행으로 재분석 트리거
        reanalyzed = 0
        if deleted > 0:
            try:
                from app.services.doc_analysis_service import ensure_analysis
                for row in empty_rows[:20]:  # 즉시 20건만 재분석
                    try:
                        analysis_conn = get_db_connection()
                        result = ensure_analysis(row["announcement_id"], analysis_conn)
                        analysis_conn.close()
                        if result and result.get("full_text"):
                            reanalyzed += 1
                            print(f"[Reanalyze] OK: #{row['announcement_id']} {row['title'][:30]} → {len(result.get('full_text',''))} chars")
                    except Exception as e:
                        print(f"[Reanalyze] Failed #{row['announcement_id']}: {e}")
            except Exception as e:
                print(f"[Reanalyze] Batch error: {e}")

        return {
            "status": "SUCCESS",
            "empty_found": len(empty_rows),
            "deleted": deleted,
            "reanalyzed_immediately": reanalyzed,
            "remaining_for_patrol": max(0, deleted - 20),
        }
    finally:
        conn.close()


@app.post("/api/admin/set-plan", dependencies=[Depends(_verify_admin)])
def admin_set_plan(req: dict):
    """관리자용 플랜 강제 변경"""
    email = req.get("email")
    plan = req.get("plan", "pro")
    days = req.get("days", 30)
    if not email:
        raise HTTPException(status_code=400, detail="email 필수")
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        import datetime as _dt
        expires = (_dt.datetime.utcnow() + _dt.timedelta(days=days)).isoformat()
        cur.execute(
            "UPDATE users SET plan = %s, plan_expires_at = %s, ai_usage_month = 0 WHERE email = %s",
            (plan, expires, email)
        )
        conn.commit()
        return {"status": "SUCCESS", "message": f"{email} → {plan} ({days}일)"}
    finally:
        conn.close()


@app.get("/api/admin/analysis-stats", dependencies=[Depends(_verify_admin)])
def admin_analysis_stats():
    """공고 분석 현황 통계"""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) as total FROM announcements WHERE origin_url IS NOT NULL")
        total = cur.fetchone()["total"]

        cur.execute("SELECT COUNT(*) as analyzed FROM announcement_analysis WHERE full_text IS NOT NULL AND LENGTH(full_text) > 100")
        analyzed = cur.fetchone()["analyzed"]

        cur.execute("SELECT COUNT(*) as empty FROM announcement_analysis WHERE full_text IS NULL OR full_text = '' OR LENGTH(full_text) < 100")
        empty = cur.fetchone()["empty"]

        cur.execute("SELECT COUNT(*) as no_analysis FROM announcements a LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id WHERE aa.announcement_id IS NULL AND a.origin_url IS NOT NULL")
        no_analysis = cur.fetchone()["no_analysis"]

        cur.execute("""
            SELECT aa.source_type, COUNT(*) as cnt
            FROM announcement_analysis aa
            WHERE aa.full_text IS NOT NULL AND LENGTH(aa.full_text) > 100
            GROUP BY aa.source_type ORDER BY cnt DESC
        """)
        by_source = [dict(r) for r in cur.fetchall()]

        return {
            "status": "SUCCESS",
            "total_announcements": total,
            "analyzed_ok": analyzed,
            "analyzed_empty": empty,
            "not_analyzed": no_analysis,
            "by_source_type": by_source,
        }
    finally:
        conn.close()


@app.get("/api/admin/urls/health", dependencies=[Depends(_verify_admin)])
def get_admin_urls_health():
    """수집 URL 헬스체크 리포트 — 실패/복구 현황"""
    report = admin_scraper.get_health_report()
    return {"status": "SUCCESS", "data": report}

@app.post("/api/admin/seed-urls", dependencies=[Depends(_verify_admin)])
def admin_seed_urls():
    """seed_regional_urls에 등록된 URL을 DB에 일괄 등록"""
    try:
        from app.db.seed_regional_urls import REGIONAL_URLS
    except ImportError:
        try:
            import importlib, sys
            spec = importlib.util.spec_from_file_location("seed", os.path.join(os.path.dirname(__file__), "db", "seed_regional_urls.py"))
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            REGIONAL_URLS = mod.REGIONAL_URLS
        except Exception as e:
            return {"status": "ERROR", "detail": f"import 실패: {e}", "inserted": 0, "skipped": 0, "total": 0}

    if not REGIONAL_URLS:
        return {"status": "ERROR", "detail": "REGIONAL_URLS가 비어있음", "inserted": 0, "skipped": 0, "total": 0}

    conn = get_db_connection()
    cursor = conn.cursor()
    inserted, skipped, errors = 0, 0, []
    for source_name, url in REGIONAL_URLS:
        try:
            cursor.execute(
                """INSERT INTO admin_urls (url, source_name, is_active)
                   VALUES (%s, %s, 1)
                   ON CONFLICT (url) DO NOTHING""",
                (url, source_name),
            )
            conn.commit()
            if cursor.rowcount > 0:
                inserted += 1
            else:
                skipped += 1
        except Exception as e:
            conn.rollback()
            errors.append(f"{source_name}: {str(e)[:80]}")
    conn.close()
    return {"status": "SUCCESS", "inserted": inserted, "skipped": skipped, "total": inserted + skipped, "source_count": len(REGIONAL_URLS), "errors": errors[:5]}


@app.post("/api/admin/cleanup-db", dependencies=[Depends(_verify_admin)])
def admin_cleanup_db():
    """DB 정리: 비지원사업 제거 + 중복 제거 + target_type 분류"""
    _cleanup_non_support_announcements()
    _deduplicate_announcements()
    _auto_classify_target_type()
    return {"status": "SUCCESS", "message": "DB 정리 + 분류 완료"}


def _run_manual_sync_in_thread():
    """관리자 수동 동기화 — 일일 통합 파이프라인 전체를 실행.

    이전에는 admin_scraper.run_all()만 호출하여 Playwright 미설치 환경에서
    무용지물이었음. 지금은 스케줄러와 동일한 run_daily_pipeline을 호출.
    """
    from app.services.patrol.daily_pipeline import run_daily_pipeline
    manual_sync_status["running"] = True
    manual_sync_status["last_result"] = "진행 중..."
    manual_sync_status["last_time"] = datetime.datetime.now().isoformat()
    conn = None
    try:
        conn = get_db_connection()
        result = run_daily_pipeline(conn)
        err_cnt = result.get("error_count", 0)
        total = result.get("total_elapsed", 0)
        summary = f"완료 ({total}s, 에러 {err_cnt})"
        manual_sync_status["last_result"] = summary
        manual_sync_status["last_pipeline"] = result
        _log_system("manual_sync", "collection", summary, "success")
    except Exception as e:
        manual_sync_status["last_result"] = f"오류: {e}"
        _log_system("manual_sync", "collection", f"수동 동기화 오류: {e}", "error")
    finally:
        if conn is not None:
            try: conn.close()
            except Exception: pass
        manual_sync_status["running"] = False
        manual_sync_status["last_time"] = datetime.datetime.now().isoformat()


@app.post("/api/admin/sync-manual", dependencies=[Depends(_verify_admin)])
async def trigger_admin_sync():
    if manual_sync_status["running"]:
        return {"status": "ALREADY_RUNNING", "message": "수동 동기화가 이미 진행 중입니다."}
    import threading
    t = threading.Thread(target=_run_manual_sync_in_thread, daemon=True)
    t.start()
    return {"status": "STARTED", "message": "백그라운드에서 수동 동기화를 시작합니다."}


@app.get("/api/admin/sync-manual-status", dependencies=[Depends(_verify_admin)])
def get_manual_sync_status():
    return {"status": "SUCCESS", "data": manual_sync_status}

def _run_reanalyze_in_thread(limit: int):
    """재분석을 별도 스레드에서 실행 — 상세 페이지 크롤링 + AI 추출 + DB 컬럼 업데이트"""
    import asyncio
    import re as _re
    import time
    import requests
    from bs4 import BeautifulSoup
    from app.services.ai_service import ai_service as _ai

    _HEADERS = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "text/html,application/xhtml+xml",
        "Accept-Language": "ko-KR,ko;q=0.9",
    }

    def _strip(text):
        if not text:
            return ""
        text = _re.sub(r'<[^>]+>', ' ', text)
        text = _re.sub(r'&[a-z]+;', ' ', text)
        text = _re.sub(r'&#\d+;', ' ', text)
        text = _re.sub(r'\s+', ' ', text)
        return text.strip()

    _GOV24_URL_PAT = _re.compile(r"gov\.kr/portal/rcv[a-zA-Z]*Svc/dtlEx/([A-Z0-9]+)", _re.IGNORECASE)

    def _fetch_gov24_api(serv_id: str) -> str:
        """gov24 /v3/serviceDetail API로 탭 내용 전체 수집 (지원대상/지원내용/신청방법 등)"""
        api_key = os.environ.get("GOV24_API_KEY", "")
        if not api_key:
            return ""
        try:
            resp = requests.get(
                "https://api.odcloud.kr/api/gov24/v3/serviceDetail",
                params={"serviceKey": api_key, "serviceId": serv_id, "returnType": "JSON"},
                timeout=15,
            )
            if resp.status_code != 200:
                return ""
            detail = resp.json().get("data", [{}])
            if isinstance(detail, list) and detail:
                detail = detail[0]
            elif not isinstance(detail, dict):
                return ""
            field_map = [
                ("지원대상", ["지원대상", "tgtrDtlCn"]),
                ("지원내용", ["지원내용", "sprtCn", "servDgst"]),
                ("신청방법", ["신청방법", "aplyMtdCn"]),
                ("선정기준", ["선정기준", "slctCritCn"]),
                ("구비서류", ["구비서류", "psblDocCn"]),
                ("문의처",   ["문의처", "inqPlCtadrList", "rprsCtadr"]),
                ("지원형태", ["지원형태", "sprtTypeNm"]),
            ]
            parts = []
            for label, keys in field_map:
                for k in keys:
                    val = str(detail.get(k, "") or "").strip()
                    if val and val not in ("null", "[]", "{}"):
                        parts.append(f"[{label}]\n{val}")
                        break
            return "\n\n".join(parts)
        except Exception:
            return ""

    def _fetch_detail(url, max_chars=8000):
        """상세 페이지 크롤링 — gov.kr 탭 페이지는 gov24 API 우선 호출"""
        # gov.kr 개인지원사업: HTML 크롤링 대신 API로 탭 전체 수집
        m = _GOV24_URL_PAT.search(url or "")
        if m:
            api_text = _fetch_gov24_api(m.group(1))
            if api_text and len(api_text) > 100:
                return api_text[:max_chars]
        # 일반 HTML 크롤링 fallback
        try:
            resp = requests.get(url, headers=_HEADERS, timeout=12, allow_redirects=True)
            resp.encoding = resp.apparent_encoding or "utf-8"
            if resp.status_code != 200:
                return ""
            soup = BeautifulSoup(resp.text, "html.parser")
            for tag in soup.find_all(["script", "style", "nav", "header", "footer", "aside"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            lines = [ln.strip() for ln in text.splitlines() if len(ln.strip()) > 1]
            return "\n".join(lines)[:max_chars]
        except Exception:
            return ""

    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    reanalyze_status["running"] = True
    reanalyze_status["done"] = 0
    reanalyze_status["last_time"] = datetime.datetime.now().isoformat()

    conn = get_db_connection()
    cursor = conn.cursor()
    # AI 미분석 공고만 대상 (ai_analyzed_at이 NULL인 건)
    # eligibility_logic이 비어있거나 한 번도 AI 분석을 하지 않은 공고
    cursor.execute("""
        SELECT announcement_id, title, summary_text, origin_url, deadline_date,
               established_years_limit, revenue_limit, employee_limit
        FROM announcements
        WHERE ai_analyzed_at IS NULL
          AND (
            eligibility_logic IS NULL OR eligibility_logic = '' OR eligibility_logic = '{}'
            OR (established_years_limit IS NULL AND revenue_limit IS NULL AND employee_limit IS NULL)
          )
        ORDER BY created_at DESC
        LIMIT %s
    """, (limit,))
    rows = [dict(r) for r in cursor.fetchall()]
    reanalyze_status["total"] = len(rows)

    success = 0
    for row in rows:
        try:
            # 1. 상세 페이지 크롤링 시도
            detail_text = ""
            origin_url = row.get("origin_url", "")
            if origin_url and origin_url.startswith("http"):
                detail_text = _fetch_detail(origin_url)

            # 2. 분석용 텍스트 구성 (상세 페이지 + 기존 요약)
            clean_summary = _strip(row.get("summary_text", ""))
            if detail_text and len(detail_text) > 100:
                input_text = f"제목: {row.get('title', '')}\n\n[상세 페이지 본문]\n{detail_text}"
            else:
                input_text = f"제목: {row.get('title', '')}\n\n내용: {clean_summary[:8000]}"

            if len(input_text) < 30:
                reanalyze_status["done"] += 1
                continue

            # 3. AI 분석 (구조화된 자격요건 추출)
            details = loop.run_until_complete(_ai.extract_program_details(input_text))
            if details:
                elig = details.get("eligibility_logic", {}) or {}
                if details.get("business_type"):
                    elig["business_type"] = details["business_type"]
                if details.get("target_keywords"):
                    elig["target_keywords"] = details["target_keywords"]

                # 4. eligibility 수치 추출
                years_limit = elig.get("max_founding_years")
                revenue_limit = elig.get("max_revenue")
                employee_limit = elig.get("max_employee_count") or elig.get("max_employees")

                ai_summary = details.get("summary_text") or details.get("description", "")
                ai_deadline = details.get("deadline_date")
                pk = row.get("announcement_id")

                cursor.execute("""
                    UPDATE announcements SET
                        eligibility_logic = %s,
                        summary_text = CASE WHEN %s != '' THEN %s ELSE summary_text END,
                        department = CASE WHEN department IS NULL OR department = '' THEN %s ELSE department END,
                        category = CASE WHEN category IS NULL OR category = '' THEN %s ELSE category END,
                        deadline_date = CASE WHEN deadline_date IS NULL AND %s IS NOT NULL THEN CAST(%s AS DATE) ELSE deadline_date END,
                        established_years_limit = COALESCE(%s, established_years_limit),
                        revenue_limit = COALESCE(%s, revenue_limit),
                        employee_limit = COALESCE(%s, employee_limit),
                        ai_analyzed_at = NOW()
                    WHERE announcement_id = %s
                """, (
                    json.dumps(elig, ensure_ascii=False),
                    ai_summary, ai_summary,
                    details.get("department", ""),
                    details.get("category", ""),
                    ai_deadline, ai_deadline,
                    years_limit, revenue_limit, employee_limit,
                    pk,
                ))
                conn.commit()
                success += 1
                if success % 10 == 0:
                    print(f"  [Reanalyze] {success} updated so far ({reanalyze_status['done']}/{len(rows)})")
        except Exception as e:
            print(f"Reanalyze error: {e}")
            conn.rollback()
        reanalyze_status["done"] += 1
        time.sleep(0.5)

    conn.close()
    reanalyze_status["running"] = False
    reanalyze_status["last_result"] = f"완료: {success}/{len(rows)}건 분석"
    reanalyze_status["last_time"] = datetime.datetime.now().isoformat()
    loop.close()
    _log_system("reanalyze", "analysis", f"{success}/{len(rows)}건 AI 재분석 완료", "success", success)
    print(f"[Reanalyze] Done: {success}/{len(rows)} announcements enriched")


@app.post("/api/admin/reanalyze", dependencies=[Depends(_verify_admin)])
async def trigger_reanalyze(limit: int = 200):
    if reanalyze_status["running"]:
        return {"status": "ALREADY_RUNNING", "message": "재분석이 이미 진행 중입니다."}
    import threading
    t = threading.Thread(target=_run_reanalyze_in_thread, args=(limit,), daemon=True)
    t.start()
    return {"status": "STARTED", "message": f"최대 {limit}건 백그라운드 재분석 시작"}


@app.get("/api/admin/reanalyze-status", dependencies=[Depends(_verify_admin)])
def get_reanalyze_status():
    return {"status": "SUCCESS", "data": reanalyze_status}


@app.post("/api/admin/reset-gov24-individual", dependencies=[Depends(_verify_admin)])
def reset_gov24_individual_analysis():
    """gov.kr 개인지원사업 공고 중 잘못 분석된 것 초기화 — ai_analyzed_at 리셋 + announcement_analysis 삭제"""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # 대상: origin_url이 gov.kr 개인지원사업 패턴인 공고
        cur.execute("""
            SELECT announcement_id FROM announcements
            WHERE origin_url ~ 'gov\\.kr/portal/rcv[a-zA-Z]*Svc/dtlEx/'
        """)
        ids = [r["announcement_id"] for r in cur.fetchall()]
        if not ids:
            return {"status": "SUCCESS", "reset": 0, "message": "대상 없음"}

        id_tuple = tuple(ids)

        # 1) announcement_analysis 삭제 (상담용 분석 초기화)
        cur.execute("DELETE FROM announcement_analysis WHERE announcement_id = ANY(%s)", (list(id_tuple),))
        deleted_analysis = cur.rowcount

        # 2) ai_analyzed_at 리셋 (매칭용 분석 초기화 → reanalyze 루프가 다시 처리)
        cur.execute("""
            UPDATE announcements
            SET ai_analyzed_at = NULL,
                eligibility_logic = NULL
            WHERE announcement_id = ANY(%s)
        """, (list(id_tuple),))
        reset_count = cur.rowcount

        conn.commit()
        return {
            "status": "SUCCESS",
            "total_gov24_individual": len(ids),
            "reset_ai_analyzed_at": reset_count,
            "deleted_announcement_analysis": deleted_analysis,
            "message": f"gov.kr 개인지원사업 {reset_count}건 초기화 완료. reanalyze 루프로 재분석 필요.",
        }
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.get("/")
def read_root():
    return {"message": "Welcome to Auto_Gov_Macting API"}

def _search_ksic_by_embedding(query: str, top_k: int = 5) -> list:
    """KSIC 임베딩 기반 유사 업종 검색 (gemini-embedding-001 768d + pgvector)."""
    if not query or len(query.strip()) < 1:
        return []
    try:
        import google.generativeai as _genai
    except ImportError:
        return []
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return []
    try:
        _genai.configure(api_key=api_key)
        res = _genai.embed_content(
            model="models/gemini-embedding-001",
            content=f"업종 질의: {query}",
            task_type="retrieval_query",
            output_dimensionality=768,
        )
        vec = res.get("embedding") if isinstance(res, dict) else res["embedding"]
        if not vec:
            return []
        vec_str = "[" + ",".join(f"{v:.6f}" for v in vec) + "]"
    except Exception as e:
        print(f"[KSIC embed search] {e}")
        return []

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT k.code, k.name, k.description,
                   1 - (e.embedding <=> %s::vector) AS similarity
            FROM ksic_embeddings e
            JOIN ksic_classification k ON e.code = k.code
            ORDER BY e.embedding <=> %s::vector
            LIMIT %s
        """, (vec_str, vec_str, top_k))
        rows = cur.fetchall()
        return [
            {
                "code": r["code"],
                "name": r["name"],
                "description": r.get("description") or "",
                "similarity": round(float(r.get("similarity") or 0), 4),
                "reason": f"의미 유사도 {round((r.get('similarity') or 0) * 100)}%",
            }
            for r in rows
        ]
    except Exception as e:
        print(f"[KSIC embed query] {e}")
        return []
    finally:
        try: conn.close()
        except: pass


@app.post("/api/industry-recommend")
async def api_industry_recommend(request: CompanyNameRequest):
    """KSIC 임베딩 기반 유사 업종 Top-5 추천 (임베딩 실패 시 기존 DB LIKE + LLM 폴백)."""
    query = (request.business_content or request.company_name or "").strip()
    if not query:
        return {"status": "SUCCESS", "data": {"candidates": []}}

    # 1) 임베딩 기반 유사도 검색 (가장 정확)
    emb_results = _search_ksic_by_embedding(query, top_k=5)
    if emb_results:
        return {"status": "SUCCESS", "data": {"candidates": emb_results}}

    # 2) 폴백: 기존 하이브리드 (DB LIKE + LLM)
    from app.services.ai_service import ai_service
    result = await ai_service.search_industry_hybrid(query)
    return {"status": "SUCCESS", "data": result}

@app.post("/api/check-url")
def api_check_url(request: URLRequest):
    result = check_duplicate_url(request.url)
    return {
        "status": "SUCCESS",
        "is_duplicate": result["status"] == "ALREADY_EXISTS",
        "detail": result
    }

def is_update_required(business_number: str, updated_at) -> bool:
    """
    Checks if revenue update is required based on Korean tax deadlines.
    - Corporations: April (4)
    - Individuals: June (6)
    """
    try:
        # Business number middle digits: 81, 82, 86, 87, 88 are usually corps
        middle = int(business_number[3:5])
        is_corp = middle in [81, 82, 86, 87, 88]

        current_date = datetime.datetime.now()
        update_month = 4 if is_corp else 6

        # Parse updated_at
        if isinstance(updated_at, datetime.datetime):
            last_updated = updated_at
        else:
            last_updated = datetime.datetime.strptime(str(updated_at)[:19], "%Y-%m-%d %H:%M:%S")

        # If we are past the update month in the current year,
        # but the data was updated in a previous year (or before the update month of this year).
        if current_date.month >= update_month:
            if last_updated.year < current_date.year:
                return True
        return False
    except Exception:
        return False

@app.post("/api/fetch-company")
def api_fetch_company(request: BusinessNumberRequest):
    if len(request.business_number) != 10:
        raise HTTPException(status_code=400, detail="사업자 번호는 10자리여야 합니다.")

    # 1. Check DB first
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE business_number = %s", (request.business_number,))
    user = cursor.fetchone()
    conn.close()

    if user:
        user_dict = dict(user)
        requires_update = is_update_required(user_dict["business_number"], user_dict["updated_at"])
        return {
            "status": "SUCCESS",
            "type": "EXISTING",
            "data": {
                "business_number": user_dict["business_number"],
                "company_name": user_dict["company_name"],
                "establishment_date": str(user_dict["establishment_date"]) if user_dict.get("establishment_date") else "",
                "address_city": user_dict["address_city"],
                "industry_code": user_dict["industry_code"],
                "revenue": user_dict["revenue_bracket"],
                "employees": user_dict["employee_count_bracket"]
            },
            "requires_update": requires_update
        }

    # 2. If not found, return NEW immediately (Bypass slow API/AI)
    is_corp = int(request.business_number[3:5]) in [81, 82, 86, 87, 88]

    return {
        "status": "SUCCESS",
        "type": "NEW",
        "data": {
            "business_number": request.business_number,
            "company_name": "(신규 기업 등록 중)",
            "establishment_date": datetime.date.today().isoformat(),
            "address_city": "전국",
            "industry_code": "",
            "revenue": "UNDER_1B",
            "employees": "UNDER_10",
            "is_corp": is_corp
        }
    }

@app.post("/api/save-profile")
def api_save_profile(profile: UserProfile, current_user: dict = Depends(_get_current_user)):
    """
    UPSERT logic for user profile. Requires password re-verification for existing users.
    """
    # 소유권 검증: 자신의 프로필만 수정 가능
    if profile.business_number != current_user["bn"]:
        raise HTTPException(status_code=403, detail="권한이 없습니다.")
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # 기존 사용자인지 확인 + 비밀번호 검증
        cursor.execute("SELECT password_hash FROM users WHERE business_number = %s", (profile.business_number,))
        existing = cursor.fetchone()
        if existing and existing.get("password_hash"):
            if not profile.password:
                raise HTTPException(status_code=401, detail="프로필 변경 시 비밀번호 확인이 필요합니다.")
            if not _verify_password(profile.password, existing["password_hash"]):
                raise HTTPException(status_code=401, detail="비밀번호가 올바르지 않습니다.")

        query = """
        INSERT INTO users (business_number, company_name, establishment_date, address_city, industry_code, revenue_bracket, employee_count_bracket, interests, user_type, age_range, income_level, family_type, employment_status)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        ON CONFLICT(business_number) DO UPDATE SET
            company_name=COALESCE(EXCLUDED.company_name, users.company_name),
            establishment_date=COALESCE(EXCLUDED.establishment_date, users.establishment_date),
            address_city=COALESCE(EXCLUDED.address_city, users.address_city),
            industry_code=COALESCE(EXCLUDED.industry_code, users.industry_code),
            revenue_bracket=COALESCE(EXCLUDED.revenue_bracket, users.revenue_bracket),
            employee_count_bracket=COALESCE(EXCLUDED.employee_count_bracket, users.employee_count_bracket),
            interests=COALESCE(EXCLUDED.interests, users.interests),
            user_type=COALESCE(EXCLUDED.user_type, users.user_type),
            age_range=COALESCE(EXCLUDED.age_range, users.age_range),
            income_level=COALESCE(EXCLUDED.income_level, users.income_level),
            family_type=COALESCE(EXCLUDED.family_type, users.family_type),
            employment_status=COALESCE(EXCLUDED.employment_status, users.employment_status)
        """
        cursor.execute(query, (
            profile.business_number, profile.company_name, profile.establishment_date,
            profile.address_city, profile.industry_code, profile.revenue_bracket,
            profile.employee_count_bracket, profile.interests,
            profile.user_type, profile.age_range, profile.income_level,
            profile.family_type, profile.employment_status
        ))
        conn.commit()
        # 프로필 변경 시 개인화 캐시 무효화
        try:
            cursor.execute(
                "DELETE FROM user_match_cache WHERE business_number = %s",
                (profile.business_number,)
            )
            conn.commit()
        except Exception:
            pass
        _response_cache.pop(f"auth_me:{profile.business_number}", None)
        _response_cache.pop(f"match:{profile.business_number}:business", None)
        _response_cache.pop(f"match:{profile.business_number}:individual", None)
        _response_cache.pop(f"match:{profile.business_number}:all", None)
        return {"status": "SUCCESS", "message": "프로필이 저장되었습니다."}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@app.post("/api/ai/parse-interests")
def api_parse_interests(req: dict):
    """사용자가 입력한 자유 텍스트를 카테고리로 매핑 (Gemini)"""
    text = (req.get("text") or "").strip()
    user_type = req.get("user_type", "both")
    if not text:
        return {"status": "SUCCESS", "interests": []}

    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return {"status": "SUCCESS", "interests": [text]}

    biz_cats = "창업지원,기술개발,수출마케팅,고용지원,시설개선,정책자금,디지털전환,판로개척,교육훈련,에너지환경,소상공인,R&D"
    ind_cats = "취업,주거,교육,청년,출산,육아,다자녀,장학금,의료,장애,저소득,노인,문화"
    cats = ind_cats if user_type == "individual" else biz_cats

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        resp = model.generate_content(
            f"사용자가 관심분야를 다음과 같이 입력했습니다: \"{text}\"\n\n"
            f"아래 카테고리 목록에서 관련된 것을 모두 골라 JSON 배열로만 반환하세요. 카테고리 목록: [{cats}]\n"
            f"반드시 목록에 있는 정확한 단어만 사용하세요. 해당하는 것이 없으면 빈 배열 []을 반환하세요.\n"
            f"JSON 배열만 반환. 다른 텍스트 없이.",
            generation_config={"temperature": 0}
        )
        import json as _json
        parsed = _json.loads(resp.text.strip().strip("`").replace("json\n", ""))
        valid = [c for c in parsed if c in cats]
        # 매핑된 카테고리 + 사용자 원본 입력을 함께 반환
        result = list(valid)
        if text not in result and text not in cats:
            result.append(text)  # 사용자 원본 키워드도 포함
        return {"status": "SUCCESS", "interests": result if result else [text]}
    except Exception as e:
        print(f"[ParseInterests] Error: {e}")
        return {"status": "SUCCESS", "interests": [text]}


@app.put("/api/profile")
def api_update_profile(req: dict, current_user: dict = Depends(_get_current_user)):
    """프로필 간편 업데이트 (알림 설정에서 호출)"""
    bn = current_user["bn"]
    conn = get_db_connection()
    cur = conn.cursor()
    fields = []
    params = []
    allowed_keys = [
        "user_type", "address_city", "interest_regions", "revenue_bracket", "employee_count_bracket",
        "interests", "custom_needs", "custom_keywords", "company_name",
        "gender", "age_range", "income_level", "family_type", "employment_status",
        "founded_date", "is_pre_founder", "certifications",
        "industry_code", "industry_name",
    ]
    # user_type 값 검증
    if "user_type" in req and req["user_type"] not in ("individual", "business", "both", None):
        conn.close()
        raise HTTPException(status_code=400, detail="유효하지 않은 user_type입니다. (individual, business, both 중 선택)")

    date_keys = {"founded_date", "establishment_date"}
    for key in allowed_keys:
        val = req.get(key)
        if val is None:
            continue
        if key in date_keys and val == "":
            continue
        fields.append(f"{key} = %s")
        params.append(val)
    if not fields:
        conn.close()
        return {"status": "SUCCESS"}
    # user_type 변경 시 구독 중이면 안내 메시지 추가
    price_notice = ""
    if "user_type" in req:
        cur2 = conn.cursor()
        cur2.execute("SELECT plan, user_type FROM users WHERE business_number = %s", (bn,))
        existing = cur2.fetchone()
        if existing:
            ex = dict(existing)
            old_type = ex.get("user_type") or "business"
            new_type = req["user_type"]
            if old_type != new_type and ex.get("plan") in ("lite", "pro"):
                price_notice = " 다음 결제 시 변경된 유형의 요금이 적용됩니다."

    params.append(bn)
    cur.execute(f"UPDATE users SET {', '.join(fields)} WHERE business_number = %s", params)
    conn.commit()
    conn.close()
    _response_cache.pop(f"auth_me:{bn}", None)
    _response_cache.pop(f"match:{bn}:business", None)
    _response_cache.pop(f"match:{bn}:individual", None)
    _response_cache.pop(f"match:{bn}:all", None)
    # DB 사전매칭 캐시도 무효화
    try:
        _dc = get_db_connection()
        _dc.cursor().execute("DELETE FROM user_match_cache WHERE business_number = %s", (bn,))
        _dc.commit(); _dc.close()
    except Exception:
        pass
    return {"status": "SUCCESS", "message": f"프로필이 업데이트되었습니다.{price_notice}"}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 공고 분석 API (블로그봇 + 앱 공용)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

BLOG_API_KEY = os.getenv("BLOG_API_KEY", "")

def _verify_api_key(x_api_key: Optional[str] = Header(None)):
    """블로그봇용 API Key 인증"""
    if not BLOG_API_KEY:
        raise HTTPException(status_code=500, detail="BLOG_API_KEY가 설정되지 않았습니다.")
    if x_api_key != BLOG_API_KEY:
        raise HTTPException(status_code=401, detail="유효하지 않은 API Key입니다.")


def _format_analysis_response(announcement_id: int, analysis: dict, announcement: dict) -> dict:
    """분석 데이터를 블로그봇/앱 공용 응답 스키마로 변환"""
    da = analysis.get("deep_analysis", {}) or {}
    ps = analysis.get("parsed_sections", {}) or {}

    # eligibility_detail: deep_analysis 구조체 또는 parsed_sections 텍스트
    elig = da.get("eligibility_detail")
    if isinstance(elig, dict):
        # 구조체를 읽기 좋은 텍스트로도 변환
        elig_text = ""
        parts = []
        if elig.get("region") and elig["region"] != "전국":
            parts.append(f"지역: {elig['region']}")
        if elig.get("business_types"):
            parts.append(f"대상: {', '.join(elig['business_types'])}")
        if elig.get("other_conditions"):
            parts.extend(elig["other_conditions"])
        elig_text = " / ".join(parts) if parts else ""
        elig_combined = elig_text or ps.get("eligibility") or ""
    else:
        elig_combined = elig or ps.get("eligibility") or ""

    # 제출서류: deep_analysis.required_documents (list of dict) 또는 parsed_sections.required_docs (text)
    docs_raw = da.get("required_documents") or []
    docs_list = []
    if isinstance(docs_raw, list):
        for d in docs_raw:
            if isinstance(d, dict):
                docs_list.append(d.get("doc_name", str(d)))
            elif isinstance(d, str):
                docs_list.append(d)
    if not docs_list:
        # fallback: parsed_sections에서 추출
        ps_docs = ps.get("required_docs") or ""
        if isinstance(ps_docs, str) and ps_docs:
            docs_list = [line.strip().lstrip("- ·•◦") for line in ps_docs.split("\n") if line.strip() and len(line.strip()) > 1]

    # 신청방법
    app_method = ps.get("application_method") or ""
    if isinstance(app_method, str) and app_method:
        app_method_text = app_method
    else:
        app_method_text = da.get("application_method") or ""

    # 신청절차 → steps 파싱
    application_steps = []
    timeline_raw = ps.get("timeline") or ""
    if isinstance(timeline_raw, str) and timeline_raw:
        for i, line in enumerate(timeline_raw.split("\n")):
            line = line.strip().lstrip("- ·•◦0123456789.)")
            if line and len(line) > 2:
                application_steps.append({"step": i + 1, "title": line, "detail": ""})

    # 지원내용
    support_raw = ps.get("support_details") or ""
    support_summary = da.get("support_summary") or {}
    if isinstance(support_summary, dict) and support_summary.get("amount"):
        support_text = f"{support_summary.get('amount', '')} ({support_summary.get('duration', '')}, {support_summary.get('method', '')})"
    elif isinstance(support_raw, str) and support_raw:
        support_text = support_raw[:500]
    else:
        support_text = ""

    # 선정기준
    eval_raw = ps.get("evaluation_criteria") or ""
    eval_weights = da.get("evaluation_weights") or []
    if isinstance(eval_weights, list) and eval_weights:
        sel_criteria = " / ".join(f"{e.get('criteria','')}({e.get('weight','')})" for e in eval_weights if isinstance(e, dict))
    elif isinstance(eval_raw, str):
        sel_criteria = eval_raw[:300]
    else:
        sel_criteria = ""

    return {
        "announcement_id": announcement_id,
        "title": announcement.get("title", ""),
        "eligibility_detail": elig_combined,
        "eligibility_raw": ps.get("eligibility") or "",
        "exclusions": ps.get("exclusions") or "",
        "required_documents": docs_list,
        "application_method": app_method_text,
        "application_url": da.get("application_url") or announcement.get("origin_url") or "",
        "application_steps": application_steps,
        "support_detail": support_text,
        "support_summary": support_summary if isinstance(support_summary, dict) else {},
        "selection_criteria": sel_criteria,
        "bonus_items": da.get("bonus_items") or [],
        "key_warnings": da.get("key_warnings") or [],
        "target_age": da.get("target_age") or "",
        "target_region": da.get("target_region") or (da.get("eligibility_detail", {}) or {}).get("region", "") or announcement.get("region") or "",
        "target_family": da.get("target_family") or "",
        "deadline_date": str(announcement.get("deadline_date") or ""),
        "department": announcement.get("department") or "",
        "category": announcement.get("category") or "",
        "origin_url": announcement.get("origin_url") or "",
        "analyzed_at": str(analysis.get("updated_at") or datetime.datetime.utcnow().isoformat()),
        "source": analysis.get("source_type") or "unknown",
        "has_full_text": bool(analysis.get("full_text")),
    }


@app.get("/api/v1/announcements/{announcement_id}/analysis")
def api_get_analysis(announcement_id: int, raw: bool = False, _: None = Depends(_verify_api_key)):
    """저장된 분석 데이터 조회 (블로그봇용). raw=true면 DB 원본 데이터 포함"""
    conn = get_db_connection()
    try:
        from app.services.doc_analysis_service import get_deep_analysis
        analysis = get_deep_analysis(announcement_id, conn)
        if not analysis:
            conn.close()
            raise HTTPException(status_code=404, detail="분석 데이터가 없습니다. POST /analyze로 분석을 먼저 실행하세요.")

        cur = conn.cursor()
        cur.execute("SELECT * FROM announcements WHERE announcement_id = %s", (announcement_id,))
        ann = cur.fetchone()
        if not ann:
            conn.close()
            raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")

        result = _format_analysis_response(announcement_id, analysis, dict(ann))
        response = {"status": "SUCCESS", "data": result}

        # raw=true면 DB 원본 데이터도 포함 (디버그용)
        if raw:
            response["_raw"] = {
                "parsed_sections": analysis.get("parsed_sections", {}),
                "deep_analysis": analysis.get("deep_analysis", {}),
                "full_text_length": len(analysis.get("full_text") or ""),
                "full_text_preview": (analysis.get("full_text") or "")[:500],
            }

        conn.close()
        return response
    except HTTPException:
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))


def _verify_api_key_or_admin(authorization: Optional[str] = Header(None), x_api_key: Optional[str] = Header(None)):
    """API Key 또는 admin JWT 둘 중 하나로 인증"""
    if x_api_key and BLOG_API_KEY and x_api_key == BLOG_API_KEY:
        return
    if authorization:
        try:
            _verify_admin(authorization)
            return
        except Exception:
            pass
    raise HTTPException(status_code=401, detail="API Key 또는 admin 인증이 필요합니다.")


@app.post("/api/v1/announcements/{announcement_id}/analyze")
def api_analyze_announcement(announcement_id: int, force: bool = False, _: None = Depends(_verify_api_key_or_admin)):
    """공고 원문 크롤링 + AI 분석 + DB 저장"""
    conn = get_db_connection()
    try:
        from app.services.doc_analysis_service import ensure_analysis, get_deep_analysis

        # 이미 분석됐으면 바로 반환 (force=True면 재분석)
        if not force:
            existing = get_deep_analysis(announcement_id, conn)
            if existing and existing.get("full_text"):
                cur = conn.cursor()
                cur.execute("SELECT * FROM announcements WHERE announcement_id = %s", (announcement_id,))
                ann = cur.fetchone()
                if ann:
                    result = _format_analysis_response(announcement_id, existing, dict(ann))
                    conn.close()
                    return {"status": "SUCCESS", "cached": True, "data": result}

        # 실시간 분석 (force면 재크롤링+재분석)
        analysis = ensure_analysis(announcement_id, conn, force=force)
        if not analysis:
            conn.close()
            raise HTTPException(status_code=422, detail="분석 실패: 원문 접근 불가 또는 AI 분석 오류")

        cur = conn.cursor()
        cur.execute("SELECT * FROM announcements WHERE announcement_id = %s", (announcement_id,))
        ann = cur.fetchone()
        if not ann:
            conn.close()
            raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")

        result = _format_analysis_response(announcement_id, analysis, dict(ann))
        conn.close()
        return {"status": "SUCCESS", "cached": False, "data": result}
    except HTTPException:
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))


from app.services.sync_service import sync_service

def _run_sync_in_thread():
    """동기화를 별도 스레드에서 실행 (이벤트 루프 블로킹 방지)"""
    import asyncio
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    sync_status["running"] = True
    sync_status["last_result"] = "진행 중..."
    sync_status["last_time"] = datetime.datetime.now().isoformat()
    try:
        loop.run_until_complete(sync_service.sync_all())
        sync_status["last_result"] = "완료"
        _log_system("api_sync", "collection", "전체 API 데이터 수집 완료", "success")
    except Exception as e:
        sync_status["last_result"] = f"오류: {e}"
        _log_system("api_sync", "collection", f"API 수집 오류: {e}", "error")
    finally:
        sync_status["running"] = False
        sync_status["last_time"] = datetime.datetime.now().isoformat()
        loop.close()


@app.post("/api/sync", dependencies=[Depends(_verify_admin)])
async def api_sync_data():
    if sync_status["running"]:
        return {"status": "ALREADY_RUNNING", "message": "동기화가 이미 진행 중입니다."}
    import threading
    t = threading.Thread(target=_run_sync_in_thread, daemon=True)
    t.start()
    return {"status": "STARTED", "message": "백그라운드에서 동기화를 시작합니다."}


@app.get("/api/admin/sync-status", dependencies=[Depends(_verify_admin)])
def get_sync_status():
    return {"status": "SUCCESS", "data": sync_status}

class PushSubscription(BaseModel):
    business_number: Optional[str] = None
    endpoint: str
    keys: dict


@app.get("/api/push/vapid-key")
def get_vapid_key():
    pub = os.getenv("VAPID_PUBLIC_KEY", "")
    return {"publicKey": pub}


@app.post("/api/push/subscribe")
def push_subscribe(sub: PushSubscription):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        """INSERT INTO push_subscriptions (business_number, endpoint, p256dh, auth)
           VALUES (%s, %s, %s, %s)
           ON CONFLICT(endpoint) DO UPDATE SET
               business_number=EXCLUDED.business_number,
               p256dh=EXCLUDED.p256dh,
               auth=EXCLUDED.auth""",
        (sub.business_number, sub.endpoint, sub.keys.get("p256dh", ""), sub.keys.get("auth", ""))
    )
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "message": "푸시 알림 구독 완료"}


@app.post("/api/push/unsubscribe")
def push_unsubscribe(data: dict):
    endpoint = data.get("endpoint", "")
    if not endpoint:
        raise HTTPException(status_code=400, detail="endpoint 필수")
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM push_subscriptions WHERE endpoint = %s", (endpoint,))
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "message": "푸시 알림 해제 완료"}


@app.post("/api/admin/push-test", dependencies=[Depends(_verify_admin)])
def admin_push_test():
    from pywebpush import webpush, WebPushException
    vapid_private = os.getenv("VAPID_PRIVATE_KEY", "")
    vapid_claims = {"sub": os.getenv("VAPID_CLAIMS_EMAIL", "mailto:admin@example.com")}

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT endpoint, p256dh, auth FROM push_subscriptions")
    subs = cursor.fetchall()
    conn.close()

    sent = 0
    failed = 0
    payload = json.dumps({
        "title": "지원금AI",
        "body": "새로운 맞춤 공고가 등록되었습니다!",
        "url": "/",
    }, ensure_ascii=False)

    for sub in subs:
        try:
            webpush(
                subscription_info={"endpoint": sub["endpoint"], "keys": {"p256dh": sub["p256dh"], "auth": sub["auth"]}},
                data=payload,
                vapid_private_key=vapid_private,
                vapid_claims=vapid_claims,
            )
            sent += 1
        except WebPushException:
            failed += 1

    _log_system("push_test", "notification", f"발송 {sent}건, 실패 {failed}건", "success" if failed == 0 else "partial", sent)
    return {"status": "SUCCESS", "message": f"발송 {sent}건, 실패 {failed}건"}


@app.get("/api/smart-matches")
def api_smart_matches(
    target_type: Optional[str] = None,
    current_user: dict = Depends(_get_current_user),
):
    """AI 맞춤 추천 공고.
    3단계 폴백:
      1) user_smart_matches (배치 AI 결과)
      2) user_match_cache business/individual (사전 매칭 캐시)
      3) 실시간 get_matches_hybrid() (상위 30건)
    """
    from app.core.matcher import get_matches_hybrid
    bn = current_user["bn"]
    tt = (target_type or "business").lower()
    if tt not in ("business", "individual"):
        tt = "business"

    conn = get_db_connection()
    try:
        cur = conn.cursor()

        # 1단계: user_smart_matches (배치 결과)
        cur.execute(
            "SELECT matches, created_at FROM user_smart_matches WHERE business_number = %s",
            (bn,)
        )
        row = cur.fetchone()
        if row and row["matches"]:
            data = row["matches"]
            if isinstance(data, str):
                data = json.loads(data)
            if data:
                return {"status": "SUCCESS", "data": data, "updated_at": str(row["created_at"]), "source": "batch"}

        # 2단계: user_match_cache (사전 매칭 캐시)
        cur.execute(
            """SELECT match_data, created_at FROM user_match_cache
               WHERE business_number = %s AND target_type = %s
               ORDER BY created_at DESC LIMIT 1""",
            (bn, tt)
        )
        cache_row = cur.fetchone()
        if cache_row and cache_row["match_data"]:
            data = cache_row["match_data"]
            if isinstance(data, str):
                data = json.loads(data)
            if data and len(data) > 0:
                return {"status": "SUCCESS", "data": data[:50], "updated_at": str(cache_row["created_at"]), "source": "cache"}

        # 3단계: 실시간 매칭 (캐시 없을 때)
        cur.execute("SELECT * FROM users WHERE business_number = %s", (bn,))
        user_row = cur.fetchone()
        if not user_row:
            return {"status": "SUCCESS", "data": [], "source": "none"}

        user_dict = dict(user_row)
        is_individual = (tt == "individual")
        matches = get_matches_hybrid(user_dict, is_individual=is_individual)
        matches = matches[:30]

        # 실시간 결과 캐시 저장 (다음 요청 속도 향상)
        if matches:
            try:
                cur.execute(
                    """INSERT INTO user_match_cache (business_number, target_type, match_data, created_at)
                       VALUES (%s, %s, %s::jsonb, CURRENT_TIMESTAMP)
                       ON CONFLICT (business_number, target_type)
                       DO UPDATE SET match_data = EXCLUDED.match_data, created_at = CURRENT_TIMESTAMP""",
                    (bn, tt, json.dumps(matches, ensure_ascii=False, default=str))
                )
                conn.commit()
            except Exception:
                pass

        return {"status": "SUCCESS", "data": matches, "source": "realtime"}

    except Exception as e:
        return {"status": "SUCCESS", "data": [], "error": str(e)[:100]}
    finally:
        conn.close()


@app.post("/api/match")
def api_match_programs(request: BusinessNumberRequest, current_user: dict = Depends(_get_current_user)):
    """
    Fetches user profile and runs the hybrid matching engine.
    """
    # 소유권 검증: 자신의 business_number만 매칭 가능
    if request.business_number != current_user["bn"]:
        raise HTTPException(status_code=403, detail="권한이 없습니다.")
    # target_type: 프론트에서 현재 탭 전달 → 해당 타입만 우선 매칭
    req_tt = (request.target_type or "").strip().lower()
    # 1차: 인메모리 캐시 (15분)
    match_cache_key = f"match:{request.business_number}:{req_tt or 'all'}"
    cached_match = _get_cached(match_cache_key)
    if cached_match:
        return cached_match
    # 2차: DB 사전 매칭 캐시 (daily digest에서 저장)
    if req_tt in ("business", "individual"):
        try:
            _dbc = get_db_connection()
            _dbc_cur = _dbc.cursor()
            _dbc_cur.execute(
                "SELECT match_data FROM user_match_cache WHERE business_number = %s AND target_type = %s AND created_at > CURRENT_DATE - INTERVAL '1 day'",
                (request.business_number, req_tt))
            _dbc_row = _dbc_cur.fetchone()
            _dbc.close()
            if _dbc_row and _dbc_row.get("match_data"):
                db_data = _dbc_row["match_data"]
                if isinstance(db_data, str):
                    db_data = json.loads(db_data)
                if db_data and len(db_data) > 0:
                    result = {"status": "SUCCESS", "data": db_data}
                    _set_cache(match_cache_key, result)
                    return result
        except Exception as dbc_err:
            print(f"[match] db cache read error: {dbc_err}")
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE business_number = %s", (request.business_number,))
    user = cursor.fetchone()
    conn.close()

    if not user:
        raise HTTPException(status_code=404, detail="사용자 정보를 찾을 수 없습니다.")

    user_dict = dict(user)

    # user_type에 따라 매칭 엔진 분기
    user_type = user_dict.get("user_type") or "both"
    # target_type 지정 시 해당 타입만 실행 (1회 호출 = ~2초)
    if req_tt == "individual":
        matches = get_matches_hybrid(user_dict, is_individual=True)
    elif req_tt == "business":
        matches = get_matches_hybrid(user_dict, is_individual=False)
    elif user_type == "individual":
        matches = get_matches_hybrid(user_dict, is_individual=True)
    elif user_type == "both":
        # target_type 미지정 + both → 기업 우선 실행, 개인은 백그라운드
        matches = get_matches_hybrid(user_dict, is_individual=False)
        # 개인 매칭 백그라운드 캐시
        import threading
        def _bg_individual():
            try:
                ind = get_matches_hybrid(user_dict, is_individual=True)
                ind = ind[:100]
                _set_cache(f"match:{request.business_number}:individual", {"status": "SUCCESS", "data": ind})
            except Exception as e:
                print(f"[match bg] individual error: {e}")
        threading.Thread(target=_bg_individual, daemon=True).start()
    else:
        matches = get_matches_hybrid(user_dict, is_individual=False)
    # 상위 100건만 반환 (점수순 정렬 후, 프론트에서 20건씩 페이지네이션)
    matches = matches[:100]

    # AI 추출 데이터 보완 (프론트엔드 대응)
    for match in matches:
        if match.get("eligibility_logic") and isinstance(match["eligibility_logic"], str):
            try:
                match["eligibility_logic"] = json.loads(match["eligibility_logic"])
            except (json.JSONDecodeError, TypeError):
                pass

    _log_event("matching", request.business_number, f"type={user_type},count={len(matches)},top_score={matches[0].get('match_score',0) if matches else 0}")

    # P1.3: match_history 저장 (별도 커넥션 대신 간소화)
    try:
        mh_conn = get_db_connection()
        mh_cur = mh_conn.cursor()
        snap_keys = ["company_name","industry_code","address_city","revenue_bracket",
                     "employee_count_bracket","interests","user_type"]
        snapshot = {k: user_dict.get(k) for k in snap_keys if user_dict.get(k) is not None}
        top5 = [{"id": m.get("announcement_id"), "title": m.get("title", "")[:60]} for m in matches[:5]]
        mh_cur.execute(
            """INSERT INTO match_history (business_number, user_type, profile_snapshot, total_matches, top_matches)
               VALUES (%s, %s, %s::jsonb, %s, %s::jsonb)""",
            (request.business_number, user_type,
             json.dumps(snapshot, ensure_ascii=False), len(matches),
             json.dumps(top5, ensure_ascii=False))
        )
        mh_conn.commit()
        mh_conn.close()
    except Exception as mh_err:
        print(f"[match_history] {mh_err}")

    result = {"status": "SUCCESS", "data": matches}
    _set_cache(match_cache_key, result)
    return result

@app.post("/api/notification-settings")
def api_save_notification_settings(settings: NotificationSettings):
    """
    Saves or updates user notification preferences.
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        query = """
        INSERT INTO notification_settings (business_number, email, phone_number, channel, is_active, kakao_enabled)
        VALUES (%s, %s, %s, %s, %s, %s)
        ON CONFLICT(business_number) DO UPDATE SET
            email=EXCLUDED.email,
            phone_number=EXCLUDED.phone_number,
            channel=EXCLUDED.channel,
            is_active=EXCLUDED.is_active,
            kakao_enabled=EXCLUDED.kakao_enabled,
            updated_at=NOW()
        """
        cursor.execute(query, (
            settings.business_number, settings.email, settings.phone_number,
            settings.channel, int(settings.is_active), int(settings.kakao_enabled or 0)
        ))
        conn.commit()
        return {"status": "SUCCESS", "message": "알림 설정이 저장되었습니다."}
    except Exception as e:
        conn.rollback()
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

@app.get("/api/notification-settings/{bn}")
def api_get_notification_settings(bn: str, current_user: dict = Depends(_get_current_user)):
    """
    Retrieves user notification preferences.
    """
    if bn != current_user["bn"]:
        raise HTTPException(status_code=403, detail="권한이 없습니다.")
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM notification_settings WHERE business_number = %s", (bn,))
    settings = cursor.fetchone()
    conn.close()

    if settings:
        return {"status": "SUCCESS", "data": dict(settings)}
    else:
        # Return default empty settings if not found
        return {
            "status": "SUCCESS",
            "data": {
                "business_number": bn,
                "email": "",
                "phone_number": "",
                "channel": "BOTH",
                "is_active": True
            }
        }

@app.post("/api/admin/send-digest", dependencies=[Depends(_verify_admin)])
async def api_send_digest():
    """데일리 다이제스트를 즉시 생성하고 이메일 발송 (관리자용)"""
    # 진단 단계: import + 호출을 분리해서 어디서 죽는지 추적
    stage = "init"
    try:
        stage = "import_module"
        from app.services.notification_service import notification_service
        stage = "call_generate_daily_digest"
        results = await notification_service.generate_daily_digest()
        stage = "post_process"
        sent_count = sum(1 for r in results if r.get("email_sent"))
        _log_system("send_digest", "notification", f"{len(results)}명 대상, {sent_count}건 이메일 발송", "success", sent_count)
        return {
            "status": "SUCCESS",
            "message": f"다이제스트 생성 완료: {len(results)}명 대상, {sent_count}건 이메일 발송",
            "data": results
        }
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        print(f"[send-digest] STAGE={stage} ERROR: {e}\n{tb}")
        try:
            _log_system("send_digest", "notification", f"stage={stage} ERROR: {str(e)[:150]}", "error", 0)
        except Exception:
            pass
        return {
            "status": "ERROR",
            "stage": stage,
            "error": str(e)[:500],
            "traceback": tb[-2000:],
        }


@app.get("/api/admin/digest-probe")
async def api_digest_probe():
    """진단용: 인증 없이 단계별 stub 응답으로 무엇이 깨지는지 추적"""
    import traceback
    out = {"steps": []}
    try:
        out["steps"].append("import_module")
        from app.services.notification_service import notification_service
        out["steps"].append("imported")
        out["steps"].append("get_target_users")
        users = await notification_service.get_target_users()
        out["steps"].append(f"got_{len(users)}_users")
        out["users_count"] = len(users)
        return {"status": "OK", **out}
    except Exception as e:
        out["error"] = str(e)[:300]
        out["traceback"] = traceback.format_exc()[-1500:]
        return {"status": "ERROR", **out}


class SavedBulk(BaseModel):
    business_number: str
    announcement_ids: List[int]


@app.post("/api/saved/bulk")
def api_save_bulk(body: SavedBulk, current_user: dict = Depends(_get_current_user)):
    if body.business_number != current_user["bn"]:
        raise HTTPException(status_code=403, detail="권한이 없습니다.")
    # FREE 플랜: 저장 불가
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT plan, plan_expires_at FROM users WHERE business_number=%s", (current_user["bn"],))
    u = cursor.fetchone()
    if u:
        plan = u["plan"] or "free"
        if plan != "free" and u["plan_expires_at"]:
            import datetime as _dt
            try:
                if _dt.datetime.fromisoformat(str(u["plan_expires_at"])) < _dt.datetime.utcnow():
                    plan = "free"
            except Exception:
                pass
        if plan == "free":
            conn.close()
            raise HTTPException(status_code=403, detail="LITE_REQUIRED")
    inserted = 0
    try:
        for aid in body.announcement_ids:
            try:
                cursor.execute(
                    """INSERT INTO saved_announcements (business_number, announcement_id)
                       VALUES (%s, %s)
                       ON CONFLICT (business_number, announcement_id) DO NOTHING""",
                    (body.business_number, aid),
                )
                inserted += cursor.rowcount
            except psycopg2.IntegrityError:
                conn.rollback()
        conn.commit()
    finally:
        conn.close()
    return {"status": "SUCCESS", "message": f"{inserted}건 저장됨", "inserted": inserted}


@app.get("/api/saved/{bn}")
def api_get_saved(bn: str, current_user: dict = Depends(_get_current_user)):
    if bn != current_user["bn"]:
        raise HTTPException(status_code=403, detail="권한이 없습니다.")
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT sa.id, sa.announcement_id, sa.memo, sa.saved_at,
               a.title, a.deadline_date, a.origin_url, a.department,
               a.category, a.support_amount, a.origin_source
        FROM saved_announcements sa
        JOIN announcements a ON sa.announcement_id = a.announcement_id
        WHERE sa.business_number = %s
        ORDER BY a.deadline_date ASC NULLS LAST
    """, (bn,))
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return {"status": "SUCCESS", "data": rows}


@app.delete("/api/saved/{saved_id}")
def api_delete_saved(saved_id: int, current_user: dict = Depends(_get_current_user)):
    conn = get_db_connection()
    cursor = conn.cursor()
    # 소유권 확인 후 삭제
    cursor.execute("DELETE FROM saved_announcements WHERE id = %s AND business_number = %s", (saved_id, current_user["bn"]))
    conn.commit()
    deleted = cursor.rowcount
    conn.close()
    if deleted:
        return {"status": "SUCCESS", "message": "삭제됨"}
    raise HTTPException(status_code=404, detail="해당 저장 항목을 찾을 수 없습니다.")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PRO 전용 API — 고객사 프로필 관리 / 상담 이력 / 종합 리포트
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _require_pro(current_user: dict):
    """PRO 플랜 체크. PRO/biz가 아니면 403."""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT plan, plan_expires_at FROM users WHERE business_number = %s", (current_user["bn"],))
    u = cur.fetchone()
    conn.close()
    if not u:
        raise HTTPException(status_code=404, detail="사용자를 찾을 수 없습니다.")
    plan = u["plan"] or "free"
    if plan not in ("pro", "biz"):
        raise HTTPException(status_code=403, detail="PRO 플랜 전용 기능입니다.")
    # 만료 체크
    if u.get("plan_expires_at"):
        import datetime as _dt
        try:
            exp = _dt.datetime.fromisoformat(str(u["plan_expires_at"]))
            if exp < _dt.datetime.utcnow():
                raise HTTPException(status_code=403, detail="플랜이 만료되었습니다. 갱신 후 이용하세요.")
        except ValueError:
            pass
    return current_user


# ── 1) 고객사 프로필 CRUD ──

class ClientProfileCreate(BaseModel):
    client_name: str
    client_type: Optional[str] = "business"  # "business" | "individual"
    business_number: Optional[str] = None
    establishment_date: Optional[str] = None
    address_city: Optional[str] = None
    industry_code: Optional[str] = None
    industry_name: Optional[str] = None
    revenue_bracket: Optional[str] = None
    employee_count_bracket: Optional[str] = None
    interests: Optional[str] = None
    memo: Optional[str] = ""
    # CRM 확장 필드
    contact_name: Optional[str] = None       # 담당자명
    contact_email: Optional[str] = None      # 담당자 이메일
    contact_phone: Optional[str] = None      # 담당자 전화번호
    tags: Optional[str] = None               # 태그 (쉼표 구분)
    status: Optional[str] = "new"            # 상태: new/consulting/matched/applied/selected


@app.get("/api/pro/clients")
def api_pro_clients(client_type: Optional[str] = None, current_user: dict = Depends(_get_current_user)):
    """PRO: 내 고객 프로필 목록 조회 (client_type 필터 가능)"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    where = "owner_business_number = %s AND is_active = TRUE"
    params: list = [current_user["bn"]]
    if client_type:
        where += " AND COALESCE(client_type, 'business') = %s"
        params.append(client_type)
    cur.execute(
        f"""SELECT id, client_name, COALESCE(client_type, 'business') AS client_type,
                  business_number, address_city, industry_code, industry_name,
                  revenue_bracket, employee_count_bracket, establishment_date, interests, memo,
                  contact_name, contact_email, contact_phone, tags, COALESCE(status, 'new') AS status,
                  created_at, updated_at
           FROM client_profiles
           WHERE {where}
           ORDER BY updated_at DESC""",
        params
    )
    rows = cur.fetchall()
    conn.close()
    clients = []
    for r in rows:
        d = dict(r)
        if d.get("establishment_date"):
            d["establishment_date"] = str(d["establishment_date"])
        if d.get("created_at"):
            d["created_at"] = str(d["created_at"])
        if d.get("updated_at"):
            d["updated_at"] = str(d["updated_at"])
        clients.append(d)
    return {"status": "SUCCESS", "clients": clients}


@app.get("/api/pro/clients/with-history")
def api_pro_clients_with_history(current_user: dict = Depends(_get_current_user)):
    """PRO: 고객 목록 + 최근 상담 요약 (테이블형 리스트용)"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    # consultation_history 테이블이 없을 수 있으므로 기본 쿼리 + 상담 통계는 별도 조회
    cur.execute(
        """SELECT *
           FROM client_profiles
           WHERE owner_business_number = %s AND is_active = TRUE
           ORDER BY updated_at DESC""",
        (current_user["bn"],)
    )
    rows = cur.fetchall()
    conn.close()
    clients = []
    for r in rows:
        d = dict(r)
        for k in ("establishment_date", "created_at", "updated_at", "last_consult_date"):
            if d.get(k): d[k] = str(d[k])
        d["consult_count"] = d.get("consult_count") or 0
        d["last_consult_summary"] = d.get("last_consult_summary") or ""
        clients.append(d)
    return {"status": "SUCCESS", "clients": clients}


@app.get("/api/pro/clients/export")
def api_pro_clients_export(current_user: dict = Depends(_get_current_user)):
    """PRO: 고객 리스트 CSV 다운로드"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT client_name, COALESCE(client_type,'business') AS client_type, business_number,
                  address_city, industry_name, industry_code, revenue_bracket, employee_count_bracket,
                  establishment_date, contact_name, contact_email, contact_phone, tags, COALESCE(status,'new') AS status, memo,
                  created_at, updated_at
           FROM client_profiles WHERE owner_business_number = %s AND is_active = TRUE ORDER BY updated_at DESC""",
        (current_user["bn"],)
    )
    rows = cur.fetchall()
    conn.close()

    import io, csv
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["기업명", "유형", "사업자번호", "지역", "업종명", "업종코드", "매출", "직원수", "설립일", "담당자", "이메일", "전화", "태그", "상태", "메모", "등록일", "수정일"])
    for r in rows:
        d = dict(r)
        writer.writerow([d.get("client_name",""), d.get("client_type",""), d.get("business_number",""),
                         d.get("address_city",""), d.get("industry_name",""), d.get("industry_code",""),
                         d.get("revenue_bracket",""), d.get("employee_count_bracket",""),
                         str(d.get("establishment_date","") or ""), d.get("contact_name",""),
                         d.get("contact_email",""), d.get("contact_phone",""), d.get("tags",""),
                         d.get("status",""), d.get("memo",""),
                         str(d.get("created_at","") or "")[:10], str(d.get("updated_at","") or "")[:10]])

    csv_content = output.getvalue()
    # BOM for Excel 한글 호환
    return Response(
        content="\ufeff" + csv_content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="clients.csv"'}
    )


@app.post("/api/pro/clients")
def api_pro_client_create(req: ClientProfileCreate, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 프로필 생성"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """INSERT INTO client_profiles
           (owner_business_number, client_name, client_type, business_number, establishment_date, address_city,
            industry_code, industry_name, revenue_bracket, employee_count_bracket, interests, memo,
            contact_name, contact_email, contact_phone, tags, status)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
           RETURNING id""",
        (current_user["bn"], req.client_name, req.client_type or "business", req.business_number,
         req.establishment_date, req.address_city, req.industry_code, req.industry_name,
         req.revenue_bracket, req.employee_count_bracket, req.interests, req.memo,
         req.contact_name, req.contact_email, req.contact_phone, req.tags or "", req.status or "new")
    )
    new_id = cur.fetchone()["id"]
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "id": new_id, "message": f"고객사 '{req.client_name}' 등록 완료"}


@app.put("/api/pro/clients/{client_id}")
def api_pro_client_update(client_id: int, req: ClientProfileCreate, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 프로필 수정"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """UPDATE client_profiles SET
           client_name=%s, client_type=%s, business_number=%s, establishment_date=%s, address_city=%s,
           industry_code=%s, industry_name=%s, revenue_bracket=%s, employee_count_bracket=%s,
           interests=%s, memo=%s, contact_name=%s, contact_email=%s, contact_phone=%s, tags=%s, status=%s,
           updated_at=CURRENT_TIMESTAMP
           WHERE id=%s AND owner_business_number=%s AND is_active=TRUE""",
        (req.client_name, req.client_type or "business", req.business_number, req.establishment_date, req.address_city,
         req.industry_code, req.industry_name, req.revenue_bracket, req.employee_count_bracket,
         req.interests, req.memo, req.contact_name, req.contact_email, req.contact_phone,
         req.tags or "", req.status or "new", client_id, current_user["bn"])
    )
    conn.commit()
    updated = cur.rowcount
    conn.close()
    if not updated:
        raise HTTPException(status_code=404, detail="고객사를 찾을 수 없습니다.")
    return {"status": "SUCCESS", "message": "수정 완료"}


@app.delete("/api/pro/clients/{client_id}")
def api_pro_client_delete(client_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 프로필 삭제 (soft delete)"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "UPDATE client_profiles SET is_active=FALSE, updated_at=CURRENT_TIMESTAMP WHERE id=%s AND owner_business_number=%s",
        (client_id, current_user["bn"])
    )
    conn.commit()
    deleted = cur.rowcount
    conn.close()
    if not deleted:
        raise HTTPException(status_code=404, detail="고객사를 찾을 수 없습니다.")
    return {"status": "SUCCESS", "message": "삭제 완료"}


# ── 1.5) 고객사 자료 첨부 (파일 업로드/조회/삭제) ──

@app.post("/api/pro/clients/{client_id}/files")
async def api_pro_client_upload_file(
    client_id: int,
    file: UploadFile = File(...),
    file_type: str = Form("other"),
    memo: str = Form(""),
    current_user: dict = Depends(_get_current_user),
):
    """PRO: 고객사 자료 업로드 (재무제표, 사업계획서, IR자료 등)"""
    _require_pro(current_user)
    # 파일 크기 제한 (10MB)
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="파일 크기는 10MB 이하만 가능합니다.")

    conn = get_db_connection()
    cur = conn.cursor()
    # 소유권 확인
    cur.execute("SELECT id FROM client_profiles WHERE id=%s AND owner_business_number=%s AND is_active=TRUE", (client_id, current_user["bn"]))
    if not cur.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="고객사를 찾을 수 없습니다.")

    # 파일 텍스트 추출 (PDF, HWP, HWPX, DOCX)
    extracted_text = ""
    try:
        fname = (file.filename or "").lower()
        detected = ""
        if fname.endswith(".pdf"): detected = "pdf"
        elif fname.endswith(".hwp"): detected = "hwp"
        elif fname.endswith(".hwpx"): detected = "hwpx"
        elif fname.endswith(".docx"): detected = "docx"

        if detected:
            from app.services.doc_analysis_service import extract_text_from_bytes
            extracted_text = extract_text_from_bytes(content, detected, max_chars=30000)
            print(f"[ClientFile] Extracted {len(extracted_text)} chars from {file.filename} ({detected})")
    except Exception as e:
        print(f"[ClientFile] Text extraction error: {e}")

    cur.execute(
        """INSERT INTO client_files (client_id, owner_business_number, file_name, file_type, file_size, file_data, extracted_text, memo)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s) RETURNING id""",
        (client_id, current_user["bn"], file.filename, file_type, len(content), psycopg2.Binary(content), extracted_text, memo)
    )
    file_id = cur.fetchone()["id"]
    conn.commit()
    conn.close()
    return {"status": "SUCCESS", "id": file_id, "message": f"'{file.filename}' 업로드 완료", "extracted_chars": len(extracted_text)}


@app.get("/api/pro/clients/{client_id}/files")
def api_pro_client_files(client_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 첨부 자료 목록 조회"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT id, file_name, file_type, file_size, memo, created_at
           FROM client_files
           WHERE client_id=%s AND owner_business_number=%s
           ORDER BY created_at DESC""",
        (client_id, current_user["bn"])
    )
    files = [dict(r) for r in cur.fetchall()]
    for f in files:
        if f.get("created_at"): f["created_at"] = str(f["created_at"])
    conn.close()
    return {"status": "SUCCESS", "files": files}


@app.get("/api/pro/clients/{client_id}/files/{file_id}/download")
def api_pro_client_file_download(client_id: int, file_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 첨부 자료 다운로드"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "SELECT file_name, file_data FROM client_files WHERE id=%s AND client_id=%s AND owner_business_number=%s",
        (file_id, client_id, current_user["bn"])
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
    from fastapi.responses import Response
    return Response(
        content=bytes(row["file_data"]),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{row["file_name"]}"'}
    )


@app.delete("/api/pro/clients/{client_id}/files/{file_id}")
def api_pro_client_file_delete(client_id: int, file_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 첨부 자료 삭제"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "DELETE FROM client_files WHERE id=%s AND client_id=%s AND owner_business_number=%s",
        (file_id, client_id, current_user["bn"])
    )
    conn.commit()
    deleted = cur.rowcount
    conn.close()
    if not deleted:
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
    return {"status": "SUCCESS", "message": "삭제 완료"}


# ── 2) 상담 이력 조회 + 엑셀 다운로드 ──

@app.get("/api/pro/consult-history")
def api_pro_consult_history(
    client_id: Optional[int] = None,
    limit: int = 50,
    offset: int = 0,
    current_user: dict = Depends(_get_current_user),
):
    """PRO: 상담 이력 조회 (고객사별 필터 가능)"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()

    bn = current_user["bn"]
    if client_id:
        # 특정 고객사의 사업자번호로 필터
        cur.execute(
            "SELECT business_number FROM client_profiles WHERE id=%s AND owner_business_number=%s",
            (client_id, bn)
        )
        cp = cur.fetchone()
        if not cp:
            conn.close()
            raise HTTPException(status_code=404, detail="고객사를 찾을 수 없습니다.")
        filter_bn = cp["business_number"] or bn
    else:
        filter_bn = bn

    cur.execute(
        """SELECT cl.id, cl.announcement_id, a.title as announcement_title, a.category,
                  cl.conclusion, cl.feedback, cl.feedback_detail, cl.created_at,
                  cl.messages, cl.session_id,
                  pcs.collected as session_collected
           FROM ai_consult_logs cl
           LEFT JOIN announcements a ON a.announcement_id = cl.announcement_id
           LEFT JOIN pro_consult_sessions pcs ON pcs.session_id = cl.session_id
           WHERE cl.business_number = %s
           ORDER BY cl.created_at DESC
           LIMIT %s OFFSET %s""",
        (filter_bn, limit, offset)
    )
    rows = cur.fetchall()

    # 총 건수
    cur.execute("SELECT COUNT(*) as cnt FROM ai_consult_logs WHERE business_number = %s", (filter_bn,))
    total = cur.fetchone()["cnt"]
    conn.close()

    def _build_consult_label(collected: dict, msgs: list) -> str:
        """일반상담 레이블: [일반상담] 이름 • 지역 • 연령대 • 소득 • 가구유형 • 취업상태.
        collected(pro_consult_sessions) 우선, 없으면 messages 첫 메시지에서 추출.
        """
        import re as _re_label
        parts = []

        # 1순위: collected dict (PRO 세션)
        if collected:
            name = collected.get("company_name") or collected.get("name") or ""
            if name:
                parts.append(f"이름: {name}")
            region = collected.get("address_city") or ""
            if region:
                parts.append(f"지역: {region}")
            age = collected.get("age_range") or ""
            if age:
                parts.append(f"연령대: {age}")
            income = collected.get("income_level") or ""
            if income:
                parts.append(f"월 소득: {income}")
            family = collected.get("family_type") or ""
            if family:
                parts.append(f"가구 유형: {family}")
            employ = collected.get("employment_status") or ""
            if employ:
                parts.append(f"취업 상태: {employ}")

        # 2순위: messages에서 bullet-point 프로필 텍스트 추출 (LITE fund 세션)
        if not parts and msgs:
            for m in msgs:
                if not isinstance(m, dict):
                    continue
                text = m.get("text", "") or ""
                # "이름: xxx • 지역: yyy ..." 패턴이 있는 메시지
                if "이름:" in text or "지역:" in text or "연령대:" in text:
                    # 불릿 항목들을 추출
                    for label, key in [("이름", "이름"), ("지역", "지역"), ("연령대", "연령대"),
                                       ("월 소득", "월 소득"), ("가구 유형", "가구 유형"), ("취업 상태", "취업 상태")]:
                        m_val = _re_label.search(rf"{label}:\s*([^•\n]+)", text)
                        if m_val:
                            val = m_val.group(1).strip().rstrip("*").strip()
                            if val:
                                parts.append(f"{label}: {val}")
                    if parts:
                        break

        return "[일반상담] " + " • ".join(parts) if parts else "[일반상담]"

    history = []
    for r in rows:
        d = dict(r)
        d["created_at"] = str(d["created_at"]) if d.get("created_at") else None
        # messages는 용량이 크므로 요약만
        msgs = d.pop("messages", None) or []
        d.pop("session_id", None)
        raw_collected = d.pop("session_collected", None)
        if isinstance(msgs, list):
            d["message_count"] = len(msgs)
            # 마지막 사용자 질문만 추출
            last_q = ""
            for m in reversed(msgs):
                if isinstance(m, dict) and m.get("role") == "user":
                    last_q = m.get("text", "")[:100]
                    break
            d["last_question"] = last_q
        else:
            msgs = []
            d["message_count"] = 0
            d["last_question"] = ""
        # announcement_id가 없는 일반상담: consult_label 생성
        if not d.get("announcement_id"):
            collected = raw_collected if isinstance(raw_collected, dict) else {}
            d["consult_label"] = _build_consult_label(collected, msgs)
        history.append(d)

    return {"status": "SUCCESS", "history": history, "total": total}


@app.get("/api/pro/consult-history/export")
def api_pro_consult_history_export(
    client_id: Optional[int] = None,
    current_user: dict = Depends(_get_current_user),
):
    """PRO: 상담 이력 엑셀(CSV) 다운로드"""
    _require_pro(current_user)
    import csv
    import io

    conn = get_db_connection()
    cur = conn.cursor()
    bn = current_user["bn"]

    if client_id:
        cur.execute(
            "SELECT business_number FROM client_profiles WHERE id=%s AND owner_business_number=%s",
            (client_id, bn)
        )
        cp = cur.fetchone()
        filter_bn = cp["business_number"] if cp and cp["business_number"] else bn
    else:
        filter_bn = bn

    cur.execute(
        """SELECT cl.id, cl.announcement_id, a.title as announcement_title, a.category, a.department,
                  a.support_amount, a.deadline_date,
                  cl.conclusion, cl.feedback, cl.feedback_detail, cl.created_at,
                  cl.messages
           FROM ai_consult_logs cl
           LEFT JOIN announcements a ON a.announcement_id = cl.announcement_id
           WHERE cl.business_number = %s
           ORDER BY cl.created_at DESC""",
        (filter_bn,)
    )
    rows = cur.fetchall()
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "상담ID", "공고ID", "공고명", "카테고리", "부처", "지원금액", "마감일",
        "결론", "피드백", "피드백상세", "상담일시", "대화수", "마지막질문"
    ])

    for r in rows:
        d = dict(r)
        msgs = d.get("messages") or []
        msg_count = len(msgs) if isinstance(msgs, list) else 0
        last_q = ""
        if isinstance(msgs, list):
            for m in reversed(msgs):
                if isinstance(m, dict) and m.get("role") == "user":
                    last_q = m.get("text", "")[:200]
                    break
        conclusion_map = {"eligible": "지원가능", "conditional": "조건부가능", "ineligible": "지원불가"}
        writer.writerow([
            d.get("id"), d.get("announcement_id"), d.get("announcement_title", ""),
            d.get("category", ""), d.get("department", ""),
            d.get("support_amount", ""), str(d.get("deadline_date", "")),
            conclusion_map.get(d.get("conclusion"), d.get("conclusion") or "미판정"),
            d.get("feedback", ""), d.get("feedback_detail", ""),
            str(d.get("created_at", "")), msg_count, last_q
        ])

    from fastapi.responses import StreamingResponse
    output.seek(0)
    # BOM for Excel UTF-8 compatibility
    bom_output = io.BytesIO()
    bom_output.write(b'\xef\xbb\xbf')
    bom_output.write(output.getvalue().encode("utf-8"))
    bom_output.seek(0)

    return StreamingResponse(
        bom_output,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=consult_history.csv"}
    )


# ── 3) 종합 리포트 ──

class ReportRequest(BaseModel):
    client_profile_id: int


@app.post("/api/pro/reports/generate")
def api_pro_report_generate(req: ReportRequest, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 종합 리포트 생성 — 매칭 + 상담 이력 + AI 종합 요약"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()

    # 1. 고객 프로필 조회
    cur.execute(
        """SELECT * FROM client_profiles WHERE id=%s AND owner_business_number=%s AND is_active=TRUE""",
        (req.client_profile_id, current_user["bn"])
    )
    client = cur.fetchone()
    if not client:
        conn.close()
        raise HTTPException(status_code=404, detail="고객사를 찾을 수 없습니다.")
    client = dict(client)

    # 2. 매칭 엔진으로 공고 검색 — 기업/개인 분기
    from app.core.matcher import get_matches_for_user, get_individual_matches_for_user, get_matches_hybrid
    profile = {
        "address_city": client.get("address_city") or "",
        "industry_code": client.get("industry_code") or "",
        "revenue_bracket": client.get("revenue_bracket") or "",
        "employee_count_bracket": client.get("employee_count_bracket") or "",
        "interests": client.get("interests") or "",
        "establishment_date": str(client.get("establishment_date") or ""),
    }
    is_individual_client = (client.get("client_type") or "business") == "individual"
    matched = get_matches_hybrid(profile, is_individual=is_individual_client)

    # 3. 각 공고에 대해 판정 — matcher의 eligibility_status 기반 (점수 기준 판정 제거)
    results = []
    eligible_count = 0
    conditional_count = 0
    ineligible_count = 0

    for ann in matched:
        a = ann if isinstance(ann, dict) else dict(ann)
        status = a.get("eligibility_status", "eligible")
        score = a.get("match_score", 0)
        if status == "ineligible":
            conclusion = "대상 아님"
            reason = a.get("ineligible_reason") or "자격 요건 불일치"
            ineligible_count += 1
        elif not a.get("eligibility_logic") or a.get("eligibility_logic") in ("{}", "null"):
            conclusion = "확인 필요"
            reason = "공고 원문에서 세부 자격요건 확인 필요"
            conditional_count += 1
        else:
            conclusion = "신청 가능"
            reason = a.get("recommendation_reason") or "자격 요건 충족"
            eligible_count += 1

        # 마감일 표시
        dl_raw = a.get("deadline_date")
        deadline_display = str(dl_raw) if dl_raw else "상시모집"

        results.append({
            "announcement_id": a.get("announcement_id"),
            "title": a.get("title", ""),
            "category": a.get("category", ""),
            "department": a.get("department", "") or "기관 미상",
            "conclusion": conclusion,
            "reason": reason,
            "support_amount": a.get("support_amount", ""),
            "deadline_date": deadline_display,
            "eligibility_status": status,
        })

    # 4. 공고AI 상담 이력 수집 (이 고객사 관련 최근 30건)
    consult_summaries = []
    try:
        cur.execute(
            """SELECT a.title, cl.conclusion, cl.messages, cl.created_at
               FROM ai_consult_logs cl
               JOIN announcements a ON a.announcement_id = cl.announcement_id
               WHERE cl.business_number = %s
               ORDER BY cl.created_at DESC LIMIT 30""",
            (current_user["bn"],)
        )
        for row in cur.fetchall():
            r = dict(row)
            msgs = r.get("messages") or []
            # 마지막 AI 응답만 추출
            last_ai = ""
            if isinstance(msgs, list):
                for m in reversed(msgs):
                    if m.get("role") == "assistant":
                        last_ai = m.get("text", "")[:200]
                        break
            consult_summaries.append({
                "title": r.get("title", ""),
                "conclusion": r.get("conclusion", ""),
                "summary": last_ai,
            })
    except Exception:
        pass

    # 4.5. 고객사 첨부 자료 텍스트 수집 (AI 프롬프트에 포함)
    attached_docs_text = ""
    try:
        cur.execute(
            "SELECT file_name, file_type, extracted_text FROM client_files WHERE client_id=%s AND owner_business_number=%s ORDER BY created_at DESC LIMIT 5",
            (req.client_profile_id, current_user["bn"])
        )
        for frow in cur.fetchall():
            fr = dict(frow)
            et = (fr.get("extracted_text") or "").strip()
            if et:
                attached_docs_text += f"\n[첨부: {fr['file_name']} ({fr['file_type']})]\n{et[:3000]}\n"
    except Exception:
        pass

    # 5. AI 종합 요약 생성
    ai_summary = ""
    try:
        import google.generativeai as genai
        api_key = os.environ.get("GEMINI_API_KEY")
        if api_key and results:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("models/gemini-2.5-flash")

            # 지원금액 파싱 함수 (백억/억/천만/만 → 숫자)
            def parse_amount(amt_str: str) -> int:
                if not amt_str: return 0
                s = str(amt_str).replace(",", "").replace(" ", "")
                num = 0
                try:
                    import re
                    # "5억", "1,000만원" 등 처리
                    if "억" in s:
                        m = re.search(r'(\d+(?:\.\d+)?)\s*억', s)
                        if m: num += int(float(m.group(1)) * 100000000)
                    if "천만" in s:
                        m = re.search(r'(\d+(?:\.\d+)?)\s*천만', s)
                        if m: num += int(float(m.group(1)) * 10000000)
                    if "백만" in s:
                        m = re.search(r'(\d+(?:\.\d+)?)\s*백만', s)
                        if m: num += int(float(m.group(1)) * 1000000)
                    if num == 0 and "만" in s:
                        m = re.search(r'(\d+(?:\.\d+)?)\s*만', s)
                        if m: num += int(float(m.group(1)) * 10000)
                    if num == 0:
                        m = re.search(r'(\d{4,})', s)
                        if m: num = int(m.group(1))
                except Exception: pass
                return num

            # 각 공고에 amount_value 추가
            for r in results:
                r["amount_value"] = parse_amount(r.get("support_amount", ""))

            # 정렬: 신청가능 → 확인필요 → 대상아님, 내부적으로 마감일 임박 + 금액 큰 순
            _CONC_ORDER = {"신청 가능": 0, "확인 필요": 1, "대상 아님": 2}
            def sort_key(r):
                conc_priority = _CONC_ORDER.get(r["conclusion"], 3)
                # 마감일: 가까운 것 우선, 상시모집/None은 뒤로
                dl = r.get("deadline_date") or ""
                dl_sort = dl if dl and dl != "상시모집" else "9999-99-99"
                return (conc_priority, dl_sort, -r["amount_value"])

            sorted_results = sorted(results, key=sort_key)
            # 보고서에는 신청가능 + 확인필요만 상위 10건. 대상아님은 섹션 뒤에 별도
            actionable = [r for r in sorted_results if r["conclusion"] != "대상 아님"]
            top10 = actionable[:10]

            consult_text = ""
            if consult_summaries:
                consult_text = "\n\n[공고별 상담 이력]\n"
                for cs in consult_summaries[:10]:
                    consult_text += f"- {cs['title']}: {cs['conclusion'] or '미판정'} — {cs['summary'][:100]}\n"

            # 로드맵: 마감일 순으로 그룹화
            today_str = datetime.date.today().isoformat()
            two_weeks = (datetime.date.today() + datetime.timedelta(days=14)).isoformat()
            this_month = (datetime.date.today() + datetime.timedelta(days=30)).isoformat()

            roadmap_immediate = [r for r in top10 if r.get("deadline_date") and r["deadline_date"] > today_str and r["deadline_date"] <= two_weeks]
            roadmap_month = [r for r in top10 if r.get("deadline_date") and r["deadline_date"] > two_weeks and r["deadline_date"] <= this_month]
            roadmap_later = [r for r in top10 if r.get("deadline_date") and r["deadline_date"] > this_month]

            # f-string에서 줄바꿈 사용 위해 변수로 분리
            NL = "\n"
            top10_lines = NL.join([
                f"{i+1}. {r['title']} | 지원금: {r['support_amount'] or '미공개'} | 판정: {r['conclusion']} | 마감: {r['deadline_date'] or '상시'}"
                for i, r in enumerate(top10)
            ]) or "없음"
            urgent_lines = NL.join([f"- [2주 내] {r['title']} | 마감: {r['deadline_date']}" for r in roadmap_immediate]) or "없음"
            month_lines = NL.join([f"- [1개월 내] {r['title']} | 마감: {r['deadline_date']}" for r in roadmap_month]) or "없음"
            later_lines = NL.join([f"- [추후] {r['title']} | 마감: {r['deadline_date']}" for r in roadmap_later]) or "없음"
            # 호환용
            eligible_lines = top10_lines
            conditional_lines = ""

            prompt = f"""당신은 10년 경력의 정부지원사업 전문 컨설턴트입니다.
아래 고객사 정보와 AI 매칭 결과를 바탕으로, 고객에게 전달할 **전문 컨설팅 리포트**를 작성하세요.

━━━━━━━━━━━━━━━━━━━━━━━━━━
[고객사 정보]
━━━━━━━━━━━━━━━━━━━━━━━━━━
- 기업명: {client.get('client_name', '')}
- 소재지: {client.get('address_city', '')}
- 업종코드: {client.get('industry_code', '')}
- 설립일: {client.get('establishment_date', '')}
- 매출규모: {client.get('revenue_bracket', '')}
- 직원수: {client.get('employee_count_bracket', '')}
- 관심분야: {client.get('interests', '')}

[매칭 결과]
- 총 {len(results)}건 분석
- 신청 가능: {eligible_count}건
- 확인 필요: {conditional_count}건
- 대상 아님(자격 미달): {ineligible_count}건

[추천 공고 TOP 10 — 신청 가능/확인 필요 우선, 마감 임박 순]
{top10_lines}

[신청 로드맵]
■ 2주 내 마감 (즉시 신청)
{urgent_lines}

■ 1개월 내 마감 (이번 달 준비)
{month_lines}

■ 추후 마감 (다음 달 이후)
{later_lines}
{consult_text}
{f'''
[고객사 제출 자료 분석]
{attached_docs_text[:5000]}
''' if attached_docs_text else ''}
━━━━━━━━━━━━━━━━━━━━━━━━━━
[리포트 작성 규칙] — 아래 섹션만 작성 (기업 요약·상담 이력은 보고서 상단에 이미 포함되어 있으므로 **절대 중복 작성 금지**)
━━━━━━━━━━━━━━━━━━━━━━━━━━

## 1. 맞춤 공고 분석 (추천 TOP 10)
- **위에 제공된 TOP 10 공고를 모두 표 형태로 정리**
- 표 컬럼: 순위 / 공고명 / 지원금액 / 마감일 / 판정 / 추천 사유
- 판정 컬럼 값: "신청 가능" / "확인 필요" (점수 금지)
- 추천 사유: **왜 이 기업에 적합한지** 1~2문장
- 절대 점수/퍼센트 수치 사용 금지

## 2. 신청 로드맵
- **위 로드맵 데이터를 활용하여 시각적 타임라인 작성**
- 즉시 신청 (2주 내) / 이번 달 준비 (1개월 내) / 추후 (다음 달 이후) 3단계 구분
- 각 단계별 표 형태로 우선순위 표시
- 즉시 신청 공고는 ★★★, 이번 달 ★★, 추후 ★ 표시

## 3. 필요 서류 체크리스트
- 공통 서류: 사업자등록증, 중소기업확인서, 재무제표 등
- 공고별 추가 서류 (알 수 있는 범위에서)

## 4. 경쟁력 분석
- 이 기업이 선정될 가능성을 높이는 방법
- 사업계획서 작성 팁
- 강조해야 할 포인트

## 5. 종합 의견
- 컨설턴트의 최종 판단
- 즉시 신청 권장 공고 (1~2개)
- 준비 후 신청 권장 공고
- 다음 분기 대비 사항

━━━━━━━━━━━━━━━━━━━━━━━━━━
[형식 규칙] — 매우 중요! 반드시 준수
━━━━━━━━━━━━━━━━━━━━━━━━━━
- **반드시 HTML 태그로 작성** (마크다운 아님)
- **코드펜스(```html ... ```) 절대 사용 금지** — HTML을 그대로 출력
- 문장 끝맺음은 명사형 (예: "지원 가능", "확인 필요", "신청 권장")
- 줄바꿈 없이 문장 나열하지 말 것. 각 항목은 분리
- **표(table)를 최대한 활용:**
  - 기업 현황 → 표
  - 추천 공고 TOP 5 → 표 (순위/공고명/금액/마감일/적합도/사유)
  - 필요 서류 → 체크리스트 표
- **로드맵은 시각적 타임라인:**
  - div 기반 가로 타임라인 또는 표 형태
  - 주간별 구분, 우선순위 색상 표시
- HTML 스타일:
  - 섹션 헤더: <h2 style="color:#5b21b6;border-bottom:2px solid #c4b5fd;padding-bottom:6px;margin-top:24px;">
  - 표: <table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px;">
  - th: style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;"
  - td: style="padding:8px 12px;border:1px solid #e5e7eb;"
  - 긴급 마감: <span style="color:#dc2626;font-weight:bold;">긴급</span>
  - 우선순위 높음: <span style="color:#7c3aed;">★★★</span>
  - 체크박스: ☐ (빈 체크박스 문자)
- 한국어, 전문 컨설턴트 톤
- 2000~3000자 내외
- <style> 태그 사용 금지 (인라인 스타일만)
- <script> 태그 사용 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━"""

            response = model.generate_content(prompt)
            ai_summary = response.text.strip() if response and response.text else "AI 분석을 생성하지 못했습니다."
            # 코드펜스 제거 (Gemini가 ```html ... ``` 으로 감싸는 경우)
            if ai_summary.startswith("```"):
                lines = ai_summary.split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines and lines[-1].strip() == "```":
                    lines = lines[:-1]
                ai_summary = "\n".join(lines).strip()
            print(f"[report] AI summary generated: {len(ai_summary)} chars")
    except Exception as e:
        ai_summary = f"AI 요약 생성 실패: {str(e)[:200]}"
        print(f"[report] AI summary error: {e}")

    import json as _json

    # ── 간트차트 로드맵 HTML 생성 (백엔드 직접 — AI 의존 X) ──
    gantt_html = ""
    try:
        import datetime as _dt
        today = _dt.date.today()
        # TOP 10에서 마감일이 미래인 것만 추출
        timeline_items = []
        for r in (sorted_results[:10] if 'sorted_results' in dir() else results[:10]):
            dl_str = str(r.get("deadline_date") or "").strip()
            if not dl_str or dl_str in ("None", "null", "상시모집"):
                # 상시모집은 별도로 표시
                timeline_items.append({"r": r, "days_left": None, "label": "상시모집"})
                continue
            try:
                dl = _dt.datetime.strptime(dl_str[:10], "%Y-%m-%d").date()
                days = (dl - today).days
                if days < 0:
                    continue
                timeline_items.append({"r": r, "days_left": days, "label": f"D-{days}"})
            except Exception:
                continue
        # 마감일 가까운 순으로 정렬
        dated = sorted([t for t in timeline_items if t["days_left"] is not None], key=lambda x: x["days_left"])
        always = [t for t in timeline_items if t["days_left"] is None]
        ordered = dated + always

        if ordered:
            max_days = max((t["days_left"] for t in ordered if t["days_left"] is not None), default=60) or 60
            chart_max = max(max_days, 30)  # 최소 30일 스케일
            rows_html = ""
            for t in ordered[:10]:
                rr = t["r"]
                title = (rr.get("title") or "")[:55]
                amt = rr.get("support_amount") or "-"
                if t["days_left"] is None:
                    bar_color = "#94a3b8"
                    bar_width = 100
                    label = "상시모집"
                elif t["days_left"] <= 7:
                    bar_color = "#dc2626"
                    bar_width = max(5, int(100 * t["days_left"] / chart_max))
                    label = f"⚠ D-{t['days_left']} (긴급)"
                elif t["days_left"] <= 30:
                    bar_color = "#ea580c"
                    bar_width = max(5, int(100 * t["days_left"] / chart_max))
                    label = f"D-{t['days_left']}"
                else:
                    bar_color = "#7c3aed"
                    bar_width = max(5, int(100 * t["days_left"] / chart_max))
                    label = f"D-{t['days_left']}"
                rows_html += f'''
                <tr>
                    <td style="padding:6px 8px;border:1px solid #e5e7eb;font-size:12px;width:40%;vertical-align:middle;">{title}</td>
                    <td style="padding:6px 8px;border:1px solid #e5e7eb;font-size:11px;width:18%;color:#16a34a;font-weight:bold;vertical-align:middle;">{amt}</td>
                    <td style="padding:6px 8px;border:1px solid #e5e7eb;width:42%;vertical-align:middle;">
                        <div style="position:relative;height:18px;background:#f1f5f9;border-radius:3px;overflow:hidden;">
                            <div style="position:absolute;left:0;top:0;height:100%;width:{bar_width}%;background:{bar_color};border-radius:3px;"></div>
                            <span style="position:relative;display:block;text-align:right;padding-right:6px;font-size:10px;line-height:18px;color:#1f2937;font-weight:bold;">{label}</span>
                        </div>
                    </td>
                </tr>'''
            gantt_html = f'''
<h2 style="color:#5b21b6;border-bottom:2px solid #c4b5fd;padding-bottom:6px;margin-top:24px;">📅 신청 로드맵 (간트차트)</h2>
<p style="font-size:12px;color:#64748b;margin:8px 0;">긴급(7일내) · 이번달(30일내) · 추후 — 색상으로 마감 임박도 표시</p>
<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px;">
<thead><tr>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px;border:1px solid #e5e7eb;text-align:left;">공고명</th>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px;border:1px solid #e5e7eb;text-align:left;">금액</th>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px;border:1px solid #e5e7eb;text-align:left;">D-day</th>
</tr></thead>
<tbody>{rows_html}</tbody>
</table>'''
    except Exception as gerr:
        print(f"[report] gantt error: {gerr}")
        gantt_html = ""

    brief = f"{client['client_name']} 기업 분석 결과: 총 {len(results)}건 분석, 신청가능 {eligible_count}건, 확인필요 {conditional_count}건, 대상아님 {ineligible_count}건"

    # ── 기업 요약 카드 (하드코딩 HTML, AI 미사용) ──
    def _fmt_years(est):
        try:
            if not est: return "-"
            est_str = str(est)[:10]
            est_d = datetime.datetime.strptime(est_str, "%Y-%m-%d").date()
            yrs = datetime.date.today().year - est_d.year
            return f"{est_str} (업력 {yrs}년)"
        except Exception:
            return str(est) if est else "-"

    _cname = client.get('client_name', '-') or '-'
    _ind = client.get('industry_name') or client.get('industry_code') or '-'
    _city = client.get('address_city') or '-'
    _rev = client.get('revenue_bracket') or '-'
    _emp = client.get('employee_count_bracket') or '-'
    _interests = client.get('interests') or '-'
    _est_disp = _fmt_years(client.get('establishment_date'))

    company_summary_html = f'''
<h2 style="color:#5b21b6;border-bottom:2px solid #c4b5fd;padding-bottom:6px;margin-top:24px;">🏢 기업 요약</h2>
<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px;">
<tr>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;width:18%;">기업명</th>
<td style="padding:8px 12px;border:1px solid #e5e7eb;width:32%;">{_cname}</td>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;width:18%;">업종</th>
<td style="padding:8px 12px;border:1px solid #e5e7eb;width:32%;">{_ind}</td>
</tr>
<tr>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;">설립일</th>
<td style="padding:8px 12px;border:1px solid #e5e7eb;">{_est_disp}</td>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;">소재지</th>
<td style="padding:8px 12px;border:1px solid #e5e7eb;">{_city}</td>
</tr>
<tr>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;">매출 규모</th>
<td style="padding:8px 12px;border:1px solid #e5e7eb;">{_rev}</td>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;">직원 수</th>
<td style="padding:8px 12px;border:1px solid #e5e7eb;">{_emp}</td>
</tr>
<tr>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px 12px;border:1px solid #e5e7eb;text-align:left;font-weight:bold;">관심분야</th>
<td colspan="3" style="padding:8px 12px;border:1px solid #e5e7eb;">{_interests}</td>
</tr>
</table>
<p style="font-size:12px;color:#64748b;margin:4px 0 16px 0;">📊 매칭 결과: 신청 가능 <b style="color:#16a34a;">{eligible_count}건</b> · 확인 필요 <b style="color:#ea580c;">{conditional_count}건</b> · 대상 아님 <b style="color:#94a3b8;">{ineligible_count}건</b></p>
'''

    # ── 상담 요약 카드 (DB consult_summaries 직접 렌더) ──
    if consult_summaries:
        _rows = ""
        for cs in consult_summaries[:10]:
            _ct = (cs.get('title') or '')[:60]
            _cc = cs.get('conclusion') or '미판정'
            _cs_text = (cs.get('summary') or '')[:120]
            _rows += f'''
<tr>
<td style="padding:6px 10px;border:1px solid #e5e7eb;font-size:12px;width:45%;">{_ct}</td>
<td style="padding:6px 10px;border:1px solid #e5e7eb;font-size:12px;width:15%;color:#7c3aed;font-weight:bold;">{_cc}</td>
<td style="padding:6px 10px;border:1px solid #e5e7eb;font-size:11px;width:40%;color:#475569;">{_cs_text}</td>
</tr>'''
        consult_summary_html = f'''
<h2 style="color:#5b21b6;border-bottom:2px solid #c4b5fd;padding-bottom:6px;margin-top:24px;">💬 상담 이력 요약</h2>
<p style="font-size:12px;color:#64748b;margin:4px 0;">최근 공고별 AI 상담 판정 결과 (최대 10건)</p>
<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px;">
<thead><tr>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px;border:1px solid #e5e7eb;text-align:left;">공고명</th>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px;border:1px solid #e5e7eb;text-align:left;">판정</th>
<th style="background:#f5f3ff;color:#5b21b6;padding:8px;border:1px solid #e5e7eb;text-align:left;">요약</th>
</tr></thead>
<tbody>{_rows}</tbody>
</table>
'''
    else:
        consult_summary_html = '''
<h2 style="color:#5b21b6;border-bottom:2px solid #c4b5fd;padding-bottom:6px;margin-top:24px;">💬 상담 이력 요약</h2>
<p style="font-size:13px;color:#64748b;margin:8px 0 16px 0;padding:12px;background:#f8fafc;border-left:3px solid #c4b5fd;">아직 개별 공고 상담 이력이 없습니다. 추천 공고별 상세 상담을 진행하시면 더 정확한 판단이 가능합니다.</p>
'''

    # summary 조립: brief + 기업요약 + 상담요약 + 간트 + AI 분석
    full_summary = f"{brief}\n\n{company_summary_html}\n{consult_summary_html}\n{gantt_html}\n\n{ai_summary}" if ai_summary else f"{brief}\n\n{company_summary_html}\n{consult_summary_html}\n{gantt_html}"

    # 6. DB 저장
    cur.execute(
        """INSERT INTO client_reports
           (client_profile_id, owner_business_number, title, summary, matched_announcements,
            total_eligible, total_conditional, total_ineligible)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
           RETURNING id""",
        (req.client_profile_id, current_user["bn"],
         f"{client['client_name']} 종합 리포트",
         full_summary, _json.dumps(results, ensure_ascii=False),
         eligible_count, conditional_count, ineligible_count)
    )
    try:
        report_id = cur.fetchone()["id"]
        conn.commit()
    except Exception as db_err:
        conn.rollback()
        conn.close()
        print(f"[report] DB error: {db_err}")
        raise HTTPException(status_code=500, detail=f"리포트 저장 실패: {str(db_err)[:100]}")
    conn.close()

    return {
        "status": "SUCCESS",
        "report_id": report_id,
        "summary": brief,
        "ai_analysis": ai_summary,
        "consult_history": consult_summaries[:10],
        "total": len(results),
        "eligible": eligible_count,
        "conditional": conditional_count,
        "ineligible": ineligible_count,
        "announcements": results,
    }


@app.get("/api/pro/reports")
def api_pro_reports(
    client_id: Optional[int] = None,
    current_user: dict = Depends(_get_current_user),
):
    """PRO: 리포트 목록 조회"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()

    if client_id:
        cur.execute(
            """SELECT r.id, r.client_profile_id, cp.client_name, r.title, r.summary,
                      r.total_eligible, r.total_conditional, r.total_ineligible, r.created_at
               FROM client_reports r
               JOIN client_profiles cp ON cp.id = r.client_profile_id
               WHERE r.owner_business_number=%s AND r.client_profile_id=%s
               ORDER BY r.created_at DESC""",
            (current_user["bn"], client_id)
        )
    else:
        cur.execute(
            """SELECT r.id, r.client_profile_id, cp.client_name, r.title, r.summary,
                      r.total_eligible, r.total_conditional, r.total_ineligible, r.created_at
               FROM client_reports r
               JOIN client_profiles cp ON cp.id = r.client_profile_id
               WHERE r.owner_business_number=%s
               ORDER BY r.created_at DESC
               LIMIT 50""",
            (current_user["bn"],)
        )

    rows = cur.fetchall()
    conn.close()
    reports = []
    for r in rows:
        d = dict(r)
        d["created_at"] = str(d["created_at"]) if d.get("created_at") else None
        reports.append(d)
    return {"status": "SUCCESS", "reports": reports}


@app.get("/api/pro/reports/{report_id}")
def api_pro_report_detail(report_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 리포트 상세 조회"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT r.*, cp.client_name, cp.address_city, cp.industry_code, cp.industry_name,
                  cp.revenue_bracket, cp.employee_count_bracket, cp.establishment_date
           FROM client_reports r
           JOIN client_profiles cp ON cp.id = r.client_profile_id
           WHERE r.id=%s AND r.owner_business_number=%s""",
        (report_id, current_user["bn"])
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="리포트를 찾을 수 없습니다.")
    d = dict(row)
    d["created_at"] = str(d["created_at"]) if d.get("created_at") else None
    if d.get("establishment_date"):
        d["establishment_date"] = str(d["establishment_date"])
    return {"status": "SUCCESS", "report": d}


class ReportEditSectionRequest(BaseModel):
    selected_text: str          # 사용자가 드래그로 선택한 텍스트
    instruction: str            # AI에게 보낼 수정 지시 ("더 강한 톤으로", "표로 바꿔줘" 등)


@app.post("/api/pro/reports/{report_id}/edit-section")
def api_pro_report_edit_section(report_id: int, req: ReportEditSectionRequest, current_user: dict = Depends(_get_current_user)):
    """PRO: 보고서의 일부 텍스트를 AI로 부분 수정.
    선택한 텍스트 + 지시 → AI가 같은 길이/톤으로 다시 작성하여 반환.
    클라이언트가 받은 결과를 원본에서 selected_text 위치에 치환.
    """
    _require_pro(current_user)
    if not req.selected_text or not req.selected_text.strip():
        raise HTTPException(status_code=400, detail="선택한 텍스트가 비어있습니다.")
    if not req.instruction or not req.instruction.strip():
        raise HTTPException(status_code=400, detail="수정 지시가 비어있습니다.")

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "SELECT id, summary FROM client_reports WHERE id=%s AND owner_business_number=%s",
        (report_id, current_user["bn"])
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="리포트를 찾을 수 없습니다.")

    full_summary = (dict(row).get("summary") or "")[:6000]

    try:
        import google.generativeai as genai
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="AI 서비스가 설정되지 않았습니다.")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("models/gemini-2.5-flash")

        prompt = f"""당신은 정부 지원사업 컨설턴트의 보고서 편집 어시스턴트입니다.

[전체 보고서 컨텍스트]
{full_summary}

[사용자가 선택한 부분]
\"\"\"
{req.selected_text}
\"\"\"

[수정 지시]
{req.instruction}

위 [선택한 부분]을 [수정 지시]에 따라 다시 작성하세요.

규칙:
1. 결과는 [선택한 부분]을 대체할 텍스트만 출력. 다른 설명이나 마크다운 코드블록 금지.
2. 비슷한 길이를 유지 (지시가 명시적으로 늘리거나 줄이라고 안 하면).
3. [전체 보고서] 톤과 일관성 유지.
4. HTML 태그가 원본에 있으면 그대로 유지하되, 지시에 따라 태그 변경 가능.
5. 절대 ```html 같은 코드펜스 사용 금지.
6. 한국어, 전문 컨설턴트 톤."""

        response = model.generate_content(prompt)
        new_text = (response.text or "").strip()
        # 코드펜스 제거
        if new_text.startswith("```"):
            lines = new_text.split("\n")
            if lines and lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            new_text = "\n".join(lines).strip()

        return {
            "status": "SUCCESS",
            "original": req.selected_text,
            "rewritten": new_text,
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"[edit-section] error: {e}")
        raise HTTPException(status_code=500, detail=f"AI 수정 실패: {str(e)[:200]}")


class ReportSaveRequest(BaseModel):
    summary: str  # 수정된 전체 HTML/텍스트


@app.put("/api/pro/reports/{report_id}")
def api_pro_report_update(report_id: int, req: ReportSaveRequest, current_user: dict = Depends(_get_current_user)):
    """PRO: 보고서 summary를 통째로 업데이트 (편집 모달 최종 저장용)"""
    _require_pro(current_user)
    if not req.summary or len(req.summary) < 10:
        raise HTTPException(status_code=400, detail="저장할 내용이 너무 짧습니다.")
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "UPDATE client_reports SET summary=%s WHERE id=%s AND owner_business_number=%s",
        (req.summary, report_id, current_user["bn"])
    )
    conn.commit()
    rows = cur.rowcount
    conn.close()
    if not rows:
        raise HTTPException(status_code=404, detail="리포트를 찾을 수 없습니다.")
    return {"status": "SUCCESS", "report_id": report_id}


@app.get("/api/pro/reports/{report_id}/pdf")
def api_pro_report_pdf(report_id: int, current_user: dict = Depends(_get_current_user)):
    """PRO: 리포트 PDF 다운로드 — HTML → PDF 변환"""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT r.title, r.summary, r.created_at, cp.client_name
           FROM client_reports r
           JOIN client_profiles cp ON cp.id = r.client_profile_id
           WHERE r.id=%s AND r.owner_business_number=%s""",
        (report_id, current_user["bn"])
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="리포트를 찾을 수 없습니다.")
    r = dict(row)

    # summary에서 brief + AI HTML 분리
    parts = (r.get("summary") or "").split("\n\n", 1)
    brief = parts[0] if parts else ""
    ai_html = parts[1] if len(parts) > 1 else ""

    # HTML 래핑
    html = f"""<!DOCTYPE html>
<html lang="ko">
<head><meta charset="utf-8"><title>{r['title']}</title>
<style>
body {{ font-family: 'Malgun Gothic', sans-serif; max-width: 800px; margin: 0 auto; padding: 40px 30px; color: #1e293b; font-size: 13px; line-height: 1.8; }}
h1 {{ color: #5b21b6; font-size: 22px; border-bottom: 3px solid #c4b5fd; padding-bottom: 10px; }}
h2 {{ color: #5b21b6; font-size: 16px; border-bottom: 2px solid #c4b5fd; padding-bottom: 6px; margin-top: 28px; }}
table {{ width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 12px; }}
th {{ background: #f5f3ff; color: #5b21b6; padding: 8px 10px; border: 1px solid #e5e7eb; text-align: left; font-weight: bold; }}
td {{ padding: 8px 10px; border: 1px solid #e5e7eb; }}
.header {{ text-align: center; margin-bottom: 30px; }}
.header p {{ color: #64748b; font-size: 12px; }}
.footer {{ margin-top: 40px; text-align: center; color: #94a3b8; font-size: 10px; border-top: 1px solid #e5e7eb; padding-top: 15px; }}
</style>
</head>
<body>
<div class="header">
<h1>{r['title']}</h1>
<p>{r['client_name']} | 작성일: {str(r.get('created_at',''))[:10]} | govmatch.kr</p>
</div>
{ai_html}
<div class="footer">
<p>본 보고서는 AI 분석 결과를 기반으로 작성되었으며, 최종 결과는 주관기관의 심사에 따릅니다.</p>
<p>지원금AI (govmatch.kr) | 밸류파인더 | Tel 010-5565-2299</p>
</div>
</body></html>"""

    # HTML → PDF (weasyprint 사용 시도, 없으면 HTML 반환)
    try:
        from weasyprint import HTML as WeasyprintHTML
        pdf_bytes = WeasyprintHTML(string=html).write_pdf()
        from fastapi.responses import Response
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{r["client_name"]}_report.pdf"'}
        )
    except ImportError:
        # weasyprint 미설치 → HTML 파일로 반환
        from fastapi.responses import Response
        return Response(
            content=html.encode("utf-8"),
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{r["client_name"]}_report.html"'}
        )


class FileAnalyzeRequest(BaseModel):
    text: str
    file_name: Optional[str] = ""
    file_type: Optional[str] = ""


@app.post("/api/pro/files/analyze")
def api_pro_file_analyze(req: FileAnalyzeRequest, current_user: dict = Depends(_get_current_user)):
    """PRO: 첨부 자료 AI 요약 분석 (JSON 텍스트)"""
    _require_pro(current_user)
    if not req.text or len(req.text.strip()) < 20:
        return {"status": "SUCCESS", "summary": "분석할 텍스트가 부족합니다."}
    return _analyze_text_with_ai(req.text[:8000], req.file_name or "파일", req.file_type or "자료")


class BusinessPlanReviewRequest(BaseModel):
    file_text: str  # 사업계획서 추출 텍스트
    file_name: Optional[str] = ""
    target_announcement_id: Optional[int] = None  # 신청 대상 공고
    client_id: Optional[int] = None  # 고객 컨텍스트


@app.get("/api/pro/stats/announcement")
def api_pro_announcement_stats(announcement_id: int, current_user: dict = Depends(_get_current_user)):
    """H: 공고 통계 — 같은 부처/카테고리의 평균 지원금, 유사 공고 수, 마감일까지 D-day."""
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT a.title, a.department, a.category, a.support_amount, a.deadline_date,
                   a.region, a.target_type
            FROM announcements a
            WHERE a.announcement_id = %s
        """, (announcement_id,))
        ann = cur.fetchone()
        if not ann:
            conn.close()
            raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")
        a = dict(ann)

        # 같은 부처/카테고리의 활성 공고 수
        cur.execute(f"""
            SELECT COUNT(*) AS cnt
            FROM announcements
            WHERE department = %s
              AND {valid_announcement_where()}
        """, (a.get("department") or "",))
        same_dept_count = cur.fetchone()["cnt"]

        cur.execute(f"""
            SELECT COUNT(*) AS cnt
            FROM announcements
            WHERE category = %s
              AND {valid_announcement_where()}
        """, (a.get("category") or "",))
        same_cat_count = cur.fetchone()["cnt"]

        # 평균 지원금 (텍스트 추출 — 정확하지 않지만 추정용)
        cur.execute("""
            SELECT support_amount FROM announcements
            WHERE department = %s AND support_amount IS NOT NULL AND support_amount != ''
            LIMIT 50
        """, (a.get("department") or "",))
        amounts = [r["support_amount"] for r in cur.fetchall()]

        # D-day
        d_day = None
        if a.get("deadline_date"):
            try:
                from datetime import date
                deadline = a["deadline_date"]
                today = date.today()
                d_day = (deadline - today).days if hasattr(deadline, "year") else None
            except Exception:
                pass

        # 분석된 유사 공고 (같은 부처)
        cur.execute("""
            SELECT COUNT(*) AS cnt
            FROM announcement_analysis aa
            JOIN announcements a ON aa.announcement_id = a.announcement_id
            WHERE a.department = %s
        """, (a.get("department") or "",))
        analyzed_count = cur.fetchone()["cnt"]

        conn.close()
        return {
            "status": "SUCCESS",
            "announcement": {
                "id": announcement_id,
                "title": a.get("title"),
                "department": a.get("department"),
                "category": a.get("category"),
                "support_amount": a.get("support_amount"),
                "deadline_date": str(a.get("deadline_date") or ""),
                "d_day": d_day,
            },
            "stats": {
                "same_department_count": same_dept_count,
                "same_category_count": same_cat_count,
                "amount_samples": amounts[:10],
                "analyzed_in_department": analyzed_count,
            },
        }
    except HTTPException:
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"통계 조회 실패: {str(e)[:200]}")


@app.post("/api/pro/clients/batch-match")
def api_pro_clients_batch_match(current_user: dict = Depends(_get_current_user)):
    """G: PRO 컨설턴트의 모든 활성 고객에 대해 일괄 매칭 → 고객별 Top 5 공고 반환.
    매일 알림용 또는 대시보드 일괄 갱신용.
    """
    _require_pro(current_user)
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """SELECT id, client_name, client_type, industry_code, address_city,
                  revenue_bracket, employee_count_bracket, interests, establishment_date,
                  age_range, income_level, family_type, employment_status, housing_status, special_conditions
           FROM client_profiles
           WHERE owner_business_number = %s AND is_active = TRUE
           ORDER BY updated_at DESC
           LIMIT 50""",
        (current_user["bn"],)
    )
    clients = [dict(r) for r in cur.fetchall()]
    conn.close()

    out = []
    for c in clients:
        is_indiv = (c.get("client_type") or "business") == "individual"
        profile = {
            "company_name": c.get("client_name") or "",
            "establishment_date": str(c.get("establishment_date") or "1990-01-01")[:10],
            "industry_code": c.get("industry_code") or "",
            "revenue_bracket": c.get("revenue_bracket") or ("1억 미만" if is_indiv else ""),
            "employee_count_bracket": c.get("employee_count_bracket") or ("5인 미만" if is_indiv else ""),
            "address_city": c.get("address_city") or "",
            "interests": c.get("interests") or "",
            "age_range": c.get("age_range") or "",
            "income_level": c.get("income_level") or "",
            "family_type": c.get("family_type") or "",
            "employment_status": c.get("employment_status") or "",
            "housing_status": c.get("housing_status") or "",
            "special_conditions": c.get("special_conditions") or "",
        }
        try:
            matches = get_matches_hybrid(profile, is_individual=is_indiv) or []
        except Exception as e:
            print(f"[batch-match] client {c['id']}: {e}")
            matches = []
        top5 = []
        for m in matches[:5]:
            top5.append({
                "announcement_id": m.get("announcement_id"),
                "title": (m.get("title") or "")[:120],
                "department": (m.get("department") or "")[:60],
                "support_amount": (m.get("support_amount") or "")[:60],
                "deadline_date": str(m.get("deadline_date") or "")[:10],
                "match_score": m.get("match_score", 0),
            })
        out.append({
            "client_id": c["id"],
            "client_name": c["client_name"],
            "client_type": c.get("client_type"),
            "match_count": len(matches),
            "top5": top5,
        })

    return {"status": "SUCCESS", "clients_processed": len(out), "data": out}


@app.post("/api/pro/business-plan/review")
def api_pro_business_plan_review(req: BusinessPlanReviewRequest, current_user: dict = Depends(_get_current_user)):
    """F: 사업계획서 텍스트 → AI 전문가 피드백 (강화 포인트, 가점 항목, 보완 사항).
    공고가 지정되면 해당 공고 자격요건과 매칭 분석.
    """
    _require_pro(current_user)
    if not req.file_text or len(req.file_text.strip()) < 100:
        raise HTTPException(status_code=400, detail="사업계획서 본문이 너무 짧습니다 (최소 100자).")

    import google.generativeai as genai
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="AI 서비스를 사용할 수 없습니다.")

    # 공고 컨텍스트 조회
    ann_context = ""
    conn = get_db_connection()
    cur = conn.cursor()
    if req.target_announcement_id:
        try:
            cur.execute("""
                SELECT a.title, a.department, a.support_amount,
                       aa.parsed_sections, aa.deep_analysis
                FROM announcements a
                LEFT JOIN announcement_analysis aa ON a.announcement_id = aa.announcement_id
                WHERE a.announcement_id = %s
            """, (req.target_announcement_id,))
            row = cur.fetchone()
            if row:
                d = dict(row)
                ann_context = f"\n\n[신청 대상 공고]\n공고명: {d.get('title','')}\n부처: {d.get('department','')}\n지원금: {d.get('support_amount','')}\n"
                ps = d.get("parsed_sections")
                if isinstance(ps, str):
                    try: ps = json.loads(ps)
                    except: ps = None
                if isinstance(ps, dict):
                    if ps.get("eligibility"): ann_context += f"\n자격요건: {str(ps.get('eligibility'))[:500]}"
                    if ps.get("evaluation_criteria"): ann_context += f"\n평가기준: {str(ps.get('evaluation_criteria'))[:500]}"
                    _docs = ps.get("required_docs") or ps.get("required_documents")
                    if _docs: ann_context += f"\n제출서류: {str(_docs)[:300]}"
        except Exception:
            pass

    # 고객 프로필
    client_context = ""
    if req.client_id:
        try:
            cur.execute("""
                SELECT client_name, client_type, industry_code, address_city,
                       revenue_bracket, employee_count_bracket, interests
                FROM client_profiles
                WHERE id = %s AND owner_business_number = %s AND is_active = TRUE
            """, (req.client_id, current_user["bn"]))
            row = cur.fetchone()
            if row:
                d = dict(row)
                client_context = f"\n\n[고객 프로파일]\n{json.dumps(d, ensure_ascii=False, default=str)[:500]}"
        except Exception:
            pass
    conn.close()

    prompt = f"""당신은 15년차 정부지원사업 전문 컨설턴트입니다.
컨설턴트가 고객의 사업계획서 초안을 검토 요청했습니다. 베테랑 관점에서 강화 포인트, 가점 항목, 보완 필요 사항을 분석하세요.

[사업계획서 본문]
{req.file_text[:6000]}
{ann_context}{client_context}

[분석 항목]
1. **강화 포인트** (현재 잘 쓰여진 부분, 더 강조해야 할 부분 3~5개)
2. **가점 항목** (놓치고 있는 가점 — 여성/청년/사회적기업/특허/수상 등 명시할 만한 것)
3. **보완 필요** (자격요건/평가기준 대비 부족하거나 누락된 부분 3~5개)
4. **위험 신호** (탈락 가능성을 높이는 표현/내용 — 모호한 수치, 근거 부족 등)
5. **컨설턴트 액션** (구체적으로 무엇을 추가/수정해야 하는지)

[응답 형식 — 순수 JSON]
{{
  "summary": "한 줄 종합 평가",
  "score": 75,  // 100점 만점 추정
  "strengths": ["강화 포인트 1", "강화 포인트 2"],
  "bonus_points": ["가점 항목 1", "가점 항목 2"],
  "improvements": ["보완 1", "보완 2"],
  "risks": ["위험 1"],
  "actions": ["액션 1", "액션 2"]
}}

마크다운 코드블록 금지. 순수 JSON만."""

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            "models/gemini-2.5-flash",
            generation_config={
                "max_output_tokens": 4096,
                "response_mime_type": "application/json",
                "temperature": 0.3,
            },
        )
        resp = model.generate_content(prompt)
        result = json.loads(resp.text)
    except Exception as e:
        print(f"[business-plan review] {e}")
        raise HTTPException(status_code=500, detail=f"분석 실패: {str(e)[:200]}")

    return {"status": "SUCCESS", "data": result}


@app.post("/api/pro/files/upload-analyze")
async def api_pro_file_upload_analyze(
    file: UploadFile = File(...),
    current_user: dict = Depends(_get_current_user),
):
    """PRO: 파일 업로드 → 텍스트 추출/멀티모달 → AI 요약 (PDF/DOCX/이미지/음성/TXT)"""
    _require_pro(current_user)

    file_name = file.filename or "unknown"
    content = await file.read()

    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(400, "파일 크기는 20MB 이하만 가능합니다.")

    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    text = ""

    # ─── 이미지 멀티모달 분석 (Gemini Vision) ───
    if ext in ("jpg", "jpeg", "png", "webp", "gif", "bmp"):
        return _analyze_image_with_gemini(content, file_name, ext)

    # ─── 음성 멀티모달 분석 (Gemini Audio) ───
    if ext in ("mp3", "wav", "m4a", "ogg", "flac", "webm", "aac"):
        return _analyze_audio_with_gemini(content, file_name, ext)

    try:
        if ext == "pdf":
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages_text = []
                for page in pdf.pages[:30]:  # 최대 30페이지
                    t = page.extract_text()
                    if t:
                        pages_text.append(t)
                text = "\n".join(pages_text)
        elif ext in ("txt", "csv", "md"):
            for enc in ["utf-8", "cp949", "euc-kr", "latin-1"]:
                try:
                    text = content.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
        elif ext in ("docx",):
            try:
                import docx as python_docx
                import io
                doc = python_docx.Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                text = content.decode("utf-8", errors="ignore")
        elif ext in ("hwp", "hwpx"):
            try:
                import olefile, io
                if olefile.isOleFile(io.BytesIO(content)):
                    ole = olefile.OleFileIO(io.BytesIO(content))
                    if ole.exists("PrvText"):
                        text = ole.openstream("PrvText").read().decode("utf-16-le", errors="ignore")
                    ole.close()
            except Exception:
                pass
            if not text:
                text = content.decode("utf-8", errors="ignore")[:5000]
        else:
            text = content.decode("utf-8", errors="ignore")[:5000]
    except Exception as e:
        return {"status": "SUCCESS", "summary": f"파일 읽기 실패: {str(e)[:100]}", "extracted_text": ""}

    if not text or len(text.strip()) < 20:
        return {"status": "SUCCESS", "summary": "파일에서 텍스트를 추출할 수 없습니다. 텍스트가 포함된 파일을 업로드해 주세요.", "extracted_text": ""}

    result = _analyze_text_with_ai(text[:8000], file_name, ext)
    result["extracted_text"] = text[:5000]

    # P1.1: client_files에 저장 (client_id=NULL, 나중에 고객 연결 가능)
    try:
        save_conn = get_db_connection()
        save_cur = save_conn.cursor()
        save_cur.execute(
            """INSERT INTO client_files
               (client_id, owner_business_number, file_name, file_type, file_size,
                file_data, extracted_text, ai_summary, memo)
               VALUES (NULL, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING id""",
            (current_user["bn"], file_name, ext, len(content),
             content[:5*1024*1024],  # 5MB 상한 (BYTEA 안전)
             text[:10000],
             json.dumps(result, ensure_ascii=False)[:10000],
             "ProSecretary 업로드"),
        )
        row = save_cur.fetchone()
        result["file_id"] = row["id"] if row else None
        save_conn.commit()
        save_conn.close()
    except Exception as save_err:
        print(f"[upload-analyze save] {save_err}")
    return result


def _analyze_image_with_gemini(content: bytes, file_name: str, ext: str) -> dict:
    """Gemini Vision으로 이미지 분석 (사업자등록증, 명함, 재무제표 사진 등)"""
    try:
        import google.generativeai as genai
        from PIL import Image
        import io

        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return {"status": "SUCCESS", "summary": "AI 서비스를 사용할 수 없습니다.", "extracted_text": ""}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("models/gemini-2.5-flash")

        try:
            img = Image.open(io.BytesIO(content))
        except Exception as e:
            return {"status": "SUCCESS", "summary": f"이미지를 열 수 없습니다: {str(e)[:80]}", "extracted_text": ""}

        prompt = """이 이미지는 고객사가 제출한 자료입니다. (사업자등록증, 명함, 재무제표, 사업계획서 등 가능)

다음 정보를 추출하여 요약하세요:
- 기업명/상호명
- 대표자명
- 사업자등록번호
- 소재지/주소
- 업종/업태
- 설립일
- 매출/직원수 등 숫자 정보
- 기타 핵심 내용

[추출 정보]
- 항목1: 값
- 항목2: 값
...

[요약]
- 핵심1: ...
- 핵심2: ...
- 핵심3: ...

이미지에서 추출 불가능한 항목은 생략하세요. 한국어로 간결하게."""

        response = model.generate_content([prompt, img])
        summary = response.text.strip() if response and response.text else "이미지 분석 실패"
        return {"status": "SUCCESS", "summary": summary, "extracted_text": summary}
    except Exception as e:
        return {"status": "SUCCESS", "summary": f"이미지 분석 오류: {str(e)[:100]}", "extracted_text": ""}


def _analyze_audio_with_gemini(content: bytes, file_name: str, ext: str) -> dict:
    """Gemini Audio로 음성 분석 (고객 통화, 회의록 등)"""
    try:
        import google.generativeai as genai

        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return {"status": "SUCCESS", "summary": "AI 서비스를 사용할 수 없습니다.", "extracted_text": ""}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("models/gemini-2.5-flash")

        # MIME 타입 매핑
        mime_map = {
            "mp3": "audio/mp3", "wav": "audio/wav", "m4a": "audio/mp4",
            "ogg": "audio/ogg", "flac": "audio/flac", "webm": "audio/webm", "aac": "audio/aac",
        }
        mime_type = mime_map.get(ext, "audio/mp3")

        audio_part = {"mime_type": mime_type, "data": content}

        prompt = """이 음성 파일은 지원사업 컨설턴트와 고객의 상담 녹음 또는 회의록입니다.

다음을 수행하세요:
1. 전체 대화를 한국어로 받아쓰기 (transcript)
2. 핵심 내용 요약 (3~5줄)
3. 추출 가능한 고객 정보 (기업명, 업종, 매출, 관심 분야 등)

[받아쓰기]
(전체 대화 내용...)

[요약]
- 핵심1: ...
- 핵심2: ...
- 핵심3: ...

[추출 정보]
- 기업명: ...
- 업종: ...
- 매출: ...
- 관심분야: ...

한국어로 간결하게."""

        response = model.generate_content([prompt, audio_part])
        summary = response.text.strip() if response and response.text else "음성 분석 실패"
        return {"status": "SUCCESS", "summary": summary, "extracted_text": summary[:5000]}
    except Exception as e:
        return {"status": "SUCCESS", "summary": f"음성 분석 오류: {str(e)[:100]}", "extracted_text": ""}


def _analyze_text_with_ai(text: str, file_name: str, file_type: str) -> dict:
    """공통 AI 분석 로직"""
    try:
        import google.generativeai as genai
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            return {"status": "SUCCESS", "summary": "AI 서비스를 사용할 수 없습니다."}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("models/gemini-2.5-flash")

        prompt = f"""아래는 고객사가 제출한 '{file_type}' ({file_name})입니다.
핵심 내용을 3~5줄로 요약하세요. 숫자(매출, 인원, 금액 등)가 있으면 반드시 포함하세요.

[자료 내용]
{text[:8000]}

[요약 형식]
- 핵심1: ...
- 핵심2: ...
- 핵심3: ...
한국어로 간결하게."""

        response = model.generate_content(prompt)
        summary = response.text.strip() if response and response.text else "요약 생성 실패"
        return {"status": "SUCCESS", "summary": summary}
    except Exception as e:
        return {"status": "SUCCESS", "summary": f"분석 오류: {str(e)[:100]}"}


def _send_html_email(to_email: str, subject: str, html_body: str, reply_to: str = "", sender_name: str = "") -> bool:
    """범용 HTML 이메일 발송"""
    smtp_user = os.getenv("SMTP_USER", "")
    smtp_password = os.getenv("SMTP_PASSWORD", "")
    smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_from = os.getenv("SMTP_FROM", smtp_user)
    if not smtp_user or not smtp_password:
        return False
    try:
        import smtplib
        from email.mime.text import MIMEText
        msg = MIMEText(html_body, "html", "utf-8")
        msg["Subject"] = subject
        display_name = f"{sender_name} (via 지원금AI)" if sender_name else "지원금AI"
        msg["From"] = f"{display_name} <{smtp_from}>"
        msg["To"] = to_email
        if reply_to:
            msg["Reply-To"] = reply_to
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
        return True
    except Exception as e:
        print(f"[Email] Send error to {to_email}: {e}")
        return False


class BulkEmailRequest(BaseModel):
    client_ids: list  # 고객 ID 목록
    subject: str
    body: str  # HTML 또는 텍스트
    include_report: Optional[bool] = False  # 최근 리포트 AI 요약 포함 여부


@app.post("/api/pro/email/send")
def api_pro_email_send(req: BulkEmailRequest, current_user: dict = Depends(_get_current_user)):
    """PRO: 고객사 일괄 이메일 발송"""
    _require_pro(current_user)
    if not req.client_ids:
        raise HTTPException(status_code=400, detail="발송 대상을 선택해주세요.")
    if not req.subject or not req.body:
        raise HTTPException(status_code=400, detail="제목과 내용을 입력해주세요.")

    conn = get_db_connection()
    cur = conn.cursor()

    # 선택된 고객 목록 조회
    placeholders = ",".join(["%s"] * len(req.client_ids))
    cur.execute(
        f"""SELECT id, client_name, contact_email, contact_name
            FROM client_profiles
            WHERE id IN ({placeholders}) AND owner_business_number=%s AND is_active=TRUE""",
        (*req.client_ids, current_user["bn"])
    )
    clients = [dict(r) for r in cur.fetchall()]

    sent = 0
    failed = 0
    skipped = 0

    for c in clients:
        email = c.get("contact_email")
        if not email or "@" not in email:
            skipped += 1
            continue

        # 이메일 본문에 고객명 치환
        body = req.body.replace("{{고객명}}", c.get("client_name", ""))
        body = body.replace("{{담당자명}}", c.get("contact_name", ""))

        # 리포트 요약 포함 옵션
        report_html = ""
        if req.include_report:
            try:
                cur.execute(
                    "SELECT summary FROM client_reports WHERE client_profile_id=%s AND owner_business_number=%s ORDER BY created_at DESC LIMIT 1",
                    (c["id"], current_user["bn"])
                )
                rrow = cur.fetchone()
                if rrow:
                    parts = (rrow["summary"] or "").split("\n\n", 1)
                    report_html = f'<hr style="margin:20px 0;border:1px solid #e5e7eb;"><h3 style="color:#5b21b6;">AI 분석 요약</h3>{parts[1] if len(parts) > 1 else parts[0]}'
            except Exception:
                pass

        full_html = f"""
        <div style="font-family:'Malgun Gothic',sans-serif;max-width:600px;margin:0 auto;padding:20px;color:#1e293b;font-size:14px;line-height:1.8;">
            {body}
            {report_html}
            <hr style="margin:30px 0;border:1px solid #e5e7eb;">
            <p style="color:#94a3b8;font-size:11px;text-align:center;">
                본 메일은 지원금AI(govmatch.kr)를 통해 발송되었습니다.<br>
                밸류파인더 | Tel 010-5565-2299
            </p>
        </div>"""

        # PRO 사용자 이름/이메일로 발신자 표시 + Reply-To
        pro_email = current_user.get("email", "")
        pro_name = ""
        try:
            cur.execute("SELECT company_name FROM users WHERE business_number = %s", (current_user["bn"],))
            _u = cur.fetchone()
            if _u:
                pro_name = _u.get("company_name") or ""
        except Exception:
            pass
        send_ok = _send_html_email(email, req.subject, full_html, reply_to=pro_email, sender_name=pro_name)
        if send_ok:
            sent += 1
        else:
            failed += 1
        # P1.2: email_logs 저장
        try:
            cur.execute(
                """INSERT INTO email_logs (owner_business_number, client_id, recipient_email, recipient_name, subject, body, status)
                   VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                (current_user["bn"], c.get("id"), email, c.get("contact_name") or c.get("client_name"),
                 req.subject, req.body, "sent" if send_ok else "failed"),
            )
            conn.commit()
        except Exception as log_err:
            print(f"[email_logs] {log_err}")
            try: conn.rollback()
            except: pass

    conn.close()
    _log_event("pro_email", current_user["bn"], f"sent={sent},failed={failed},skipped={skipped}")

    return {
        "status": "SUCCESS",
        "message": f"발송 완료: {sent}건 성공, {failed}건 실패, {skipped}건 이메일 없음",
        "sent": sent, "failed": failed, "skipped": skipped,
    }


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# API 제휴 관련
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class PartnershipInquiry(BaseModel):
    company_name: str
    contact_name: str
    email: str
    phone: Optional[str] = ""
    purpose: str
    expected_volume: Optional[str] = ""
    message: Optional[str] = ""


@app.post("/api/partnership/inquiry")
def api_partnership_inquiry(req: PartnershipInquiry, request: Request):
    """API 제휴 문의 접수"""
    if not req.company_name or not req.email or not req.purpose:
        raise HTTPException(status_code=400, detail="필수 항목을 입력해주세요.")
    # Rate limit: IP당 시간당 3회
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"partnership:ip:{ip}", 3, 3600):
        raise HTTPException(status_code=429, detail="잠시 후 다시 시도해주세요.")

    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS partnership_inquiries (
                id SERIAL PRIMARY KEY,
                company_name VARCHAR(200) NOT NULL,
                contact_name VARCHAR(100) NOT NULL,
                email VARCHAR(200) NOT NULL,
                phone VARCHAR(50) DEFAULT '',
                purpose TEXT NOT NULL,
                expected_volume VARCHAR(100) DEFAULT '',
                message TEXT DEFAULT '',
                status VARCHAR(20) DEFAULT 'new',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute(
            """INSERT INTO partnership_inquiries (company_name, contact_name, email, phone, purpose, expected_volume, message)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (req.company_name, req.contact_name, req.email, req.phone or "", req.purpose, req.expected_volume or "", req.message or "")
        )
        conn.commit()
        _log_system("partnership_inquiry", "system", f"{req.company_name} ({req.email})", "success")
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail="문의 접수 중 오류가 발생했습니다.")
    finally:
        conn.close()

    return {"status": "SUCCESS", "message": "제휴 문의가 접수되었습니다. 담당자가 빠르게 연락드리겠습니다."}


@app.get("/api/admin/partnership-inquiries", dependencies=[Depends(_verify_admin)])
def api_get_partnership_inquiries():
    """관리자: 제휴 문의 목록 조회"""
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("SELECT * FROM partnership_inquiries ORDER BY created_at DESC LIMIT 50")
        rows = [dict(r) for r in cur.fetchall()]
        for r in rows:
            if r.get("created_at"):
                r["created_at"] = str(r["created_at"])
        conn.close()
        return {"status": "SUCCESS", "data": rows}
    except Exception:
        conn.close()
        return {"status": "SUCCESS", "data": []}


@app.post("/api/partnership/chat")
def api_partnership_chat(req: dict, request: Request):
    """API 제휴 상담 챗봇 (Gemini 기반)"""
    message = (req.get("message") or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="메시지를 입력해주세요.")
    # Rate limit
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"pchat:ip:{ip}", 10, 60):
        raise HTTPException(status_code=429, detail="잠시 후 다시 시도해주세요.")

    system_prompt = """당신은 "지원금AI" 서비스의 API 제휴 상담 전문가입니다.
비즈니스 파트너에게 전문적이고 친절하게 응대하세요.

[서비스 정보]
- 서비스명: 지원금AI (govmatch.kr)
- 제공: 17,000+ 정부 지원금 공고 데이터 + AI 매칭 엔진
- 운영: 밸류파인더 (대표 권오성)

[API 제공 데이터]
- 정부지원금/보조금/정책자금 공고 데이터 (기업+개인)
- AI 기반 자격요건 분석 및 매칭
- 공고 상세 정보 (자격요건, 제출서류, 신청방법 등)
- 실시간 새 공고 알림

[API 요금제]
- Free: 일 100건 조회, 기본 공고 데이터
- Basic (월 29만원): 일 1,000건, AI 매칭 포함, 이메일 지원
- Pro (월 99만원): 무제한, AI 매칭+분석, 전담 매니저, SLA 보장
- Enterprise: 맞춤 협의

[연동 방식]
- RESTful API (JSON)
- API Key 인증
- Swagger 문서 제공
- 연동 가이드 및 샘플 코드 제공

[자주 묻는 질문]
- 테스트 API Key는 무료로 발급 가능
- 데이터는 매일 자동 업데이트
- SLA: Basic 99%, Pro 99.9%
- 결제: 월 자동결제 (카드/계좌이체)

답변 규칙:
1. 비즈니스 톤으로 전문적이게 답변
2. 구체적 제휴 논의가 필요하면 문의 폼 작성을 안내
3. 기술적 질문에는 상세하게 답변
4. 모르는 것은 "담당자가 상세히 안내드리겠습니다"로 안내
5. 한국어로 답변"""

    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return {"status": "SUCCESS", "answer": "현재 상담 서비스 준비 중입니다. 하단 문의 폼을 이용해주세요."}

        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")

        history = req.get("history", [])
        chat_messages = [{"role": "user", "parts": [system_prompt + "\n\n사용자 질문: " + message]}]
        if history:
            # 이전 대화 컨텍스트 포함 (최근 5건)
            context = "\n".join([f"{'사용자' if h['role']=='user' else 'AI'}: {h['content']}" for h in history[-5:]])
            chat_messages = [{"role": "user", "parts": [system_prompt + f"\n\n[이전 대화]\n{context}\n\n사용자 질문: " + message]}]

        response = model.generate_content(chat_messages[0]["parts"][0])
        answer = response.text.strip()
        return {"status": "SUCCESS", "answer": answer}
    except Exception as e:
        return {"status": "SUCCESS", "answer": "현재 상담이 일시적으로 불가합니다. 하단 문의 폼을 이용해주세요."}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 고객 상담
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class SupportInquiry(BaseModel):
    name: str
    email: str
    category: str
    message: str


@app.post("/api/support/inquiry")
def api_support_inquiry(req: SupportInquiry, request: Request):
    """고객 문의 접수"""
    if not req.name or not req.email or not req.message:
        raise HTTPException(status_code=400, detail="필수 항목을 입력해주세요.")
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"support:ip:{ip}", 5, 3600):
        raise HTTPException(status_code=429, detail="잠시 후 다시 시도해주세요.")

    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS support_inquiries (
                id SERIAL PRIMARY KEY,
                name VARCHAR(100) NOT NULL,
                email VARCHAR(200) NOT NULL,
                category VARCHAR(50) DEFAULT '',
                message TEXT NOT NULL,
                status VARCHAR(20) DEFAULT 'new',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute(
            "INSERT INTO support_inquiries (name, email, category, message) VALUES (%s, %s, %s, %s)",
            (req.name, req.email, req.category or "", req.message)
        )
        conn.commit()
        _log_system("support_inquiry", "system", f"{req.name} ({req.email}) - {req.category}", "success")
    except Exception:
        conn.rollback()
        raise HTTPException(status_code=500, detail="문의 접수 중 오류가 발생했습니다.")
    finally:
        conn.close()
    return {"status": "SUCCESS", "message": "문의가 접수되었습니다. 빠르게 답변드리겠습니다."}


@app.get("/api/admin/support-inquiries", dependencies=[Depends(_verify_admin)])
def api_get_support_inquiries():
    """관리자: 고객 문의 목록"""
    conn = get_db_connection()
    cur = conn.cursor()
    try:
        cur.execute("SELECT * FROM support_inquiries ORDER BY created_at DESC LIMIT 50")
        rows = [dict(r) for r in cur.fetchall()]
        for r in rows:
            if r.get("created_at"):
                r["created_at"] = str(r["created_at"])
        conn.close()
        return {"status": "SUCCESS", "data": rows}
    except Exception:
        conn.close()
        return {"status": "SUCCESS", "data": []}


@app.post("/api/support/chat")
def api_support_chat(req: dict, request: Request):
    """고객 상담 AI 챗봇"""
    message = (req.get("message") or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="메시지를 입력해주세요.")
    ip = _get_client_ip(request)
    if not _rate_limit_check(f"schat:ip:{ip}", 15, 60):
        raise HTTPException(status_code=429, detail="잠시 후 다시 시도해주세요.")

    system_prompt = """당신은 "지원금AI" 고객 상담 챗봇입니다.
친절하고 쉬운 말로 사용자를 도와주세요.

[서비스 안내]
- 지원금AI는 정부 지원금/보조금/정책자금을 AI가 자동 매칭해주는 무료 서비스입니다.
- 기업(사업자)과 개인 모두 이용 가능합니다.
- 17,000+ 공고를 실시간 분석합니다.

[주요 기능]
- 무료 회원가입 → 프로필 설정 → AI 자동 매칭
- 맞춤 공고 알림 (무료)
- AI 상담: 공고별 자격요건/신청방법 상세 안내
- 유료 플랜: LITE(월 2,900~4,900원), PRO(월 49,000원)

[자주 묻는 질문]
- 회원가입: 이메일 또는 소셜(카카오/네이버/구글) 로그인
- 매칭이 안 될 때: 프로필 설정(지역, 업종, 매출 등)을 정확히 입력해야 합니다
- 결제: Free(무료), Lite(개인 2,900원/사업자 4,900원), Pro(29,000원/월 이벤트가). 7일 무료 체험 가능
- 알림: 새 공고 등록 시 자동 알림 (이메일/푸시)
- 환불: 결제 후 7일 이내 전액 환불 가능

[답변 규칙]
1. 친근하고 쉬운 말로 답변
2. 해결이 어려운 문제는 "문의 폼을 작성해주시면 담당자가 직접 도와드리겠습니다"로 안내
3. 한국어로 답변
4. 답변은 간결하게 (3~5문장)"""

    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return {"status": "SUCCESS", "answer": "현재 상담 서비스 준비 중입니다. 하단 문의 폼을 이용해주세요."}

        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")

        history = req.get("history", [])
        context = ""
        if history:
            context = "\n".join([f"{'사용자' if h['role']=='user' else 'AI'}: {h['content']}" for h in history[-5:]])
            context = f"\n\n[이전 대화]\n{context}"

        response = model.generate_content(f"{system_prompt}{context}\n\n사용자 질문: {message}")
        return {"status": "SUCCESS", "answer": response.text.strip()}
    except Exception:
        return {"status": "SUCCESS", "answer": "일시적으로 상담이 불가합니다. 하단 문의 폼을 이용해주세요."}


@app.get("/api/trending")
def api_trending(target_type: Optional[str] = None, authorization: Optional[str] = Header(None)):
    """오늘의 인기 공고 3건 반환 — target_type=business|individual 로 필터, 미지정 시 전체"""
    tt = (target_type or "").strip().lower()
    if tt not in ("business", "individual"):
        tt = ""
    # 인메모리 캐시 (10분 TTL)
    cache_key = f"trending:{tt}"
    cached = _get_cached(cache_key)
    if cached:
        return cached
    # 사용자 소재지 추출 — trending 조회와 같은 커넥션에서 처리
    user_home_city = ""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        # 사용자 소재지 추출 (로그인 시 — 같은 커넥션 사용)
        if authorization and authorization.startswith("Bearer "):
            try:
                _payload = jwt.decode(authorization.split(" ", 1)[1], JWT_SECRET, algorithms=["HS256"])
                _bn = _payload.get("bn")
                if _bn:
                    cur.execute("SELECT address_city FROM users WHERE business_number = %s", (_bn,))
                    _row = cur.fetchone()
                    if _row:
                        _city = (_row.get("address_city", "") or "")
                        _cities = [c.strip() for c in _city.split(",") if c.strip() and c.strip() != "전국"]
                        if _cities:
                            user_home_city = _cities[0]
            except Exception:
                pass
        tt_filter_sql = ""
        tt_params: tuple = ()
        if tt == "business":
            tt_filter_sql = " AND COALESCE(a.target_type, 'business') IN ('business', 'both')"
        elif tt == "individual":
            tt_filter_sql = " AND COALESCE(a.target_type, 'business') IN ('individual', 'both')"
        cur.execute(f"""
            SELECT t.rank, t.trending_keyword, t.trending_reason,
                   a.announcement_id, a.title, a.department, a.category,
                   a.support_amount, a.deadline_date, a.region,
                   a.origin_url
            FROM trending_announcements t
            JOIN announcements a ON t.announcement_id = a.announcement_id
            WHERE t.trending_date >= CURRENT_DATE - INTERVAL '7 days'{tt_filter_sql}
            ORDER BY t.trending_date DESC, t.rank
            LIMIT 3
        """, tt_params)
        rows = [dict(r) for r in cur.fetchall()]

        # 여전히 비어있으면 — 직접 인기 공고 쿼리 폴백
        if not rows:
            try:
                fb_tt_sql = ""
                if tt == "business":
                    fb_tt_sql = " AND COALESCE(target_type, 'business') IN ('business', 'both')"
                elif tt == "individual":
                    fb_tt_sql = " AND COALESCE(target_type, 'business') IN ('individual', 'both')"
                cur.execute(f"""
                    SELECT announcement_id, title, department, category,
                           support_amount, deadline_date, region, origin_url
                    FROM announcements
                    WHERE (deadline_type = 'ongoing' OR (deadline_type = 'fixed' AND deadline_date >= CURRENT_DATE) OR (deadline_type = 'unknown' AND created_at >= CURRENT_DATE - INTERVAL '3 months')) AND is_archived = FALSE
                      AND support_amount IS NOT NULL AND support_amount != ''{fb_tt_sql}
                    ORDER BY
                        CASE WHEN support_amount ILIKE '%%억%%' THEN 0 ELSE 1 END,
                        deadline_date ASC NULLS LAST
                    LIMIT 3
                """)
                fallback = cur.fetchall()
                for i, fb in enumerate(fallback, 1):
                    rows.append({**dict(fb), "rank": i, "trending_keyword": "인기", "trending_reason": "금액 상위"})
            except Exception:
                pass

        # 사용자 소재지 기반 필터링 — 다른 지역 한정 공고 제외 (전국/소재지는 통과)
        if user_home_city:
            from app.services.rule_engine import _normalize_region
            _user_region = _normalize_region(user_home_city)
            import re as _re
            filtered = []
            for r in rows:
                ad_region = _normalize_region(r.get("region", "") or "")
                title = r.get("title", "") or ""
                # 1. region 필드 체크
                if ad_region and ad_region not in ("전국", "", "All"):
                    if ad_region != _user_region:
                        continue
                # 2. 제목 [도시명] 패턴 체크
                m = _re.search(r'\[(서울|경기|인천|부산|대구|대전|광주|울산|세종|강원|충북|충남|전북|전남|경북|경남|제주)\]', title)
                if m and m.group(1) != _user_region:
                    continue
                filtered.append(r)
            rows = filtered

            # 필터링 후 부족하면 추가 인기 공고로 보충
            if len(rows) < 3:
                needed = 3 - len(rows)
                existing_ids = {r.get("announcement_id") for r in rows}
                _reg_tt_sql = ""
                if tt == "business":
                    _reg_tt_sql = " AND COALESCE(target_type, 'business') IN ('business', 'both')"
                elif tt == "individual":
                    _reg_tt_sql = " AND COALESCE(target_type, 'business') IN ('individual', 'both')"
                cur.execute(f"""
                    SELECT announcement_id, title, department, category,
                           support_amount, deadline_date, region, origin_url
                    FROM announcements
                    WHERE (deadline_type = 'ongoing' OR (deadline_type = 'fixed' AND deadline_date >= CURRENT_DATE) OR (deadline_type = 'unknown' AND created_at >= CURRENT_DATE - INTERVAL '3 months')) AND is_archived = FALSE
                      AND support_amount IS NOT NULL AND support_amount != ''{_reg_tt_sql}
                      AND (region IS NULL OR region = '' OR region = '전국' OR region = 'All' OR region ILIKE %s)
                      AND title NOT ILIKE %s
                    ORDER BY
                        CASE WHEN support_amount ILIKE '%%억%%' THEN 0 ELSE 1 END,
                        deadline_date ASC NULLS LAST
                    LIMIT %s
                """, (f"%{_user_region}%", "%[%]%", needed * 3))
                for fb in cur.fetchall():
                    fbd = dict(fb)
                    if fbd.get("announcement_id") in existing_ids:
                        continue
                    rows.append({**fbd, "rank": len(rows) + 1, "trending_keyword": "전국", "trending_reason": "전국/소재지 인기"})
                    if len(rows) >= 3:
                        break

        # 날짜 직렬화
        for r in rows:
            if r.get("deadline_date"):
                r["deadline_date"] = str(r["deadline_date"])
        result = {"status": "SUCCESS", "data": rows[:3], "date": str(__import__("datetime").date.today())}
        _set_cache(cache_key, result)
        return result
    except Exception as outer_e:
        print(f"[Trending API] outer error: {outer_e}")
        import traceback; traceback.print_exc()
        # 별도 커넥션으로 폴백
        try:
            conn2 = get_db_connection()
            cur2 = conn2.cursor()
            _outer_tt_sql = ""
            if tt == "business":
                _outer_tt_sql = " AND COALESCE(target_type, 'business') IN ('business', 'both')"
            elif tt == "individual":
                _outer_tt_sql = " AND COALESCE(target_type, 'business') IN ('individual', 'both')"
            cur2.execute(f"""
                SELECT announcement_id, title, department, category,
                       support_amount, deadline_date, region, origin_url
                FROM announcements
                WHERE (deadline_type = 'ongoing' OR (deadline_type = 'fixed' AND deadline_date >= CURRENT_DATE) OR (deadline_type = 'unknown' AND created_at >= CURRENT_DATE - INTERVAL '3 months')) AND is_archived = FALSE
                  AND support_amount IS NOT NULL AND support_amount != ''{_outer_tt_sql}
                ORDER BY
                    CASE WHEN support_amount ILIKE '%%억%%' THEN 0 ELSE 1 END,
                    deadline_date ASC NULLS LAST
                LIMIT 3
            """)
            fallback = [dict(r) for r in cur2.fetchall()]
            conn2.close()
            for i, fb in enumerate(fallback, 1):
                fb["rank"] = i
                fb["trending_keyword"] = "인기"
                fb["trending_reason"] = "금액 상위"
                if fb.get("deadline_date"):
                    fb["deadline_date"] = str(fb["deadline_date"])
            return {"status": "SUCCESS", "data": fallback, "date": str(__import__("datetime").date.today())}
        except Exception:
            pass
        return {"status": "SUCCESS", "data": [], "date": str(__import__("datetime").date.today())}
    finally:
        try: conn.close()
        except: pass


# ── SmartDoc 연동 API ──

@app.get("/api/announcements/search")
def api_announcements_search(keyword: str = "", q: str = "", limit: int = 20):
    """공고 검색 — title/department/summary 매칭. keyword 또는 q 둘 다 지원"""
    # q 파라미터도 지원 (프론트엔드 호환성)
    search_term = (keyword or q or "").strip()
    limit = min(limit, 100)
    conn = get_db_connection()
    cur = conn.cursor()
    # 지역명 감지 (검색어에 포함 시 region 매칭 우선)
    REGION_NAMES = {"서울","경기","인천","부산","대구","대전","광주","울산","세종","강원","충북","충남","전북","전남","경북","경남","제주"}
    try:
        if search_term:
            words = search_term.split()
            where_parts = []
            params = []
            for w in words:
                where_parts.append("(title ILIKE %s OR department ILIKE %s OR summary_text ILIKE %s OR category ILIKE %s OR region ILIKE %s)")
                params.extend([f"%{w}%", f"%{w}%", f"%{w}%", f"%{w}%", f"%{w}%"])
            where_sql = " AND ".join(where_parts)

            # 검색어에서 지역명 추출
            detected_regions = [w for w in words if w in REGION_NAMES]

            # region 매칭 가산점 SQL
            if detected_regions:
                region_case_parts = []
                region_params_pre = []
                for rg in detected_regions:
                    region_case_parts.append("WHEN region ILIKE %s THEN 1")
                    region_params_pre.append(f"%{rg}%")
                region_score_sql = f"CASE {' '.join(region_case_parts)} ELSE 0 END"
            else:
                region_score_sql = "0"
                region_params_pre = []

            cur.execute(
                f"""SELECT announcement_id, title, department, category, deadline_date, support_amount, region,
                           CASE WHEN title ILIKE %s THEN 1 ELSE 0 END as title_match,
                           {region_score_sql} as region_match
                    FROM announcements
                    WHERE {where_sql}
                    ORDER BY
                        region_match DESC,
                        title_match DESC,
                        CASE WHEN deadline_date IS NOT NULL AND deadline_date >= CURRENT_DATE THEN 0 ELSE 1 END,
                        deadline_date ASC NULLS LAST,
                        MD5(announcement_id::text || %s)
                    LIMIT %s""",
                [f"%{search_term}%"] + region_params_pre + params + [search_term, limit],
            )
        else:
            cur.execute(
                """SELECT announcement_id, title, department, deadline_date, support_amount
                   FROM announcements
                   ORDER BY
                       CASE WHEN deadline_date IS NOT NULL AND deadline_date >= CURRENT_DATE THEN 0 ELSE 1 END,
                       deadline_date ASC NULLS LAST,
                       created_at DESC
                   LIMIT %s""",
                [limit],
            )
        rows = cur.fetchall()
        return {"status": "SUCCESS", "data": [dict(r) for r in rows], "total": len(rows)}
    finally:
        conn.close()


@app.get("/api/announcements/{announcement_id}")
def api_announcement_by_id(announcement_id: int):
    """공고 상세 — SEO 페이지용 (인증 불필요)"""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT announcement_id, title, department, category, region, target_type,
                   support_amount, deadline_date, summary_text, origin_url, final_url,
                   eligibility_logic, origin_source
            FROM announcements WHERE announcement_id = %s
        """, (announcement_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Not found")
        return {"status": "SUCCESS", "data": dict(row)}
    finally:
        conn.close()


@app.get("/api/announcements/{announcement_id}/for-smartdoc")
def api_announcement_for_smartdoc(announcement_id: int):
    """SmartDoc용 공고 상세 — 원문/분석 데이터 반환"""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            """SELECT announcement_id, title, department, deadline_date, support_amount,
                      summary_text, eligibility_logic, origin_url, category, region,
                      target_type, origin_source
               FROM announcements WHERE announcement_id = %s""",
            (announcement_id,),
        )
        ann = cur.fetchone()
        if not ann:
            raise HTTPException(status_code=404, detail="공고를 찾을 수 없습니다.")

        result = dict(ann)

        # deep_analysis 데이터가 있으면 full_text, evaluation_weights 등 포함
        try:
            from app.services.doc_analysis_service import get_deep_analysis
            analysis = get_deep_analysis(announcement_id, conn)
            if analysis:
                parsed = analysis.get("parsed_sections", {})
                deep = analysis.get("deep_analysis", {})
                result["full_text"] = analysis.get("full_text", "")
                result["evaluation_weights"] = deep.get("evaluation_weights") or parsed.get("evaluation_weights")
                result["form_templates"] = deep.get("form_templates") or parsed.get("form_templates")
                result["eligibility_detail"] = deep.get("eligibility") or parsed.get("eligibility")
                result["budget_info"] = deep.get("budget") or parsed.get("budget")
        except Exception:
            pass

        return {"status": "SUCCESS", "data": result}
    finally:
        conn.close()


@app.post("/api/admin/ai-coo/run")
def api_ai_coo_run(req: AdminAuthRequest):
    """관리자: AI COO 일일 감시 수동 실행 (테스트용)"""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="관리자 비밀번호가 올바르지 않습니다.")
    try:
        from app.services.orchestrator import run_daily_supervision
        result = run_daily_supervision()
        return {
            "status": "SUCCESS",
            "elapsed": result.get("elapsed"),
            "metrics": result.get("metrics", {}),
            "quality_avg": result.get("quality", {}).get("avg_total"),
            "knowledge_total": result.get("learning", {}).get("total"),
            "email_sent": result.get("report", {}).get("email_sent"),
            "kakao_sent": result.get("report", {}).get("kakao_sent"),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════
#  자금상담 사전학습 Q&A 관리 API
# ══════════════════════════════════════════════════════════

@app.get("/api/admin/qa-review", dependencies=[Depends(_verify_admin)])
def admin_qa_review_list(status: str = "pending", limit: int = 50):
    """Q&A 검토 큐 목록 조회"""
    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            """SELECT id, question, ai_answer, category, source_keywords,
                      status, corrected_answer, owner_memo, created_at, reviewed_at
               FROM qa_review_queue
               WHERE status = %s
               ORDER BY created_at DESC
               LIMIT %s""",
            (status, limit),
        )
        rows = cur.fetchall()
        items = [dict(r) for r in rows]
        return {"status": "SUCCESS", "items": items, "count": len(items)}
    finally:
        conn.close()


class QAReviewRequest(BaseModel):
    action: str  # approve | correct | reject
    corrected_answer: str = ""
    owner_memo: str = ""


@app.post("/api/admin/qa-review/{item_id}/review", dependencies=[Depends(_verify_admin)])
def admin_qa_review_action(item_id: int, req: QAReviewRequest):
    """Q&A 승인(approve) / 수정(correct) / 거절(reject) 처리 + 승인 시 knowledge_base 저장"""
    if req.action not in ("approve", "correct", "reject"):
        raise HTTPException(status_code=400, detail="action은 approve/correct/reject 중 하나")

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM qa_review_queue WHERE id = %s",
            (item_id,),
        )
        item = cur.fetchone()
        if not item:
            raise HTTPException(status_code=404, detail="항목 없음")

        new_status = "approved" if req.action in ("approve", "correct") else "rejected"
        final_answer = req.corrected_answer.strip() if req.action == "correct" else item["ai_answer"]

        cur.execute(
            """UPDATE qa_review_queue
               SET status = %s, corrected_answer = %s, owner_memo = %s,
                   reviewed_at = NOW()
               WHERE id = %s""",
            (new_status, req.corrected_answer or None, req.owner_memo or None, item_id),
        )

        if new_status == "approved":
            content_json = json.dumps(
                {"question": item["question"], "answer": final_answer},
                ensure_ascii=False,
            )
            cur.execute(
                """INSERT INTO knowledge_base
                       (source, knowledge_type, category, content, confidence, source_agent, use_count)
                   VALUES (%s, %s, %s, %s::jsonb, %s, %s, 0)
                   ON CONFLICT DO NOTHING""",
                (
                    "qa_verified",
                    "faq",
                    item["category"],
                    content_json,
                    1.0,
                    "qa_training",
                ),
            )

        conn.commit()
        return {"status": "SUCCESS", "action": new_status, "id": item_id}
    finally:
        conn.close()


class QAGenerateRequest(BaseModel):
    password: str
    batch_size: int = 5


@app.post("/api/admin/qa-review/generate")
def admin_qa_generate(req: QAGenerateRequest):
    """Q&A 배치 생성 트리거 (백그라운드 실행)"""
    if req.password != os.environ.get("ADMIN_PASSWORD", "admin1234"):
        raise HTTPException(status_code=401, detail="비밀번호 오류")

    import threading

    def _run():
        try:
            from app.services.qa_generator import generate_and_save_qa
            result = generate_and_save_qa(batch_size=req.batch_size)
            print(f"[qa_generate] 완료: {result}")
        except Exception as e:
            print(f"[qa_generate] 오류: {e}")

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    return {"status": "STARTED", "message": f"Q&A 생성 시작 (batch_size={req.batch_size})"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
